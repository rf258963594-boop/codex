from __future__ import annotations

import hashlib
import json
from collections import defaultdict
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(r"D:\RSIN GROUP Dropbox\RSIN GROUP TEAM\网站项目文件\Rbiz 智能自助线上业务操作平台-开发\签字文件")
OUT = Path(r"C:\Users\25896\Documents\Codex\2026-05-25\new-chat\outputs")

SIG_TERMS = [
    "signature",
    "signed",
    "sign",
    "director",
    "shareholder",
    "member",
    "secretary",
    "witness",
    "client",
    "beneficial owner",
    "nominee",
    "签名",
    "签署",
    "董事",
    "股东",
    "秘书",
    "客户",
    "见证",
]

FIELD_TERMS = [
    "Company Name",
    "公司名称",
    "UEN",
    "FYE",
    "Client Name",
    "Name of Company",
    "Director",
    "Shareholder",
    "Secretary",
    "Passport",
    "NRIC",
    "FIN",
    "Address",
    "Nationality",
    "Date",
    "Share",
    "Paid",
    "Capital",
    "SSIC",
    "Business Activity",
    "Registered Office",
    "Financial Year",
    "AGM",
    "RORC",
    "Controller",
]


def sha(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def docx_text(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " | ") for cell in row.cells]
            if any(cells):
                parts.append(" || ".join(cells))
    return "\n".join(parts)


def pdf_text(path: Path) -> tuple[str, int]:
    reader = PdfReader(str(path))
    texts: list[str] = []
    for index, page in enumerate(reader.pages):
        try:
            texts.append(page.extract_text() or "")
        except Exception as exc:  # pypdf can fail on odd legacy PDFs.
            texts.append(f"[EXTRACT_ERROR page {index + 1}: {exc}]")
    return "\n".join(texts), len(reader.pages)


def classify(path: Path) -> str:
    rel = str(path.relative_to(ROOT)).lower()
    name = path.name.lower()
    if "年审" in rel or "agm" in name:
        return "annual_review_agm"
    if "注册" in rel:
        return "incorporation"
    if "签名记录" in name:
        return "signature_audit_trail"
    if "清单" in name:
        return "file_checklist"
    return "unknown"


def file_role(path: Path) -> str:
    name = path.name.lower()
    if "first minutes" in name or "resolution" in name:
        return "first_directors_resolution"
    if "form 24" in name:
        return "allotment_or_share_statement_form24"
    if "form45b" in name:
        return "secretary_consent_form45b"
    if "form 45" in name:
        return "director_consent_form45"
    if "share certificate" in name:
        return "share_certificate"
    if "rorc" in name:
        return "rorc"
    if "secretary agreement" in name:
        return "secretary_agreement"
    if "nominee director" in name or "委任挂名董事" in name:
        return "nominee_director_agreement"
    if "client acceptance" in name:
        return "client_acceptance_form"
    if "背景调查" in name:
        return "background_check"
    if "agm" in name:
        return "agm_package"
    if "签名记录" in name:
        return "signature_record_attachment"
    if "清单" in name:
        return "signature_file_list"
    return "unknown"


def extract_records() -> tuple[list[dict], list[dict], list[dict]]:
    files = [path for path in ROOT.rglob("*") if path.is_file()]
    hash_groups: dict[str, list[Path]] = defaultdict(list)
    for path in files:
        hash_groups[sha(path)].append(path)

    duplicates = [
        {"hash": digest, "files": [str(path.relative_to(ROOT)) for path in paths]}
        for digest, paths in hash_groups.items()
        if len(paths) > 1
    ]

    records: list[dict] = []
    for path in files:
        ext = path.suffix.lower()
        text = ""
        pages = None
        status = "ok"
        if ext == ".docx":
            try:
                text = docx_text(path)
            except Exception as exc:
                status = f"docx_extract_error: {exc}"
        elif ext == ".pdf":
            try:
                text, pages = pdf_text(path)
            except Exception as exc:
                status = f"pdf_extract_error: {exc}"
        elif ext in {".doc", ".xls"}:
            status = "legacy_binary_not_extracted"
        else:
            status = "unsupported"

        lower = text.lower()
        field_hits = sorted({term for term in FIELD_TERMS if term.lower() in lower})
        sig_hits = sorted({term for term in SIG_TERMS if term.lower() in lower})
        signature_lines = []
        for line in text.splitlines():
            compact = " ".join(line.split())
            if compact and any(term.lower() in compact.lower() for term in SIG_TERMS):
                signature_lines.append(compact[:220])
            if len(signature_lines) >= 10:
                break
        first_lines = [" ".join(line.split())[:220] for line in text.splitlines() if line.strip()][:12]
        records.append(
            {
                "relative_path": str(path.relative_to(ROOT)),
                "name": path.name,
                "extension": ext,
                "size": path.stat().st_size,
                "hash": sha(path),
                "category": classify(path),
                "role": file_role(path),
                "status": status,
                "pages": pages,
                "text_length": len(text),
                "field_hits": field_hits,
                "signature_hits": sig_hits,
                "signature_lines": signature_lines,
                "first_lines": first_lines,
            }
        )

    role_groups: dict[tuple[str, str], list[dict]] = defaultdict(list)
    for record in records:
        role_groups[(record["category"], record["role"])].append(record)

    canonical: list[dict] = []
    for key, group in sorted(role_groups.items()):
        if key[1] == "unknown":
            canonical.extend(group)
            continue

        def rank(record: dict) -> tuple[int, int, str]:
            ext_rank = {".docx": 0, ".doc": 1, ".pdf": 2, ".xls": 3}.get(record["extension"], 9)
            no_remark_penalty = 1 if "没有备注" in record["relative_path"] else 0
            return ext_rank, no_remark_penalty, record["name"]

        canonical.append(sorted(group, key=rank)[0])
    return records, duplicates, canonical


def write_report(records: list[dict], duplicates: list[dict], canonical: list[dict]) -> tuple[Path, Path]:
    OUT.mkdir(parents=True, exist_ok=True)
    json_path = OUT / "签字文件模板审计数据.json"
    json_path.write_text(
        json.dumps(
            {
                "root": str(ROOT),
                "file_count": len(records),
                "duplicates": duplicates,
                "records": records,
                "canonical": canonical,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    priority_map = {
        "first_directors_resolution": ("P1", "注册基础董事决议", "多人同签一份：董事/挂名董事按规则签"),
        "director_consent_form45": ("P1", "董事同意/声明", "每个董事一份"),
        "secretary_consent_form45b": ("P1", "秘书同意/声明", "秘书签，通常每个秘书一份"),
        "share_certificate": ("P1", "股票证书", "每个股东一份，通常董事/秘书签发，不一定股东签"),
        "secretary_agreement": ("P1", "秘书服务协议", "客户/公司与秘书服务方签，需人工确认签署方"),
        "nominee_director_agreement": ("P1", "挂名董事协议", "客户/公司与挂名董事签，常用人员库强相关"),
        "client_acceptance_form": ("P2", "客户接受/KYC表", "客户签，可能每个客户/UBO一份"),
        "background_check": ("P2", "背景调查/KYC", "客户/UBO填写或签署"),
        "rorc": ("P2", "RORC 控制人登记", "控制人/公司记录，签字规则需确认"),
        "allotment_or_share_statement_form24": ("P2", "股份/配股声明 legacy Form 24", "建议改名为 Share Allotment/Statement；是否仍用需复核"),
        "agm_package": ("P2", "AGM/年审包", "董事/股东按 AGM 文件签"),
        "signature_record_attachment": ("P1", "签名记录附件", "每份签名文件后附带，不是业务文件"),
    }

    lines: list[str] = []
    lines.extend(
        [
            "# 签字文件模板初步审计报告",
            "",
            f"- 原始目录：`{ROOT}`",
            f"- 文件总数：{len(records)}",
            f"- 重复文件组：{len(duplicates)}（主要是“没有备注”PDF副本）",
            "- 说明：本报告只读取原文件，未修改原目录。老 `.doc` / `.xls` 是二进制旧格式，本轮主要通过同名 PDF 或 `.docx` 判断内容。",
            "",
            "## 文件包判断",
            "",
            "| 文件包 | 结论 |",
            "|---|---|",
            "| 公司注册 | 可作为第一版注册文件包接入，约 10 类文件。需要把 FORM 24/45/45B 等过时命名改成内部模板名称或标记为 legacy。 |",
            "| 年审 AGM | 可以作为第二个文件包接入，但目前只有 AGM 一套，字段和签字规则要单独建。 |",
            "| 变更 | 当前目录没有变更文件，不应从这批文件推断变更规则。 |",
            "",
            "## 建议的模板接入清单",
            "",
            "| 优先级 | 角色/文件 | 现有文件 | 建议接入方式 | 签字判断 |",
            "|---|---|---|---|---|",
        ]
    )

    for record in canonical:
        role = record["role"]
        if role in {"unknown", "signature_file_list"}:
            continue
        priority, label, signing = priority_map.get(role, ("P3", role, "需人工判断"))
        mode = "可直接做模板" if record["extension"] in {".docx", ".doc"} else "只有 PDF，需找可编辑模板或重建"
        lines.append(f"| {priority} | {label} | `{record['relative_path']}` | {mode} | {signing} |")

    lines.extend(
        [
            "",
            "## 关键逻辑问题",
            "",
            "1. 注册文件包里有 `FORM 24 / FORM 45 / FORM45B` 这类旧命名。系统里不建议直接用旧名做业务逻辑，应映射成内部标准名称，例如 `director_consent`、`secretary_consent`、`share_statement`。",
            "2. `公司注册-客户签名文件` 与 `公司注册-客户签名文件 - 没有备注` 大多是重复 PDF。系统只需要保留一套模板来源，另一个可作为输出预览参考。",
            "3. `每份签名文件后需附带的签名记录.docx` 应作为通用附件模板，不应混在注册/年审业务文件判断里。",
            "4. 年审目前只有 AGM 模板，不能代表 Annual Return 全流程。建议作为 `AGM_PACKAGE` 单独接，不要和注册文件包混在一起。",
            "5. 老 `.doc` 文件后续最好转换成 `.docx` 或重新建 docx 模板，否则自动字段替换、段落循环和签字区布局会比较难维护。",
            "",
            "## 字段适配方向",
            "",
            "| 标准字段组 | 可能对应模板内容 | 来源表 |",
            "|---|---|---|",
            "| company.name | Company Name / Name of Company | 注册模板 Company.company_name |",
            "| company.uen | UEN / Registration No. | BizFile 或变更模板 Company_Current.uen |",
            "| company.registered_office_address | Registered Office / Address | Company.registered_office_address |",
            "| company.fye / first_fye | FYE / Financial Year End | Company.fye / first_fye |",
            "| officer.directors[] | Director / Form 45 / First Resolution signers | People 中 is_director=Yes |",
            "| officer.secretary[] | Secretary / Form45B / Secretary Agreement | People 中 is_secretary=Yes |",
            "| shareholder.list[] | Share Certificate / Form 24 / Register | Shareholders 表 |",
            "| beneficial_owners/controllers[] | RORC / Background Check | 需要后续在注册模板加 UBO/controller 输入 |",
            "| signature.audit | 每份签名文件后需附带的签名记录 | 通用附件模板 |",
            "",
            "## 抽取摘要",
            "",
        ]
    )

    for record in canonical:
        if record["role"] == "signature_file_list":
            continue
        lines.append(f"### {record['role']} — {record['name']}")
        lines.append(
            f"- 类型：{record['category']} / {record['extension']} / 状态：{record['status']} / 页数：{record['pages'] or ''}"
        )
        if record["field_hits"]:
            lines.append(f"- 识别到字段词：{', '.join(record['field_hits'][:18])}")
        if record["signature_lines"]:
            lines.append("- 签字相关片段：")
            for snippet in record["signature_lines"][:4]:
                lines.append(f"  - {snippet}")
        elif record["first_lines"]:
            lines.append(f"- 文本开头：{'; '.join(record['first_lines'][:3])}")
        lines.append("")

    md_path = OUT / "签字文件模板初步审计报告.md"
    md_path.write_text("\n".join(lines), encoding="utf-8")
    return md_path, json_path


def main() -> None:
    records, duplicates, canonical = extract_records()
    md_path, json_path = write_report(records, duplicates, canonical)
    print(md_path)
    print(json_path)
    print(f"files={len(records)} duplicates={len(duplicates)} canonical={len(canonical)}")


if __name__ == "__main__":
    main()

