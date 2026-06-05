from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DIR = ROOT / "app" / "doc_templates" / "p2_standard_v1"
OUTPUT_DIR = ROOT / "outputs" / "P2_standard_templates_v1"
TEMPLATE_NAME = "M01_combined_directors_resolution_standard.docx"
FIELD_MAP_NAME = "M01_field_map.md"
MANIFEST_NAME = "M01_manifest.json"

BASE_FONT = "Calibri"
EAST_ASIA_FONT = "Microsoft YaHei"
INK = RGBColor(0x11, 0x18, 0x27)
MUTED = RGBColor(0x63, 0x6B, 0x78)
ACCENT = RGBColor(0x1F, 0x4D, 0x78)
LIGHT_FILL = "F2F4F7"
TABLE_BORDER = "B8C1CC"


def ensure_dirs() -> None:
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for directory in (TEMPLATE_DIR, OUTPUT_DIR):
        for path in directory.glob("M01_*"):
            if path.is_file():
                path.unlink()


def style_font(style, size: float, *, bold: bool = False, color: RGBColor | None = None, after: float = 6, before: float = 0, line_spacing: float = 1.10) -> None:
    style.font.name = BASE_FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def set_run_font(run, size: float | None = None, *, bold: bool | None = None, italic: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = BASE_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    style_font(doc.styles["Normal"], 11, color=INK, after=6, line_spacing=1.10)
    style_font(doc.styles["Title"], 12.5, bold=True, color=INK, after=12, line_spacing=1.0)
    style_font(doc.styles["Heading 1"], 12, bold=True, color=INK, before=10, after=5, line_spacing=1.05)
    style_font(doc.styles["Heading 2"], 11, bold=True, color=ACCENT, before=8, after=4, line_spacing=1.05)
    style_font(doc.styles["Heading 3"], 10.5, bold=True, color=ACCENT, before=7, after=3, line_spacing=1.05)

    for style_name, size, color, italic in [
        ("Legal Clause", 10.8, INK, False),
        ("Template Marker", 8.2, MUTED, True),
        ("Signature Label", 10.5, INK, False),
        ("Small Note", 8.8, MUTED, False),
    ]:
        if style_name not in doc.styles:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = doc.styles[style_name]
        style_font(style, size, color=color, after=4, line_spacing=1.10)
        style.font.italic = italic
        style.paragraph_format.space_after = Pt(4)

    doc.core_properties.title = "M01 Combined Directors' Resolution"
    doc.core_properties.subject = "P2 ordinary company maintenance directors' resolution master template"
    doc.core_properties.author = "RSIN template rebuild"
    doc.core_properties.comments = "M01 P2 master template. Conditional markers are intended to be removed by the document generator."
    add_footer(doc)


def add_field(paragraph, code: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = code
    run._r.append(instr)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)


def add_footer(doc: Document) -> None:
    paragraph = doc.sections[0].footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Page ")
    set_run_font(run, 8.5, color=MUTED)
    add_field(paragraph, "PAGE")
    run = paragraph.add_run(" of ")
    set_run_font(run, 8.5, color=MUTED)
    add_field(paragraph, "NUMPAGES")


def add_para(
    doc: Document,
    text: str = "",
    *,
    style: str | None = None,
    align: WD_ALIGN_PARAGRAPH | None = None,
    size: float | None = None,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    before: float | None = None,
    after: float | None = None,
    keep_with_next: bool = False,
):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    if before is not None:
        paragraph.paragraph_format.space_before = Pt(before)
    if after is not None:
        paragraph.paragraph_format.space_after = Pt(after)
    if keep_with_next:
        paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return paragraph


def add_company_header(doc: Document) -> None:
    add_para(doc, "{{company.company_name}}", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12.2, after=1, keep_with_next=True)
    add_para(doc, "Registration No. {{company.uen}}", align=WD_ALIGN_PARAGRAPH.CENTER, size=10.2, after=1, keep_with_next=True)
    add_para(doc, "(Incorporated in the Republic of Singapore)", align=WD_ALIGN_PARAGRAPH.CENTER, size=10.2, after=1, keep_with_next=True)
    add_para(doc, "(the \"Company\")", align=WD_ALIGN_PARAGRAPH.CENTER, size=10.2, after=12, keep_with_next=True)
    add_para(
        doc,
        "DIRECTORS' RESOLUTIONS IN WRITING MADE PURSUANT TO THE COMPANY'S CONSTITUTION",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        size=12.4,
        after=14,
        keep_with_next=True,
    )


def marker(doc: Document, text: str) -> None:
    add_para(doc, text, style="Template Marker", after=2)


def clause_title(doc: Document, title: str) -> None:
    add_para(doc, title, style="Heading 1", keep_with_next=True)


def clause_para(doc: Document, text: str) -> None:
    add_para(doc, text, style="Legal Clause")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_table_borders(table, color: str = TABLE_BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for tag in ("w:top", "w:left", "w:bottom", "w:right", "w:insideH", "w:insideV"):
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top: int = 90, bottom: int = 90, start: int = 120, end: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for tag, value in (("w:top", top), ("w:bottom", bottom), ("w:start", start), ("w:end", end)):
        element = margins.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        row_pr = row._tr.get_or_add_trPr()
        if row_pr.find(qn("w:cantSplit")) is None:
            row_pr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    row_pr = row._tr.get_or_add_trPr()
    tbl_header = row_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        row_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 9.8, align: WD_ALIGN_PARAGRAPH | None = None, fill: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=INK)
    if fill:
        shade_cell(cell, fill)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_width(table, widths)
    set_table_borders(table)
    set_cell_margins(table)
    set_repeat_table_header(table.rows[0])
    for col, header in enumerate(headers):
        set_cell_text(table.cell(0, col), header, bold=True, size=9.4, align=WD_ALIGN_PARAGRAPH.CENTER, fill=LIGHT_FILL)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            set_cell_text(table.cell(r_idx, c_idx), value, size=9.4, align=WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else None)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_business_activity_table(doc: Document) -> None:
    add_table(
        doc,
        ["Item", "Updated ACRA record"],
        [
            ["Primary SSIC / activity", "{{m01.primary_ssic_new}}\n{{m01.primary_activity_new}}"],
            ["Secondary SSIC / activity", "{{m01.secondary_ssic_new}}\n{{m01.secondary_activity_new}}"],
        ],
        [2600, 6760],
    )


def add_officer_particulars_table(doc: Document) -> None:
    marker(doc, "[[REPEAT m01.officer_particular_updates[]]]")
    add_table(
        doc,
        ["Particular", "Existing record", "Updated record"],
        [
            [
                "{{officer_update.field_label}}",
                "{{officer_update.old_value}}",
                "{{officer_update.new_value}}",
            ],
        ],
        [2200, 3580, 3580],
    )


def add_people_table(doc: Document, loop_name: str, capacity_label: str) -> None:
    marker(doc, f"[[REPEAT m01.{loop_name}[]]]")
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [6760, 2600])
    set_table_borders(table)
    set_cell_margins(table)
    set_cell_text(table.cell(0, 0), "{{person.full_name}}", size=9.8)
    set_cell_text(table.cell(0, 1), "{{person.effective_date}}", size=9.8, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_share_transfer_table(doc: Document) -> None:
    marker(doc, "[[REPEAT m01.share_transfers[]]]")
    add_table(
        doc,
        ["Transferor", "Transferee", "Class", "No. of shares", "Certificate / remarks"],
        [
            [
                "{{transfer.transferor_name}}",
                "{{transfer.transferee_name}}",
                "{{transfer.share_class}}",
                "{{transfer.shares_transferred}}",
                "{{transfer.certificate_reference}}",
            ],
        ],
        [2040, 2040, 1320, 1540, 2420],
    )


def add_share_allotment_table(doc: Document) -> None:
    marker(doc, "[[REPEAT m01.share_allotments[]]]")
    add_table(
        doc,
        ["Allottee", "Class", "No. of shares", "Amount paid", "Allotment date"],
        [
            [
                "{{allotment.allottee_name}}",
                "{{allotment.share_class}}",
                "{{allotment.shares_allotted}}",
                "{{allotment.total_paid}}",
                "{{allotment.allotment_date}}",
            ],
        ],
        [2900, 1600, 1620, 1620, 1620],
    )


def add_signature_table(doc: Document) -> None:
    add_para(doc, "DIRECTOR(S)", style="Signature Label", bold=True, before=12, after=4, keep_with_next=True)
    marker(doc, "[[REPEAT m01.director_signers[]]]")
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [4680, 4680])
    set_table_borders(table, "FFFFFF")
    set_cell_margins(table, top=80, bottom=120, start=0, end=180)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    set_cell_text(left, "\n\n______________________________\n{{signer.full_name}}\nDirector", size=10.5)
    set_cell_text(right, "", size=10.5)


def build_docx() -> Document:
    doc = Document()
    configure_doc(doc)
    add_company_header(doc)
    clause_para(
        doc,
        "The undersigned, being the director(s) of the Company, hereby pass the following resolutions in writing pursuant to the Constitution of the Company.",
    )
    clause_para(doc, "RESOLVED: -")

    marker(doc, "[[IF m01.include_registered_office]]")
    clause_title(doc, "CHANGE OF REGISTERED OFFICE")
    clause_para(
        doc,
        "That, with effect from {{m01.registered_office_effective_date}}, the registered office of the Company be changed from {{m01.registered_office_old}} to {{m01.registered_office_new}}.",
    )
    marker(doc, "[[END IF]]")

    marker(doc, "[[IF m01.include_office_hours]]")
    clause_title(doc, "OPENING HOURS OF REGISTERED OFFICE")
    clause_para(
        doc,
        "That the opening hours of the registered office of the Company be recorded as {{m01.office_hours_new}} with effect from {{m01.office_hours_effective_date}}.",
    )
    marker(doc, "[[END IF]]")

    marker(doc, "[[IF m01.include_business_activity]]")
    clause_title(doc, "CHANGE OF BUSINESS ACTIVITIES")
    clause_para(
        doc,
        "That the business activities and/or SSIC code(s) of the Company recorded with the Accounting and Corporate Regulatory Authority be amended as follows with effect from {{m01.business_activity_effective_date}}:",
    )
    add_business_activity_table(doc)
    marker(doc, "[[END IF]]")

    marker(doc, "[[IF m01.include_fye]]")
    clause_title(doc, "CHANGE OF FINANCIAL YEAR END")
    clause_para(
        doc,
        "That the financial year end of the Company be changed from {{m01.fye_old}} to {{m01.fye_new}}, and that the next set of accounts shall be made up for the period from {{m01.next_accounts_period_start}} to {{m01.next_accounts_period_end}}.",
    )
    marker(doc, "[[END IF]]")

    marker(doc, "[[IF m01.include_officer_particulars]]")
    clause_title(doc, "UPDATE OF OFFICER'S PARTICULARS")
    clause_para(
        doc,
        "That the particulars of {{m01.officer_name}} in the records of the Company and/or ACRA be updated as follows:",
    )
    add_officer_particulars_table(doc)
    marker(doc, "[[END IF]]")

    marker(doc, "[[IF m01.include_director_appointments]]")
    clause_title(doc, "APPOINTMENT OF DIRECTOR")
    clause_para(doc, "That the following person(s) be and are hereby appointed as director(s) of the Company:")
    add_people_table(doc, "director_appointments", "Director")
    marker(doc, "[[END IF]]")

    marker(doc, "[[IF m01.include_director_resignations]]")
    clause_title(doc, "RESIGNATION OF DIRECTOR")
    clause_para(doc, "That the resignation of the following director(s) be and is hereby accepted:")
    add_people_table(doc, "director_resignations", "Director")
    marker(doc, "[[END IF]]")

    marker(doc, "[[IF m01.include_secretary_appointments]]")
    clause_title(doc, "APPOINTMENT OF COMPANY SECRETARY")
    clause_para(doc, "That the following person or corporate secretarial provider be and is hereby appointed as secretary of the Company:")
    add_people_table(doc, "secretary_appointments", "Company Secretary")
    marker(doc, "[[END IF]]")

    marker(doc, "[[IF m01.include_secretary_resignations]]")
    clause_title(doc, "RESIGNATION OF COMPANY SECRETARY")
    clause_para(doc, "That the resignation of the following secretary of the Company be and is hereby accepted:")
    add_people_table(doc, "secretary_resignations", "Company Secretary")
    marker(doc, "[[END IF]]")

    marker(doc, "[[IF m01.include_share_transfer_approval]]")
    clause_title(doc, "APPROVAL OF TRANSFER OF SHARES")
    clause_para(doc, "That the transfer of the following shares in the Company be and is hereby approved:")
    add_share_transfer_table(doc)
    clause_para(
        doc,
        "That the relevant share certificate(s) be cancelled and/or issued as required, and that the Register of Members be updated accordingly.",
    )
    marker(doc, "[[END IF]]")

    marker(doc, "[[IF m01.include_share_allotment_approval]]")
    clause_title(doc, "ALLOTMENT OF SHARES")
    clause_para(
        doc,
        "That, pursuant to the authority of the members where applicable, the allotment and issue of the following shares be and is hereby approved:",
    )
    add_share_allotment_table(doc)
    clause_para(
        doc,
        "That the required statutory return(s), including Form 24 where applicable, be lodged with ACRA and that the allotment be recorded in the Register of Members.",
    )
    marker(doc, "[[END IF]]")

    marker(doc, "[[IF m01.include_bizfile_authorization]]")
    clause_title(doc, "ACRA / BIZFILE LODGEMENT")
    clause_para(
        doc,
        "That any director of the Company and/or the company secretary and/or {{provider.name}} be authorised to prepare, sign, lodge, amend and submit all necessary documents, notices and electronic transactions with ACRA and any other relevant authority to give effect to the above resolutions.",
    )
    marker(doc, "[[END IF]]")

    add_para(doc, "Dated this {{signature.day_ordinal}} day of {{signature.month_year}}", style="Signature Label", before=10, after=10)
    add_signature_table(doc)
    return doc


def field_map_markdown() -> str:
    return """# M01 Combined Directors' Resolution - Field Map

Purpose: ordinary directors' written resolution for P2 company maintenance matters. The final generator should keep only triggered sections, repeat rows/blocks marked with `[[REPEAT ...[]]]`, and remove all `[[IF ...]]`, `[[END IF]]`, and `[[REPEAT ...]]` markers.

## Event To Clause Mapping

| Event / source | M01 section | Default signer | Notes |
|---|---|---|---|
| change_registered_office | CHANGE OF REGISTERED OFFICE | director_signers | Usually combined with ACRA/BizFile lodgement clause. |
| change_office_hours | OPENING HOURS OF REGISTERED OFFICE | director_signers | Normally default office hours; only include when explicitly changed. |
| change_business_activity | CHANGE OF BUSINESS ACTIVITIES | director_signers | Use DR wording; do not convert to EGM/Special Resolution unless a separate shareholder action is required. |
| change_fye | CHANGE OF FINANCIAL YEAR END | director_signers | Use old FYE, new FYE, and next accounts period fields. |
| update_director_particulars / update_secretary_particulars / update_shareholder_particulars | UPDATE OF OFFICER'S PARTICULARS | director_signers | Covers ID type/no., address, phone/email if recorded internally or for filing. |
| appoint_director / resign_director | APPOINTMENT / RESIGNATION OF DIRECTOR | director_signers | Removal by shareholders belongs to M02, not M01. |
| appoint_secretary / resign_secretary | APPOINTMENT / RESIGNATION OF COMPANY SECRETARY | director_signers | Optional resignation letter remains a separate output when selected. |
| share_transfer | APPROVAL OF TRANSFER OF SHARES | director_signers | M01 only approves transfer; M03 generates instrument, certificate and register checklist. No Form 24. |
| share_allotment | ALLOTMENT OF SHARES | director_signers | M02 should handle S161/member authority where required; M03 handles Form 24/share certificate. |
| bizfile_authorization | ACRA / BIZFILE LODGEMENT | director_signers | Usually included whenever any filing is needed. |

## Core Fields

| Field | Description |
|---|---|
| company.company_name | Company legal name. |
| company.uen | Company registration number / UEN. |
| signature.day / signature.month_year | Resolution date split for legal wording. |
| provider.name | Filing provider or secretary firm, default RSIN GROUP PTE. LTD. |
| m01.director_signers[] | One or more directors signing the DR. |

## Clause Fields

| Field group | Required fields |
|---|---|
| registered office | m01.registered_office_old, m01.registered_office_new, m01.registered_office_effective_date |
| office hours | m01.office_hours_new, m01.office_hours_effective_date |
| business activity | m01.primary_ssic_new, m01.primary_activity_new, m01.secondary_ssic_new, m01.secondary_activity_new |
| FYE | m01.fye_old, m01.fye_new, m01.next_accounts_period_start, m01.next_accounts_period_end |
| officer particulars | m01.officer_name, m01.officer_particular_updates[] |
| people changes | m01.director_appointments[], m01.director_resignations[], m01.secretary_appointments[], m01.secretary_resignations[] |
| share transfer | m01.share_transfers[] |
| share allotment | m01.share_allotments[] |

## Repeat Markers

| Marker | Meaning |
|---|---|
| [[REPEAT m01.officer_particular_updates[]]] | Repeat the following table row once per changed officer particular. |
| [[REPEAT m01.director_appointments[]]] | Repeat the following table row once per appointed director. |
| [[REPEAT m01.director_resignations[]]] | Repeat the following table row once per resigned director. |
| [[REPEAT m01.secretary_appointments[]]] | Repeat the following table row once per appointed secretary. |
| [[REPEAT m01.secretary_resignations[]]] | Repeat the following table row once per resigned secretary. |
| [[REPEAT m01.share_transfers[]]] | Repeat the following table row once per transfer item. |
| [[REPEAT m01.share_allotments[]]] | Repeat the following table row once per allotment item. |
| [[REPEAT m01.director_signers[]]] | Repeat the signature block once per director signer. |

## Exclusions

M01 should not handle company name change, director removal by shareholders, compulsory takeover authority, constitution change, strike-off shareholder consent, annual review AGM/AR, RORC/ROND/RONS registers, or standalone transfer instruments. Those belong to M02-M05 or registers packs.
"""


def write_manifest(docx_path: Path) -> None:
    manifest = {
        "template_id": "M01_combined_directors_resolution",
        "file_name": TEMPLATE_NAME,
        "display_name": "M01 Combined Directors' Resolution",
        "phase": "P2",
        "status": "template_built_pending_generator_integration",
        "preset": "standard_business_brief",
        "named_override": "formal secretarial legal template: centered company heading, compact legal clauses, conditional block markers, fixed-width tables and signature blocks.",
        "source_samples": [
            "DR - Change of Registered Office Address - SUNSEA",
            "DR - Change of Business Activity - TODRINK",
            "DR - Change of Officer Particular - TODRINK",
            "DR - Change of FYE - LIGHT CONE",
            "TRANSFER OF SHARES - TRI-X",
            "DR - Authority to Issue Shares - FQ AI",
        ],
        "output": str(docx_path),
        "field_map": FIELD_MAP_NAME,
    }
    (OUTPUT_DIR / MANIFEST_NAME).write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> None:
    ensure_dirs()
    doc = build_docx()
    output_path = OUTPUT_DIR / TEMPLATE_NAME
    doc.save(output_path)
    shutil.copy2(output_path, TEMPLATE_DIR / TEMPLATE_NAME)
    (OUTPUT_DIR / FIELD_MAP_NAME).write_text(field_map_markdown(), encoding="utf-8")
    write_manifest(output_path)
    print(output_path)
    print(TEMPLATE_DIR / TEMPLATE_NAME)
    print(OUTPUT_DIR / FIELD_MAP_NAME)


if __name__ == "__main__":
    main()
