from __future__ import annotations

import json
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_ROOT = Path(
    "D:/RSIN GROUP Dropbox/RSIN GROUP TEAM/"
    "\u7f51\u7ad9\u9879\u76ee\u6587\u4ef6/"
    "Rbiz \u667a\u80fd\u81ea\u52a9\u7ebf\u4e0a\u4e1a\u52a1\u64cd\u4f5c\u5e73\u53f0-\u5f00\u53d1/"
    "\u7b7e\u5b57\u6587\u4ef6"
)
SOURCE_REG = SOURCE_ROOT / "\u516c\u53f8\u6ce8\u518c-\u5ba2\u6237\u7b7e\u540d\u6587\u4ef6"

TEMPLATE_DIR = ROOT / "app" / "doc_templates" / "p1_preserved_v2"
OUTPUT_DIR = ROOT / "outputs" / "P1_preserved_templates_v2"
ZIP_PATH = ROOT / "outputs" / "P1_preserved_templates_v2.zip"

V1_DIRS = [
    ROOT / "app" / "doc_templates" / "p1_rebuilt",
    ROOT / "outputs" / "P1_rebuilt_templates",
]


MANIFEST = [
    {
        "template_id": "first_directors_resolution",
        "file_name": "01_first_directors_resolution_preserved.docx",
        "display_name": "First Directors Resolution",
        "source": "1. Company Name + First minutes resolution.docx",
        "build_method": "source_docx_fieldized",
        "repeat": "one_per_company",
    },
    {
        "template_id": "director_consent",
        "file_name": "02_director_consent_form45_preserved.docx",
        "display_name": "Consent to Act as Director / Form 45",
        "source": "3. Company Name + Form 45.doc/pdf + ACRA Form 45 reference",
        "build_method": "pdf_reconstructed_with_current_acra_wording",
        "repeat": "one_per_director",
    },
    {
        "template_id": "secretary_consent",
        "file_name": "03_secretary_consent_form45b_preserved.docx",
        "display_name": "Consent to Act as Secretary / Form 45B",
        "source": "4. Company Name + Form45B.doc/pdf + ACRA Form 45B reference",
        "build_method": "pdf_reconstructed_with_current_acra_wording",
        "repeat": "one_per_secretary",
    },
    {
        "template_id": "share_certificate",
        "file_name": "04_share_certificate_preserved.docx",
        "display_name": "Share Certificate",
        "source": "5. Company Name + Share Certificate-2.docx",
        "build_method": "source_docx_fieldized",
        "repeat": "one_per_shareholder_or_certificate",
    },
    {
        "template_id": "secretary_service_agreement",
        "file_name": "05_secretary_service_agreement_preserved.docx",
        "display_name": "Secretary Service Agreement",
        "source": "7.Secretary Agreement.docx",
        "build_method": "source_docx_fieldized_minor_text_cleanup",
        "repeat": "one_per_company",
    },
    {
        "template_id": "nominee_director_agreement",
        "file_name": "06_nominee_director_agreement_preserved.docx",
        "display_name": "Nominee Director Agreement",
        "source": "8 Agreement for Appointment of Nominee Director.doc/pdf",
        "build_method": "pdf_reconstructed_preserving_full_clause_structure",
        "repeat": "one_per_nominee_director_relationship",
        "manual_review": True,
    },
    {
        "template_id": "signature_record_attachment",
        "file_name": "07_signature_record_attachment_preserved.docx",
        "display_name": "Signature Record Attachment",
        "source": "\u6bcf\u4efd\u7b7e\u540d\u6587\u4ef6\u540e\u9700\u9644\u5e26\u7684\u7b7e\u540d\u8bb0\u5f55.docx",
        "build_method": "source_docx_fieldized_sample_data_removed",
        "repeat": "append_to_each_signed_document",
    },
]


FIELD_NOTES = {
    "company.company_name": "Company legal name.",
    "company.uen": "Company registration number/UEN; can be blank before incorporation.",
    "company.incorporation_date": "Date of incorporation or appointment effective date.",
    "company.registered_office_address": "Registered office address.",
    "company.register_location": "Place where register of members and index is kept; default can be registered office.",
    "company.first_directors_names": "Derived list of first director names for old resolution sentence.",
    "company.subscriber_share_lines": "Derived multi-line subscriber/share table text for old resolution layout.",
    "company.share_par_value": "Legacy wording value per share; default usually 1.",
    "company.fye": "Financial year end.",
    "company.first_financial_period_start": "First accounts period start.",
    "company.first_financial_period_end": "First accounts period end.",
    "director.*": "One director context for per-director documents.",
    "secretary.*": "One secretary context for per-secretary documents.",
    "shareholder.*": "One shareholder/share certificate context.",
    "provider.*": "Website setting for RSIN/service provider details.",
    "client_signatory.*": "Client authorised signatory context.",
    "nominee_director.*": "Nominee director context.",
    "signature.*": "Signature date fields, including day/month/year split for old legal wording.",
    "signature_record.*": "Post-signing e-signature audit data.",
}


def ensure_dirs() -> None:
    for directory in (TEMPLATE_DIR, OUTPUT_DIR):
        directory.mkdir(parents=True, exist_ok=True)
        for path in directory.iterdir():
            if path.is_file():
                path.unlink()


def mark_v1_as_simplified() -> None:
    note = (
        "# Simplified Draft Notice\n\n"
        "This folder was the first rebuilt draft. It is useful for field planning, but it intentionally simplified several agreement clauses. "
        "For production template work, use `P1_preserved_templates_v2` instead.\n"
    )
    for directory in V1_DIRS:
        if directory.exists():
            (directory / "README_SIMPLIFIED_DRAFT.md").write_text(note, encoding="utf-8")


def configure_business_doc(doc: Document, title: str, compact: bool = False) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72 if compact else 0.8)
    section.bottom_margin = Inches(0.72 if compact else 0.8)
    section.left_margin = Inches(0.78 if compact else 0.85)
    section.right_margin = Inches(0.78 if compact else 0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(9 if compact else 10)
    normal.paragraph_format.space_after = Pt(3 if compact else 5)
    normal.paragraph_format.line_spacing = 1.03 if compact else 1.08
    doc.core_properties.title = title
    doc.core_properties.author = "RSIN template rebuild v2"
    doc.core_properties.comments = "Preserved/fieldized P1 template v2. Original source files were not modified."


def set_font(run, size: float | int | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic


def set_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        run = paragraph.add_run()
        set_font(run)
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def add_para(
    doc: Document,
    text: str = "",
    *,
    align: WD_ALIGN_PARAGRAPH | None = None,
    size: float | int | None = None,
    bold: bool = False,
    italic: bool = False,
    before: float | int | None = None,
    after: float | int | None = None,
) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)


def add_title(doc: Document, text: str, size: float | int = 11) -> None:
    add_para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, size=size, bold=True, after=2)


def add_clause_heading(doc: Document, english: str, chinese: str | None = None) -> None:
    add_para(doc, english, bold=True, before=6, after=1)
    if chinese:
        add_para(doc, chinese, bold=True, after=2)


def add_signature_pair(doc: Document, left: list[str], right: list[str]) -> None:
    table = doc.add_table(rows=4, cols=2)
    table.autofit = False
    for cell in table.columns[0].cells:
        cell.width = Inches(3.0)
    for cell in table.columns[1].cells:
        cell.width = Inches(3.0)
    for idx, text in enumerate(["\n".join(left[:1]), "\n\n", "\n".join(left[1:]), ""]):
        table.cell(idx, 0).text = text
    for idx, text in enumerate(["\n".join(right[:1]), "\n\n", "\n".join(right[1:]), ""]):
        table.cell(idx, 1).text = text


def replace_in_document(doc: Document, replacements: list[tuple[str | re.Pattern, str]]) -> None:
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        original = paragraph.text
        updated = original
        for pattern, repl in replacements:
            if isinstance(pattern, str):
                updated = updated.replace(pattern, repl)
            else:
                updated = pattern.sub(repl, updated)
        if updated != original:
            set_paragraph_text(paragraph, updated)


def reveal_white_text(doc: Document) -> None:
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        for run in paragraph.runs:
            if run.font.color and run.font.color.rgb == RGBColor(0xFF, 0xFF, 0xFF):
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def remove_embedded_drawings(doc: Document) -> None:
    for drawing in doc.element.xpath(".//w:drawing"):
        drawing.getparent().remove(drawing)


def load_source_docx(path: Path, title: str) -> Document:
    doc = Document(path)
    configure_business_doc(doc, title)
    return doc


def save(doc: Document, file_name: str) -> None:
    doc.save(TEMPLATE_DIR / file_name)
    doc.save(OUTPUT_DIR / file_name)


def build_first_directors_resolution() -> None:
    doc = load_source_docx(SOURCE_REG / "1. Company Name + First minutes resolution.docx", "First Directors Resolution")
    p = doc.paragraphs
    if len(p) > 92:
        p[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_text(p[0], "{{company.company_name}}")
        set_paragraph_text(p[2], "Company Registration No.: {{company.uen}}")
        set_paragraph_text(p[17], "The Certificate numbered {{company.incorporation_certificate_no}} confirming Incorporation of Company on {{company.incorporation_date}} together with a copy of the Constitution of the Company.")
        set_paragraph_text(p[23], "That {{company.first_directors_names}} be and are hereby confirmed as the First Directors of the Company.")
        set_paragraph_text(p[29], "That {{secretary.full_name}} is hereby appointed as Secretary of the Company with immediate effect.")
        set_paragraph_text(p[37], "That the Registered Office of the Company situated at {{company.registered_office_address}}")
        set_paragraph_text(p[41], "That the Place where Register of Members and Index is kept situated at {{company.register_location}} be and is hereby confirmed.")
        set_paragraph_text(p[48], "")
        set_paragraph_text(p[33], "4.   CONFIRMATION OF REGISTERED OFFICE AND PLACE WHERE REGISTER OF MEMBERS AND INDEX IS KEPT ")
        set_paragraph_text(p[50], "Company Registration No.: {{company.uen}}")
        set_paragraph_text(p[73], "That the signatories to the Constitution be registered as members in respect of the shares for which they subscribed in full, in cash at {{company.share_currency}} {{company.share_par_value}} /- per share, namely: -")
        set_paragraph_text(p[75], "Name of Subscriber                                                                    No. of share of {{company.share_currency}} {{company.share_par_value}} /- each")
        set_paragraph_text(p[76], "{{company.subscriber_share_lines}}")
        set_paragraph_text(p[84], "That the financial year end of the Company be fixed on {{company.fye}} and the first set of accounts will be for the period from {{company.first_financial_period_start}} to {{company.first_financial_period_end}}.")
        set_paragraph_text(p[87], "{{director.full_name}}\nDirector")
        set_paragraph_text(p[92], "Dated: {{signature.date}}")
    remove_embedded_drawings(doc)
    save(doc, "01_first_directors_resolution_preserved.docx")


def build_share_certificate() -> None:
    doc = load_source_docx(SOURCE_REG / "5. Company Name + Share Certificate-2.docx", "Share Certificate")
    p = doc.paragraphs
    if len(p) > 40:
        set_paragraph_text(p[2], "Certificate No.: {{shareholder.certificate_no}}")
        set_paragraph_text(p[4], "{{company.company_name}}")
        set_paragraph_text(p[5], "")
        set_paragraph_text(p[6], "REG. HOLDER {{shareholder.shareholder_name}}                                                                                         DATE {{signature.date}}")
        set_paragraph_text(p[7], "FOLIO IN MEMBERS' REG. {{shareholder.folio_no}}                                                     ALLOTMENT / TRANSFER NO. {{shareholder.allotment_transfer_no}}")
        set_paragraph_text(p[8], "                                                                                              Remarks {{shareholder.remarks}}")
        set_paragraph_text(p[9], "                                                                                                                                                                                                              No. of Shares: {{shareholder.shares}}")
        set_paragraph_text(p[10], "{{director.full_name}}                                    {{secretary_or_director.full_name}}                                                Distinctive No.: {{shareholder.distinctive_numbers}}")
        set_paragraph_text(p[13], "Certificate No.: {{shareholder.certificate_no}}")
        set_paragraph_text(p[15], "{{company.company_name}}")
        set_paragraph_text(p[16], "")
        set_paragraph_text(p[20], "{{shareholder.shares}} SHARES IN THE COMPANY NUMBERED AS ON THE FACE THEREOF.")
        set_paragraph_text(p[29], "THIS IS TO CERTIFY THAT {{shareholder.shareholder_name}}")
        set_paragraph_text(p[30], "of {{shareholder.shareholder_address}}")
        set_paragraph_text(p[31], "is/are the registered holder(s)")
        set_paragraph_text(p[32], "of {{shareholder.shares}}")
        set_paragraph_text(p[33], "{{shareholder.share_class}} Share(s) fully paid in the Company subject to the Memorandum & Articles of Association of the Company.")
        set_paragraph_text(p[34], "Dated this {{signature.day}} day of {{signature.month_year}}")
        set_paragraph_text(p[39], "  Director {{director.full_name}}                                                                 Director / Secretary {{secretary_or_director.full_name}}")
    save(doc, "04_share_certificate_preserved.docx")


def build_secretary_agreement() -> None:
    doc = load_source_docx(SOURCE_REG / "7.Secretary Agreement.docx", "Secretary Service Agreement")
    replacements: list[tuple[str | re.Pattern, str]] = [
        (re.compile(r"COMPANY NAME PTE\. LT\s+D\."), "{{company.company_name}}"),
        ("COMPANY NAME PTE. LTD.", "{{company.company_name}}"),
        ("Rsin Group Pte. Ltd. & related companies & related companies", "{{provider.name}} & related companies"),
        ("Rsin Group Pte. Ltd. & related companies", "{{provider.name}} & related companies"),
        ("Rsin Group Pte. Ltd.", "{{provider.name}}"),
        ("www.ibizfile.com", "{{provider.portal_url}}"),
        ("{{provider.name}} & related companies\u2019s", "{{provider.name}} & related companies'"),
        ("{{provider.name}} & related companies's", "{{provider.name}} & related companies'"),
        ("will always due and payable", "will always be due and payable"),
        ("including but not related to corporate secretarial services", "including but not limited to corporate secretarial services"),
        ("We aware that", "We are aware that"),
        ("such other agency as may be gazette by", "such other agency as may be gazetted by"),
        ("designated agents/staffs", "designated agents/staff"),
        ("share certificates numbering will commence from number \u201c31\u201d", "share certificates numbering will commence from number \u201c{{provider.default_share_certificate_start_no}}\u201d"),
        ("Date: ", "Date: {{signature.date}}"),
        ("To:\t{{provider.name}} ", "To:\t{{provider.name}}"),
    ]
    replace_in_document(doc, replacements)
    reveal_white_text(doc)
    p = doc.paragraphs
    if len(p) > 78:
        set_paragraph_text(p[64], "Signed by: {{provider.authorised_signatory_name}}")
        set_paragraph_text(p[65], "For and on behalf of {{provider.name}}:")
        set_paragraph_text(p[71], "Signed by: {{client_signatory.full_name}}")
        set_paragraph_text(p[72], "Name: {{client_signatory.full_name}}")
        set_paragraph_text(p[74], "Signed by: {{client_signatory_2.full_name}}")
        set_paragraph_text(p[78], "Signed by: {{client_signatory_3.full_name}}")
    save(doc, "05_secretary_service_agreement_preserved.docx")


def build_signature_record() -> None:
    source = SOURCE_ROOT / "\u6bcf\u4efd\u7b7e\u540d\u6587\u4ef6\u540e\u9700\u9644\u5e26\u7684\u7b7e\u540d\u8bb0\u5f55.docx"
    doc = load_source_docx(source, "Signature Record Attachment")
    p = doc.paragraphs
    if len(p) > 100:
        replacements = {
            0: "{{signature_record.envelope_id}}",
            4: "{{signer.email}} Party ID: {{signer.party_id}} IP Address: {{signer.ip_address}}",
            12: "{{signer.digital_fingerprint_checksum}}",
            15: "{{signer_2.full_name}}",
            16: "Party ID: {{signer_2.party_id}} IP Address: {{signer_2.ip_address}}",
            24: "{{signer_2.digital_fingerprint_checksum}}",
            28: "Timestamp\tAudit",
            29: "{{signature_record.audit_line_1}}",
            30: "{{signature_record.audit_line_2}}",
            31: "{{signature_record.audit_line_3}}",
            33: "{{signature_record.audit_line_4}}",
            36: "{{signature_record.audit_line_5}}",
            53: "{{signature_record.envelope_id_page_2}}",
            56: "{{signer.full_name}}",
            57: "Party ID: {{signer.party_id}} IP Address: {{signer.ip_address}}",
            58: "{{signer.email}}",
            67: "{{signer.digital_fingerprint_checksum}}",
            72: "{{signature_record.audit_line_1}}",
            73: "{{signature_record.audit_line_2}}",
            74: "{{signature_record.audit_line_3}}",
            75: "{{signature_record.audit_line_4}}",
        }
        for idx, text in replacements.items():
            set_paragraph_text(p[idx], text)
    save(doc, "07_signature_record_attachment_preserved.docx")


def build_director_consent() -> None:
    doc = Document()
    configure_business_doc(doc, "Consent to Act as Director / Form 45", compact=True)
    add_title(doc, "THE COMPANIES ACT", 10)
    add_title(doc, "(CHAPTER 50)", 9)
    add_title(doc, "SECTION 173C(a)", 9)
    add_title(doc, "CONSENT TO ACT AS DIRECTOR AND", 10)
    add_title(doc, "STATEMENT OF NON DISQUALIFICATION TO ACT AS DIRECTOR", 10)
    add_para(doc, "FORM 45", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)
    add_para(doc, "Name of Company: {{company.company_name}}")
    add_para(doc, "Company No: {{company.uen}}")
    add_para(doc, "I, the undermentioned person, hereby consent to act as a director of the abovenamed company with effect from {{director.appointment_date}} (date) and declare that:")
    add_para(doc, "(a) I am not disqualified from acting as a director, in that:", after=2)
    form45_items = [
        "(i) I am not below 18 years of age and that I am otherwise of full legal capacity.",
        "(ii) Within a period of 3 years preceding the date of this statement I have not had any disqualification order made by the General Division of the High Court of Singapore against me under section 149A(1) of the Companies Act (\"the Act\").",
        "(iii) Within a period of 5 years preceding the date of this statement I have not had any disqualification order made by the General Division of the High Court of Singapore against me under section 149(1) or 154(2) of the Act.",
        "*(iv) That within a period of 5 years preceding 12th November 1993 I have not been convicted whether within or without Singapore, of any offence - (A) in connection with the promotion, formation or management of a corporation; (B) involving fraud or dishonesty punishable on conviction with imprisonment for 3 months or more; or (C) under section 157 (failure to act honestly and diligently as a director or making improper use of company information for gain) or under section 339 (failure to keep proper company accounts books) of the Act.",
        "*(v) That within a period of 5 years preceding the date of this statement I have not been convicted, in Singapore or elsewhere, of any offence involving fraud or dishonesty punishable on conviction with imprisonment for 3 months or more.",
        "(vi) That - (A) I have not been convicted of 3 or more offences under the Act in relation to the requirements on the filing of returns, accounts or other documents with the Registrar of Companies and have not had 3 or more orders of the General Division of the High Court of Singapore made against me under section 13 or 399 of the Act in relation to such requirements; (B) the last of any such conviction did not take place or the last of any such order was not made during the period of 5 years preceding the date of this statement; and (C) I am not an undischarged bankrupt under section 148(1) of the Act.",
        "(vii) By virtue of the foregoing I am not disqualified from acting as a director of the abovenamed company.",
    ]
    for item in form45_items:
        add_para(doc, item)
    doc.add_page_break()
    add_para(doc, "Name of Company: {{company.company_name}}")
    add_para(doc, "Company No: {{company.uen}}")
    add_para(doc, "(b) I am aware of and undertake to abide by my duties, responsibilities and liabilities specified in the Act as well as under the common law where applicable, including the following key administrative and substantive duties, that is, to:")
    duties = [
        "(i) discharge my responsibilities in the company;",
        "(ii) ensure that I have a reasonable degree of skill and knowledge to handle the affairs of the company;",
        "(iii) act honestly and be reasonably diligent in discharging my duties and act in the interest of the company without putting myself in a position of conflict of interest;",
        "(iv) employ the powers and assets that I am entrusted with for the proper purposes of the company and not for any collateral purpose;",
        "(v) ensure that the company and I comply with all the requirements and obligations under the Act including those in respect of meetings, requisitions, resolutions, accounts, reports, statements, records and other documents on the company, filing and notices and any other prerequisites; and",
        "(vi) account to the shareholders for my conduct of the affairs of the company and make such disclosures that are incumbent upon me under the Act.",
    ]
    for item in duties:
        add_para(doc, item)
    add_para(doc, "(c) That -")
    add_para(doc, "*(i) I have read and understood the above statements; or")
    add_para(doc, "*(ii) the above statements were interpreted to me in {{director.interpretation_language}} (state language/dialect) by {{director.interpreter_name}} NRIC NO: {{director.interpreter_id_number}} before I executed this form and I confirm that the statements are true. I am also aware that I can be prosecuted in Court if I willfully give any information on this form which is false.")
    add_para(doc, "Name: {{director.full_name}}        EMAIL Address: {{director.email}}        Contact Number: {{director.phone}}")
    add_para(doc, "Address: {{director.residential_address}}")
    add_para(doc, "*{{director.id_type}} No: {{director.id_number}}        Nationality: {{director.nationality}}")
    add_para(doc, "Signature: ................................")
    add_para(doc, "Dated this {{signature.day}} day of {{signature.month_year}}")
    add_para(doc, "* Delete where inapplicable.", italic=True, size=8)
    save(doc, "02_director_consent_form45_preserved.docx")


def build_secretary_consent() -> None:
    doc = Document()
    configure_business_doc(doc, "Consent to Act as Secretary / Form 45B", compact=True)
    add_title(doc, "THE COMPANIES ACT", 10)
    add_title(doc, "(CHAPTER 50)", 9)
    add_title(doc, "SECTION 173C(b)", 9)
    add_title(doc, "CONSENT TO ACT AS SECRETARY", 11)
    add_para(doc, "FORM 45B", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)
    add_para(doc, "Name of company: {{company.company_name}}")
    add_para(doc, "Company No: {{company.uen}}")
    add_para(doc, "1. I, {{secretary.full_name}}, the undermentioned person, hereby consent to act as a secretary of the abovenamed company with effect from {{secretary.appointment_date}} (date).")
    add_para(doc, "\u20202. I am a qualified person under section 171(1AA) of the Companies Act by virtue of my being -")
    secretary_items = [
        "*(i) a secretary of a company for at least 3 years of the 5 years immediately preceding the abovementioned date of my appointment as secretary of the abovenamed company.",
        "*(ii) a qualified person under the Legal Profession Act (Cap. 161).",
        "*(iii) a public accountant.",
        "*(iiia) a member of the Institute of Singapore Chartered Accountants (formerly known as the Institute of Certified Public Accountants of Singapore).",
        "*(iv) a member of the Chartered Secretaries Institute of Singapore.",
        "*(v) a member of the Association of International Accountants (Singapore Branch).",
        "*(vi) a member of The Institute of Company Accountants, Singapore.",
    ]
    for item in secretary_items:
        add_para(doc, item)
    add_para(doc, "Name: {{secretary.full_name}}")
    add_para(doc, "Address: {{secretary.residential_address}}")
    add_para(doc, "*{{secretary.id_type}} No: {{secretary.id_number}}        Nationality: {{secretary.nationality}}")
    add_para(doc, "Signature: ................................")
    add_para(doc, "Dated this {{signature.day}} day of {{signature.month_year}}")
    add_para(doc, "\u2020 To be completed by secretaries of public companies only or by secretaries of private companies appointed under section 171(1AB) of the Act.", italic=True, size=8)
    add_para(doc, "* Delete where inapplicable.", italic=True, size=8)
    save(doc, "03_secretary_consent_form45b_preserved.docx")


def build_nominee_director_agreement() -> None:
    doc = Document()
    configure_business_doc(doc, "Nominee Director Agreement", compact=True)
    add_para(doc, "THIS AGREEMENT is dated {{signature.date}}", after=1)
    add_para(doc, "\u672c\u534f\u8bae\u8ba2\u7acb\u4e8e {{signature.date}}", after=6)
    add_title(doc, "AGREEMENT FOR APPOINTMENT OF NOMINEE DIRECTOR", 11)
    add_title(doc, "\u59d4\u4efb\u6302\u540d\u8463\u4e8b\u534f\u8bae", 11)
    add_para(doc, "Between", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "{{nominee_director.full_name}} ({{nominee_director.id_number}})", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_para(doc, "and", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "{{company.company_name}} (Company Registration No. {{company.uen}}), a company incorporated in the Republic of Singapore and having its registered address at {{company.registered_office_address}} (\"Company\"); {{company.company_name_chinese}}(\u516c\u53f8\u6ce8\u518c\u53f7\uff1a{{company.uen}})\uff0c\u4e00\u5bb6\u4e8e\u65b0\u52a0\u5761\u6210\u7acb\u7684\u6709\u9650\u8d23\u4efb\u516c\u53f8\uff0c\u6ce8\u518c\u5730\u5740\u4e3a {{company.registered_office_address_chinese}}\uff08\u7b80\u79f0\u201c\u516c\u53f8\u201d\uff09\u3002")
    add_clause_heading(doc, "WHEREAS", "\u9274\u4e8e")
    add_para(doc, "The Company has requested us to provide the services of a nominee who shall be appointed as a director of the Company (\"Nominee Director\") upon the terms and conditions as stated herein.")
    add_para(doc, "\u4f9d\u636e\u672c\u534f\u8bae\u7ea6\u5b9a\u7684\u6761\u6b3e\u548c\u6761\u4ef6\uff0c\u516c\u53f8\u8981\u6c42\u63d0\u4f9b\u6302\u540d\u8463\u4e8b\u670d\u52a1\uff08\u7b80\u79f0\u201c\u6302\u540d\u8463\u4e8b\u201d\uff09\u3002")
    add_para(doc, "NOW IT IS HEREBY AGREED AS FOLLOWS:", bold=True)
    add_para(doc, "\u53cc\u65b9\u540c\u610f\u5982\u4e0b\uff1a", bold=True)

    nominee_clauses = [
        ("A. APPOINTMENT AND DUTIES", "\u59d4\u4efb\u548c\u8d23\u4efb", [
            ("1. We shall provide the Company with the services of the Nominee Director. Upon the notification of the particulars of the Nominee Director, the Company shall take such steps as are necessary for the Nominee Director to be appointed as a director of the Company.", "\u6211\u4eec\u5c06\u63d0\u4f9b\u516c\u53f8\u6302\u540d\u8463\u4e8b\u670d\u52a1\u3002\u516c\u53f8\u88ab\u901a\u77e5\u6302\u540d\u8463\u4e8b\u7684\u4e2a\u4eba\u4fe1\u606f\u540e\uff0c\u5e94\u4e3a\u8be5\u6302\u540d\u8463\u4e8b\u91c7\u53d6\u5fc5\u8981\u7684\u6b65\u9aa4\u4f7f\u5176\u59d4\u4efb\u4e3a\u516c\u53f8\u4e4b\u8463\u4e8b\u3002"),
            ("2. The duties of the Nominee Director shall be limited to the signing of statutory forms and board resolutions in order to comply with only the requirements of the Companies Act, Cap. 50 and its regulations and the Company's Constitution.", "\u6302\u540d\u8463\u4e8b\u4e4b\u8d23\u4efb\u4ec5\u9650\u4e8e\u7b7e\u7f72\u6cd5\u5b9a\u8868\u683c\u548c\u8463\u4e8b\u4f1a\u51b3\u8bae\uff0c\u4ee5\u7b26\u5408\u65b0\u52a0\u5761\u516c\u53f8\u6cd5\u548c\u5176\u89c4\u5219\u53ca\u516c\u53f8\u7ae0\u7a0b\u7684\u8981\u6c42\u3002"),
            ("3. For the avoidance of doubt, the Nominee Director shall not be required to participate, in any manner whatsoever, in the management or decision-making of the Company and shall not be required to sign or execute any agreement, assurance, guarantee or other documents which would incur or in any way impose any personal liability on the Nominee Director, whether as a principal or guarantor, notwithstanding that the same is purportedly for the benefit of the Company or any other third party. The Nominee Director shall also not be required to comply with any instructions which may be unlawful under the laws of Singapore as well as all other laws which may be applicable.", "\u6302\u540d\u8463\u4e8b\u5c06\u4e0d\u4f1a\u4ee5\u4efb\u4f55\u5f62\u5f0f\u53c2\u4e0e\u516c\u53f8\u7ba1\u7406\u5c42\u6216\u516c\u53f8\u51b3\u7b56\uff0c\u5c06\u4e0d\u4f1a\u88ab\u8981\u6c42\u7b7e\u7f72\u6216\u6267\u884c\u4efb\u4f55\u534f\u8bae\u3001\u4fdd\u8bc1\u3001\u62c5\u4fdd\u6216\u5176\u4ed6\u53ef\u80fd\u4ea7\u751f\u7684\u6587\u4ef6\uff0c\u6216\u4efb\u4f55\u4f1a\u5f15\u8d77\u8be5\u6302\u540d\u8463\u4e8b\u4e2a\u4eba\u8d23\u4efb\u7684\u6587\u4ef6\uff0c\u5c3d\u7ba1\u8fd9\u4e9b\u6587\u4ef6\u636e\u79f0\u662f\u4e3a\u4e86\u516c\u53f8\u6216\u4efb\u4f55\u7b2c\u4e09\u65b9\u7684\u5229\u76ca\u3002\u6302\u540d\u8463\u4e8b\u5c06\u4e5f\u4e0d\u9700\u9075\u5b88\u4efb\u4f55\u8fdd\u53cd\u65b0\u52a0\u5761\u516c\u53f8\u6cd5\u6216\u5176\u4ed6\u6210\u6587\u6cd5\u89c4\u5b9a\u7684\u6307\u793a\u3002"),
            ("4. Statement on matters regarding assistance bank account opening:", "\u5173\u4e8e\u534f\u52a9\u94f6\u884c\u5f00\u6237\u4e8b\u9879\u58f0\u660e\uff1a"),
            ("a) In the event of requirement, the Nominee Director will provide assistance pertaining to matters of local bank account opening by providing valid personal identification documents to the bank. The company retains the legal right to claim compensation and accountability from the Client should the Client submit the Nominee Director's personal identification for any other uses beside bank account opening.", "\u5728\u6709\u9700\u8981\u7684\u60c5\u51b5\u4e0b\uff0c\u65b0\u52a0\u5761\u6302\u540d\u8463\u4e8b\u5c06\u5bf9\u4e8e\u516c\u53f8\u5728\u65b0\u52a0\u5761\u672c\u5730\u94f6\u884c\u5f00\u6237\u7684\u670d\u52a1\u8fdb\u884c\u534f\u52a9\uff0c\u5e76\u4e14\u63d0\u4f9b\u6709\u6548\u7684\u4e2a\u4eba\u8eab\u4efd\u8bc1\u4ef6\u7ed9\u4e88\u94f6\u884c\u8fdb\u884c\u5f00\u6237\u7684\u3002\u4e00\u65e6\u53d1\u73b0\u65b0\u52a0\u5761\u6302\u540d\u8463\u4e8b\u7684\u8eab\u4efd\u8bc1\u4ef6\u5728\u672a\u544a\u77e5\u7684\u60c5\u51b5\u4e0b\u7528\u4f5c\u5176\u4ed6\u7528\u9014\uff0c\u6211\u4eec\u5c06\u4fdd\u7559\u6cd5\u5f8b\u8ffd\u7a76\u4ee5\u53ca\u8981\u6c42\u8d54\u507f\u7684\u6743\u5229\u3002"),
            ("b) Assistance for opening bank account will only be provided by the Nominee Director in the event that the bank in question is a local (Singapore) bank. The Nominee Director reserves the rights to reject the Client should there be a request for bank account opening at an overseas bank.", "\u65b0\u52a0\u5761\u6302\u540d\u8463\u4e8b\u534f\u52a9\u5f00\u6237\u7684\u4e49\u52a1\u5c3d\u6b62\u4e8e\u65b0\u52a0\u5761\u672c\u5730\u94f6\u884c\uff0c\u82e5\u516c\u53f8\u8981\u6c42\u65b0\u52a0\u5761\u6302\u540d\u8463\u4e8b\u534f\u52a9\u5916\u56fd\u94f6\u884c\u7684\u5f00\u6237\uff0c\u6302\u540d\u8463\u4e8b\u6709\u6743\u63d0\u51fa\u62d2\u7edd\u3002"),
        ]),
        ("B. TERMINATION", "\u7ec8\u6b62", [
            ("1. Either party may terminate this Agreement at any time by providing the other party with prior one (1) month's notice in writing.", "\u534f\u8bae\u53cc\u65b9\u6709\u6743\u5728\u63d0\u4f9b\u5bf9\u65b9\u63d0\u524d 1 \u4e2a\u6708\u7684\u4e66\u9762\u901a\u77e5\u7684\u524d\u63d0\u4e0b\uff0c\u5728\u4efb\u4f55\u65f6\u5019\u7ec8\u6b62\u672c\u534f\u8bae\u3002"),
            ("2. In the event of ceasing (for whatever reason) to provide the services of the Nominee Director, the board of directors of the Company shall accept the resignation of the Nominee Director and where required under the Companies Act, Cap. 50 and its regulations or the Company's Constitution, shall take immediate steps to appoint another person in place of the Nominee Director as a director of the Company on or before the date of termination and/or cessation.", "\u82e5\u505c\u6b62\u63d0\u4f9b\u6302\u540d\u8463\u4e8b\u670d\u52a1\uff08\u65e0\u8bba\u4f55\u79cd\u539f\u56e0\uff09\uff0c\u516c\u53f8\u8463\u4e8b\u4f1a\u5e94\u63a5\u53d7\u6302\u540d\u8463\u4e8b\u4e4b\u8bf7\u8f9e\u5e76\u5728\u7ec8\uff08\u505c\uff09\u6b62\u65e5\u524d\u6216\u5f53\u5929\u4f9d\u636e\u65b0\u52a0\u5761\u516c\u53f8\u6cd5\u53ca\u5176\u89c4\u5b9a\u6216\u516c\u53f8\u7ae0\u7a0b\uff0c\u91c7\u53d6\u7acb\u5373\u7684\u63aa\u65bd\u59d4\u4efb\u53e6\u4e00\u4e2a\u4eba\u4e3a\u516c\u53f8\u8463\u4e8b\uff0c\u6765\u4ee3\u66ff\u8be5\u6302\u540d\u8463\u4e8b\u3002"),
            ("3. The Company shall provide with satisfactory documentary evidence that such changes have been effected and notified to the ACRA.", "\u516c\u53f8\u63d0\u4f9b\u7b26\u5408\u8981\u6c42\u7684\u4e66\u9762\u6587\u4ef6\uff0c\u8bc1\u660e\u8be5\u66ff\u6362\u5df2\u7ecf\u751f\u6548\u5e76\u5df2\u7ecf\u901a\u77e5\u65b0\u52a0\u5761\u5546\u4e1a\u6ce8\u518c\u5c40\u3002"),
        ]),
        ("C. WARRANTIES AND COVENANTS", "\u4fdd\u8bc1\u548c\u534f\u7ea6", [
            ("1. The Company hereby undertakes and warrants that all acts required to be done by the Nominee Director shall be lawful under the laws of Singapore as well as all other laws which may be applicable and that all statements and documents, which the Nominee Director is requested to sign, shall be true, accurate and not misleading.", "\u516c\u53f8\u636e\u6b64\u627f\u8bfa\u548c\u4fdd\u8bc1\u8981\u6c42\u6302\u540d\u8463\u4e8b\u5f00\u5c55\u7684\u6d3b\u52a8\u7b26\u5408\u65b0\u52a0\u5761\u6cd5\u5f8b\u548c\u5176\u4ed6\u9002\u7528\u7684\u6cd5\u5f8b\uff0c\u6240\u6709\u8981\u6c42\u6302\u540d\u8463\u4e8b\u7b7e\u7f72\u7684\u58f0\u660e\u548c\u6587\u4ef6\u90fd\u662f\u771f\u5b9e\u3001\u6b63\u786e\u4e14\u672a\u6709\u4efb\u4f55\u8bef\u5bfc\u3002"),
            ("2. The Company hereby undertakes and warrants that the Company will keep proper accounts and will file all income tax returns and all papers and documents which are required by law to be filed with any authority and that all taxes required to be paid by the Company will be duly paid.", "\u516c\u53f8\u636e\u6b64\u627f\u8bfa\u548c\u4fdd\u8bc1\u516c\u53f8\u5c06\u4f1a\u9075\u5b88\u4f1a\u8ba1\u51c6\u5219\uff0c\u7533\u62a5\u516c\u53f8\u6240\u5f97\u7a0e\uff0c\u5411\u6709\u5173\u673a\u5173\u7533\u62a5\u6cd5\u5f8b\u8981\u6c42\u7533\u62a5\u7684\u6240\u6709\u6587\u4ef6\uff0c\u516c\u53f8\u6240\u9700\u8981\u652f\u4ed8\u7684\u7a0e\u6536\u90fd\u5c06\u4f1a\u6309\u65f6\u652f\u4ed8\u3002"),
        ]),
        ("D. GENERAL", "\u7efc\u5408", [
            ("1. If any of the provisions of this Agreement becomes invalid, illegal or unenforceable in any respect under any law, the validity, legality and enforceability of the remaining provisions shall not in any way be affected or impaired.", "\u82e5\u672c\u534f\u8bae\u4e2d\u7684\u4efb\u4f55\u6761\u6b3e\u65e0\u6548\uff0c\u975e\u6cd5\u6216\u4e0d\u80fd\u6267\u884c\uff0c\u5e76\u4e0d\u5f71\u54cd\u6216\u635f\u5bb3\u672c\u534f\u8bae\u5176\u4ed6\u6761\u6b3e\u7684\u6709\u6548\u6027\u3001\u5408\u6cd5\u6027\u548c\u6267\u884c\u6027\u3002"),
            ("2. This Agreement embodies the entire understanding of the parties and there are no promises, terms and conditions or disclosures or obligations, oral or written, express or implied other than those contained herein, save and except those mutually agreed to in writing between the parties after the due execution of this Agreement and each party acknowledge that it/he has not relied on any oral or written representations made to it/him by the other or its/his employees or agents.", "\u9664\u53cc\u65b9\u7b7e\u7f72\u672c\u534f\u8bae\u540e\u4e66\u9762\u5171\u540c\u540c\u610f\uff0c\u4e14\u6bcf\u4e00\u65b9\u786e\u8ba4\u5176\u6ca1\u6709\u4f9d\u8d56\u4efb\u4f55\u5bf9\u65b9\u3001\u5bf9\u65b9\u96c7\u5458\u6216\u4ee3\u7406\u53e3\u5934\u6216\u4e66\u9762\u7684\u58f0\u660e\u5916\uff0c\u672c\u534f\u8bae\u5305\u542b\u53cc\u65b9\u7684\u5168\u90e8\u7406\u89e3\uff0c\u9664\u672c\u534f\u8bae\u5916\uff0c\u53cc\u65b9\u4e4b\u95f4\u65e0\u4efb\u4f55\u627f\u8bfa\u3001\u6761\u6b3e\u548c\u6761\u4ef6\uff0c\u6216\u62ab\u9732\u6216\u4e49\u52a1\uff0c\u65e0\u8bba\u662f\u53e3\u5934\u8fd8\u662f\u4e66\u9762\uff0c\u660e\u793a\u6216\u6697\u793a\u7684\u3002"),
            ("3. Time wherever mentioned in this Agreement shall be of the essence of this Agreement.", "\u672c\u534f\u8bae\u63d0\u5230\u7684\u65f6\u95f4\u89c6\u4e3a\u534f\u8bae\u4e4b\u8981\u7d20\u3002"),
            ("4. Subject as otherwise provided herein, all notices, demands or other communications required or permitted to be given or made hereunder shall be in writing and delivered personally or sent by prepaid registered post or certificate of posting or by facsimile message addressed to the intended recipient thereof at its address or its facsimile number set out herein (or to such other address or facsimile number as any party may from time to time notify the others for the purpose of this clause) or by other means which parties may agree in writing thereafter. Any notice, demand or communication shall be deemed to have been duly served: (a) if delivered personally, on the day of delivery; (b) if sent by facsimile, on the day of the conclusion of transmission; (c) if sent by letter to a local address, [48 hours] after posting; and (d) if sent by letter to an overseas address, [3 days] after posting.", "\u9664\u53e6\u6709\u89c4\u5b9a\u5916\uff0c\u6240\u6709\u7684\u901a\u77e5\u3001\u8bf7\u6c42\u6216\u8005\u5176\u4ed6\u8981\u6c42\u6216\u5141\u8bb8\u63d0\u4f9b\u7684\u901a\u4fe1\u5e94\u91c7\u7528\u4e66\u9762\u5f62\u5f0f\uff0c\u7531\u672c\u4eba\u4eb2\u81ea\u9012\u9001\uff0c\u6216\u4ee5\u9884\u4ed8\u6b3e\u6302\u53f7\u4fe1\u6216\u90ae\u5bc4\u8bc1\u660e\u7684\u5f62\u5f0f\u5bc4\u9001\uff0c\u6216\u4ee5\u4f20\u771f\u81f3\u6536\u4fe1\u4eba\u5730\u5740\u6216\u4f20\u771f\u53f7\u7801\u7684\u5f62\u5f0f\uff0c\u6216\u53cc\u65b9\u4e66\u9762\u540c\u610f\u7684\u5176\u4ed6\u65b9\u5f0f\u3002"),
            ("5. This Agreement shall be binding on each of the parties and their respective assignees, executors and administrators.", "\u672c\u534f\u8bae\u5bf9\u53cc\u65b9\uff0c\u4ed6\u4eec\u7684\u627f\u7ee7\u8005\uff0c\u5236\u5b9a\u9057\u5631\u6267\u884c\u4eba\u548c\u7ba1\u7406\u4eba\u5177\u6709\u7ea6\u675f\u529b\u3002"),
            ("6. A person who is not a party to this Agreement has no rights under the Contracts (Rights of Third Parties) Act (Cap. 53B) to enforce any term of this Agreement.", "\u975e\u534f\u8bae\u4e4b\u5f53\u4e8b\u4eba\u65e0\u6743\u4f9d\u636e\u5408\u540c\u6cd5\uff08\u7b2c\u4e09\u65b9\u6743\u5229\uff09\u6765\u6267\u884c\u672c\u534f\u8bae\u7684\u4efb\u4f55\u6761\u6b3e\u3002"),
            ("7. Termination of this Agreement for whatever reason shall not release either party from any obligations which, either expressly or by implication, are intended to survive the termination of this Agreement, nor will it affect the accrued rights and liabilities of the parties as at the date of termination.", "\u65e0\u8bba\u56e0\u4f55\u79cd\u539f\u56e0\u7ec8\u6b62\u672c\u534f\u8bae\u5c06\u4e0d\u4f1a\u8c41\u514d\u534f\u8bae\u4efb\u4f55\u4e00\u65b9\u7684\u4e49\u52a1\uff0c\u4e5f\u4e0d\u4f1a\u5f71\u54cd\u5230\u81f3\u7ec8\u6b62\u65e5\u6b62\u534f\u8bae\u53cc\u65b9\u4ea7\u751f\u7684\u6743\u5229\u548c\u503a\u52a1\u3002"),
            ("8. Words denoting the singular number only shall include the plural number and vice versa.", "\u5355\u6570\u5f62\u5f0f\u5c06\u5305\u62ec\u590d\u6570\u5f62\u5f0f\uff0c\u53cd\u4e4b\u4ea6\u7136\u3002"),
            ("9. Words denoting the masculine gender only shall include the feminine and neuter genders.", "\u7537\u6027\u4eba\u79f0\u5c06\u5305\u62ec\u5973\u6027\u4eba\u79f0\u548c\u4e2d\u6027\u4eba\u79f0\u3002"),
            ("10. References to persons shall be deemed to include bodies incorporated or unincorporated.", "\u4eba\u79f0\u5c06\u89c6\u4e3a\u5305\u62ec\u6210\u7acb\u6216\u672a\u6210\u7acb\u7684\u7ec4\u7ec7\u3002"),
            ("11. The Recitals to this Agreement shall be and form an integral part of this Agreement.", "\u534f\u8bae\u7684\u524d\u8a00\u90e8\u5206\u4e3a\u534f\u8bae\u4e4b\u7ec4\u6210\u90e8\u5206\u3002"),
            ("12. Headings in this Agreement are for convenient reference only and shall not be used to construe or interpret this Agreement.", "\u534f\u8bae\u4e4b\u6807\u9898\u53ea\u4e3a\u65b9\u4fbf\u53c2\u8003\u4e4b\u4f7f\u7528\uff0c\u5e76\u4e0d\u662f\u7528\u6765\u5206\u6790\u6216\u6f14\u7ece\u534f\u8bae\u4e4b\u7528\u3002"),
            ("13. Any reference to a statutory provision shall include such provision and any regulations made in pursuance thereof as from time to time modified or re-enacted whether before or after the date of this Agreement.", "\u5f15\u7528\u6cd5\u6761\u5c06\u5305\u62ec\u8be5\u6cd5\u6761\u548c\u4efb\u4f55\u4f9d\u636e\u8be5\u6761\u5236\u8ba2\u7684\u89c4\u5219\uff0c\u65e0\u8bba\u662f\u534f\u8bae\u8ba2\u7acb\u524d\u6216\u540e\u4e0d\u65f6\u4fee\u6539\u6216\u8005\u91cd\u65b0\u5236\u5b9a\u3002"),
            ("14. \"Dollars\" and the sign \"$\" shall mean the lawful currency of Singapore.", "\u65b0\u5143\u6216\u8005\u7b26\u53f7\u201c$\u201d\u4ee3\u8868\u65b0\u52a0\u5761\u7684\u6cd5\u5b9a\u8d27\u5e01\u3002"),
            ("15. This Agreement shall be governed by and construed in accordance with the laws of the Republic of Singapore and parties agree to submit to the non-exclusive jurisdiction of the courts of the Republic of Singapore.", "\u672c\u534f\u8bae\u53d7\u65b0\u52a0\u5761\u6cd5\u5f8b\u7ba1\u8f96\uff0c\u65b0\u52a0\u5761\u6cd5\u9662\u5bf9\u672c\u534f\u8bae\u5177\u6709\u975e\u6392\u5916\u7ba1\u8f96\u6743\u3002"),
        ]),
    ]
    for heading, zh, clauses in nominee_clauses:
        add_clause_heading(doc, heading, zh)
        for en, cn in clauses:
            add_para(doc, en)
            add_para(doc, cn)
    doc.add_page_break()
    add_para(doc, "IN WITNESS whereof the parties have duly executed this Agreement as of the date and year written above.")
    add_para(doc, "\u534f\u8bae\u53cc\u65b9\u5728\u4e0a\u8ff0\u7684\u5e74\u4efd\u548c\u65e5\u671f\u5408\u6cd5\u6709\u6548\u5730\u7b7e\u7f72\u672c\u534f\u8bae\u3002")
    add_signature_pair(
        doc,
        ["The Signature of", "Nominee Director", "\u4ee3\u7406\u8463\u4e8b", "{{nominee_director.full_name}}"],
        ["The Stamp / Signature of", "{{company.company_name}}", "{{client_signatory.full_name}}", "Director / \u8463\u4e8b"],
    )
    save(doc, "06_nominee_director_agreement_preserved.docx")


def write_guides() -> None:
    review_notes = {
        "principle": "v2 preserves old text/order as much as possible, fieldizes sample data, and only applies clear professionalism/source updates.",
        "preserved_from_source_docx": [
            "First Directors Resolution",
            "Share Certificate",
            "Secretary Service Agreement",
            "Signature Record Attachment",
        ],
        "reconstructed_from_pdf_or_reference": [
            "Director Consent / Form 45",
            "Secretary Consent / Form 45B",
            "Nominee Director Agreement",
        ],
        "updates_applied": [
            "Removed sample names, addresses, IDs, email addresses, dates and company numbers; replaced them with placeholders.",
            "Form 45 wording uses ACRA reference wording where the old PDF was already close but sample-filled; notably 'General Division of the High Court of Singapore'.",
            "Form 45B institute names follow the ACRA reference wording, including Institute of Singapore Chartered Accountants and Chartered Secretaries Institute of Singapore.",
            "Secretary Agreement keeps the long bilingual clauses from the source DOCX, but fixes obvious duplicate provider wording and visible grammar typos such as 'will always be due and payable', 'We are aware', and 'including but not limited to'.",
            "First Directors Resolution keeps the old layout and fixes visible heading/grammar typos such as 'MEMBERS AND INDEX' spacing and plural signatories/members wording, while removing old sample arrow/page-marker artifacts.",
            "Recovered old Secretary Agreement clauses that were present in the source file but formatted as white text.",
            "Nominee Director Agreement keeps the full clause structure from the source PDF instead of summary clauses.",
        ],
        "manual_review_flags": [
            "These are operational templates, not legal advice. Final production use should be checked by a Singapore-qualified professional or your internal reviewer.",
            "Old templates still refer to Companies Act (Chapter 50) / Cap. 50 in places because the ACRA Form 45/Form 45B PDFs also display that legacy label. Do not bulk change this without legal review.",
            "The Nominee Director Agreement includes broad bilingual terms from the old PDF. Commercial terms such as fees, deposit, notice, local bank support and liability scope should be confirmed before use.",
        ],
        "official_references": [
            "ACRA Form 45 - Consent to Act as Director and Statement of Non Disqualification to Act as Director",
            "ACRA Form 45B - Consent to Act as Secretary",
            "Singapore Statutes Online - Companies Act 1967 sections 149, 154, 171 and 173A/173E context",
        ],
        "templates": MANIFEST,
        "fields": FIELD_NOTES,
    }
    readme_lines = [
        "# P1 Preserved Template Pack v2",
        "",
        "This version replaces the first simplified rebuild. It keeps the old template content and order as much as possible, while replacing sample data with clear placeholders.",
        "",
        "## Files",
        "",
    ]
    for item in MANIFEST:
        readme_lines.append(f"- `{item['file_name']}` - {item['display_name']} ({item['build_method']})")
    readme_lines.extend(
        [
            "",
            "## Upgrade Rules",
            "",
            "- Do not summarize or shorten agreement clauses.",
            "- Replace sample data with placeholders.",
            "- Keep old display names such as Form 45/Form 45B, but use stable template IDs internally.",
            "- Put legal/commercial uncertainty in review notes instead of silently changing signed terms.",
            "",
            "## Main Placeholders",
            "",
        ]
    )
    for key, note in FIELD_NOTES.items():
        readme_lines.append(f"- `{{{{{key}}}}}`: {note}")
    readme_lines.extend(
        [
            "",
            "## Review Notes",
            "",
        ]
    )
    for note in review_notes["manual_review_flags"]:
        readme_lines.append(f"- {note}")

    for directory in (TEMPLATE_DIR, OUTPUT_DIR):
        (directory / "template_manifest_v2.json").write_text(json.dumps(review_notes, ensure_ascii=False, indent=2), encoding="utf-8")
        (directory / "README_v2_fields_and_review.md").write_text("\n".join(readme_lines), encoding="utf-8")


def zip_output() -> None:
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(OUTPUT_DIR.iterdir()):
            if path.is_file():
                zf.write(path, arcname=path.name)


def main() -> None:
    ensure_dirs()
    mark_v1_as_simplified()
    build_first_directors_resolution()
    build_director_consent()
    build_secretary_consent()
    build_share_certificate()
    build_secretary_agreement()
    build_nominee_director_agreement()
    build_signature_record()
    write_guides()
    zip_output()
    print(f"Generated preserved v2 templates in {TEMPLATE_DIR}")
    print(f"Copied package to {OUTPUT_DIR}")
    print(f"Zip package: {ZIP_PATH}")


if __name__ == "__main__":
    main()
