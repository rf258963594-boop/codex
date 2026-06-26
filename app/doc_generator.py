from __future__ import annotations

import json
import calendar
import re
import zipfile
from copy import deepcopy
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Any

from docx import Document
from docx.table import Table
from docx.table import _Row
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from config import DOC_TEMPLATE_DIR, GENERATED_DIR, P1_VERSION
from doc_render import convert_docx_to_pdf
from signatures import SIGNATURE_DOCX_WIDTH_IN, apply_auto_signatures, resolve_signature_path


P1_TEMPLATE_DIR = DOC_TEMPLATE_DIR / "p1_standard_v3_part1"
P2_TEMPLATE_DIR = DOC_TEMPLATE_DIR / "p2_standard_v1"
M01_TEMPLATE_NAME = "M01_combined_directors_resolution_standard.docx"
M01_VERSION = "P2_M01_v0.1"
M02_VERSION = "P2_M02_v0.2"
M03_VERSION = "P2_M03_v0.1"
M04_VERSION = "P2_M04_v0.1"
M05_VERSION = "P2_M05_v0.3"
M02_TEMPLATE_NAMES = {
    "resolution_package": "M02_resolution_package_transfer_in_standard.docx",
    "handover_resignation_package": "M02_handover_and_resignation_package_standard.docx",
}
M03_TEMPLATE_NAMES = {
    "resolution": "M03_share_transfer_directors_resolution_standard.docx",
    "instrument": "M03_instrument_of_transfer_standard.docx",
    "certificate": "M03_updated_share_certificate_standard.docx",
    "register": "M03_register_of_members_update_standard.docx",
    "checklist": "M03_register_and_stamp_duty_checklist_standard.docx",
}
M04_TEMPLATE_NAMES = {
    "authority": "M04_s161_members_authority_standard.docx",
    "resolution": "M04_allotment_directors_resolution_standard.docx",
    "application": "M04_share_application_standard.docx",
    "certificate": "M04_share_certificate_standard.docx",
    "form24": "M04_return_of_allotment_form24_standard.docx",
    "register": "M04_register_of_members_update_standard.docx",
    "checklist": "M04_register_update_checklist_standard.docx",
}
M05_TEMPLATE_NAMES = {
    "agm_package": "M05_agm_documents_package_standard.docx",
    "annual_return_package": "M05_annual_return_authorisation_package_standard.docx",
    "checklist": "M05_annual_review_checklist_standard.docx",
}
NO_VALUES = {"no", "n", "false", "0", "否", "不", "无需", "不用"}
M02_TRANSFER_EVENT_TYPES = {"transfer_in", "transfer_in_cooperative", "transfer_in_non_cooperative"}
DEFAULT_PREVIOUS_SECRETARIAL_PROVIDER = "the previous corporate secretarial service provider and/or the former company secretary"
M01_DR_EVENT_TYPES = {
    "change_registered_office",
    "change_office_hours",
    "change_business_activity",
    "change_fye",
    "update_officer_particulars",
    "update_director_particulars",
    "update_secretary_particulars",
    "update_shareholder_particulars",
    "appoint_director",
    "resign_director",
    "appoint_secretary",
    "resign_secretary",
    "share_transfer_approval",
    "share_allotment_approval",
    "bizfile_authorization",
}
M01_REPEAT_VARS = {
    "m01.officer_particular_updates": "officer_update",
    "m01.director_appointments": "person",
    "m01.director_resignations": "person",
    "m01.secretary_appointments": "person",
    "m01.secretary_resignations": "person",
    "m01.share_transfers": "transfer",
    "m01.share_allotments": "allotment",
    "m01.director_signers": "signer",
    "m03.transfers": "transfer",
    "m03.director_signature_rows": "sigrow",
    "m03.transfer_parties": "party",
    "m03.checklist_items": "check",
    "m04.allotments": "allotment",
    "m04.member_signature_rows": "sigrow",
    "m04.director_signature_rows": "sigrow",
    "m04.checklist_items": "check",
    "m05.member_signature_rows": "sigrow",
    "m05.director_signature_rows": "sigrow",
    "m05.attendance_rows": "attendee",
    "m05.checklist_items": "check",
}


def generate_p1_package(parsed: dict[str, Any], job_code: str) -> Path:
    """Generate the current P1 incorporation DOCX package and return the zip path."""
    if parsed.get("task_type") != "incorporation":
        raise ValueError("Current P1 document generation only supports incorporation Excel files.")
    if not P1_TEMPLATE_DIR.exists():
        raise FileNotFoundError(f"P1 template folder not found: {P1_TEMPLATE_DIR}")

    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    package_dir = GENERATED_DIR / f"{safe_filename(job_code)}_P1_docs"
    package_dir.mkdir(parents=True, exist_ok=True)
    for old_path in package_dir.iterdir():
        if old_path.is_file():
            old_path.unlink()

    context = build_context(parsed)
    generated: list[str] = []

    company_templates = [
        ("01_first_directors_resolution_standard.docx", "01_first_directors_resolution.docx"),
        ("05_secretary_service_agreement_standard.docx", "05_secretary_service_agreement.docx"),
        ("07_return_of_allotment_form24_standard.docx", "07_return_of_allotment_form24.docx"),
        ("09_register_of_members_standard.docx", "09_register_of_members.docx"),
        ("10_paid_up_capital_confirmation_standard.docx", "10_paid_up_capital_confirmation.docx"),
    ]
    for template_name, output_name in company_templates:
        render_docx(P1_TEMPLATE_DIR / template_name, package_dir / output_name, context)
        generated.append(output_name)

    for idx, director in enumerate(context["directors"], start=1):
        child = context_with(context, director=director, signature=context["signature"])
        name = f"02_director_consent_form45_{idx:02d}_{safe_filename(director.get('full_name') or 'director')}.docx"
        render_docx(P1_TEMPLATE_DIR / "02_director_consent_form45_standard.docx", package_dir / name, child)
        generated.append(name)

    for idx, secretary in enumerate(context["secretaries"], start=1):
        child = context_with(context, secretary=secretary, signature=context["signature"])
        name = f"03_secretary_consent_form45b_{idx:02d}_{safe_filename(secretary.get('full_name') or 'secretary')}.docx"
        render_docx(P1_TEMPLATE_DIR / "03_secretary_consent_form45b_standard.docx", package_dir / name, child)
        generated.append(name)

    for idx, shareholder in enumerate(context["shareholders"], start=1):
        child = context_with(context, shareholder=shareholder, director=context["certificate_director"])
        name = f"04_share_certificate_{idx:02d}_{safe_filename(shareholder.get('shareholder_name') or 'shareholder')}.docx"
        render_docx(P1_TEMPLATE_DIR / "04_share_certificate_standard.docx", package_dir / name, child)
        generated.append(name)

    nominees = context["nominee_directors"] or context["local_directors"]
    for idx, nominee in enumerate(nominees, start=1):
        child = context_with(context, nominee_director=nominee)
        name = f"06_nominee_director_agreement_{idx:02d}_{safe_filename(nominee.get('full_name') or 'nominee_director')}.docx"
        render_docx(P1_TEMPLATE_DIR / "06_nominee_director_agreement_standard.docx", package_dir / name, child)
        generated.append(name)

    for idx, shareholder in enumerate(context["registrable_controllers"], start=1):
        child = context_with(context, shareholder=shareholder)
        name = f"08_rorc_notice_controller_{idx:02d}_{safe_filename(shareholder.get('shareholder_name') or 'controller')}.docx"
        render_docx(P1_TEMPLATE_DIR / "08_rorc_notice_controller_standard.docx", package_dir / name, child)
        generated.append(name)

    summary_path = package_dir / "generation_summary.json"
    summary_path.write_text(
        json.dumps(
            {
                "job_code": job_code,
                "p1_version": P1_VERSION,
                "company": context["company"].get("company_name", ""),
                "generated_files": generated,
                "counts": {
                    "directors": len(context["directors"]),
                    "secretaries": len(context["secretaries"]),
                    "shareholders": len(context["shareholders"]),
                    "registrable_controllers": len(context["registrable_controllers"]),
                },
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    zip_path = GENERATED_DIR / f"{safe_filename(job_code)}_P1_docx_package.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(package_dir.iterdir()):
            if path.is_file():
                zf.write(path, arcname=path.name)
    return zip_path


def generate_p1_pdf_package(parsed: dict[str, Any], job_code: str) -> Path:
    """Generate the current P1 package and return a PDF-only zip for users."""
    docx_zip = generate_p1_package(parsed, job_code)
    docx_zip.unlink(missing_ok=True)
    code = safe_filename(job_code)
    return build_pdf_zip_from_docx_dir(
        GENERATED_DIR / f"{code}_P1_docs",
        GENERATED_DIR / f"{code}_P1_pdf",
        GENERATED_DIR / f"{code}_P1_pdf_package.zip",
    )


def generate_p2_m01_package(parsed: dict[str, Any], job_code: str) -> Path:
    """Generate the first P2 M01 DOCX package for ordinary directors' resolutions."""
    if parsed.get("task_type") != "maintenance":
        raise ValueError("M01 generation only supports maintenance / P2 Excel files.")
    template_path = P2_TEMPLATE_DIR / M01_TEMPLATE_NAME
    if not template_path.exists():
        raise FileNotFoundError(f"M01 template not found: {template_path}")

    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    package_dir = GENERATED_DIR / f"{safe_filename(job_code)}_P2_M01_docs"
    package_dir.mkdir(parents=True, exist_ok=True)
    for old_path in package_dir.iterdir():
        if old_path.is_file():
            old_path.unlink()

    context = build_m01_context(parsed)
    if not m01_has_content(context["m01"]):
        raise ValueError("This task has no M01 ordinary directors' resolution content.")

    company_name = safe_filename(context["company"].get("company_name") or "company")
    generated: list[str] = []
    output_name = f"01_M01_combined_directors_resolution_{company_name}.docx"
    render_m01_docx(template_path, package_dir / output_name, context)
    generated.append(output_name)
    generated.extend(render_m01_appointment_consents(package_dir, context))

    summary_path = package_dir / "generation_summary.json"
    summary_path.write_text(
        json.dumps(
            {
                "job_code": job_code,
                "version": M01_VERSION,
                "company": context["company"].get("company_name", ""),
                "generated_files": generated,
                "ordinary_dr_sections": m01_section_names(context["m01"]),
                "director_consents": [
                    row.get("full_name", "") for row in context["m01"].get("director_appointments", [])
                ],
                "secretary_consents": [
                    row.get("full_name", "") for row in context["m01"].get("secretary_appointments", [])
                ],
                "director_signers": [row.get("full_name", "") for row in context["m01"].get("director_signers", [])],
                "notes": [
                    "M01 is the first P2 generated document type.",
                    "Default M01 signing logic is all current directors identified from the People sheet or one-page director signer list.",
                    "Specific signer fields are used only when m01_director_signing_mode is specified_signers.",
                    "Director and secretary appointment events also generate Form 45 / Form 45B consent documents.",
                    "M02, M03 and M04 documents are generated by their own dedicated packages.",
                ],
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    zip_path = GENERATED_DIR / f"{safe_filename(job_code)}_P2_M01_docx_package.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(package_dir.iterdir()):
            if path.is_file():
                zf.write(path, arcname=path.name)
    return zip_path


def render_m01_appointment_consents(package_dir: Path, context: dict[str, Any]) -> list[str]:
    """Generate appointment consent documents that belong with an M01 appointment event."""
    generated: list[str] = []
    m01 = context.get("m01", {})

    for idx, director in enumerate(m01.get("director_appointments", []), start=1):
        consent_person = appointment_consent_person(director, "Director")
        child = context_with(context, director=consent_person, signature=context["signature"])
        name = f"02_M01_director_consent_form45_{idx:02d}_{safe_filename(consent_person.get('full_name') or 'director')}.docx"
        render_docx(P1_TEMPLATE_DIR / "02_director_consent_form45_standard.docx", package_dir / name, child)
        generated.append(name)

    for idx, secretary in enumerate(m01.get("secretary_appointments", []), start=1):
        consent_person = appointment_consent_person(secretary, "Company Secretary")
        child = context_with(context, secretary=consent_person, signature=context["signature"])
        name = f"03_M01_secretary_consent_form45b_{idx:02d}_{safe_filename(consent_person.get('full_name') or 'secretary')}.docx"
        render_docx(P1_TEMPLATE_DIR / "03_secretary_consent_form45b_standard.docx", package_dir / name, child)
        generated.append(name)

    return generated


def generate_p2_m01_pdf_package(parsed: dict[str, Any], job_code: str) -> Path:
    """Generate the first P2 ordinary DR package and return a PDF-only zip."""
    docx_zip = generate_p2_m01_package(parsed, job_code)
    docx_zip.unlink(missing_ok=True)
    code = safe_filename(job_code)
    return build_pdf_zip_from_docx_dir(
        GENERATED_DIR / f"{code}_P2_M01_docs",
        GENERATED_DIR / f"{code}_P2_M01_pdf",
        GENERATED_DIR / f"{code}_P2_M01_pdf_package.zip",
    )


def generate_p2_m02_package(parsed: dict[str, Any], job_code: str) -> Path:
    """Generate the P2 M02 transfer-in DOCX package."""
    if parsed.get("task_type") != "maintenance":
        raise ValueError("M02 generation only supports maintenance / P2 Excel files.")
    missing_templates = [
        name for name in M02_TEMPLATE_NAMES.values()
        if not (P2_TEMPLATE_DIR / name).exists()
    ]
    if missing_templates:
        raise FileNotFoundError(f"M02 template(s) not found: {', '.join(missing_templates)}")

    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    package_dir = GENERATED_DIR / f"{safe_filename(job_code)}_P2_M02_docs"
    package_dir.mkdir(parents=True, exist_ok=True)
    for old_path in package_dir.iterdir():
        if old_path.is_file():
            old_path.unlink()

    context = build_m02_context(parsed)
    if not m02_has_content(context["m02"]):
        raise ValueError("This task has no M02 transfer-in content.")

    company_name = safe_filename(context["company"].get("company_name") or "company")
    generated: list[str] = []
    base_outputs = [
        ("resolution_package", f"01_M02_EGM_and_board_resolutions_{company_name}.docx"),
        ("handover_resignation_package", f"02_M02_handover_and_resignation_package_{company_name}.docx"),
    ]
    for template_key, output_name in base_outputs:
        render_m01_docx(P2_TEMPLATE_DIR / M02_TEMPLATE_NAMES[template_key], package_dir / output_name, context)
        if template_key == "handover_resignation_package":
            append_m02_resignation_letters(package_dir / output_name, context)
        generated.append(output_name)

    summary_path = package_dir / "generation_summary.json"
    summary_path.write_text(
        json.dumps(
            {
                "job_code": job_code,
                "version": M02_VERSION,
                "company": context["company"].get("company_name", ""),
                "generated_files": generated,
                "transfer_in_mode": context["m02"].get("mode_label", ""),
                "existing_service_provider": context["m02"].get("old_secretary_company", ""),
                "new_service_provider": context["m02"].get("new_secretary_company", ""),
                "member_signers": [row.get("full_name", "") for row in context["m02"].get("member_signers", [])],
                "director_signers": [row.get("full_name", "") for row in context["m02"].get("director_signers", [])],
                "related_personnel_changes": context["m02"].get("personnel_change_lines", []),
                "resignation_letter_sections": [row.get("full_name", "") for row in context["m02"].get("resigning_persons", [])],
                "notes": [
                    "M02 is the P2 transfer-in package.",
                    "The output is grouped into two signing PDFs: a resolution package and a handover/resignation package.",
                    "The new service provider is described as corporate secretarial service provider, not as the individual statutory company secretary.",
                    "Optional resignation letter sections are appended only when resigning persons are identified from personnel action rows and resignation letters are requested.",
                ],
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    zip_path = GENERATED_DIR / f"{safe_filename(job_code)}_P2_M02_docx_package.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(package_dir.iterdir()):
            if path.is_file():
                zf.write(path, arcname=path.name)
    return zip_path


def append_m02_resignation_letters(docx_path: Path, context: dict[str, Any]) -> None:
    resigning_people = context.get("m02", {}).get("resigning_persons", [])
    if not resigning_people:
        return

    doc = Document(docx_path)
    company = context.get("company", {})
    for person in resigning_people:
        doc.add_page_break()
        add_m02_appended_paragraph(doc, clean(person.get("full_name")), bold=True, after=2)
        address = clean(person.get("address") or person.get("residential_address"))
        if address:
            add_m02_appended_paragraph(doc, address, bold=True, after=12)
        resignation_date = date_text(person.get("effective_date_raw") or person.get("effective_date"))
        add_m02_appended_paragraph(doc, resignation_date, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, after=14)
        add_m02_appended_paragraph(doc, "The Board of Directors", after=2)
        add_m02_appended_paragraph(doc, clean(company.get("company_name")), bold=True, after=2)
        add_m02_appended_paragraph(doc, f"Registration No. {bold_field(clean(company.get('uen')))}", after=12)
        add_m02_appended_paragraph(doc, "Dear Sirs", after=10)
        capacity = clean(person.get("capacity")) or "Officer"
        add_m02_appended_paragraph(doc, f"RE: NOTICE OF RESIGNATION AS {capacity.upper()}", bold=True, after=8)
        add_m02_appended_paragraph(
            doc,
            f"I, {bold_field(clean(person.get('full_name')))}, hereby tender my resignation as {bold_field(capacity)} of {bold_field(clean(company.get('company_name')))} (the \"Company\") with effect from {bold_field(resignation_date)}.",
        )
        add_m02_appended_paragraph(
            doc,
            f"I confirm that I have no claim against the Company in respect of my resignation as {bold_field(capacity)}, except for any accrued rights or obligations which cannot be waived by law.",
        )
        add_m02_appended_paragraph(doc, "Please arrange for the Company's statutory records and any relevant lodgements to be updated accordingly.", after=18)
        add_m02_appended_paragraph(doc, "Yours faithfully,", after=28)
        add_m02_appended_signature(doc, person, after=12)
        add_m02_appended_paragraph(doc, f"Name: {bold_field(clean(person.get('full_name')))}", after=2)
    doc.save(docx_path)


def add_m02_appended_signature(doc: Document, person: dict[str, Any], *, after: float = 12) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.10
    label = paragraph.add_run("Signature: ")
    label.font.name = "Calibri"
    label._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    label.font.size = Pt(10.8)
    signature_path = resolve_signature_path(person.get("signature_image_path"))
    if signature_path:
        paragraph.add_run().add_picture(str(signature_path), width=Inches(SIGNATURE_DOCX_WIDTH_IN))
    else:
        fallback = paragraph.add_run("______________________________")
        fallback.font.name = "Calibri"
        fallback._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        fallback.font.size = Pt(10.8)


def add_m02_appended_paragraph(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH | None = None,
    before: float = 0,
    after: float = 6,
    size: float = 10.8,
) -> None:
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.10
    segments: list[tuple[str, bool]] = []
    append_marked_segments(segments, text, bold)
    for part, is_bold in merge_text_segments(segments):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(size)
        run.font.bold = is_bold


def generate_p2_m02_pdf_package(parsed: dict[str, Any], job_code: str) -> Path:
    """Generate the P2 transfer-in package and return a PDF-only zip."""
    docx_zip = generate_p2_m02_package(parsed, job_code)
    docx_zip.unlink(missing_ok=True)
    code = safe_filename(job_code)
    return build_pdf_zip_from_docx_dir(
        GENERATED_DIR / f"{code}_P2_M02_docs",
        GENERATED_DIR / f"{code}_P2_M02_pdf",
        GENERATED_DIR / f"{code}_P2_M02_pdf_package.zip",
    )


def generate_p2_m03_package(parsed: dict[str, Any], job_code: str) -> Path:
    """Generate the P2 M03 share transfer DOCX package."""
    if parsed.get("task_type") != "maintenance":
        raise ValueError("M03 generation only supports maintenance / P2 Excel files.")
    missing_templates = [
        name for name in M03_TEMPLATE_NAMES.values()
        if not (P2_TEMPLATE_DIR / name).exists()
    ]
    if missing_templates:
        raise FileNotFoundError(f"M03 template(s) not found: {', '.join(missing_templates)}")

    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    package_dir = GENERATED_DIR / f"{safe_filename(job_code)}_P2_M03_docs"
    package_dir.mkdir(parents=True, exist_ok=True)
    for old_path in package_dir.iterdir():
        if old_path.is_file():
            old_path.unlink()

    context = build_m03_context(parsed)
    if not m03_has_content(context["m03"]):
        raise ValueError("This task has no M03 share transfer content.")

    company_name = safe_filename(context["company"].get("company_name") or "company")
    generated: list[str] = []
    base_outputs = [
        ("resolution", f"01_M03_share_transfer_directors_resolution_{company_name}.docx"),
        ("instrument", f"02_M03_instrument_of_transfer_{company_name}.docx"),
    ]
    for template_key, output_name in base_outputs:
        render_m01_docx(P2_TEMPLATE_DIR / M03_TEMPLATE_NAMES[template_key], package_dir / output_name, context)
        generated.append(output_name)

    for idx, certificate in enumerate(context["m03"].get("certificates", []), start=1):
        holder_name = safe_filename(certificate.get("holder_name") or f"shareholder_{idx:02d}")
        output_name = f"03_M03_updated_share_certificate_{idx:02d}_{holder_name}.docx"
        render_m01_docx(
            P2_TEMPLATE_DIR / M03_TEMPLATE_NAMES["certificate"],
            package_dir / output_name,
            context_with(context, certificate=certificate),
        )
        generated.append(output_name)

    register_name = f"04_M03_register_of_members_update_{company_name}.docx"
    render_m01_docx(P2_TEMPLATE_DIR / M03_TEMPLATE_NAMES["register"], package_dir / register_name, context)
    generated.append(register_name)

    checklist_name = f"05_M03_register_and_stamp_duty_checklist_{company_name}.docx"
    render_m01_docx(P2_TEMPLATE_DIR / M03_TEMPLATE_NAMES["checklist"], package_dir / checklist_name, context)
    generated.append(checklist_name)

    summary_path = package_dir / "generation_summary.json"
    summary_path.write_text(
        json.dumps(
            {
                "job_code": job_code,
                "version": M03_VERSION,
                "company": context["company"].get("company_name", ""),
                "generated_files": generated,
                "share_transfer_count": len(context["m03"].get("transfers", [])),
                "certificate_count": len(context["m03"].get("certificates", [])),
                "director_signers": [row.get("full_name", "") for row in context["m03"].get("director_signers", [])],
                "manual_review_flags": context["m03"].get("manual_review_flags", []),
                "notes": [
                    "M03 is the P2 ordinary share transfer package.",
                    "Ordinary share transfers do not generate Form 24; Form 24 remains reserved for allotment / capital increase workflows.",
                    "Stamp duty, NAV and actual consideration remain manual review points.",
                    "Updated share certificate numbers should be checked before signing and release.",
                ],
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    zip_path = GENERATED_DIR / f"{safe_filename(job_code)}_P2_M03_docx_package.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(package_dir.iterdir()):
            if path.is_file():
                zf.write(path, arcname=path.name)
    return zip_path


def generate_p2_m03_pdf_package(parsed: dict[str, Any], job_code: str) -> Path:
    """Generate the P2 share transfer package and return a PDF-only zip."""
    docx_zip = generate_p2_m03_package(parsed, job_code)
    docx_zip.unlink(missing_ok=True)
    code = safe_filename(job_code)
    return build_pdf_zip_from_docx_dir(
        GENERATED_DIR / f"{code}_P2_M03_docs",
        GENERATED_DIR / f"{code}_P2_M03_pdf",
        GENERATED_DIR / f"{code}_P2_M03_pdf_package.zip",
    )


def generate_p2_m04_package(parsed: dict[str, Any], job_code: str) -> Path:
    """Generate the P2 M04 share allotment DOCX package."""
    if parsed.get("task_type") != "maintenance":
        raise ValueError("M04 generation only supports maintenance / P2 Excel files.")
    missing_templates = [
        name for name in M04_TEMPLATE_NAMES.values()
        if not (P2_TEMPLATE_DIR / name).exists()
    ]
    if missing_templates:
        raise FileNotFoundError(f"M04 template(s) not found: {', '.join(missing_templates)}")

    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    package_dir = GENERATED_DIR / f"{safe_filename(job_code)}_P2_M04_docs"
    package_dir.mkdir(parents=True, exist_ok=True)
    for old_path in package_dir.iterdir():
        if old_path.is_file():
            old_path.unlink()

    context = build_m04_context(parsed)
    if not m04_has_content(context["m04"]):
        raise ValueError("This task has no M04 share allotment content.")

    company_name = safe_filename(context["company"].get("company_name") or "company")
    generated: list[str] = []
    base_outputs = [
        ("authority", f"01_M04_s161_members_authority_{company_name}.docx"),
        ("resolution", f"02_M04_allotment_directors_resolution_{company_name}.docx"),
    ]
    for template_key, output_name in base_outputs:
        render_m01_docx(P2_TEMPLATE_DIR / M04_TEMPLATE_NAMES[template_key], package_dir / output_name, context)
        generated.append(output_name)

    for idx, allotment in enumerate(context["m04"].get("allotments", []), start=1):
        allottee_name = safe_filename(allotment.get("allottee_name") or f"allottee_{idx:02d}")
        output_name = f"03_M04_share_application_{idx:02d}_{allottee_name}.docx"
        render_m01_docx(
            P2_TEMPLATE_DIR / M04_TEMPLATE_NAMES["application"],
            package_dir / output_name,
            context_with(context, allotment=allotment),
        )
        generated.append(output_name)

    for idx, certificate in enumerate(context["m04"].get("certificates", []), start=1):
        holder_name = safe_filename(certificate.get("holder_name") or f"shareholder_{idx:02d}")
        output_name = f"04_M04_share_certificate_{idx:02d}_{holder_name}.docx"
        render_m01_docx(
            P2_TEMPLATE_DIR / M04_TEMPLATE_NAMES["certificate"],
            package_dir / output_name,
            context_with(context, certificate=certificate),
        )
        generated.append(output_name)

    form24_name = f"05_M04_return_of_allotment_form24_{company_name}.docx"
    render_m01_docx(P2_TEMPLATE_DIR / M04_TEMPLATE_NAMES["form24"], package_dir / form24_name, context)
    generated.append(form24_name)

    register_name = f"06_M04_register_of_members_update_{company_name}.docx"
    render_m01_docx(P2_TEMPLATE_DIR / M04_TEMPLATE_NAMES["register"], package_dir / register_name, context)
    generated.append(register_name)

    checklist_name = f"07_M04_register_update_checklist_{company_name}.docx"
    render_m01_docx(P2_TEMPLATE_DIR / M04_TEMPLATE_NAMES["checklist"], package_dir / checklist_name, context)
    generated.append(checklist_name)

    summary_path = package_dir / "generation_summary.json"
    summary_path.write_text(
        json.dumps(
            {
                "job_code": job_code,
                "version": M04_VERSION,
                "company": context["company"].get("company_name", ""),
                "generated_files": generated,
                "allotment_count": len(context["m04"].get("allotments", [])),
                "certificate_count": len(context["m04"].get("certificates", [])),
                "member_signers": [row.get("full_name", "") for row in context["m04"].get("member_signers", [])],
                "director_signers": [row.get("full_name", "") for row in context["m04"].get("director_signers", [])],
                "manual_review_flags": context["m04"].get("manual_review_flags", []),
                "notes": [
                    "M04 is the P2 share allotment / capital increase package.",
                    "Form 24 is generated here, not in ordinary share transfer M03.",
                    "ACRA / BizFile lodgement and final capital position remain manual review points.",
                    "Share certificate numbers and paid-up status should be checked before signing and release.",
                ],
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    zip_path = GENERATED_DIR / f"{safe_filename(job_code)}_P2_M04_docx_package.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(package_dir.iterdir()):
            if path.is_file():
                zf.write(path, arcname=path.name)
    return zip_path


def generate_p2_m04_pdf_package(parsed: dict[str, Any], job_code: str) -> Path:
    """Generate the P2 share allotment package and return a PDF-only zip."""
    docx_zip = generate_p2_m04_package(parsed, job_code)
    docx_zip.unlink(missing_ok=True)
    code = safe_filename(job_code)
    return build_m04_pdf_zip_from_docx_dir(
        GENERATED_DIR / f"{code}_P2_M04_docs",
        GENERATED_DIR / f"{code}_P2_M04_pdf",
        GENERATED_DIR / f"{code}_P2_M04_pdf_package.zip",
    )


def generate_p2_m05_package(parsed: dict[str, Any], job_code: str) -> Path:
    """Generate the P2 M05 annual review DOCX package."""
    if parsed.get("task_type") != "maintenance":
        raise ValueError("M05 generation only supports maintenance / P2 Excel files.")
    missing_templates = [
        name for name in M05_TEMPLATE_NAMES.values()
        if not (P2_TEMPLATE_DIR / name).exists()
    ]
    if missing_templates:
        raise FileNotFoundError(f"M05 template(s) not found: {', '.join(missing_templates)}")

    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    package_dir = GENERATED_DIR / f"{safe_filename(job_code)}_P2_M05_docs"
    package_dir.mkdir(parents=True, exist_ok=True)
    for old_path in package_dir.iterdir():
        if old_path.is_file():
            old_path.unlink()

    context = build_m05_context(parsed)
    if not m05_has_content(context["m05"]):
        raise ValueError("This task has no M05 annual review content.")

    company_name = safe_filename(context["company"].get("company_name") or "company")
    generated: list[str] = []
    outputs = [
        ("agm_package", f"01_M05_agm_documents_package_{company_name}.docx"),
        ("annual_return_package", f"02_M05_annual_return_authorisation_package_{company_name}.docx"),
        ("checklist", f"99_M05_annual_review_internal_checklist_{company_name}.docx"),
    ]
    for template_key, output_name in outputs:
        render_m01_docx(P2_TEMPLATE_DIR / M05_TEMPLATE_NAMES[template_key], package_dir / output_name, context)
        generated.append(output_name)

    summary_path = package_dir / "generation_summary.json"
    summary_path.write_text(
        json.dumps(
            {
                "job_code": job_code,
                "version": M05_VERSION,
                "company": context["company"].get("company_name", ""),
                "generated_files": generated,
                "fye_date": context["m05"].get("fye_date", ""),
                "agm_date": context["m05"].get("agm_date", ""),
                "accounts_status": context["m05"].get("accounts_status_label", ""),
                "agm_mode": context["m05"].get("agm_mode_label", ""),
                "financial_statements_required": "No" if context["m05"].get("no_financial_statements_required") else "Yes",
                "director_signers": [row.get("full_name", "") for row in context["m05"].get("director_signers", [])],
                "member_signers": [row.get("full_name", "") for row in context["m05"].get("member_signers", [])],
                "manual_review_flags": context["m05"].get("manual_review_flags", []),
                "notes": [
                    "M05 is the P2 annual review / AGM / Annual Return authorisation package.",
                    "The package prepares signing documents and authorisations; it does not confirm that ACRA / BizFile filing has been lodged.",
                    "Financial statements, audit exemption and Annual Return data remain manual review points.",
                ],
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    zip_path = GENERATED_DIR / f"{safe_filename(job_code)}_P2_M05_docx_package.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(package_dir.iterdir()):
            if path.is_file():
                zf.write(path, arcname=path.name)
    return zip_path


def generate_p2_m05_pdf_package(parsed: dict[str, Any], job_code: str) -> Path:
    """Generate the P2 annual review package and return a PDF-only zip."""
    docx_zip = generate_p2_m05_package(parsed, job_code)
    docx_zip.unlink(missing_ok=True)
    code = safe_filename(job_code)
    return build_m05_pdf_zip_from_docx_dir(
        GENERATED_DIR / f"{code}_P2_M05_docs",
        GENERATED_DIR / f"{code}_P2_M05_pdf",
        GENERATED_DIR / f"{code}_P2_M05_pdf_package.zip",
    )


def build_pdf_zip_from_docx_dir(docx_dir: Path, pdf_dir: Path, zip_path: Path) -> Path:
    if not docx_dir.exists():
        raise FileNotFoundError(f"Generated DOCX folder not found: {docx_dir}")
    pdf_dir.mkdir(parents=True, exist_ok=True)
    for old_path in pdf_dir.iterdir():
        if old_path.is_file():
            old_path.unlink()
    pdf_paths: list[Path] = []
    for docx_path in sorted(docx_dir.glob("*.docx")):
        pdf_paths.append(convert_docx_to_pdf(docx_path, pdf_dir))
    if not pdf_paths:
        raise ValueError("No DOCX files were generated for PDF conversion.")
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for pdf_path in pdf_paths:
            zf.write(pdf_path, arcname=pdf_path.name)
    return zip_path


def build_m04_pdf_zip_from_docx_dir(docx_dir: Path, pdf_dir: Path, zip_path: Path) -> Path:
    if not docx_dir.exists():
        raise FileNotFoundError(f"Generated DOCX folder not found: {docx_dir}")
    pdf_dir.mkdir(parents=True, exist_ok=True)
    final_dir = pdf_dir / "final"
    final_dir.mkdir(parents=True, exist_ok=True)
    for old_path in list(pdf_dir.glob("*.pdf")) + list(final_dir.glob("*.pdf")):
        old_path.unlink()

    converted: list[Path] = []
    for docx_path in sorted(docx_dir.glob("*.docx")):
        converted.append(convert_docx_to_pdf(docx_path, pdf_dir))
    if not converted:
        raise ValueError("No DOCX files were generated for PDF conversion.")

    final_paths: list[Path] = []

    authority = first_pdf_with_prefix(converted, "01_M04_s161_members_authority_")
    company_suffix = suffix_after_prefix(authority, "01_M04_s161_members_authority_") if authority else "company"
    if authority:
        target = final_dir / f"01_M04_member_authority_package_{company_suffix}.pdf"
        copy_pdf(authority, target)
        final_paths.append(target)

    resolution = first_pdf_with_prefix(converted, "02_M04_allotment_directors_resolution_")
    form24 = first_pdf_with_prefix(converted, "05_M04_return_of_allotment_form24_")
    if resolution or form24:
        target = final_dir / f"02_M04_directors_allotment_and_form24_package_{company_suffix}.pdf"
        merge_pdf_files([path for path in [resolution, form24] if path], target)
        final_paths.append(target)

    for source in sorted(path for path in converted if path.name.startswith("03_M04_share_application_")):
        target = final_dir / source.name
        copy_pdf(source, target)
        final_paths.append(target)

    for source in sorted(path for path in converted if path.name.startswith("04_M04_share_certificate_")):
        target = final_dir / source.name
        copy_pdf(source, target)
        final_paths.append(target)

    checklist = first_pdf_with_prefix(converted, "07_M04_register_update_checklist_")
    if checklist:
        target = final_dir / f"99_M04_internal_register_update_checklist_{suffix_after_prefix(checklist, '07_M04_register_update_checklist_')}.pdf"
        copy_pdf(checklist, target)
        final_paths.append(target)

    used_names = {path.name for path in final_paths}
    for source in converted:
        if source.name in used_names:
            continue
        if any(
            source.name.startswith(prefix)
            for prefix in [
                "01_M04_s161_members_authority_",
                "02_M04_allotment_directors_resolution_",
                "03_M04_share_application_",
                "04_M04_share_certificate_",
                "05_M04_return_of_allotment_form24_",
                "07_M04_register_update_checklist_",
            ]
        ):
            continue
        target = final_dir / source.name
        copy_pdf(source, target)
        final_paths.append(target)

    if not final_paths:
        raise ValueError("No M04 PDF files were prepared for the final package.")
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for pdf_path in sorted(final_paths):
            zf.write(pdf_path, arcname=pdf_path.name)
    return zip_path


def build_m05_pdf_zip_from_docx_dir(docx_dir: Path, pdf_dir: Path, zip_path: Path) -> Path:
    if not docx_dir.exists():
        raise FileNotFoundError(f"Generated DOCX folder not found: {docx_dir}")
    pdf_dir.mkdir(parents=True, exist_ok=True)
    final_dir = pdf_dir / "final"
    final_dir.mkdir(parents=True, exist_ok=True)
    for old_path in list(pdf_dir.glob("*.pdf")) + list(final_dir.glob("*.pdf")):
        old_path.unlink()

    converted: list[Path] = []
    for docx_path in sorted(docx_dir.glob("*.docx")):
        converted.append(convert_docx_to_pdf(docx_path, pdf_dir))
    if not converted:
        raise ValueError("No DOCX files were generated for PDF conversion.")

    final_paths: list[Path] = []

    agm_package = first_pdf_with_prefix(converted, "01_M05_agm_documents_package_")
    annual_return_package = first_pdf_with_prefix(converted, "02_M05_annual_return_authorisation_package_")
    company_suffix = (
        suffix_after_prefix(agm_package, "01_M05_agm_documents_package_")
        if agm_package
        else suffix_after_prefix(annual_return_package, "02_M05_annual_return_authorisation_package_")
    )

    signing_sources = [path for path in [agm_package, annual_return_package] if path]
    if signing_sources:
        target = final_dir / f"01_M05_annual_review_signing_package_{company_suffix}.pdf"
        merge_pdf_files(signing_sources, target)
        final_paths.append(target)

    checklist = first_pdf_with_prefix(converted, "99_M05_annual_review_internal_checklist_")
    if checklist:
        target = final_dir / f"99_M05_annual_review_internal_checklist_{suffix_after_prefix(checklist, '99_M05_annual_review_internal_checklist_')}.pdf"
        copy_pdf(checklist, target)
        final_paths.append(target)

    used_names = {path.name for path in final_paths}
    for source in converted:
        if source.name in used_names:
            continue
        if any(
            source.name.startswith(prefix)
            for prefix in [
                "01_M05_agm_documents_package_",
                "02_M05_annual_return_authorisation_package_",
                "99_M05_annual_review_internal_checklist_",
            ]
        ):
            continue
        target = final_dir / source.name
        copy_pdf(source, target)
        final_paths.append(target)

    if not final_paths:
        raise ValueError("No M05 PDF files were prepared for the final package.")
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for pdf_path in sorted(final_paths):
            zf.write(pdf_path, arcname=pdf_path.name)
    return zip_path


def first_pdf_with_prefix(paths: list[Path], prefix: str) -> Path | None:
    return next((path for path in sorted(paths) if path.name.startswith(prefix)), None)


def suffix_after_prefix(path: Path | None, prefix: str) -> str:
    if not path:
        return "company"
    return path.stem.removeprefix(prefix)


def copy_pdf(source: Path, target: Path) -> None:
    target.write_bytes(source.read_bytes())


def merge_pdf_files(source_paths: list[Path], target: Path) -> None:
    from pypdf import PdfReader, PdfWriter

    writer = PdfWriter()
    for source in source_paths:
        reader = PdfReader(str(source))
        for page in reader.pages:
            writer.add_page(page)
    with target.open("wb") as handle:
        writer.write(handle)


def build_m01_context(parsed: dict[str, Any]) -> dict[str, Any]:
    company = normalize_m01_company(parsed.get("company", {}))
    people = [normalize_m01_person(row) for row in parsed.get("people", [])]
    events = [row for row in parsed.get("change_events", []) if active_m01_event(row)]
    transfers = (
        [normalize_m01_transfer(row) for row in parsed.get("share_transfers", []) if active_m01_row(row, "generate", ["transferor_name", "transferee_name", "shares_transferred"])]
        if is_yes(company.get("include_share_transfer_in_m01"))
        else []
    )
    allotments = (
        [normalize_m01_allotment(row) for row in parsed.get("share_allotments", []) if active_m01_row(row, "generate", ["allottee_name", "shares_allotted", "total_paid"])]
        if is_yes(company.get("include_share_allotment_in_m01"))
        else []
    )

    m01: dict[str, Any] = {
        "include_registered_office": False,
        "include_office_hours": False,
        "include_business_activity": False,
        "include_fye": False,
        "include_officer_particulars": False,
        "include_director_appointments": False,
        "include_director_resignations": False,
        "include_secretary_appointments": False,
        "include_secretary_resignations": False,
        "include_share_transfer_approval": bool(transfers),
        "include_share_allotment_approval": bool(allotments),
        "include_bizfile_authorization": False,
        "registered_office_old": "",
        "registered_office_new": "",
        "registered_office_effective_date": "",
        "office_hours_new": "",
        "office_hours_effective_date": "",
        "primary_ssic_old": "",
        "primary_activity_old": "",
        "primary_ssic_new": "",
        "primary_activity_new": "",
        "secondary_ssic_old": "",
        "secondary_activity_old": "",
        "secondary_ssic_new": "",
        "secondary_activity_new": "",
        "business_activity_effective_date": "",
        "fye_old": "",
        "fye_new": "",
        "next_accounts_period_start": "",
        "next_accounts_period_end": "",
        "officer_name": "",
        "officer_particular_updates": [],
        "director_appointments": [],
        "director_resignations": [],
        "secretary_appointments": [],
        "secretary_resignations": [],
        "share_transfers": transfers,
        "share_allotments": allotments,
        "director_signers": director_signers_for_m01(people, company),
    }

    for row in events:
        event_type = clean(row.get("event_type"))
        if event_type not in M01_DR_EVENT_TYPES and not is_yes(row.get("combine_in_dr")):
            continue
        apply_m01_event(m01, row, people, company)

    m01["include_bizfile_authorization"] = m01_has_content(m01)
    if not m01["director_signers"]:
        m01["director_signers"] = [{"full_name": "", "capacity": "Director"}]

    signature_date = company.get("default_document_date") or first_filled(
        [row.get("effective_date") for row in events],
        [row.get("transfer_date") for row in transfers],
        [row.get("allotment_date") for row in allotments],
    )
    return {
        "company": company,
        "people": people,
        "provider": provider_context(),
        "signature": signature_context(signature_date),
        "m01": m01,
    }


def build_m02_context(parsed: dict[str, Any]) -> dict[str, Any]:
    company = normalize_m01_company(parsed.get("company", {}))
    people = [normalize_m01_person(row) for row in parsed.get("people", [])]
    all_events = parsed.get("change_events", [])
    events = [row for row in all_events if active_m02_event(row)]
    transfer = next((row for row in events if clean(row.get("event_type")) in M02_TRANSFER_EVENT_TYPES), {})
    registered_office_change = registered_office_change_for_m02(all_events)
    provider = provider_context()

    effective_date_raw = clean(transfer.get("effective_date") or company.get("default_document_date"))
    old_provider = clean(
        transfer.get("old_secretary_company")
        or transfer.get("old_value")
        or company.get("old_secretary_company")
        or company.get("current_secretary_company")
    ) or DEFAULT_PREVIOUS_SECRETARIAL_PROVIDER
    new_provider = clean(
        transfer.get("new_secretary_company")
        or transfer.get("new_value")
        or company.get("new_secretary_company")
        or provider["name"]
    )
    new_registered_office = clean(registered_office_change.get("new_registered_office_address") or registered_office_change.get("new_value"))
    meeting_place = clean(
        company.get("egm_meeting_place")
        or company.get("meeting_place")
        or transfer.get("meeting_place")
        or new_registered_office
        or provider["registered_address"]
    )
    mode = clean(transfer.get("event_type")) or "transfer_in_cooperative"
    member_signers = member_signers_for_m02(people, company)
    director_signers = director_signers_for_m01(people, company)
    if not director_signers:
        director_signers = [{"full_name": "", "capacity": "Director"}]
    if not member_signers:
        member_signers = client_side_signer_fallback(people, company, director_signers)

    client_signatory = client_signatory_for_m02(people, company, member_signers, director_signers)
    notice_issuer = notice_issuer_for_m02(people, company, director_signers, client_signatory)
    personnel_changes = personnel_changes_for_m02(events, people)
    resignation_people = resignation_persons_for_m02(events, people, transfer, company)

    m02: dict[str, Any] = {
        "include_transfer_in": bool(transfer),
        "mode": mode,
        "mode_label": "Transfer-in",
        "manual_review_required": "No",
        "effective_date_raw": effective_date_raw,
        "effective_date": date_text(effective_date_raw),
        "old_secretary_company": old_provider,
        "new_secretary_company": new_provider,
        "current_registered_office": company.get("registered_office_address", ""),
        "new_registered_office_address": new_registered_office,
        "meeting_place": meeting_place,
        "member_signers": member_signers,
        "director_signers": director_signers,
        "client_signatory": client_signatory,
        "notice_issuer": notice_issuer,
        "personnel_changes": personnel_changes,
        "personnel_change_lines": [row.get("summary", "") for row in personnel_changes if row.get("summary")],
        "personnel_change_summary": personnel_change_summary_for_m02(personnel_changes),
        "resigning_persons": resignation_people,
        "member_signature_blocks": signature_blocks(member_signers, "Member / Authorised Signatory", effective_date_raw),
        "director_signature_blocks": signature_blocks(director_signers, "Director", effective_date_raw),
        "client_signature_block": signature_block(client_signatory, "Authorised Signatory", effective_date_raw),
        "notice_issuer_signature_block": signature_block(notice_issuer, "Director / Authorised Signatory", effective_date_raw),
        "handover_request_items": handover_request_items(),
    }
    return {
        "company": company,
        "people": people,
        "provider": provider,
        "signature": signature_context(effective_date_raw),
        "m02": m02,
    }


def build_m03_context(parsed: dict[str, Any]) -> dict[str, Any]:
    company = normalize_m01_company(parsed.get("company", {}))
    people = [normalize_m01_person(row) for row in parsed.get("people", [])]
    raw_transfers = [
        row for row in parsed.get("share_transfers", [])
        if active_m01_row(row, "generate", ["transferor_name", "transferee_name", "shares_transferred"])
    ]
    transfers = [normalize_m03_transfer(row, company, idx) for idx, row in enumerate(raw_transfers, start=1)]
    director_signers = director_signers_for_m01(people, company)
    if not director_signers:
        director_signers = [{"full_name": "", "capacity": "Director"}]
    transfer_parties = m03_transfer_parties(transfers)
    certificates = [m03_certificate_from_transfer(row) for row in transfers if row.get("generate_new_certificate")]
    signature_date = company.get("default_document_date") or first_filled([row.get("transfer_date_raw") for row in transfers])
    manual_review_flags = m03_manual_review_flags(transfers)
    m03 = {
        "include_share_transfer": bool(transfers),
        "transfers": transfers,
        "director_signers": director_signers,
        "director_signature_rows": m03_signature_rows(director_signers, "Director"),
        "transfer_parties": transfer_parties,
        "certificates": certificates,
        "checklist_items": m03_checklist_items(transfers, certificates, manual_review_flags),
        "manual_review_flags": manual_review_flags,
    }
    return {
        "company": company,
        "people": people,
        "provider": provider_context(),
        "signature": signature_context(signature_date),
        "m03": m03,
        "register": register_context(
            company,
            m03_register_entries(transfers),
            title="REGISTER OF MEMBERS UPDATE RECORD",
            subtitle="Share transfer update record",
            transaction_type="Transfer of shares",
            effective_date=signature_date,
            note="This record summarises the member register update arising from the share transfer documents generated in this package. It should be checked against the signed transfer instrument, share certificates and final statutory records before filing or release.",
        ),
    }


def m03_has_content(m03: dict[str, Any]) -> bool:
    return bool(m03.get("transfers"))


def build_m04_context(parsed: dict[str, Any]) -> dict[str, Any]:
    company = normalize_m01_company(parsed.get("company", {}))
    people = [normalize_m01_person(row) for row in parsed.get("people", [])]
    raw_allotments = [
        row for row in parsed.get("share_allotments", [])
        if active_m01_row(row, "generate", ["allottee_name", "shares_allotted", "total_paid", "issued_share_capital"])
    ]
    allotments = [normalize_m04_allotment(row, company, idx) for idx, row in enumerate(raw_allotments, start=1)]
    director_signers = director_signers_for_m01(people, company)
    if not director_signers:
        director_signers = [{"full_name": "", "capacity": "Director"}]
    member_signers = member_signers_for_m02(people, company)
    if not member_signers:
        member_signers = client_side_signer_fallback(people, company, director_signers)
    certificates = [m04_certificate_from_allotment(row) for row in allotments if row.get("generate_certificate")]
    signature_date = company.get("default_document_date") or first_filled(
        [row.get("authority_date_raw") for row in allotments],
        [row.get("allotment_date_raw") for row in allotments],
    )
    manual_review_flags = m04_manual_review_flags(allotments, company)
    m04 = {
        "include_share_allotment": bool(allotments),
        "allotments": allotments,
        "certificates": certificates,
        "member_signers": member_signers,
        "director_signers": director_signers,
        "member_signature_rows": m03_signature_rows(member_signers, "Member / Authorised Signatory"),
        "director_signature_rows": m03_signature_rows(director_signers, "Director"),
        "authority_date": date_text(first_filled([row.get("authority_date_raw") for row in allotments], [signature_date])),
        "share_class_summary": m04_share_class_summary(allotments),
        "total_shares_allotted": format_number(sum_number(row.get("shares_allotted_raw") for row in allotments)),
        "amount_paid_per_share_summary": m04_per_share_summary(allotments, "amount_paid_per_share"),
        "amount_due_per_share_summary": m04_per_share_summary(allotments, "amount_due_per_share"),
        "post_issued_share_capital_text": m04_post_capital_text(company, allotments, "post_allotment_issued_share_capital", "issued_share_capital_raw"),
        "post_paid_up_capital_text": m04_post_capital_text(company, allotments, "post_allotment_paid_up_capital", "paid_up_share_capital_raw"),
        "capital_review_note": m04_capital_review_note(company, allotments),
        "checklist_items": m04_checklist_items(allotments, certificates, manual_review_flags),
        "manual_review_flags": manual_review_flags,
    }
    return {
        "company": company,
        "people": people,
        "provider": provider_context(),
        "signature": signature_context(signature_date),
        "m04": m04,
        "register": register_context(
            company,
            m04_register_entries(allotments),
            title="REGISTER OF MEMBERS UPDATE RECORD",
            subtitle="Share allotment update record",
            transaction_type="Allotment of shares",
            effective_date=signature_date,
            note="This record summarises the member register update arising from the share allotment documents generated in this package. It should be checked against the signed allotment authority, Form 24, share certificates and final ACRA / BizFile records before filing or release.",
        ),
    }


def build_m05_context(parsed: dict[str, Any]) -> dict[str, Any]:
    company = normalize_m01_company(parsed.get("company", {}))
    people = [normalize_m01_person(row) for row in parsed.get("people", [])]
    annual = normalize_m05_annual(parsed.get("annual_review", {}), company)

    director_signers = m05_director_signers(annual, people, company)
    if not director_signers:
        director_signers = [{"full_name": "", "capacity": "Director"}]
    member_signers = m05_member_signers(annual, people, company, director_signers)
    if not member_signers:
        member_signers = [{"full_name": "", "capacity": "Member / Authorised Signatory"}]
    ar_signer = m05_ar_signer(annual, people, director_signers, member_signers)
    chairperson = director_signers[0] if director_signers else ar_signer

    signature_date = annual.get("agm_date_raw") or company.get("default_document_date")
    manual_review_flags = m05_manual_review_flags(annual, company, director_signers, member_signers)
    m05 = {
        **annual,
        "include_annual_review": m05_annual_required(annual),
        "director_signers": director_signers,
        "member_signers": member_signers,
        "ar_signer": ar_signer,
        "ar_signer_name": ar_signer.get("full_name", ""),
        "ar_signer_capacity": ar_signer.get("capacity") or "Director / Authorised Signatory",
        "chairperson_name": chairperson.get("full_name", ""),
        "notice_issuer_name": (director_signers[0] if director_signers else ar_signer).get("full_name", ""),
        "notice_issuer_capacity": (director_signers[0] if director_signers else ar_signer).get("capacity") or "Director",
        "director_signature_rows": m03_signature_rows(director_signers, "Director"),
        "member_signature_rows": m03_signature_rows(member_signers, "Member / Authorised Signatory"),
        "attendance_rows": m05_attendance_rows(director_signers, member_signers),
        "director_names_text": names_text(director_signers),
        "member_names_text": names_text(member_signers),
        "manual_review_flags": manual_review_flags,
        "checklist_items": m05_checklist_items(annual, manual_review_flags),
    }
    return {
        "company": company,
        "people": people,
        "provider": provider_context(),
        "signature": signature_context(signature_date),
        "m05": m05,
    }


def m04_has_content(m04: dict[str, Any]) -> bool:
    return bool(m04.get("allotments"))


def m05_has_content(m05: dict[str, Any]) -> bool:
    return bool(m05.get("include_annual_review"))


def active_m02_event(row: dict[str, Any]) -> bool:
    event_type = clean(row.get("event_type"))
    if not event_type:
        return False
    if event_type in M02_TRANSFER_EVENT_TYPES:
        return active_m01_row(
            row,
            "generate",
            ["effective_date", "old_value", "new_value", "old_secretary_company", "new_secretary_company"],
        )
    if event_type in {"appoint_director", "appoint_secretary", "resign_director", "resign_secretary"}:
        return active_m01_row(row, "generate", ["target_person_id", "target_name", "effective_date", "resignation_letter"])
    return False


def registered_office_change_for_m02(events: list[dict[str, Any]]) -> dict[str, Any]:
    for row in events:
        if clean(row.get("event_type")) != "change_registered_office":
            continue
        if active_m01_row(row, "generate", ["new_registered_office_address", "new_value"]):
            return row
    return {}


def m02_has_content(m02: dict[str, Any]) -> bool:
    return bool(m02.get("include_transfer_in"))


def member_signers_for_m02(people: list[dict[str, Any]], company: dict[str, Any]) -> list[dict[str, Any]]:
    tokens = split_signer_tokens(
        company.get("member_signer_person_id")
        or company.get("member_signer_person_ids")
        or company.get("member_signer_name")
        or company.get("member_signer_names")
        or company.get("shareholder_signer_name")
        or company.get("shareholder_signer_names")
        or company.get("client_signatory_name")
        or company.get("client_signatory_person_id")
    )
    signers = signers_from_tokens(tokens, people, "Member / Authorised Signatory")
    if signers:
        return force_signer_capacity(signers, "Member")
    signers = [
        {**person, "capacity": "Member"}
        for person in people
        if is_yes(person.get("is_shareholder")) or is_yes(person.get("is_client_signatory"))
    ]
    return dedupe_signers(force_signer_capacity(signers, "Member"))


def client_side_signer_fallback(people: list[dict[str, Any]], company: dict[str, Any], director_signers: list[dict[str, Any]]) -> list[dict[str, Any]]:
    tokens = split_signer_tokens(company.get("client_signatory_name") or company.get("client_signatory_person_id"))
    signers = signers_from_tokens(tokens, people, "Authorised Signatory")
    if signers:
        return signers
    client_people = [person for person in people if is_yes(person.get("is_client_signatory"))]
    if client_people:
        return [{**client_people[0], "capacity": client_people[0].get("capacity") or "Authorised Signatory"}]
    if director_signers:
        return [{**director_signers[0], "capacity": "Authorised Signatory"}]
    return [{"full_name": "", "capacity": "Authorised Signatory"}]


def client_signatory_for_m02(
    people: list[dict[str, Any]],
    company: dict[str, Any],
    member_signers: list[dict[str, Any]],
    director_signers: list[dict[str, Any]],
) -> dict[str, Any]:
    fallback = client_side_signer_fallback(people, company, director_signers)
    signer = (member_signers or fallback or [{}])[0]
    return {**signer, "capacity": "Authorised Signatory"}


def notice_issuer_for_m02(
    people: list[dict[str, Any]],
    company: dict[str, Any],
    director_signers: list[dict[str, Any]],
    client_signatory: dict[str, Any],
) -> dict[str, Any]:
    tokens = split_signer_tokens(
        company.get("notice_issuer_name")
        or company.get("notice_issuer_person_id")
        or company.get("egm_notice_issuer_name")
        or company.get("egm_notice_issuer_person_id")
    )
    signers = signers_from_tokens(tokens, people, "Director / Authorised Signatory")
    if signers:
        return signers[0]
    if director_signers:
        return {**director_signers[0], "capacity": "Director"}
    return {**client_signatory, "capacity": client_signatory.get("capacity") or "Authorised Signatory"}


def signers_from_tokens(tokens: list[str], people: list[dict[str, Any]], default_capacity: str) -> list[dict[str, Any]]:
    signers: list[dict[str, Any]] = []
    for token in tokens:
        matched = find_person_by_reference(people, token)
        if matched:
            signers.append({**matched, "capacity": matched.get("capacity") or default_capacity})
        else:
            signers.append({"full_name": token, "capacity": default_capacity})
    return dedupe_signers(signers)


def force_signer_capacity(signers: list[dict[str, Any]], capacity: str) -> list[dict[str, Any]]:
    return [{**signer, "capacity": capacity} for signer in signers]


def dedupe_signers(signers: list[dict[str, Any]]) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    seen: set[str] = set()
    for signer in signers:
        key = (clean(signer.get("full_name")).lower(), clean(signer.get("capacity")).lower())
        if key in seen or not key[0]:
            continue
        seen.add(key)
        out.append(signer)
    return out


def dedupe_text(items: list[str]) -> list[str]:
    out: list[str] = []
    seen: set[str] = set()
    for item in items:
        key = clean(item)
        if not key or key in seen:
            continue
        seen.add(key)
        out.append(key)
    return out


def personnel_changes_for_m02(events: list[dict[str, Any]], people: list[dict[str, Any]]) -> list[dict[str, Any]]:
    changes: list[dict[str, Any]] = []
    for row in events:
        event_type = clean(row.get("event_type"))
        if event_type not in {"appoint_director", "appoint_secretary", "resign_director", "resign_secretary"}:
            continue
        if event_type.endswith("director"):
            capacity = "Director"
        else:
            capacity = "Company Secretary"
        action = "Appointment of" if event_type.startswith("appoint_") else "Resignation / cessation of"
        person = m01_target_person(row, people)
        item = person_change_row(row, person, capacity)
        item["capacity"] = capacity
        item["action"] = action
        item["summary"] = f"{action} {item.get('full_name', '')} as {capacity} with effect from {item.get('effective_date', '')}."
        if item.get("full_name"):
            changes.append(item)
    return changes


def personnel_change_summary_for_m02(changes: list[dict[str, Any]]) -> str:
    lines = []
    for idx, row in enumerate(changes, start=1):
        summary = clean(row.get("summary"))
        if not summary:
            continue
        name = clean(row.get("full_name"))
        capacity = clean(row.get("capacity"))
        effective_date = clean(row.get("effective_date"))
        if name:
            summary = summary.replace(name, bold_field(name), 1)
        if capacity:
            summary = summary.replace(f"as {capacity}", f"as {bold_field(capacity)}", 1)
        if effective_date:
            summary = summary.replace(effective_date, bold_field(effective_date), 1)
        lines.append(f"{idx}. {summary}")
    if lines:
        return "\n".join(lines)
    return "No separate director or company secretary appointment, resignation, removal or cessation item is included in this transfer-in package."


def resignation_persons_for_m02(
    events: list[dict[str, Any]],
    people: list[dict[str, Any]],
    transfer: dict[str, Any],
    company: dict[str, Any],
) -> list[dict[str, Any]]:
    transfer_requests_letters = (
        is_yes(transfer.get("resignation_letter"))
        or is_yes(transfer.get("generate_resignation_letter"))
        or is_yes(company.get("generate_resignation_letter"))
    )
    resigning: list[dict[str, Any]] = []
    for row in events:
        event_type = clean(row.get("event_type"))
        if event_type not in {"resign_director", "resign_secretary"}:
            continue
        if not (transfer_requests_letters or is_yes(row.get("resignation_letter"))):
            continue
        fallback_capacity = "Director" if event_type == "resign_director" else "Company Secretary"
        person = m01_target_person(row, people)
        item = person_change_row(row, person, fallback_capacity)
        item["capacity"] = fallback_capacity
        item["effective_date_raw"] = clean(row.get("effective_date") or transfer.get("effective_date"))
        item["effective_date"] = date_text(item["effective_date_raw"])
        if item.get("full_name"):
            resigning.append(item)
    return dedupe_signers(resigning)


def signature_blocks(signers: list[dict[str, Any]], default_capacity: str, signature_date: Any = "") -> str:
    return "\n\n".join(signature_block(signer, default_capacity, signature_date) for signer in signers) or signature_block({}, default_capacity, signature_date)


def signature_block(signer: dict[str, Any], default_capacity: str, signature_date: Any = "") -> str:
    name = clean(signer.get("full_name"))
    capacity = clean(signer.get("capacity")) or default_capacity
    date_line = date_text(signature_date) if clean(signature_date) else "______________________________"
    return "\n".join(
        [
            "Signature: ______________________________",
            f"Name: {bold_field(name)}",
            f"Capacity: {bold_field(capacity)}",
            f"Date: {bold_field(date_line) if clean(signature_date) else date_line}",
        ]
    )


def handover_request_items() -> str:
    return "\n".join(
        [
            "1. statutory registers and minute books;",
            "2. incorporation and constitutional records;",
            "3. ACRA / BizFile filing history and pending filing information;",
            "4. registers of directors, secretaries, members, controllers and nominee arrangements where maintained;",
            "5. share certificates, transfer instruments, allotment records and related registers where maintained;",
            "6. correspondence, resolutions and other corporate records relevant to the Company's statutory compliance.",
        ]
    )


def normalize_m01_company(raw: dict[str, Any]) -> dict[str, Any]:
    company = {str(k): clean(v) for k, v in raw.items()}
    company.setdefault("company_name", "")
    company.setdefault("uen", "")
    company.setdefault("registered_office_address", "")
    company.setdefault("register_location", company.get("registered_office_address", ""))
    company.setdefault("default_document_date", "")
    return company


def normalize_m01_person(raw: dict[str, Any]) -> dict[str, Any]:
    item = normalize_person(raw)
    item["person_id"] = clean(raw.get("person_id"))
    item["address"] = item.get("residential_address", "")
    item["capacity"] = person_capacity(item)
    return item


def person_capacity(person: dict[str, Any]) -> str:
    roles = []
    if is_yes(person.get("is_director")):
        roles.append("Director")
    if is_yes(person.get("is_secretary")):
        roles.append("Company Secretary")
    if is_yes(person.get("is_shareholder")):
        roles.append("Shareholder")
    return " / ".join(roles) or clean(person.get("new_roles")) or clean(person.get("current_roles")) or ""


def director_signers_for_m01(people: list[dict[str, Any]], company: dict[str, Any] | None = None) -> list[dict[str, Any]]:
    company = company or {}
    signing_mode = clean(company.get("m01_director_signing_mode") or company.get("director_signing_mode")).lower()
    if signing_mode in {"specified_signers", "specific", "manual", "override"}:
        requested = split_signer_tokens(
            company.get("director_signer_person_id")
            or company.get("director_signer_person_ids")
            or company.get("director_signer_name")
            or company.get("director_signer_names")
        )
        matched: list[dict[str, Any]] = []
        for token in requested:
            token_lower = token.lower()
            person = next(
                (
                    row
                    for row in people
                    if clean(row.get("person_id")).lower() == token_lower
                    or clean(row.get("full_name")).lower() == token_lower
                ),
                None,
            )
            matched.append({**(person or {"full_name": token}), "capacity": "Director"})
        return matched
    return [
        {**person, "capacity": "Director"}
        for person in people
        if is_yes(person.get("is_director"))
    ]


def split_signer_tokens(value: Any) -> list[str]:
    return [token.strip() for token in re.split(r"[,;/\n]+", clean(value)) if token.strip()]


def apply_m01_event(m01: dict[str, Any], row: dict[str, Any], people: list[dict[str, Any]], company: dict[str, Any]) -> None:
    event_type = clean(row.get("event_type"))
    effective_date = clean(row.get("effective_date"))
    old_value = clean(row.get("old_value"))
    new_value = clean(row.get("new_value"))
    target = m01_target_person(row, people)

    if event_type == "change_registered_office":
        m01["include_registered_office"] = True
        m01["registered_office_old"] = clean(row.get("old_registered_office_address") or old_value or company.get("registered_office_address", ""))
        m01["registered_office_new"] = clean(row.get("new_registered_office_address") or new_value)
        m01["registered_office_effective_date"] = date_text(effective_date)
    elif event_type == "change_office_hours":
        m01["include_office_hours"] = True
        m01["office_hours_new"] = clean(row.get("new_office_hours") or new_value)
        m01["office_hours_effective_date"] = date_text(effective_date)
    elif event_type == "change_business_activity":
        m01["include_business_activity"] = True
        m01["primary_ssic_old"] = clean(row.get("primary_ssic_old") or row.get("old_primary_ssic"))
        m01["primary_activity_old"] = clean(row.get("primary_activity_old") or row.get("old_primary_activity") or old_value)
        m01["primary_ssic_new"] = clean(row.get("primary_ssic_new") or row.get("new_primary_ssic"))
        m01["primary_activity_new"] = clean(row.get("primary_activity_new") or row.get("new_primary_activity") or new_value)
        m01["secondary_ssic_old"] = clean(row.get("secondary_ssic_old") or row.get("old_secondary_ssic"))
        m01["secondary_activity_old"] = clean(row.get("secondary_activity_old") or row.get("old_secondary_activity"))
        m01["secondary_ssic_new"] = clean(row.get("secondary_ssic_new") or row.get("new_secondary_ssic"))
        m01["secondary_activity_new"] = clean(row.get("secondary_activity_new") or row.get("new_secondary_activity"))
        m01["business_activity_effective_date"] = date_text(effective_date)
    elif event_type == "change_fye":
        m01["include_fye"] = True
        m01["fye_old"] = old_value
        m01["fye_new"] = new_value
        m01["next_accounts_period_start"] = clean(row.get("next_accounts_period_start") or "")
        m01["next_accounts_period_end"] = clean(row.get("next_accounts_period_end") or new_value)
    elif event_type in {"update_officer_particulars", "update_director_particulars", "update_secretary_particulars", "update_shareholder_particulars"}:
        m01["include_officer_particulars"] = True
        target_name = target.get("full_name") or clean(row.get("target_name"))
        officer_names = m01.setdefault("_officer_names", [])
        if target_name and target_name not in officer_names:
            officer_names.append(target_name)
        if len(officer_names) > 1:
            m01["officer_name"] = "the following person(s)"
        elif not m01["officer_name"]:
            m01["officer_name"] = target_name
        field_label = clean(row.get("field_label")) or clean(row.get("change_field")) or clean(row.get("remarks")) or clean(row.get("event_name_cn")) or "Particulars"
        m01["officer_particular_updates"].append(
            {
                "field_label": field_label,
                "_target_name": target_name,
                "_field_label": field_label,
                "old_value": old_value,
                "new_value": new_value,
            }
        )
        refresh_m01_officer_particular_labels(m01)
    elif event_type == "appoint_director":
        m01["include_director_appointments"] = True
        m01["director_appointments"].append(person_change_row(row, target, "Director", company.get("default_document_date") or today_text()))
    elif event_type == "resign_director":
        m01["include_director_resignations"] = True
        m01["director_resignations"].append(person_change_row(row, target, "Director", company.get("default_document_date") or today_text()))
    elif event_type == "appoint_secretary":
        m01["include_secretary_appointments"] = True
        m01["secretary_appointments"].append(person_change_row(row, target, "Company Secretary", company.get("default_document_date") or today_text()))
    elif event_type == "resign_secretary":
        m01["include_secretary_resignations"] = True
        m01["secretary_resignations"].append(person_change_row(row, target, "Company Secretary", company.get("default_document_date") or today_text()))
    elif event_type == "bizfile_authorization":
        m01["include_bizfile_authorization"] = True


def m01_target_person(row: dict[str, Any], people: list[dict[str, Any]]) -> dict[str, Any]:
    target_id = clean(row.get("target_person_id")).lower()
    target_name = clean(row.get("target_name")).lower()
    matches: list[dict[str, Any]] = []
    for person in people:
        if target_id and clean(person.get("person_id")).lower() == target_id:
            matches.append(person)
        if target_name and clean(person.get("full_name")).lower() == target_name:
            matches.append(person)
    if not matches:
        return {}
    return max(matches, key=person_match_priority)


def person_match_priority(person: dict[str, Any]) -> tuple[int, int, int, int]:
    return (
        1 if clean(person.get("signature_image_path")) else 0,
        1 if clean(person.get("source")).lower() == "common" else 0,
        1 if clean(person.get("common_person_name") or person.get("common_person_display_name")) else 0,
        len([value for value in person.values() if clean(value)]),
    )


def refresh_m01_officer_particular_labels(m01: dict[str, Any]) -> None:
    names = m01.get("_officer_names", [])
    multi_officer = len(names) > 1
    for item in m01.get("officer_particular_updates", []):
        base = clean(item.get("_field_label") or item.get("field_label"))
        target_name = clean(item.get("_target_name"))
        item["field_label"] = f"{target_name} - {base}" if multi_officer and target_name else base


def person_change_row(row: dict[str, Any], person: dict[str, Any], fallback_capacity: str, default_date: Any = "") -> dict[str, Any]:
    effective_date_raw = clean(row.get("effective_date") or default_date)
    address = (
        person.get("address")
        or person.get("residential_address")
        or clean(row.get("target_address") or row.get("residential_address") or row.get("address"))
    )
    item = dict(person)
    item.update(
        {
            "full_name": person.get("full_name") or clean(row.get("target_name")),
            "id_type": person.get("id_type") or clean(row.get("target_id_type") or row.get("id_type")),
            "id_number": person.get("id_number") or clean(row.get("target_id_number") or row.get("id_number")),
            "nationality": person.get("nationality") or clean(row.get("target_nationality") or row.get("nationality")),
            "date_of_birth": person.get("date_of_birth") or clean(row.get("target_date_of_birth") or row.get("date_of_birth")),
            "residential_address": person.get("residential_address") or address,
            "address": address,
            "email": person.get("email") or clean(row.get("target_email") or row.get("email")),
            "phone": person.get("phone") or clean(row.get("target_phone") or row.get("phone")),
            "effective_date_raw": effective_date_raw,
            "effective_date": date_text(effective_date_raw),
            "effective_date_short": short_date_text(effective_date_raw),
            "appointment_date": date_text(effective_date_raw),
            "appointment_date_short": short_date_text(effective_date_raw),
            "consent_date": date_text(effective_date_raw),
            "consent_date_short": short_date_text(effective_date_raw),
            "document_date": date_text(effective_date_raw),
            "document_date_short": short_date_text(effective_date_raw),
            "signature_date": date_text(effective_date_raw),
            "signature_date_short": short_date_text(effective_date_raw),
            "capacity": person.get("capacity") or fallback_capacity,
        }
    )
    return item


def appointment_consent_person(person: dict[str, Any], fallback_capacity: str) -> dict[str, Any]:
    item = dict(person)
    date_raw = clean(item.get("effective_date_raw") or item.get("appointment_date_raw"))
    date_long = clean(item.get("appointment_date") or item.get("effective_date") or date_text(date_raw))
    date_short = clean(item.get("appointment_date_short") or item.get("effective_date_short") or short_date_text(date_raw))
    if date_long:
        for key in ["appointment_date", "effective_date", "consent_date", "document_date", "signature_date"]:
            item[key] = date_long
        for key in ["appointment_date_short", "effective_date_short", "consent_date_short", "document_date_short", "signature_date_short"]:
            item[key] = date_short
    item["residential_address"] = item.get("residential_address") or item.get("address", "")
    item["capacity"] = item.get("capacity") or fallback_capacity
    return item


def normalize_m01_transfer(row: dict[str, Any]) -> dict[str, Any]:
    return {
        "transferor_name": clean(row.get("transferor_name") or row.get("transferor_shareholder_id")),
        "transferee_name": clean(row.get("transferee_name") or row.get("transferee_shareholder_id")),
        "share_class": clean(row.get("share_class")) or "Ordinary",
        "shares_transferred": format_number(row.get("shares_transferred")),
        "transfer_date": date_text(row.get("transfer_date")),
        "certificate_reference": certificate_reference(row),
    }


def normalize_m03_transfer(row: dict[str, Any], company: dict[str, Any], index: int) -> dict[str, Any]:
    currency = clean(row.get("currency") or company.get("currency") or "SGD")
    transfer_date_raw = clean(row.get("transfer_date") or company.get("default_document_date"))
    consideration_amount = clean(row.get("consideration_amount"))
    consideration_basis = clean(row.get("consideration_basis")) or "internal_paid_up_basis"
    old_certificate_no = clean(row.get("old_certificate_no"))
    new_certificate_no = clean(row.get("new_certificate_no")) or f"To be assigned ({index:02d})"
    generate_certificate_value = clean(row.get("generate_new_certificate")).lower()
    stamp_review_value = clean(row.get("stamp_duty_review")).lower()
    item = {
        "transfer_id": clean(row.get("transfer_id")) or f"TR{index:03d}",
        "transferor_name": clean(row.get("transferor_name") or row.get("transferor_shareholder_id")),
        "transferor_id_number": clean(row.get("transferor_id_number") or row.get("transferor_id_or_reg_no")),
        "transferor_address": clean(row.get("transferor_address")),
        "transferee_name": clean(row.get("transferee_name") or row.get("transferee_shareholder_id")),
        "transferee_id_number": clean(row.get("transferee_id_number") or row.get("transferee_id_or_reg_no")),
        "transferee_address": clean(row.get("transferee_address")),
        "share_class": clean(row.get("share_class")) or "Ordinary",
        "shares_transferred": format_number(row.get("shares_transferred")),
        "shares_transferred_raw": clean(row.get("shares_transferred")),
        "transfer_date_raw": transfer_date_raw,
        "transfer_date": date_text(transfer_date_raw),
        "consideration_basis": consideration_basis,
        "consideration_basis_label": consideration_basis_label(consideration_basis),
        "consideration_amount": consideration_amount,
        "currency": currency,
        "consideration_text": consideration_text(currency, consideration_amount, consideration_basis),
        "old_certificate_no": old_certificate_no,
        "new_certificate_no": new_certificate_no,
        "transferor_remaining_shares": format_number(row.get("transferor_remaining_shares")) if clean(row.get("transferor_remaining_shares")) else "",
        "generate_new_certificate": generate_certificate_value not in NO_VALUES,
        "stamp_duty_review": is_yes(row.get("stamp_duty_review")) or consideration_basis == "stamp_duty_higher_of_price_or_nav",
        "remarks": clean(row.get("remarks")),
    }
    item["certificate_reference"] = certificate_reference(item)
    item["consideration_and_certificate_text"] = m03_consideration_and_certificate_text(item)
    return item


def consideration_basis_label(value: str) -> str:
    return {
        "internal_paid_up_basis": "Internal transfer / paid-up basis",
        "acra_paid_up_capital_basis": "ACRA paid-up capital basis",
        "stamp_duty_higher_of_price_or_nav": "Stamp duty review: higher of price or NAV/share value",
    }.get(clean(value), clean(value) or "Internal transfer / paid-up basis")


def consideration_text(currency: str, amount: str, basis: str) -> str:
    if amount:
        return f"{currency} {format_money_number(amount)}"
    if basis == "stamp_duty_higher_of_price_or_nav":
        return "Subject to stamp duty / NAV review"
    return "As agreed between the parties"


def m03_consideration_and_certificate_text(row: dict[str, Any]) -> str:
    parts = [bold_field(clean(row.get("consideration_text")))]
    if clean(row.get("old_certificate_no")):
        parts.append(f"old cert: {bold_field(clean(row.get('old_certificate_no')))}")
    if clean(row.get("new_certificate_no")):
        parts.append(f"new cert: {bold_field(clean(row.get('new_certificate_no')))}")
    if clean(row.get("remarks")):
        parts.append(clean(row.get("remarks")))
    return "; ".join(part for part in parts if part)


def m03_transfer_parties(transfers: list[dict[str, Any]]) -> list[dict[str, Any]]:
    parties: list[dict[str, Any]] = []
    for row in transfers:
        parties.append(
            {
                "full_name": row.get("transferor_name", ""),
                "capacity": "Transferor",
                "id_number": row.get("transferor_id_number", ""),
                "address": row.get("transferor_address", ""),
                "id_line": id_line(row.get("transferor_id_number", "")),
                "address_line": address_line(row.get("transferor_address", "")),
            }
        )
        parties.append(
            {
                "full_name": row.get("transferee_name", ""),
                "capacity": "Transferee",
                "id_number": row.get("transferee_id_number", ""),
                "address": row.get("transferee_address", ""),
                "id_line": id_line(row.get("transferee_id_number", "")),
                "address_line": address_line(row.get("transferee_address", "")),
            }
        )
    return dedupe_signers([party for party in parties if clean(party.get("full_name"))])


def m03_register_entries(transfers: list[dict[str, Any]]) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    for row in transfers:
        transferor = clean(row.get("transferor_name"))
        old_cert = clean(row.get("old_certificate_no"))
        remarks = []
        if transferor:
            remarks.append(f"Transfer from {transferor}")
        if old_cert:
            remarks.append(f"Old certificate: {old_cert}")
        if clean(row.get("consideration_text")):
            remarks.append(clean(row.get("consideration_text")))
        entries.append(
            {
                "folio_no": clean(row.get("transfer_id")),
                "member_name": clean(row.get("transferee_name")),
                "id_number": clean(row.get("transferee_id_number")),
                "address": clean(row.get("transferee_address")),
                "shares": clean(row.get("shares_transferred")),
                "share_class": clean(row.get("share_class")) or "Ordinary",
                "certificate_no": clean(row.get("new_certificate_no")),
                "entry_date": clean(row.get("transfer_date")),
                "transaction_reference": clean(row.get("transfer_id")),
                "paid_status_text": "fully paid",
                "paid_amount_raw": "",
                "remarks": "; ".join(part for part in remarks if part),
            }
        )
    return entries


def m03_signature_rows(signers: list[dict[str, Any]], default_capacity: str) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for idx in range(0, len(signers), 2):
        left = signers[idx] if idx < len(signers) else {}
        right = signers[idx + 1] if idx + 1 < len(signers) else {}
        rows.append(
            {
                "left_block": compact_signature_block(left, default_capacity) if clean(left.get("full_name")) else "",
                "right_block": compact_signature_block(right, default_capacity) if clean(right.get("full_name")) else "",
            }
        )
    return rows or [{"left_block": compact_signature_block({}, default_capacity), "right_block": ""}]


def compact_signature_block(signer: dict[str, Any], default_capacity: str) -> str:
    name = clean(signer.get("full_name"))
    capacity = clean(signer.get("capacity")) or default_capacity
    return f"\n______________________________\n{bold_field(name)}\n{bold_field(capacity)}"


def id_line(value: Any) -> str:
    text_value = clean(value)
    return f"ID / Reg. No.: {bold_field(text_value)}" if text_value else "ID / Reg. No.: ______________________________"


def address_line(value: Any) -> str:
    text_value = clean(value)
    return f"Address: {bold_field(text_value)}" if text_value else "Address: ______________________________"


def m03_certificate_from_transfer(row: dict[str, Any]) -> dict[str, Any]:
    return {
        "holder_name": row.get("transferee_name", ""),
        "certificate_no": row.get("new_certificate_no", "") or "To be assigned",
        "share_class": row.get("share_class", "Ordinary"),
        "shares_text": row.get("shares_transferred", ""),
        "issue_date": row.get("transfer_date", ""),
        "paid_status_text": "fully paid",
        "source_transfer_id": row.get("transfer_id", ""),
        "source_reference": f"Transfer ref.: {row.get('transfer_id', '')}" if row.get("transfer_id") else "",
    }


def m03_manual_review_flags(transfers: list[dict[str, Any]]) -> list[str]:
    flags: list[str] = []
    for row in transfers:
        if row.get("stamp_duty_review") or row.get("consideration_basis") == "stamp_duty_higher_of_price_or_nav":
            flags.append(f"{row.get('transfer_id')}: stamp duty / NAV / consideration basis requires review.")
        if looks_corporate_party(row.get("transferor_name")) or looks_corporate_party(row.get("transferee_name")):
            flags.append(f"{row.get('transfer_id')}: corporate party authorisation / authorised representative should be reviewed.")
        if not clean(row.get("new_certificate_no")) or "assigned" in clean(row.get("new_certificate_no")).lower():
            flags.append(f"{row.get('transfer_id')}: new share certificate number should be checked before signing.")
    return dedupe_text(flags)


def looks_corporate_party(value: Any) -> bool:
    text_value = clean(value).lower()
    return any(token in text_value for token in ["pte", "ltd", "limited", "llp", "corp", "company"])


def m03_checklist_items(transfers: list[dict[str, Any]], certificates: list[dict[str, Any]], flags: list[str]) -> list[dict[str, Any]]:
    transfer_count = len(transfers)
    certificate_count = len(certificates)
    items = [
        {
            "item": "Proper instrument of transfer",
            "status": "Prepare / sign",
            "note": f"{transfer_count} transfer item(s) identified. Ensure transferor and transferee signatures are complete.",
        },
        {
            "item": "Share certificate cancellation / issue",
            "status": "Review",
            "note": f"{certificate_count} updated certificate(s) prepared. Check old and new certificate numbers before release.",
        },
        {
            "item": "Register of Members / EROM",
            "status": "Update after signing",
            "note": "Update statutory register only after signed documents and required review are complete.",
        },
        {
            "item": "Stamp duty / NAV / consideration",
            "status": "Manual review",
            "note": "Do not treat the generated documents as confirmation that stamp duty or NAV review is complete.",
        },
        {
            "item": "RORC / controller review",
            "status": "Check if control changes",
            "note": "If the post-transfer shareholding changes registrable controllers, prepare the relevant RORC updates separately.",
        },
    ]
    for flag in flags:
        items.append({"item": "Manual review flag", "status": "Attention", "note": flag})
    return items


def normalize_m04_allotment(row: dict[str, Any], company: dict[str, Any], index: int) -> dict[str, Any]:
    currency = clean(row.get("currency") or company.get("currency") or company.get("share_currency") or "SGD")
    share_class = clean(row.get("share_class") or company.get("share_class_default") or company.get("share_class") or "Ordinary")
    shares_raw = clean(row.get("shares_allotted"))
    shares_number = to_number(shares_raw)
    issued_raw = clean(row.get("issued_share_capital") or row.get("allotment_issued_share_capital"))
    if not issued_raw and shares_number > 0:
        issued_per_share = to_number(row.get("issued_amount_per_share")) or to_number(company.get("share_par_value")) or 1
        issued_raw = format_money_number(shares_number * issued_per_share)
    paid_raw = clean(row.get("paid_up_share_capital") or row.get("total_paid") or row.get("paid_amount"))
    if not paid_raw and clean(row.get("amount_paid_per_share")) and shares_number > 0:
        paid_raw = format_money_number(shares_number * to_number(row.get("amount_paid_per_share")))
    if not paid_raw and issued_raw:
        paid_raw = issued_raw
    unpaid = max(to_number(issued_raw) - to_number(paid_raw), 0)
    paid_per_share = per_share_amount(to_number(paid_raw), shares_number)
    due_per_share = "-" if unpaid <= 0 else per_share_amount(unpaid, shares_number)
    allotment_date_raw = clean(row.get("allotment_date") or company.get("default_document_date"))
    authority_date_raw = clean(row.get("authority_date") or company.get("default_document_date") or allotment_date_raw)
    certificate_no = clean(row.get("certificate_no") or row.get("new_certificate_no")) or f"To be assigned ({index:02d})"
    generate_certificate_value = clean(row.get("generate_certificate") or row.get("certificate_required")).lower()
    form24_value = clean(row.get("form24_required")).lower()
    item = {
        "allotment_id": clean(row.get("allotment_id")) or f"AL{index:03d}",
        "allottee_name": clean(row.get("allottee_name") or row.get("allottee_person_id")),
        "allottee_id_number": clean(row.get("allottee_id_number") or row.get("allottee_id_or_reg_no")),
        "allottee_address": clean(row.get("allottee_address") or row.get("address")),
        "nationality": clean(row.get("nationality") or row.get("country") or row.get("country_of_incorporation")),
        "share_class": share_class,
        "shares_allotted": format_number(shares_raw),
        "shares_allotted_raw": shares_raw,
        "issued_share_capital_raw": issued_raw,
        "paid_up_share_capital_raw": paid_raw,
        "issued_share_capital": format_money_number(issued_raw),
        "paid_up_share_capital": format_money_number(paid_raw),
        "issued_share_capital_text": f"{currency} {format_money_number(issued_raw)}" if issued_raw else f"{currency} -",
        "paid_up_share_capital_text": f"{currency} {format_money_number(paid_raw)}" if paid_raw else f"{currency} -",
        "amount_paid_per_share": paid_per_share,
        "amount_paid_per_share_text": f"{currency} {paid_per_share}" if paid_per_share else "",
        "amount_due_per_share": due_per_share,
        "amount_due_per_share_text": "-" if unpaid <= 0 else f"{currency} {due_per_share}",
        "unpaid_share_capital": format_money_number(unpaid),
        "paid_status_text": "fully paid" if unpaid <= 0 else f"paid up to {currency} {paid_per_share} per share, with {currency} {due_per_share} unpaid per share",
        "currency": currency,
        "allotment_date_raw": allotment_date_raw,
        "allotment_date": date_text(allotment_date_raw),
        "authority_date_raw": authority_date_raw,
        "authority_date": date_text(authority_date_raw),
        "certificate_no": certificate_no,
        "generate_certificate": generate_certificate_value not in NO_VALUES,
        "form24_required": form24_value not in NO_VALUES,
        "post_allotment_total_shares": clean(row.get("post_allotment_total_shares")),
        "post_allotment_issued_share_capital": clean(row.get("post_allotment_issued_share_capital") or row.get("post_allotment_share_capital")),
        "post_allotment_paid_up_capital": clean(row.get("post_allotment_paid_up_capital")),
        "consideration_text": clean(row.get("consideration")) or "Cash",
        "remarks": clean(row.get("remarks")),
    }
    item["allottee_form24_name_id"] = "\n".join(
        part for part in [bold_field(item["allottee_name"]), id_line(item["allottee_id_number"]) if item["allottee_id_number"] else ""] if part
    )
    item["form24_allotment_text"] = (
        f"{bold_field(item['shares_allotted'])} {bold_field(share_class)} share(s)\n"
        f"Issued share capital: {bold_field(item['issued_share_capital_text'])}\n"
        f"Paid-up share capital: {bold_field(item['paid_up_share_capital_text'])}"
    )
    if unpaid > 0:
        item["form24_allotment_text"] += f"\nUnpaid share capital: {bold_field(f'{currency} {format_money_number(unpaid)}')}"
    return item


def m04_certificate_from_allotment(row: dict[str, Any]) -> dict[str, Any]:
    share_class = row.get("share_class", "Ordinary")
    return {
        "holder_name": row.get("allottee_name", ""),
        "holder_address": row.get("allottee_address", ""),
        "certificate_no": row.get("certificate_no", "") or "To be assigned",
        "share_class": share_class,
        "shares_text": row.get("shares_allotted", ""),
        "shares_in_words": shares_in_words(row.get("shares_allotted_raw") or row.get("shares_allotted"), share_class),
        "issue_date": row.get("allotment_date", ""),
        "paid_status_text": row.get("paid_status_text", "fully paid"),
        "paid_status_line": m04_certificate_paid_status_line(row),
        "source_allotment_id": row.get("allotment_id", ""),
        "source_reference": f"Allotment ref.: {row.get('allotment_id', '')}" if row.get("allotment_id") else "",
    }


def m04_certificate_paid_status_line(row: dict[str, Any]) -> str:
    text_value = clean(row.get("paid_status_text"))
    if text_value:
        return text_value
    return "fully paid"


def m04_register_entries(allotments: list[dict[str, Any]]) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    for row in allotments:
        remarks = []
        if clean(row.get("issued_share_capital_text")):
            remarks.append(f"Issued: {clean(row.get('issued_share_capital_text'))}")
        if clean(row.get("paid_up_share_capital_text")):
            remarks.append(f"Paid-up: {clean(row.get('paid_up_share_capital_text'))}")
        if clean(row.get("remarks")):
            remarks.append(clean(row.get("remarks")))
        entries.append(
            {
                "folio_no": clean(row.get("allotment_id")),
                "member_name": clean(row.get("allottee_name")),
                "id_number": clean(row.get("allottee_id_number")),
                "address": clean(row.get("allottee_address")),
                "shares": clean(row.get("shares_allotted")),
                "share_class": clean(row.get("share_class")) or "Ordinary",
                "certificate_no": clean(row.get("certificate_no")),
                "entry_date": clean(row.get("allotment_date")),
                "transaction_reference": clean(row.get("allotment_id")),
                "paid_status_text": clean(row.get("paid_status_text")),
                "paid_amount_raw": clean(row.get("paid_up_share_capital_raw")),
                "remarks": "; ".join(part for part in remarks if part),
            }
        )
    return entries


def m04_share_class_summary(allotments: list[dict[str, Any]]) -> str:
    classes = dedupe_text([row.get("share_class", "") for row in allotments])
    return ", ".join(classes) if classes else "Ordinary"


def m04_per_share_summary(allotments: list[dict[str, Any]], key: str) -> str:
    values = dedupe_text([row.get(key, "") for row in allotments])
    if not values:
        return "-"
    currencies = dedupe_text([row.get("currency", "") for row in allotments])
    prefix = currencies[0] if len(currencies) == 1 else ""
    if len(values) == 1:
        return f"{prefix} {values[0]}".strip()
    return "See allottee list"


def m04_post_capital_text(company: dict[str, Any], allotments: list[dict[str, Any]], post_key: str, added_key: str) -> str:
    explicit = first_filled([row.get(post_key) for row in allotments])
    currency = clean(company.get("currency") or company.get("share_currency") or (allotments[0].get("currency") if allotments else "") or "SGD")
    if explicit:
        return f"{currency} {format_money_number(explicit)}"
    company_key = "issued_share_capital" if "issued" in post_key else "paid_up_capital"
    existing = to_number(company.get(company_key) or company.get(f"current_{company_key}"))
    added = sum_number(row.get(added_key) for row in allotments)
    if existing or added:
        return f"{currency} {format_money_number(existing + added)}"
    return f"{currency} -"


def m04_capital_review_note(company: dict[str, Any], allotments: list[dict[str, Any]]) -> str:
    if any(row.get("post_allotment_issued_share_capital") or row.get("post_allotment_paid_up_capital") for row in allotments):
        return "Post-allotment capital entered in the business sheet; verify against ACRA before filing."
    if clean(company.get("issued_share_capital")) or clean(company.get("paid_up_capital")):
        return "Post-allotment capital derived from company fields plus allotment rows; verify before filing."
    return "Manual review: current capital was not provided; check final ACRA share capital position before filing."


def m04_manual_review_flags(allotments: list[dict[str, Any]], company: dict[str, Any]) -> list[str]:
    flags: list[str] = []
    if not (clean(company.get("issued_share_capital")) or clean(company.get("paid_up_capital"))):
        flags.append("Current issued / paid-up share capital is blank; Form 24 capital position should be checked manually.")
    for row in allotments:
        if not clean(row.get("allottee_id_number")):
            flags.append(f"{row.get('allotment_id')}: allottee ID / registration number is blank.")
        if not clean(row.get("allottee_address")):
            flags.append(f"{row.get('allotment_id')}: allottee address is blank.")
        if row.get("amount_due_per_share") not in {"", "-"}:
            flags.append(f"{row.get('allotment_id')}: partly paid shares / unpaid capital requires review.")
        if not clean(row.get("certificate_no")) or "assigned" in clean(row.get("certificate_no")).lower():
            flags.append(f"{row.get('allotment_id')}: share certificate number should be checked before signing.")
    return dedupe_text(flags)


def m04_checklist_items(allotments: list[dict[str, Any]], certificates: list[dict[str, Any]], flags: list[str]) -> list[dict[str, Any]]:
    items = [
        {
            "item": "Section 161 / members' authority",
            "status": "Prepare / sign",
            "note": "Confirm member approval requirements and signatories before issuing shares.",
        },
        {
            "item": "Directors' allotment approval",
            "status": "Prepare / sign",
            "note": f"{len(allotments)} allotment item(s) identified.",
        },
        {
            "item": "Return of Allotment / Form 24",
            "status": "Prepare / lodge",
            "note": "Generated as signing / lodgement support; final ACRA / BizFile filing still requires review.",
        },
        {
            "item": "Share certificate issue",
            "status": "Review",
            "note": f"{len(certificates)} certificate(s) prepared. Check numbering and paid-up status before release.",
        },
        {
            "item": "Register of Members / share capital",
            "status": "Update after signing",
            "note": "Update register and share capital records after signed documents and filing review are complete.",
        },
    ]
    for flag in flags:
        items.append({"item": "Manual review flag", "status": "Attention", "note": flag})
    return items


def normalize_m05_annual(raw: dict[str, Any], company: dict[str, Any]) -> dict[str, Any]:
    annual = {str(k): clean(v) for k, v in raw.items()}
    fye_raw = clean(annual.get("fye_date") or annual.get("financial_year_end") or company.get("fye_date") or company.get("fye"))
    agm_raw = clean(annual.get("agm_date") or company.get("default_document_date"))
    if not agm_raw:
        agm_raw = today_text()
    statement_raw = clean(annual.get("financial_statement_date") or fye_raw)
    start_raw = clean(annual.get("financial_year_start")) or default_financial_year_start(fye_raw)
    ar_as_at_raw = clean(annual.get("ar_as_at_date") or agm_raw)
    agm_place = clean(annual.get("agm_place") or company.get("registered_office_address"))
    agm_time = clean(annual.get("agm_time") or "10.00 a.m.")
    currency = clean(company.get("currency") or "SGD")
    normalized = {
        "annual_review_required": clean(annual.get("annual_review_required") or company.get("annual_review_required")),
        "fye_date_raw": fye_raw,
        "fye_date": date_text(fye_raw),
        "fye_date_upper": formal_date_upper(fye_raw),
        "fye_year_label": financial_year_label(fye_raw),
        "financial_statement_date_raw": statement_raw,
        "financial_statement_date": date_text(statement_raw),
        "financial_statement_date_upper": formal_date_upper(statement_raw),
        "financial_year_start_raw": start_raw,
        "financial_year_start": date_text(start_raw),
        "financial_year_start_upper": formal_date_upper(start_raw),
        "financial_period_text": financial_period_text(start_raw, fye_raw),
        "agm_date_raw": agm_raw,
        "agm_date": date_text(agm_raw),
        "agm_date_upper": formal_date_upper(agm_raw),
        "agm_time": agm_time,
        "agm_place": agm_place,
        "agm_route": clean(annual.get("agm_route") or "ordinary_agm"),
        "company_activity_status": clean(annual.get("company_activity_status") or annual.get("activity_status") or ""),
        "accounts_status": clean(annual.get("accounts_status") or "active"),
        "financial_statements_type": clean(annual.get("financial_statements_type") or annual.get("financial_statement_type") or ""),
        "financial_statements_required": clean(annual.get("financial_statements_required") or ""),
        "audit_exemption_status": clean(annual.get("audit_exemption_status") or ""),
        "agm_status": clean(annual.get("agm_status") or ""),
        "acra_dormant_relevant_company": clean(annual.get("acra_dormant_relevant_company") or annual.get("dormant_relevant_company") or ""),
        "total_assets_under_500k": clean(annual.get("total_assets_under_500k") or annual.get("assets_under_500k") or ""),
        "iras_tax_status": clean(annual.get("iras_tax_status") or ""),
        "ar_as_at_date_raw": ar_as_at_raw,
        "ar_as_at_date": date_text(ar_as_at_raw),
        "director_signer_name": clean(annual.get("director_signer_name") or annual.get("director_signer_names")),
        "shareholder_signer_name": clean(annual.get("shareholder_signer_name") or annual.get("shareholder_signer_names")),
        "ar_authorized_signer_name": clean(annual.get("ar_authorized_signer_name") or annual.get("ar_authorized_signer_names")),
        "directors_fee": clean(annual.get("directors_fee") or "0"),
        "directors_fee_text": m05_fee_text(annual.get("directors_fee"), currency, "directors' fees"),
        "directors_remuneration": clean(annual.get("directors_remuneration") or "0"),
        "directors_remuneration_text": m05_fee_text(annual.get("directors_remuneration"), currency, "directors' remuneration"),
        "shorter_notice_consent": clean(annual.get("shorter_notice_consent") or "Auto"),
        "management_rep_letter": clean(annual.get("management_rep_letter") or "Yes"),
        "remarks": clean(annual.get("remarks")),
    }
    normalized.update(m05_status_profile(normalized))
    return normalized


def m05_status_option(value: Any) -> str:
    return clean(value).lower().replace("-", "_").replace("/", "_").replace(" ", "_")


def m05_status_profile(annual: dict[str, Any]) -> dict[str, Any]:
    accounts = m05_status_option(annual.get("accounts_status"))
    activity = m05_status_option(annual.get("company_activity_status"))
    fs_type = m05_status_option(annual.get("financial_statements_type"))
    audit_status = m05_status_option(annual.get("audit_exemption_status"))
    agm_status = m05_status_option(annual.get("agm_status"))
    agm_route = m05_status_option(annual.get("agm_route"))

    # Keep M05 routing intentionally simple for operations:
    # accounts_status is the controlling field. Other fields are only fallbacks
    # when accounts_status is blank/Auto, so advisory defaults cannot change the
    # generated document route unexpectedly.
    controlling_values = [accounts] if accounts and accounts != "auto" else [activity, fs_type, audit_status, agm_status, agm_route]
    active_values = {"active", "non_dormant", "non_dormant_company", "normal", "ordinary", "unaudited", "small_company"}
    dormant_values = {
        "dormant",
        "dormant_company",
        "dormant_no_fs",
        "dormant_no_financial_statements",
        "dormant_financial_statements_review_required",
    }
    audited_values = {"audited", "audit_required", "audited_accounts", "audited_financial_statements"}
    is_dormant = any(value in dormant_values for value in controlling_values if value)
    is_audited = any(value in audited_values for value in controlling_values if value)
    if accounts in active_values:
        is_dormant = False
        is_audited = False
    no_fs_required = is_dormant

    agm_no_meeting_statuses = {
        "dispensed_with_agm",
        "dispensed",
        "exempt_from_agm",
        "agm_exempt",
        "exempt_private_company",
        "written_resolution",
        "written_resolutions",
        "dormant_company",
        "no_agm",
    }
    explicit_agm_held = agm_status in {"held_agm", "ordinary_agm", "agm_held"} or agm_route in {"ordinary_agm", "held_agm"}
    no_agm_meeting = agm_status in agm_no_meeting_statuses or agm_route in agm_no_meeting_statuses
    if is_dormant and agm_status in {"", "auto"}:
        no_agm_meeting = True
        explicit_agm_held = False
    use_agm_meeting_documents = explicit_agm_held and not no_agm_meeting
    use_written_annual_review_documents = not use_agm_meeting_documents

    if is_dormant:
        accounts_label = "Dormant company / no financial statements required"
    elif is_audited:
        accounts_label = "Audited financial statements"
    else:
        accounts_label = "Active company / unaudited financial statements"

    if no_fs_required:
        financial_statement_display = "Not applicable - dormant company"
    else:
        financial_statement_display = annual.get("financial_statement_date") or annual.get("fye_date")
    fye_upper = annual.get("fye_date_upper") or annual.get("fye_date") or ""
    agm_date = annual.get("agm_date") or ""
    agm_time = annual.get("agm_time") or ""
    agm_place = annual.get("agm_place") or ""
    financial_period = annual.get("financial_period_text") or fye_upper
    fye_upper_bold = bold_field(fye_upper)
    agm_date_bold = bold_field(agm_date)
    agm_time_bold = bold_field(agm_time)
    agm_place_bold = bold_field(agm_place)
    financial_period_bold = bold_field(financial_period)

    if is_dormant and no_fs_required:
        board_accounts_title = "Dormant Status and Annual Return"
        board_accounts_resolution_1 = (
            "RESOLVED - That the directors note, based on the information provided and subject to final review, "
            "that the Company is dormant for the relevant financial year and is proceeding on the basis that financial statements are not required to be prepared for that year."
        )
        board_accounts_resolution_2 = (
            "RESOLVED - That any director of the Company be authorised to confirm the Company's dormant status, statutory records and Annual Return particulars "
            "for the purpose of the Annual Return and related declarations."
        )
        agm_accounts_business = f"To note the Company's dormant status and the basis on which no financial statements are required for the financial year ended {fye_upper_bold}."
        minutes_accounts_title = "Dormant Status"
        minutes_accounts_resolution = (
            f"RESOLVED - That the Company's dormant status for the financial year ended {fye_upper_bold} be and is hereby noted, "
            "and that the Annual Return particulars be reviewed and approved for filing support."
        )
        audit_title = "Dormant Company / No Financial Statements Statement"
        audit_intro = "I/We, the undermentioned director(s) of the Company, hereby state and declare that:"
        audit_a = (
            "a) based on the information provided, the Company was dormant for the relevant financial year and is proceeding on the basis "
            "that it is not required to prepare financial statements for that year;"
        )
        audit_b = "b) the Company has not carried on business or had accounting transactions requiring full financial statements for the relevant year, subject to final review;"
        audit_c = "c) the accounting and statutory records required to be kept by the Company have been maintained or reviewed for Annual Return purposes; and"
        audit_d = "d) the directors remain responsible for confirming the Company's dormant status and any applicable statutory exemption before the Annual Return is lodged."
        mrl_title = "DORMANT COMPANY REPRESENTATION"
        mrl_intro = f"We confirm, to the best of our knowledge and belief, the following matters for the financial period ended {fye_upper_bold}:"
        mrl_lines = [
            "1. The Company was dormant for the relevant financial year, subject to final review against its accounting records and bank activity.",
            "2. No business activity, income-generating transaction or material accounting transaction has been identified for the relevant year, except as disclosed to the preparer.",
            "3. The Company's statutory records, registers and accounting records have been maintained or made available for Annual Return review.",
            "4. The directors have reviewed whether the Company qualifies as a dormant company for which financial statements are not required.",
            "5. The Company has no undisclosed assets, liabilities, contingent liabilities, guarantees, commitments or subsequent events requiring disclosure.",
            "6. The Annual Return information and supporting particulars provided for filing are complete and accurate to the best of our knowledge.",
            "7. The directors understand that ACRA and IRAS dormant status should be reviewed separately where applicable.",
        ]
    elif is_audited:
        board_accounts_title = "Audited Financial Statements"
        board_accounts_resolution_1 = (
            f"RESOLVED - That the audited financial statements of the Company for the financial year ended {fye_upper_bold}, together with the auditor's report, "
            "be and are hereby approved for submission to the members of the Company."
        )
        board_accounts_resolution_2 = (
            "RESOLVED - That any director of the Company be authorised to make such amendments to the audited financial statements and related reports "
            "as may be necessary or desirable, provided that such amendments are not inconsistent with the records approved by the Board."
        )
        agm_accounts_business = f"To receive and consider the audited financial statements of the Company and the auditor's report for the financial year ended {fye_upper_bold}."
        minutes_accounts_title = "Audited Financial Statements"
        minutes_accounts_resolution = (
            f"RESOLVED - That the audited financial statements of the Company for the financial year ended {fye_upper_bold}, together with the auditor's report, "
            "be and are hereby received and adopted."
        )
        audit_title = "AUDITED ACCOUNTS REVIEW STATEMENT"
        audit_intro = "I/We, the undermentioned director(s) of the Company, hereby state and confirm that:"
        audit_a = "a) the Company is proceeding on the basis that audited financial statements are required or have been prepared for the relevant financial year;"
        audit_b = "b) the audit exemption statement under Section 205C should not be used unless the Company separately qualifies for audit exemption;"
        audit_c = "c) the auditor's report, financial statements and Annual Return particulars should be checked against the Company's records before filing; and"
        audit_d = "d) any auditor appointment, re-appointment or resignation matter should be reviewed separately where applicable."
        mrl_title = "MANAGEMENT REPRESENTATION"
        mrl_intro = f"We confirm, to the best of our knowledge and belief, the following matters for the financial period ended {fye_upper_bold}:"
        mrl_lines = [
            "1. The financial statements and supporting schedules have been prepared from the Company's accounting records and information provided to the preparer and/or auditor.",
            "2. All assets, liabilities, contingent liabilities, guarantees, commitments and material transactions have been recorded or disclosed.",
            "3. The accounting and other records required to be kept by the Company under the Companies Act 1967 have been maintained.",
            "4. Auditor-related information and any audit report should be checked before the Annual Return is lodged.",
            "5. There have been no subsequent events requiring adjustment or disclosure in the financial statements, except as already disclosed.",
            "6. The Company is able to meet its liabilities as and when they fall due, unless otherwise disclosed to the directors and preparer.",
            "7. The Annual Return information and supporting particulars provided for filing are complete and accurate to the best of our knowledge.",
        ]
    else:
        board_accounts_title = "Report of the Directors and Financial Statements"
        board_accounts_resolution_1 = (
            "RESOLVED - That the report of the directors and the financial statements of the Company for the financial year ended "
            f"{fye_upper_bold} be and are hereby approved and signed for submission to the members of the Company."
        )
        board_accounts_resolution_2 = (
            "RESOLVED - That any director of the Company be authorised to make such amendments to the report of the directors and financial statements "
            "as may be necessary or desirable, provided that such amendments are not inconsistent with the records approved by the Board."
        )
        agm_accounts_business = f"To receive and consider the report of the directors and the financial statements of the Company for the financial year ended {fye_upper_bold}."
        minutes_accounts_title = "Report of the Directors and Financial Statements"
        minutes_accounts_resolution = (
            "RESOLVED - That the report of the directors and the financial statements of the Company for the financial year ended "
            f"{fye_upper_bold} be and are hereby received and adopted."
        )
        audit_title = "STATEMENT BY A SMALL COMPANY EXEMPT FROM AUDIT REQUIREMENTS"
        audit_intro = "I/We, the undermentioned director(s) of the Company, hereby state and declare that:"
        audit_a = (
            f"a) for the financial period {financial_period_bold}, the Company qualifies, or is expected to qualify based on the information "
            "provided, as a small company under Section 205C of the Companies Act 1967 read with the Thirteenth Schedule;"
        )
        audit_b = "b) no notice has been received from any member requiring the Company to obtain an audit of its financial statements for the said financial year;"
        audit_c = "c) the accounting and other records required to be kept by the Company under Section 199 of the Companies Act 1967 have been kept; and"
        audit_d = "d) this statement is made for the purpose of the Annual Return and any related declaration required under the Companies Act 1967."
        mrl_title = "MANAGEMENT REPRESENTATION"
        mrl_intro = f"We confirm, to the best of our knowledge and belief, the following matters for the financial period ended {fye_upper_bold}:"
        mrl_lines = [
            "1. The financial statements have been prepared from the Company's accounting records and information provided to the preparer.",
            "2. All assets, liabilities, contingent liabilities, guarantees, commitments and material transactions have been recorded or disclosed.",
            "3. The accounting and other records required to be kept by the Company under the Companies Act 1967 have been maintained.",
            "4. No notice has been received from any member requiring the Company to obtain an audit of its financial statements for the relevant year.",
            "5. There have been no subsequent events requiring adjustment or disclosure in the financial statements, except as already disclosed.",
            "6. The Company is able to meet its liabilities as and when they fall due, unless otherwise disclosed to the directors and preparer.",
            "7. The Annual Return information and supporting particulars provided for filing are complete and accurate to the best of our knowledge.",
        ]

    if use_agm_meeting_documents:
        agm_mode_label = "Annual General Meeting held"
        board_meeting_resolution = (
            f"RESOLVED - That an Annual General Meeting of the Company be convened and held on {agm_date_bold} at {agm_time_bold} "
            f"at {agm_place_bold}, or at such other place, date and time as may be agreed by the directors and members."
        )
        board_documents_resolution = (
            "RESOLVED - That the notice of Annual General Meeting, agenda, agreement to shorter notice, proxy form, attendance sheet and minutes "
            "prepared in connection with the Annual General Meeting be and are hereby approved."
        )
    else:
        agm_mode_label = "AGM exempted or dispensed; written annual review"
        board_meeting_resolution = (
            "RESOLVED - That, subject to the Company satisfying the applicable statutory requirements and member approvals, the annual review matters "
            "may be dealt with by written resolutions, member consent or the applicable exemption from holding an Annual General Meeting."
        )
        board_documents_resolution = (
            "RESOLVED - That the written annual review resolutions, member consent, Annual Return authorisation and related declarations prepared for the Company "
            "be and are hereby approved."
        )

    written_resolution_title = "MEMBERS' WRITTEN RESOLUTION / CONSENT"
    written_resolution_body = (
        f"The undersigned member(s) of the Company confirm that the annual review matters for the financial year ended {fye_upper_bold} "
        "may be dealt with by written resolution, member consent or the applicable exemption from holding an Annual General Meeting, subject to final statutory review."
    )
    written_resolution_accounts = minutes_accounts_resolution.replace("RESOLVED - ", "")

    return {
        "is_dormant": is_dormant,
        "is_audited": is_audited,
        "no_financial_statements_required": no_fs_required,
        "use_agm_meeting_documents": use_agm_meeting_documents,
        "use_written_annual_review_documents": use_written_annual_review_documents,
        "accounts_status_label": accounts_label,
        "financial_statement_date_display": financial_statement_display,
        "agm_mode_label": agm_mode_label,
        "board_accounts_title": board_accounts_title,
        "board_accounts_resolution_1": board_accounts_resolution_1,
        "board_accounts_resolution_2": board_accounts_resolution_2,
        "board_meeting_resolution": board_meeting_resolution,
        "board_documents_resolution": board_documents_resolution,
        "agm_accounts_business": agm_accounts_business,
        "minutes_accounts_title": minutes_accounts_title,
        "minutes_accounts_resolution": minutes_accounts_resolution,
        "audit_statement_title": audit_title,
        "audit_statement_intro": audit_intro,
        "audit_statement_a": audit_a,
        "audit_statement_b": audit_b,
        "audit_statement_c": audit_c,
        "audit_statement_d": audit_d,
        "management_representation_title": mrl_title,
        "management_representation_intro": mrl_intro,
        "management_representation_1": mrl_lines[0],
        "management_representation_2": mrl_lines[1],
        "management_representation_3": mrl_lines[2],
        "management_representation_4": mrl_lines[3],
        "management_representation_5": mrl_lines[4],
        "management_representation_6": mrl_lines[5],
        "management_representation_7": mrl_lines[6],
        "written_resolution_title": written_resolution_title,
        "written_resolution_body": written_resolution_body,
        "written_resolution_accounts": written_resolution_accounts,
    }


def m05_annual_required(annual: dict[str, Any]) -> bool:
    flag = clean(annual.get("annual_review_required")).lower()
    if flag in NO_VALUES:
        return False
    return is_yes(annual.get("annual_review_required")) or any(clean(annual.get(key)) for key in ["fye_date_raw", "agm_date_raw"])


def m05_director_signers(annual: dict[str, Any], people: list[dict[str, Any]], company: dict[str, Any]) -> list[dict[str, Any]]:
    tokens = split_signer_tokens(annual.get("director_signer_name") or company.get("director_signer_names") or company.get("director_signer_name"))
    signers = signers_from_tokens(tokens, people, "Director")
    if signers:
        return signers
    return director_signers_for_m01(people, company)


def m05_member_signers(
    annual: dict[str, Any],
    people: list[dict[str, Any]],
    company: dict[str, Any],
    director_signers: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    tokens = split_signer_tokens(
        annual.get("shareholder_signer_name")
        or company.get("member_signer_names")
        or company.get("shareholder_signer_names")
        or company.get("client_signatory_name")
    )
    signers = signers_from_tokens(tokens, people, "Member / Authorised Signatory")
    if signers:
        return signers
    signers = member_signers_for_m02(people, company)
    if signers:
        return signers
    return client_side_signer_fallback(people, company, director_signers)


def m05_ar_signer(
    annual: dict[str, Any],
    people: list[dict[str, Any]],
    director_signers: list[dict[str, Any]],
    member_signers: list[dict[str, Any]],
) -> dict[str, Any]:
    tokens = split_signer_tokens(annual.get("ar_authorized_signer_name"))
    signers = signers_from_tokens(tokens, people, "Director / Authorised Signatory")
    if signers:
        return signers[0]
    if director_signers:
        return {**director_signers[0], "capacity": director_signers[0].get("capacity") or "Director"}
    if member_signers:
        return {**member_signers[0], "capacity": member_signers[0].get("capacity") or "Authorised Signatory"}
    return {"full_name": "", "capacity": "Director / Authorised Signatory"}


def m05_attendance_rows(director_signers: list[dict[str, Any]], member_signers: list[dict[str, Any]]) -> list[dict[str, Any]]:
    rows = dedupe_signers(
        [
            *[{**row, "capacity": row.get("capacity") or "Director"} for row in director_signers],
            *[{**row, "capacity": row.get("capacity") or "Member / Authorised Signatory"} for row in member_signers],
        ]
    )
    return rows or [{"full_name": "", "capacity": ""}]


def m05_manual_review_flags(annual: dict[str, Any], company: dict[str, Any], director_signers: list[dict[str, Any]], member_signers: list[dict[str, Any]]) -> list[str]:
    flags: list[str] = []
    if not clean(annual.get("fye_date_raw")):
        flags.append("FYE is blank; annual review documents should be checked manually.")
    if not clean(annual.get("agm_date_raw")):
        flags.append("AGM date is blank; the system used today's/document date fallback.")
    if not clean(company.get("uen")):
        flags.append("UEN is blank; Annual Return documents should be checked manually.")
    if not director_signers or not clean(director_signers[0].get("full_name")):
        flags.append("Director signer is blank.")
    if not member_signers or not clean(member_signers[0].get("full_name")):
        flags.append("Member/shareholder signer is blank.")
    if annual.get("is_dormant"):
        flags.append("Dormant company status selected; confirm ACRA dormant relevant company status and IRAS dormant/tax waiver status separately.")
        if m05_status_option(annual.get("total_assets_under_500k")) not in {"yes", "auto"}:
            flags.append("Dormant AGM exemption may depend on asset threshold; confirm whether total assets are not more than S$500,000.")
    if annual.get("is_audited"):
        flags.append("Audited financial statements selected; check auditor information and do not use the small company audit exemption statement unless separately confirmed.")
    if clean(annual.get("audit_exemption_status")) and m05_status_option(annual.get("audit_exemption_status")) in {"manual", "manual_review"}:
        flags.append("Audit exemption status is marked manual review.")
    if clean(annual.get("agm_status")) and m05_status_option(annual.get("agm_status")) in {"manual", "manual_review"}:
        flags.append("AGM status is marked manual review.")
    return dedupe_text(flags)


def m05_checklist_items(annual: dict[str, Any], flags: list[str]) -> list[dict[str, Any]]:
    items = [
        {"item": "FYE and financial statements", "status": "Review", "note": f"FYE: {annual.get('fye_date') or '-'}; financial statements: {annual.get('financial_statement_date_display') or '-'}."},
        {"item": "Accounts / activity status", "status": "Review", "note": annual.get("accounts_status_label") or annual.get("accounts_status") or "-"},
        {"item": "Annual review method", "status": "Review", "note": annual.get("agm_mode_label") or annual.get("agm_route") or "-"},
        {"item": "AGM documents", "status": "Prepare / sign", "note": f"AGM date/time: {annual.get('agm_date') or '-'} at {annual.get('agm_time') or '-'}." if annual.get("use_agm_meeting_documents") else "AGM meeting documents are replaced by written annual review / member consent documents."},
        {"item": "Shorter notice consent", "status": "Prepare / sign" if annual.get("use_agm_meeting_documents") else "Not applicable / review", "note": "Generated as part of the AGM package." if annual.get("use_agm_meeting_documents") else "Not generated when AGM is exempted, dispensed or replaced by written annual review documents."},
        {"item": "Annual Return authorisation", "status": "Prepare / sign", "note": "Authorises preparation/lodgement support; final ACRA / BizFile filing requires manual review."},
        {"item": "Audit / dormant statement", "status": "Review", "note": annual.get("audit_statement_title") or "Confirm audit exemption or dormant company status against statutory records."},
        {"item": "Management representation", "status": "Prepare / sign", "note": f"Setting: {annual.get('management_rep_letter') or 'Yes'}."},
    ]
    for flag in flags:
        items.append({"item": "Manual review flag", "status": "Attention", "note": flag})
    return items


def m05_fee_text(value: Any, currency: str, label: str) -> str:
    amount = to_number(value)
    if amount <= 0:
        return f"It was noted that there were no {label} payable for the financial year."
    return f"It was resolved that {label} of {currency} {format_money_number(amount)} be and are hereby approved for the financial year."


def default_financial_year_start(fye_value: Any) -> str:
    try:
        fye = parse_date(clean(fye_value))
    except ValueError:
        return ""
    start = date(fye.year - 1, fye.month, fye.day) + timedelta(days=1)
    return start.strftime("%d/%m/%Y")


def financial_period_text(start_value: Any, end_value: Any) -> str:
    start = formal_date_upper(start_value)
    end = formal_date_upper(end_value)
    if start and end:
        return f"{start} to {end}"
    return end or start


def formal_date_upper(value: Any) -> str:
    raw = clean(value)
    if not raw:
        return ""
    try:
        dt = parse_date(raw)
        return dt.strftime("%d %B %Y").upper()
    except ValueError:
        return raw.upper()


def financial_year_label(value: Any) -> str:
    try:
        return str(parse_date(clean(value)).year)
    except ValueError:
        return clean(value)[-4:] if clean(value) else ""


def names_text(signers: list[dict[str, Any]]) -> str:
    return ", ".join(row.get("full_name", "") for row in signers if clean(row.get("full_name")))


def normalize_m01_allotment(row: dict[str, Any]) -> dict[str, Any]:
    total_paid = clean(row.get("total_paid"))
    if not total_paid:
        total_paid = format_money_number(to_number(row.get("shares_allotted")) * to_number(row.get("amount_paid_per_share")))
    return {
        "allottee_name": clean(row.get("allottee_name") or row.get("allottee_person_id")),
        "share_class": clean(row.get("share_class")) or "Ordinary",
        "shares_allotted": format_number(row.get("shares_allotted")),
        "amount_paid_per_share": format_money_number(row.get("amount_paid_per_share")),
        "total_paid": (clean(row.get("currency")) or "SGD") + " " + total_paid if total_paid else "",
        "allotment_date": date_text(row.get("allotment_date")),
    }


def certificate_reference(row: dict[str, Any]) -> str:
    refs = []
    for key, label in [("old_certificate_no", "old cert"), ("new_certificate_no", "new cert")]:
        value = clean(row.get(key))
        if value:
            refs.append(f"{label}: {value}")
    remarks = clean(row.get("remarks"))
    if remarks:
        refs.append(remarks)
    return "; ".join(refs)


def m01_has_content(m01: dict[str, Any]) -> bool:
    return any(
        bool(m01.get(key))
        for key in [
            "include_registered_office",
            "include_office_hours",
            "include_business_activity",
            "include_fye",
            "include_officer_particulars",
            "include_director_appointments",
            "include_director_resignations",
            "include_secretary_appointments",
            "include_secretary_resignations",
            "include_share_transfer_approval",
            "include_share_allotment_approval",
        ]
    )


def m01_section_names(m01: dict[str, Any]) -> list[str]:
    labels = [
        ("include_registered_office", "Change of Registered Office"),
        ("include_office_hours", "Opening Hours of Registered Office"),
        ("include_business_activity", "Change of Business Activities"),
        ("include_fye", "Change of Financial Year End"),
        ("include_officer_particulars", "Update of Officer's Particulars"),
        ("include_director_appointments", "Appointment of Director"),
        ("include_director_resignations", "Resignation of Director"),
        ("include_secretary_appointments", "Appointment of Company Secretary"),
        ("include_secretary_resignations", "Resignation of Company Secretary"),
        ("include_share_transfer_approval", "Approval of Transfer of Shares"),
        ("include_share_allotment_approval", "Allotment of Shares"),
        ("include_bizfile_authorization", "ACRA / BizFile Lodgement"),
    ]
    return [label for key, label in labels if m01.get(key)]


def active_m01_event(row: dict[str, Any]) -> bool:
    if not clean(row.get("event_type")):
        return False
    return active_m01_row(
        row,
        "generate",
        [
            "effective_date",
            "target_person_id",
            "target_name",
            "old_value",
            "new_value",
            "event_name_cn",
            "new_registered_office_address",
            "new_primary_activity",
            "new_primary_ssic",
            "new_secondary_activity",
            "new_secondary_ssic",
            "new_fye",
            "new_office_hours",
            "field_label",
        ],
    )


def active_m01_row(row: dict[str, Any], flag_key: str, trigger_keys: list[str]) -> bool:
    flag = clean(row.get(flag_key)).lower()
    if flag in NO_VALUES:
        return False
    if is_yes(row.get(flag_key)):
        return True
    return any(clean(row.get(key)) for key in trigger_keys)


def date_text(value: Any) -> str:
    raw = clean(value)
    if not raw:
        return ""
    return long_date_text(raw)


def first_filled(*groups) -> str:
    for group in groups:
        for value in group:
            cleaned = clean(value)
            if cleaned:
                return cleaned
    return ""


def build_context(parsed: dict[str, Any]) -> dict[str, Any]:
    company_options = dict(parsed.get("company", {}))
    merge_p1_generation_options(company_options, parsed.get("generation", {}))
    merge_p1_generation_options(company_options, parsed.get("output_options", {}))
    company = normalize_company(company_options)
    people = [normalize_person(row) for row in parsed.get("people", [])]
    for person in people:
        lock_p1_person_dates(person, company["incorporation_date_raw"])
    shareholders = [normalize_shareholder(row, people, company, idx) for idx, row in enumerate(parsed.get("shareholders", []), start=1)]
    shareholders = [row for row in shareholders if row.get("shareholder_name") and to_number(row.get("shares")) > 0]

    directors = [p for p in people if is_yes(p.get("is_director"))]
    secretaries = [p for p in people if is_yes(p.get("is_secretary"))]
    nominee_directors = [p for p in people if is_yes(p.get("is_nominee_director"))]
    local_directors = [p for p in directors if is_yes(p.get("is_local_resident_director"))]
    client_signatories = [p for p in people if is_yes(p.get("signing_required")) and not is_yes(p.get("is_secretary"))]

    total_shares = sum_number(row.get("shares") for row in shareholders) or to_number(company.get("total_issued_shares"))
    total_issued = sum_number(row.get("issued_share_capital") for row in shareholders) or to_number(company.get("issued_share_capital")) or total_shares
    total_paid = sum_number(row.get("paid_amount") for row in shareholders) or to_number(company.get("paid_up_capital")) or total_issued
    total_unpaid = max(total_issued - total_paid, 0)
    apply_distinctive_numbers(shareholders)
    company["subscriber_share_totals"] = format_number(total_shares)
    company["issued_share_capital"] = format_money_number(total_issued)
    company["paid_up_capital"] = format_money_number(total_paid)
    company["unpaid_share_capital"] = format_money_number(total_unpaid)
    company["issued_amount_per_share"] = per_share_amount(total_issued, total_shares)
    company["amount_paid_per_share"] = per_share_amount(total_paid, total_shares)
    company["amount_due_per_share"] = "-" if total_unpaid <= 0 else per_share_amount(total_unpaid, total_shares)
    company["share_par_value"] = company["amount_paid_per_share"]
    company["share_payment_review_note"] = "Manual review: partly paid shares / unpaid share capital." if total_unpaid > 0 else ""
    company["shareholder_signature_blocks"] = "\n\n".join(
        f"----------------------------------------\n{bold_field(row.get('shareholder_name', ''))}\n{bold_field('Allottee / Shareholder')}" for row in shareholders
    )
    company["director_signature_blocks"] = "\n\n".join(
        f"Signature: ______________________________\nName: {bold_field(p.get('full_name', ''))}\nCapacity: {bold_field('Director')}" for p in directors
    )
    company["first_directors_names"] = ", ".join(p.get("full_name", "") for p in directors if p.get("full_name"))
    company["subscriber_share_lines"] = "\n".join(
        f"{row.get('shareholder_name', '')}    {format_number(row.get('shares'))}" for row in shareholders
    )
    client_signatory = default_client_signatory(shareholders, people, client_signatories, directors, company)
    certificate_director = default_certificate_director(directors)
    paid_up_confirmation_signatory = default_paid_up_confirmation_signatory(directors, client_signatory)

    sig = signature_context(company.get("incorporation_date_raw"))
    return {
        "company": company,
        "people": people,
        "directors": directors,
        "secretaries": secretaries,
        "shareholders": shareholders,
        "nominee_directors": nominee_directors,
        "local_directors": local_directors,
        "client_signatories": client_signatories,
        "certificate_director": certificate_director,
        "registrable_controllers": registrable_controllers(shareholders, total_shares),
        "provider": provider_context(),
        "signature": sig,
        "register": p1_register_context(company, shareholders),
        "director": directors[0] if directors else {},
        "secretary": secretaries[0] if secretaries else {},
        "shareholder": shareholders[0] if shareholders else {},
        "nominee_director": (nominee_directors or local_directors or directors or [{}])[0],
        "client_signatory": client_signatory,
        "paid_up_confirmation_signatory": paid_up_confirmation_signatory,
        "client_signatory_2": {},
        "client_signatory_3": {},
        "secretary_or_director": (secretaries or directors or [{}])[0],
    }


def p1_register_context(company: dict[str, Any], shareholders: list[dict[str, Any]]) -> dict[str, Any]:
    return register_context(
        company,
        p1_register_entries(shareholders),
        title="REGISTER OF MEMBERS",
        subtitle="Initial register after incorporation",
        transaction_type="Subscription on incorporation",
        effective_date=company.get("incorporation_date_raw") or company.get("incorporation_date"),
        note="This initial register is prepared from the incorporation shareholding information provided for the document package. It should be checked against the Constitution, Form 24 / Return of Allotment and issued share certificates before being treated as the Company's statutory register.",
    )


def register_context(
    company: dict[str, Any],
    entries: list[dict[str, Any]],
    *,
    title: str,
    subtitle: str,
    transaction_type: str,
    effective_date: Any,
    note: str,
) -> dict[str, Any]:
    date_raw = clean(effective_date) or company.get("default_document_date") or company.get("incorporation_date_raw")
    return {
        "title": title,
        "subtitle": subtitle,
        "transaction_type": transaction_type,
        "effective_date": date_text(date_raw),
        "prepared_date": date_text(date_raw),
        "location": clean(company.get("register_location") or company.get("registered_office_address")),
        "share_class": clean(company.get("share_class") or company.get("share_class_default")) or "Ordinary",
        "currency": clean(company.get("share_currency") or company.get("currency")) or "SGD",
        "total_shares": format_number(sum_number(row.get("shares") for row in entries)),
        "total_paid": register_total_paid_text(company, entries),
        "entries": entries,
        "note": note,
    }


def register_total_paid_text(company: dict[str, Any], entries: list[dict[str, Any]]) -> str:
    currency = clean(company.get("share_currency") or company.get("currency")) or "SGD"
    explicit = clean(company.get("paid_up_capital"))
    if explicit:
        return f"{currency} {format_money_number(explicit)}"
    total = sum_number(row.get("paid_amount_raw") for row in entries)
    return f"{currency} {format_money_number(total)}" if total else f"{currency} -"


def p1_register_entries(shareholders: list[dict[str, Any]]) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    for idx, row in enumerate(shareholders, start=1):
        entries.append(
            {
                "folio_no": clean(row.get("folio_no")) or str(idx),
                "member_name": clean(row.get("shareholder_name")),
                "id_number": clean(row.get("id_number")),
                "address": clean(row.get("shareholder_address")),
                "shares": clean(row.get("shares")),
                "share_class": clean(row.get("share_class")) or "Ordinary",
                "certificate_no": clean(row.get("certificate_no")),
                "entry_date": clean(row.get("allotment_date") or row.get("issue_date") or row.get("document_date")),
                "transaction_reference": clean(row.get("allotment_transfer_no")) or "Subscription",
                "paid_status_text": clean(row.get("paid_status_text")),
                "paid_amount_raw": clean(row.get("paid_amount") or row.get("paid_up_share_capital")),
                "remarks": clean(row.get("remarks")),
            }
        )
    return entries


def merge_p1_generation_options(company: dict[str, Any], source: Any) -> None:
    if not isinstance(source, dict):
        return
    for key in [
        "client_signatory_person_id",
        "client_signer_person_id",
        "client_signatory_logic",
        "authorized_representative_person_id",
        "prepared_by",
        "output_format",
    ]:
        value = clean(source.get(key))
        if value and not clean(company.get(key)):
            company[key] = value


def normalize_company(raw: dict[str, Any]) -> dict[str, Any]:
    company = {str(k): clean(v) for k, v in raw.items()}
    set_default(company, "company_name", "")
    set_default(company, "uen", "")
    set_default(company, "registered_office_address", "")
    set_default(company, "register_location", company.get("registered_office_address", ""))
    set_default(company, "share_currency", company.get("currency") or "SGD")
    set_default(company, "share_class", company.get("share_class_default") or "Ordinary")
    set_default(company, "share_par_value", par_value(company))
    set_default(company, "incorporation_date", today_text())
    raw_incorporation_date = company["incorporation_date"]
    company["incorporation_date_raw"] = short_date_text(raw_incorporation_date)
    company["incorporation_date_short"] = short_date_text(raw_incorporation_date)
    company["incorporation_date_long"] = long_date_text(raw_incorporation_date)
    company["incorporation_date"] = company["incorporation_date_long"]
    company["document_date"] = company["incorporation_date_long"]
    company["signature_date"] = company["incorporation_date_long"]
    set_default(company, "office_hours", "Monday to Friday, 9.00 a.m. to 5.00 p.m.")

    default_first_fye = default_first_fye_date(raw_incorporation_date)
    first_period_start = company.get("first_financial_period_start") or raw_incorporation_date
    first_period_end = company.get("first_financial_period_end") or company.get("first_fye") or default_first_fye
    fye_source = company.get("fye") or first_period_end
    company["fye_month"] = fye_month_text(fye_source)
    company["first_financial_period_start"] = long_date_text(first_period_start)
    company["first_financial_period_end"] = long_date_text(first_period_end)
    return company


def set_default(item: dict[str, Any], key: str, value: Any) -> None:
    if not clean(item.get(key)):
        item[key] = value


def default_first_fye_date(incorporation_date: Any) -> str:
    """Default first FYE to the last day of the month before incorporation month."""
    try:
        dt = parse_date(clean(incorporation_date))
    except ValueError:
        return clean(incorporation_date)
    if dt.month == 1:
        year = dt.year
        month = 12
    else:
        year = dt.year + 1
        month = dt.month - 1
    last_day = calendar.monthrange(year, month)[1]
    return date(year, month, last_day).strftime("%d/%m/%Y")


def fye_month_text(value: Any) -> str:
    text = clean(value)
    if not text:
        return ""
    try:
        return parse_date(text).strftime("%B").upper()
    except ValueError:
        month_names = {name.lower(): name.upper() for name in calendar.month_name if name}
        month_names.update({name[:3].lower(): name.upper() for name in calendar.month_name if name})
        lowered = text.strip().lower()
        return month_names.get(lowered, text.upper())


def long_date_text(value: Any) -> str:
    text = clean(value)
    if not text:
        return ""
    try:
        dt = parse_date(text)
    except ValueError:
        return text
    return f"{dt.day}{ordinal_suffix(dt.day)} {dt.strftime('%B').upper()} {dt.year}"


def short_date_text(value: Any) -> str:
    text = clean(value)
    if not text:
        return ""
    try:
        return parse_date(text).strftime("%d/%m/%Y")
    except ValueError:
        return text


def ordinal_suffix(day: int) -> str:
    if 10 <= day % 100 <= 20:
        return "th"
    return {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")


def normalize_person(raw: dict[str, Any]) -> dict[str, Any]:
    item = {str(k): clean(v) for k, v in raw.items()}
    full_name = item.get("full_name") or item.get("common_person_name") or ""
    item["full_name"] = full_name
    item.setdefault("appointment_date", "")
    item.setdefault("residential_address", "")
    item.setdefault("id_type", "")
    item.setdefault("id_number", "")
    item.setdefault("nationality", "")
    item.setdefault("date_of_birth", "")
    return item


def lock_p1_person_dates(person: dict[str, Any], incorporation_date: Any) -> None:
    date_value = long_date_text(incorporation_date)
    short_value = short_date_text(incorporation_date)
    for key in ["appointment_date", "effective_date", "consent_date", "document_date", "signature_date"]:
        person[key] = date_value
        person[f"{key}_short"] = short_value


def normalize_shareholder(raw: dict[str, Any], people: list[dict[str, Any]], company: dict[str, Any], idx: int) -> dict[str, Any]:
    item = {str(k): clean(v) for k, v in raw.items()}
    matched = match_person(item, people)
    is_corporate = str(item.get("shareholder_type", "")).lower() == "corporate" or bool(item.get("corporate_name"))
    name = item.get("corporate_name") if is_corporate else item.get("person_full_name")
    name = name or matched.get("full_name") or item.get("person_full_name") or item.get("corporate_name") or ""
    id_number = item.get("corporate_registration_number") if is_corporate else item.get("person_id_number")
    id_number = id_number or matched.get("id_number", "")
    address = item.get("corporate_registered_address") if is_corporate else matched.get("residential_address", "")
    nationality = item.get("corporate_registration_country") if is_corporate else matched.get("nationality", "")
    shares = clean(item.get("shares") or "")
    issued = clean(item.get("issued_share_capital") or item.get("share_capital") or item.get("issued_capital") or "")
    if not issued and to_number(shares) > 0:
        issued = shares
    paid = clean(item.get("paid_amount") or item.get("paid_up_share_capital") or item.get("paid_up_capital") or "")
    if not paid and issued:
        paid = issued
    unpaid = max(to_number(issued) - to_number(paid), 0)
    currency = item.get("currency") or company.get("share_currency") or "SGD"
    paid_per_share = per_share_amount(to_number(paid), to_number(shares))
    due_per_share = per_share_amount(unpaid, to_number(shares))
    share_class = item.get("share_class") or company.get("share_class") or "Ordinary"
    paid_status_text = "fully paid" if unpaid <= 0 else f"paid up to {currency} {paid_per_share} per share, with {currency} {due_per_share} unpaid per share"
    form24_allotment_text = (
        f"{bold_field(format_number(shares))} {bold_field(share_class)} Shares\n"
        f"Issued share capital: {bold_field(f'{currency} {format_money_number(issued)}')}\n"
        f"Paid-up share capital: {bold_field(f'{currency} {format_money_number(paid)}')}"
    )
    if unpaid > 0:
        form24_allotment_text += f"\nUnpaid share capital: {bold_field(f'{currency} {format_money_number(unpaid)}')}"
    return {
        **item,
        "is_corporate": "Yes" if is_corporate else "",
        "is_individual": "" if is_corporate else "Yes",
        "shareholder_name": name,
        "id_number": id_number,
        "id_type": "UEN/Registration No." if is_corporate else matched.get("id_type", item.get("id_type", "")),
        "shareholder_address": address,
        "nationality": nationality,
        "date_of_birth": matched.get("date_of_birth", ""),
        "share_class": share_class,
        "shares": format_number(shares),
        "issued_share_capital": format_money_number(issued),
        "paid_amount": format_money_number(paid),
        "paid_up_share_capital": format_money_number(paid),
        "unpaid_share_capital": format_money_number(unpaid),
        "issued_amount_per_share": per_share_amount(to_number(issued), to_number(shares)),
        "amount_paid_per_share": paid_per_share,
        "amount_due_per_share": "-" if unpaid <= 0 else due_per_share,
        "paid_status_text": paid_status_text,
        "form24_allotment_text": form24_allotment_text,
        "currency": currency,
        "certificate_no": item.get("certificate_no") or item.get("certificate_no_override") or str(idx),
        "folio_no": item.get("folio_no") or str(idx),
        "allotment_transfer_no": item.get("allotment_transfer_no") or str(idx),
        "distinctive_numbers": distinctive_numbers(idx, shares),
        "allotment_date": company.get("incorporation_date_long", ""),
        "issue_date": company.get("incorporation_date_long", ""),
        "certificate_date": company.get("incorporation_date_long", ""),
        "document_date": company.get("incorporation_date_long", ""),
        "signature_date": company.get("incorporation_date_long", ""),
        "remarks": item.get("remarks", ""),
        "shares_in_words": shares_in_words(shares, share_class),
    }


def default_client_signatory(
    shareholders: list[dict[str, Any]],
    people: list[dict[str, Any]],
    client_signatories: list[dict[str, Any]],
    directors: list[dict[str, Any]],
    company: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Choose the client-side signer for service/nominee agreements."""
    company = company or {}
    explicit_ref = clean(company.get("client_signatory_person_id") or company.get("client_signer_person_id"))
    if explicit_ref:
        matched = find_person_by_reference(people, explicit_ref)
        if matched:
            return client_signatory_from_person(matched)
    logic = clean(company.get("client_signatory_logic")).lower()
    if logic.startswith("shareholder_"):
        try:
            index = int(logic.rsplit("_", 1)[1]) - 1
        except ValueError:
            index = 0
        if 0 <= index < len(shareholders):
            return client_signatory_from_shareholder(shareholders[index], people)
    if shareholders:
        return client_signatory_from_shareholder(shareholders[0], people)
    if client_signatories:
        return client_signatory_from_person(client_signatories[0])
    if directors:
        return client_signatory_from_person(directors[0])
    return {}


def find_person_by_reference(people: list[dict[str, Any]], reference: Any) -> dict[str, Any]:
    ref = clean(reference).lower()
    if not ref:
        return {}
    for person in people:
        for key in ("person_id", "full_name", "id_number", "common_person_name"):
            if ref == clean(person.get(key)).lower():
                return person
    return {}


def default_certificate_director(directors: list[dict[str, Any]]) -> dict[str, Any]:
    """Prefer a client-side director for share certificates; use nominee only as fallback."""
    shareholder_directors = [
        p for p in directors if not is_yes(p.get("is_nominee_director")) and is_yes(p.get("is_shareholder"))
    ]
    if shareholder_directors:
        return shareholder_directors[0]
    non_nominee_directors = [p for p in directors if not is_yes(p.get("is_nominee_director"))]
    if non_nominee_directors:
        return non_nominee_directors[0]
    return directors[0] if directors else {}


def default_paid_up_confirmation_signatory(
    directors: list[dict[str, Any]],
    client_signatory: dict[str, Any],
) -> dict[str, Any]:
    """Prefer a non-nominee client director for the paid-up capital confirmation."""
    non_nominee_directors = [p for p in directors if not is_yes(p.get("is_nominee_director"))]
    if non_nominee_directors:
        return {**non_nominee_directors[0], "capacity": "Director"}
    if client_signatory and "director" in clean(client_signatory.get("capacity")).lower():
        return {**client_signatory, "capacity": "Director"}
    if directors:
        return {**directors[0], "capacity": "Director"}
    return {**client_signatory, "capacity": client_signatory.get("capacity") or "Client Representative"} if client_signatory else {}


def client_signatory_from_shareholder(shareholder: dict[str, Any], people: list[dict[str, Any]]) -> dict[str, Any]:
    matched = match_person(shareholder, people)
    is_corporate = str(shareholder.get("shareholder_type", "")).lower() == "corporate" or bool(shareholder.get("corporate_name"))
    if is_corporate:
        full_name = shareholder.get("authorized_rep_full_name") or shareholder.get("shareholder_name", "")
        capacity = "Authorised Representative of Shareholder"
    else:
        full_name = shareholder.get("shareholder_name", "")
        capacity = "Director / Shareholder" if is_yes(matched.get("is_director")) else "Shareholder"
    return {
        **shareholder,
        "full_name": full_name,
        "id_number": shareholder.get("id_number", ""),
        "id_type": shareholder.get("id_type", ""),
        "residential_address": shareholder.get("shareholder_address", ""),
        "capacity": capacity,
    }


def client_signatory_from_person(person: dict[str, Any]) -> dict[str, Any]:
    capacity = "Client Representative"
    if is_yes(person.get("is_director")) and is_yes(person.get("is_shareholder")):
        capacity = "Director / Shareholder"
    elif is_yes(person.get("is_director")):
        capacity = "Director"
    elif is_yes(person.get("is_shareholder")):
        capacity = "Shareholder"
    return {**person, "capacity": capacity}


def render_docx(template_path: Path, output_path: Path, context: dict[str, Any]) -> None:
    doc = Document(template_path)
    filter_conditional_blocks(doc, context)
    expand_repeat_markers(doc, context)
    for paragraph in doc.paragraphs:
        render_paragraph(paragraph, context)
    for table in doc.tables:
        render_table(table, context)
    strip_empty_marker_paragraphs(doc)
    apply_auto_signatures(doc, context)
    doc.save(output_path)


IF_BLOCK_RE = re.compile(r"\[\[IF\s+([\w.]+)\]\]")
END_IF_RE = re.compile(r"\[\[END IF\]\]")
REPEAT_BLOCK_RE = re.compile(r"\[\[REPEAT\s+([\w.]+)\[\]\]\]")


def render_m01_docx(template_path: Path, output_path: Path, context: dict[str, Any]) -> None:
    doc = Document(template_path)
    filter_conditional_blocks(doc, context)
    expand_repeat_markers(doc, context)
    for paragraph in doc.paragraphs:
        render_paragraph(paragraph, context)
    for table in doc.tables:
        render_table(table, context)
    strip_empty_marker_paragraphs(doc)
    apply_auto_signatures(doc, context)
    doc.save(output_path)


def filter_conditional_blocks(doc: Document, context: dict[str, Any]) -> None:
    body = doc.element.body
    include_stack: list[bool] = []
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        text_value = block_text(child)
        match = IF_BLOCK_RE.search(text_value)
        if match:
            include_stack.append(bool(resolve(context, match.group(1))))
            remove_element(child)
            continue
        if END_IF_RE.search(text_value):
            if include_stack:
                include_stack.pop()
            remove_element(child)
            continue
        if include_stack and not all(include_stack):
            remove_element(child)


def expand_repeat_markers(doc: Document, context: dict[str, Any]) -> None:
    body = doc.element.body
    idx = 0
    while idx < len(body):
        child = body[idx]
        if child.tag == qn("w:sectPr"):
            idx += 1
            continue
        text_value = block_text(child)
        match = REPEAT_BLOCK_RE.search(text_value)
        if not match:
            idx += 1
            continue
        expr = match.group(1)
        remove_element(child)
        if idx >= len(body):
            continue
        next_child = body[idx]
        if next_child.tag != qn("w:tbl"):
            continue
        table = Table(next_child, doc)
        expand_repeat_table(table, context, expr)
        idx += 1


def expand_repeat_table(table: Table, context: dict[str, Any], expr: str) -> None:
    rows = resolve(context, expr)
    if not isinstance(rows, list):
        rows = []
    variable_name = M01_REPEAT_VARS.get(expr, "item")
    template_idx = 0 if expr == "m01.director_signers" or len(table.rows) == 1 else 1
    template_row = table.rows[template_idx]
    tr_parent = template_row._tr.getparent()
    insert_at = tr_parent.index(template_row._tr)
    template_tr = deepcopy(template_row._tr)
    tr_parent.remove(template_row._tr)
    for offset, item in enumerate(rows):
        new_tr = deepcopy(template_tr)
        tr_parent.insert(insert_at + offset, new_tr)
        child_context = dict(context)
        child_context[variable_name] = item
        render_row(_Row(new_tr, table), child_context)


def strip_empty_marker_paragraphs(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        text_value = paragraph.text.strip()
        if IF_BLOCK_RE.search(text_value) or END_IF_RE.search(text_value) or REPEAT_BLOCK_RE.search(text_value):
            remove_element(paragraph._p)


def block_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t")))


def remove_element(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


ROW_LOOP_START_RE = re.compile(r"{%\s*for\s+(\w+)\s+in\s+([\w.]+)\s*%}")
ROW_LOOP_END_RE = re.compile(r"{%\s*endfor\s*%}")


def render_table(table, context: dict[str, Any]) -> None:
    idx = 0
    while idx < len(table.rows):
        row = table.rows[idx]
        row_text = "\n".join(cell.text for cell in row.cells)
        match = ROW_LOOP_START_RE.search(row_text)
        if match and ROW_LOOP_END_RE.search(row_text):
            var_name, list_expr = match.groups()
            rows = resolve(context, list_expr)
            if not isinstance(rows, list):
                rows = []
            tr_parent = row._tr.getparent()
            insert_at = tr_parent.index(row._tr)
            template_tr = deepcopy(row._tr)
            tr_parent.remove(row._tr)
            for offset, item in enumerate(rows):
                new_tr = deepcopy(template_tr)
                tr_parent.insert(insert_at + offset, new_tr)
                child = dict(context)
                child[var_name] = item
                render_row(_Row(new_tr, table), child)
            idx = insert_at + len(rows)
            continue
        render_row(row, context)
        idx += 1


def render_row(row, context: dict[str, Any]) -> None:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            render_paragraph(paragraph, context)
        for table in cell.tables:
            render_table(table, context)


def render_paragraph(paragraph, context: dict[str, Any]) -> None:
    text = paragraph.text
    if "{{" not in text and "{%" not in text:
        return
    segments = render_text_segments(text, context)
    rendered = "".join(part for part, _ in segments)
    if rendered == text and not any(is_field for _, is_field in segments):
        return
    first = paragraph.runs[0] if paragraph.runs else None
    paragraph.clear()
    for part, is_field in segments:
        if not part:
            continue
        run = paragraph.add_run(part)
        apply_run_template_style(run, first)
        if is_field:
            run.bold = True


def apply_run_template_style(run, template_run) -> None:
    if template_run is None:
        return
    run.bold = template_run.bold
    run.italic = template_run.italic
    run.underline = template_run.underline
    if template_run.font.name:
        run.font.name = template_run.font.name
    if template_run.font.size:
        run.font.size = template_run.font.size
    if template_run.font.color and template_run.font.color.rgb:
        run.font.color.rgb = template_run.font.color.rgb


LOOP_RE = re.compile(r"{%\s*for\s+(\w+)\s+in\s+([\w.]+)\s*%}(.*?){%\s*endfor\s*%}", re.DOTALL)
VAR_RE = re.compile(r"{{\s*([^{}]+?)\s*}}")
BOLD_FIELD_START = "[[B]]"
BOLD_FIELD_END = "[[/B]]"
BOLD_FIELD_MARKER_RE = re.compile(r"\[\[/?B\]\]")


def render_text(text: str, context: dict[str, Any]) -> str:
    return "".join(part for part, _ in render_text_segments(text, context))


def render_text_segments(text: str, context: dict[str, Any]) -> list[tuple[str, bool]]:
    text = ROW_LOOP_START_RE.sub("", text)
    text = ROW_LOOP_END_RE.sub("", text)
    if LOOP_RE.search(text):
        segments: list[tuple[str, bool]] = []
        pos = 0
        for match in LOOP_RE.finditer(text):
            if match.start() > pos:
                segments.extend(render_text_segments(text[pos : match.start()], context))
            var_name, list_expr, body = match.groups()
            rows = resolve(context, list_expr)
            if isinstance(rows, list):
                for row in rows:
                    child = dict(context)
                    child[var_name] = row
                    segments.extend(render_text_segments(body, child))
            pos = match.end()
        if pos < len(text):
            segments.extend(render_text_segments(text[pos:], context))
        return merge_text_segments(segments)

    segments: list[tuple[str, bool]] = []
    pos = 0
    for match in VAR_RE.finditer(text):
        if match.start() > pos:
            append_marked_segments(segments, text[pos : match.start()], False)
        expr = match.group(1).strip()
        value = str(resolve(context, expr) or "")
        visible_value = strip_bold_field_markers(value)
        append_marked_segments(segments, value, should_bold_rendered_field(expr, visible_value))
        pos = match.end()
    if pos < len(text):
        append_marked_segments(segments, text[pos:], False)
    return merge_text_segments(segments)


def append_marked_segments(segments: list[tuple[str, bool]], text: str, default_bold: bool) -> None:
    if not text:
        return
    pos = 0
    current_bold = default_bold
    for match in BOLD_FIELD_MARKER_RE.finditer(text):
        if match.start() > pos:
            segments.append((text[pos : match.start()], current_bold))
        current_bold = match.group(0) == BOLD_FIELD_START or default_bold
        pos = match.end()
    if pos < len(text):
        segments.append((text[pos:], current_bold))


def strip_bold_field_markers(value: str) -> str:
    return value.replace(BOLD_FIELD_START, "").replace(BOLD_FIELD_END, "")


def bold_field(value: Any) -> str:
    text = str(value or "")
    if not text:
        return ""
    return f"{BOLD_FIELD_START}{text}{BOLD_FIELD_END}"


def merge_text_segments(segments: list[tuple[str, bool]]) -> list[tuple[str, bool]]:
    merged: list[tuple[str, bool]] = []
    for text, is_field in segments:
        if not text:
            continue
        if merged and merged[-1][1] == is_field:
            merged[-1] = (merged[-1][0] + text, is_field)
        else:
            merged.append((text, is_field))
    return merged


def should_bold_rendered_field(expr: str, value: str) -> bool:
    if not value.strip():
        return False
    key = expr.strip().lower()
    name = key.rsplit(".", 1)[-1]
    words = re.findall(r"\w+", value)
    short_value = len(words) <= 12 and "\n" not in value and not value.lstrip().upper().startswith(("RESOLVED", "NOTED"))
    value_field_markers = (
        "date",
        "amount",
        "capital",
        "share",
        "shares",
        "certificate",
        "currency",
        "uen",
        "number",
    )
    if short_value and not any(marker in name for marker in ("line", "block", "blocks")) and any(marker in name for marker in value_field_markers):
        return True
    generated_text_markers = (
        "resolution",
        "statement",
        "representation",
        "business",
        "intro",
        "body",
        "title",
        "block",
        "blocks",
        "line",
        "lines",
        "note",
        "notes",
        "remark",
        "remarks",
        "summary",
        "description",
        "declaration",
        "clause",
    )
    if any(marker in name for marker in generated_text_markers):
        return False
    if len(words) > 12 and any(mark in value for mark in [".", ";", ":"]):
        return False
    if value.lstrip().upper().startswith(("RESOLVED", "NOTED")):
        return False
    return True


def resolve(context: dict[str, Any], expr: str) -> Any:
    current: Any = context
    for part in expr.split("."):
        part = part.strip()
        if isinstance(current, dict):
            current = current.get(part, "")
        else:
            current = getattr(current, part, "")
        if current is None:
            return ""
    return current


def context_with(context: dict[str, Any], **overrides: Any) -> dict[str, Any]:
    child = deepcopy(context)
    child.update(overrides)
    return child


def provider_context() -> dict[str, str]:
    return {
        "name": "RSIN GROUP PTE. LTD.",
        "registered_address": "111 North Bridge Road, #29-06A, Peninsula Plaza, Singapore 179098",
        "lodging_officer_name": "",
        "account_no": "",
        "phone": "",
    }


def signature_context(value: Any = "") -> dict[str, str]:
    text = clean(value) or today_text()
    short_date = text
    day = ""
    day_ordinal = ""
    month_year = ""
    try:
        dt = parse_date(text)
        short_date = short_date_text(text)
        text = long_date_text(text)
        day = str(dt.day)
        day_ordinal = f"{dt.day}{ordinal_suffix(dt.day)}"
        month_year = dt.strftime("%B %Y")
    except ValueError:
        day = text[:2] if text else ""
        day_ordinal = day
        month_year = text
    return {"date": text, "date_short": short_date, "day": day, "day_ordinal": day_ordinal, "month_year": month_year}


def registrable_controllers(shareholders: list[dict[str, Any]], total_shares: float) -> list[dict[str, Any]]:
    if not shareholders:
        return []
    if total_shares <= 0:
        return shareholders
    controllers = []
    for row in shareholders:
        shares = to_number(row.get("shares"))
        if shares / total_shares >= 0.25:
            controllers.append(row)
    return controllers or shareholders


def match_person(shareholder: dict[str, Any], people: list[dict[str, Any]]) -> dict[str, Any]:
    name = clean(shareholder.get("person_full_name")).lower()
    ident = clean(shareholder.get("person_id_number")).lower()
    for person in people:
        if ident and clean(person.get("id_number")).lower() == ident:
            return person
        if name and clean(person.get("full_name")).lower() == name:
            return person
    return {}


def par_value(company: dict[str, Any]) -> str:
    shares = to_number(company.get("total_issued_shares"))
    paid = to_number(company.get("paid_up_capital"))
    if shares > 0 and paid > 0:
        return format_money_number(paid / shares)
    return "1"


def per_share_amount(amount: Any, shares: Any) -> str:
    share_count = to_number(shares)
    if share_count <= 0:
        return "0"
    return format_per_share_number(to_number(amount) / share_count)


def format_per_share_number(value: Any) -> str:
    number = to_number(value)
    if number == 0:
        return "0"
    if abs(number - round(number, 2)) < 1e-9:
        return format_money_number(number)
    return f"{number:,.6f}".rstrip("0").rstrip(".")


def distinctive_numbers(idx: int, shares: Any) -> str:
    count = int(to_number(shares))
    if count <= 0:
        return ""
    start = 1 if idx == 1 else ""
    if start == "":
        return ""
    return f"{start} - {count}"


def apply_distinctive_numbers(shareholders: list[dict[str, Any]]) -> None:
    start = 1
    for row in shareholders:
        count = int(to_number(row.get("shares")))
        if count <= 0:
            continue
        end = start + count - 1
        row["distinctive_numbers"] = f"{start} - {end}" if start != end else str(start)
        start = end + 1


def shares_in_words(value: Any, share_class: str) -> str:
    number = int(to_number(value))
    if number <= 0:
        return f"{format_number(value)} {share_class} Share(s)"
    return f"{format_number(number)} ({int_to_words(number).upper()}) {share_class} Share(s)"


def int_to_words(number: int) -> str:
    if number == 0:
        return "zero"
    units = ["", "one", "two", "three", "four", "five", "six", "seven", "eight", "nine", "ten", "eleven", "twelve", "thirteen", "fourteen", "fifteen", "sixteen", "seventeen", "eighteen", "nineteen"]
    tens = ["", "", "twenty", "thirty", "forty", "fifty", "sixty", "seventy", "eighty", "ninety"]

    def under_thousand(n: int) -> str:
        words = []
        if n >= 100:
            words.append(units[n // 100])
            words.append("hundred")
            n %= 100
        if n >= 20:
            words.append(tens[n // 10])
            n %= 10
        if n:
            words.append(units[n])
        return " ".join(words)

    parts = []
    for scale, label in ((1_000_000, "million"), (1_000, "thousand"), (1, "")):
        if number >= scale:
            chunk = number // scale
            number %= scale
            words = under_thousand(chunk)
            parts.append(f"{words} {label}".strip())
    return " ".join(parts)


def is_yes(value: Any) -> bool:
    return str(value or "").strip().lower() in {"yes", "y", "true", "1", "是", "需要"}


def clean(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (datetime, date)):
        return value.strftime("%d/%m/%Y")
    text = str(value).strip()
    if text.endswith(".0") and text[:-2].isdigit():
        return text[:-2]
    return text


def parse_date(value: str) -> datetime:
    text = clean(value)
    text = re.sub(r"(\d{1,2})(st|nd|rd|th)\b", r"\1", text, flags=re.IGNORECASE)
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%Y/%m/%d", "%d-%m-%Y", "%d %B %Y", "%d %b %Y", "%d %B, %Y", "%d %b, %Y"):
        try:
            return datetime.strptime(text, fmt)
        except ValueError:
            continue
    raise ValueError(text)


def today_text() -> str:
    return date.today().strftime("%d/%m/%Y")


def to_number(value: Any) -> float:
    text = clean(value).replace(",", "").replace("SGD", "").replace("S$", "").strip()
    if not text:
        return 0.0
    try:
        return float(text)
    except ValueError:
        return 0.0


def sum_number(values) -> float:
    return sum(to_number(value) for value in values)


def format_number(value: Any) -> str:
    number = to_number(value)
    if number == 0 and clean(value) not in {"0", "0.0"}:
        return clean(value)
    if number.is_integer():
        return f"{int(number):,}"
    return f"{number:,.2f}"


def format_money_number(value: Any) -> str:
    number = to_number(value)
    if number == 0 and clean(value) not in {"0", "0.0"}:
        return clean(value)
    if number.is_integer():
        return f"{int(number):,}"
    return f"{number:,.2f}"


def safe_filename(value: Any) -> str:
    text = clean(value) or "output"
    text = re.sub(r'[<>:"/\\|?*\r\n\t]+', "_", text)
    text = re.sub(r"\s+", "_", text).strip("._ ")
    return text[:90] or "output"
