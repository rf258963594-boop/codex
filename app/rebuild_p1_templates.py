from __future__ import annotations

import json
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DIR = ROOT / "app" / "doc_templates" / "p1_rebuilt"
OUTPUT_DIR = ROOT / "outputs" / "P1_rebuilt_templates"
ZIP_PATH = ROOT / "outputs" / "P1_rebuilt_templates.zip"


MANIFEST = [
    {
        "template_id": "first_directors_resolution",
        "file_name": "01_first_directors_resolution.docx",
        "display_name": "First Directors Resolution",
        "old_reference": "1. Company Name + First minutes resolution.docx",
        "level": "company",
        "repeat": "one_per_company",
        "signers": "All signing directors",
        "priority": "P1",
    },
    {
        "template_id": "director_consent",
        "file_name": "02_director_consent_legacy_form45.docx",
        "display_name": "Consent to Act as Director / legacy Form 45",
        "old_reference": "3. Company Name + Form 45.doc / pdf",
        "level": "person",
        "repeat": "one_per_director",
        "signers": "Relevant director",
        "priority": "P1",
    },
    {
        "template_id": "secretary_consent",
        "file_name": "03_secretary_consent_legacy_form45b.docx",
        "display_name": "Consent to Act as Secretary / legacy Form 45B",
        "old_reference": "4. Company Name + Form45B.doc / pdf",
        "level": "person",
        "repeat": "one_per_secretary",
        "signers": "Relevant secretary",
        "priority": "P1",
    },
    {
        "template_id": "share_certificate",
        "file_name": "04_share_certificate.docx",
        "display_name": "Share Certificate",
        "old_reference": "5. Company Name + Share Certificate-2.docx",
        "level": "shareholder",
        "repeat": "one_per_shareholder_or_certificate",
        "signers": "Director and director/secretary; shareholder receipt if used",
        "priority": "P1",
    },
    {
        "template_id": "secretary_service_agreement",
        "file_name": "05_secretary_service_agreement.docx",
        "display_name": "Secretary Service Agreement",
        "old_reference": "7.Secretary Agreement.docx",
        "level": "agreement",
        "repeat": "one_per_company",
        "signers": "Service provider and client signatories",
        "priority": "P1",
    },
    {
        "template_id": "nominee_director_agreement",
        "file_name": "06_nominee_director_agreement.docx",
        "display_name": "Nominee Director Agreement",
        "old_reference": "8 Company Name + Agreement for Appointment of Nominee Director.doc / pdf",
        "level": "agreement",
        "repeat": "one_per_nominee_director_relationship",
        "signers": "Nominee director/service provider and company/client signatory",
        "priority": "P1",
        "manual_review": True,
    },
    {
        "template_id": "signature_record_attachment",
        "file_name": "07_signature_record_attachment.docx",
        "display_name": "Signature Record Attachment",
        "old_reference": "每份签名文件后需附带的签名记录.docx",
        "level": "attachment",
        "repeat": "append_to_each_signed_document",
        "signers": "No pre-signature; populated after e-signing",
        "priority": "P1",
    },
]


FIELD_GROUPS = {
    "Company": [
        ("company.company_name", "Company legal name", "Company sheet: company_name"),
        ("company.uen", "Company registration number/UEN", "BizFile or post-incorporation value"),
        ("company.company_type", "Company type", "Default: Private Company Limited by Shares"),
        ("company.incorporation_date", "Date of incorporation or document date", "Company sheet"),
        ("company.registered_office_address", "Registered office", "Company sheet"),
        ("company.register_location", "Place where registers are kept", "Default: registered office"),
        ("company.office_hours", "Registered office hours", "Default working hours"),
        ("company.fye", "Financial year end", "Company sheet"),
        ("company.first_financial_period_start", "First financial period start", "Calculated or manual"),
        ("company.first_financial_period_end", "First financial period end", "Calculated or manual"),
        ("company.share_currency", "Share currency", "Default: SGD"),
        ("company.share_class", "Share class", "Default: Ordinary"),
        ("company.share_par_value", "Per-share value where used by legacy wording", "Default: 1"),
        ("company.paid_up_capital", "Total paid-up capital", "Calculated from Shareholders"),
        ("company.has_common_seal", "Whether common seal wording should be kept", "Default: No / review"),
    ],
    "People": [
        ("director.full_name", "Director full legal name", "People.full_name"),
        ("director.id_type", "Passport/NRIC/FIN", "People.id_type"),
        ("director.id_number", "ID number", "People.id_number"),
        ("director.nationality", "Nationality", "People.nationality"),
        ("director.residential_address", "Residential address", "People.residential_address"),
        ("director.email", "Email", "People.email"),
        ("director.phone", "Phone", "People.phone"),
        ("director.appointment_date", "Director appointment/effective date", "People.appointment_date"),
        ("director.signature_date", "Director signature date", "Generated or manual"),
        ("secretary.full_name", "Secretary full legal name", "People.full_name"),
        ("secretary.id_type", "Passport/NRIC/FIN", "People.id_type"),
        ("secretary.id_number", "ID number", "People.id_number"),
        ("secretary.nationality", "Nationality", "People.nationality"),
        ("secretary.residential_address", "Residential address", "People.residential_address"),
        ("secretary.email", "Secretary email", "People.email or common person setting"),
        ("secretary.phone", "Secretary phone", "People.phone or common person setting"),
        ("secretary.appointment_date", "Secretary appointment/effective date", "People.appointment_date"),
        ("secretary.signature_date", "Secretary signature date", "Generated or manual"),
        ("nominee_director.full_name", "Nominee director name", "Common people or People.full_name"),
        ("nominee_director.id_number", "Nominee director ID number", "Common people or People.id_number"),
        ("nominee_director.signature_date", "Nominee director signature date", "Generated or manual"),
        ("client_signatory.full_name", "Client signing person", "People where signing_required=Yes"),
        ("client_signatory.signature_date", "Client signatory signature date", "Generated or manual"),
        ("secretary_or_director.full_name", "Second certificate signer", "Website chooses secretary or another director"),
    ],
    "Shareholders": [
        ("shareholder.shareholder_name", "Person or corporate shareholder display name", "Derived"),
        ("shareholder.person_full_name", "Individual shareholder name", "Shareholders.person_full_name"),
        ("shareholder.person_id_number", "Individual ID number", "Shareholders.person_id_number"),
        ("shareholder.corporate_name", "Corporate shareholder name", "Shareholders.corporate_name"),
        ("shareholder.corporate_registration_number", "Corporate registration number", "Shareholders.corporate_registration_number"),
        ("shareholder.corporate_registered_address", "Corporate registered address", "Shareholders.corporate_registered_address"),
        ("shareholder.share_class", "Share class", "Shareholders.share_class"),
        ("shareholder.shares", "Number of shares", "Shareholders.shares"),
        ("shareholder.paid_amount", "Paid amount", "Shareholders.paid_amount"),
        ("shareholder.currency", "Currency", "Shareholders.currency"),
        ("shareholder.certificate_no", "Share certificate number", "Generated by website"),
        ("shareholder.folio_no", "Folio number", "Generated or manual"),
        ("shareholder.distinctive_numbers", "Distinctive share numbers", "Generated or manual"),
        ("shareholder.allotment_transfer_no", "Allotment or transfer number", "Generated or manual"),
        ("shareholder.shareholder_address", "Shareholder display address", "Derived from individual/corporate address"),
        ("shareholder.remarks", "Certificate/counterfoil remarks", "Generated or manual"),
    ],
    "Provider and signing": [
        ("provider.name", "Corporate secretary service provider", "Website setting"),
        ("provider.uen", "Provider UEN", "Website setting"),
        ("provider.address", "Provider address", "Website setting"),
        ("provider.email", "Provider email", "Website setting"),
        ("provider.authorised_signatory_name", "Provider signing person", "Website setting"),
        ("generation.prepared_by", "Internal preparer", "Generation.prepared_by"),
        ("generation.signing_mode", "Signing mode", "Generation.signing_mode"),
        ("signature.date", "Signature date", "Generated or manual"),
        ("signature.day", "Signature day for legacy wording", "Derived from signature.date"),
        ("signature.month_year", "Signature month/year for legacy wording", "Derived from signature.date"),
        ("signature_record.document_title", "Signed document title", "Post-signing data"),
        ("signature_record.envelope_id", "E-sign envelope/document ID", "Post-signing data"),
        ("signature_record.generated_at", "Audit record generated timestamp", "Post-signing data"),
        ("signature_record.checksum", "Digital fingerprint/checksum", "Post-signing data"),
        ("signature_record.signers", "Signer audit table", "Post-signing data"),
        ("signer.full_name", "Signer name inside signature audit table", "Post-signing data"),
        ("signer.email", "Signer email inside signature audit table", "Post-signing data"),
        ("signer.party_id", "Signer party ID", "Post-signing data"),
        ("signer.status", "Signer status", "Post-signing data"),
        ("signer.signed_at", "Signer completion timestamp", "Post-signing data"),
        ("event.timestamp", "Audit event timestamp", "Post-signing data"),
        ("event.action", "Audit event action", "Post-signing data"),
        ("event.ip_address", "Audit event IP address", "Post-signing data"),
        ("event.details", "Audit event details", "Post-signing data"),
    ],
}


def clean_dirs() -> None:
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for directory in (TEMPLATE_DIR, OUTPUT_DIR):
        for path in directory.glob("*"):
            if path.is_file():
                path.unlink()


def configure_doc(doc: Document, title: str = "") -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size in [("Heading 1", 13), ("Heading 2", 11), ("Heading 3", 10)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)

    if title:
        doc.core_properties.title = title
    doc.core_properties.author = "RSIN template rebuild"
    doc.core_properties.comments = "Rebuilt P1 template with clear placeholders. Original source files were not modified."


def set_run_font(run, size: float | int | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic


def para(
    doc: Document,
    text: str = "",
    *,
    align: WD_ALIGN_PARAGRAPH | None = None,
    bold: bool = False,
    italic: bool = False,
    size: float | int | None = None,
    before: float | int | None = None,
    after: float | int | None = None,
    keep: bool = False,
) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if keep:
        p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)


def centered_title(doc: Document, text: str, size: float | int = 12) -> None:
    para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=size, after=2, keep=True)


def company_heading(doc: Document) -> None:
    centered_title(doc, "{{company.company_name}}", 12)
    para(doc, "Company Registration No.: {{company.uen}}", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, after=0)
    para(doc, "(Incorporated in the Republic of Singapore)", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, after=8)


def page_no(doc: Document, label: str) -> None:
    para(doc, label, align=WD_ALIGN_PARAGRAPH.CENTER, size=9, after=6)


def section_heading(doc: Document, text: str) -> None:
    para(doc, text, bold=True, size=10, before=7, after=3, keep=True)


def resolved(doc: Document) -> None:
    para(doc, "Resolved -", italic=True, after=2, keep=True)


def table(doc: Document, rows: int, cols: int, widths: list[float] | None = None):
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    if widths:
        for row in tbl.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    return tbl


def cell_text(cell, text: str, *, bold: bool = False, size: float | int = 9, align=None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def set_borders(tbl) -> None:
    tbl_pr = tbl._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "808080")


def signature_line(doc: Document, name_field: str, role_label: str, date_field: str = "{{signature.date}}") -> None:
    para(doc, "", after=8)
    tbl = table(doc, 3, 2, [3.0, 3.0])
    set_borders(tbl)
    cell_text(tbl.cell(0, 0), "Signature: ______________________________", size=9)
    cell_text(tbl.cell(0, 1), "Date: " + date_field, size=9)
    cell_text(tbl.cell(1, 0), "Name: " + name_field, size=9)
    cell_text(tbl.cell(1, 1), "Role: " + role_label, size=9)
    cell_text(tbl.cell(2, 0), "ID: ______________________________", size=9)
    cell_text(tbl.cell(2, 1), "Witness: ______________________________", size=9)


def save_doc(doc: Document, name: str) -> None:
    template_path = TEMPLATE_DIR / name
    output_path = OUTPUT_DIR / name
    doc.save(template_path)
    doc.save(output_path)


def build_first_directors_resolution() -> None:
    doc = Document()
    configure_doc(doc, "First Directors Resolution")
    company_heading(doc)
    centered_title(doc, "DIRECTORS' RESOLUTION IN WRITING", 12)
    para(doc, "_" * 78, align=WD_ALIGN_PARAGRAPH.CENTER, size=8, after=7)

    section_heading(doc, "1. CERTIFICATE CONFIRMING THE INCORPORATION OF COMPANY")
    para(doc, "Noted -", italic=True, after=2)
    para(
        doc,
        "The certificate confirming incorporation of the Company dated {{company.incorporation_date}}, together with a copy of the Constitution of the Company, was tabled and noted.",
    )

    section_heading(doc, "2. CONFIRMATION OF FIRST DIRECTORS")
    resolved(doc)
    para(
        doc,
        "That the following persons be and are hereby confirmed as the first directors of the Company with effect from {{company.incorporation_date}}:",
    )
    tbl = table(doc, 2, 4, [2.15, 1.35, 1.55, 1.15])
    set_borders(tbl)
    for idx, label in enumerate(["Name", "ID Type", "ID Number", "Nationality"]):
        cell_text(tbl.cell(0, idx), label, bold=True)
    for idx, value in enumerate(["{{director.full_name}}", "{{director.id_type}}", "{{director.id_number}}", "{{director.nationality}}"]):
        cell_text(tbl.cell(1, idx), "{% for director in directors %}" + value if idx == 0 else value)
    para(doc, "{% endfor %}", size=1, after=1)

    section_heading(doc, "3. APPOINTMENT OF SECRETARY")
    resolved(doc)
    para(doc, "That {{secretary.full_name}} be and is hereby appointed as Secretary of the Company with immediate effect from {{secretary.appointment_date}}.")
    para(doc, "That the Accounting and Corporate Regulatory Authority, Singapore be notified accordingly.")

    section_heading(doc, "4. CONFIRMATION OF REGISTERED OFFICE AND PLACE WHERE REGISTER OF MEMBERS AND INDEX IS KEPT")
    resolved(doc)
    para(doc, "That the registered office of the Company situated at {{company.registered_office_address}} since the date of incorporation be and is hereby confirmed.")
    para(doc, "That the place where the Register of Members and Index is kept, situated at {{company.register_location}}, be and is hereby confirmed.")
    para(doc, "Office hours: {{company.office_hours}}.")

    doc.add_page_break()
    page_no(doc, "- Page 2 -")
    company_heading(doc)
    centered_title(doc, "DIRECTORS' RESOLUTION IN WRITING", 12)
    para(doc, "_" * 78, align=WD_ALIGN_PARAGRAPH.CENTER, size=8, after=7)

    section_heading(doc, "5. ADOPTION OF COMMON SEAL")
    resolved(doc)
    para(doc, "{% if company.has_common_seal %}That the seal, an impression of which is affixed to the margin of these resolutions, be and is hereby adopted as the Common Seal of the Company.{% else %}This section is not applicable unless the Company has adopted a common seal.{% endif %}")

    section_heading(doc, "6. ISSUE OF SUBSCRIBERS' SHARES")
    resolved(doc)
    para(doc, "That the signatories to the Constitution be registered as members in respect of the shares for which they subscribed, fully paid in cash at {{company.share_currency}} {{company.share_par_value}} per share, namely:")
    tbl = table(doc, 2, 4, [2.55, 1.2, 1.2, 1.45])
    set_borders(tbl)
    for idx, label in enumerate(["Name of Subscriber", "Share Class", "No. of Shares", "Paid Amount"]):
        cell_text(tbl.cell(0, idx), label, bold=True)
    for idx, value in enumerate(["{{shareholder.shareholder_name}}", "{{shareholder.share_class}}", "{{shareholder.shares}}", "{{shareholder.currency}} {{shareholder.paid_amount}}"]):
        cell_text(tbl.cell(1, idx), "{% for shareholder in shareholders %}" + value if idx == 0 else value)
    para(doc, "{% endfor %}", size=1, after=1)
    para(doc, "That share certificates be issued to the subscribers in accordance with the Company's Constitution.")

    section_heading(doc, "7. FINANCIAL YEAR END")
    resolved(doc)
    para(doc, "That the financial year end of the Company be fixed on {{company.fye}} and the first set of accounts will be for the period from {{company.first_financial_period_start}} to {{company.first_financial_period_end}}.")

    section_heading(doc, "SIGNATURES")
    para(doc, "{% for director in directors %}", size=1, after=1)
    signature_line(doc, "{{director.full_name}}", "Director", "{{director.signature_date}}")
    para(doc, "{% endfor %}", size=1, after=1)
    save_doc(doc, "01_first_directors_resolution.docx")


def build_director_consent() -> None:
    doc = Document()
    configure_doc(doc, "Consent to Act as Director")
    centered_title(doc, "THE COMPANIES ACT 1967", 12)
    centered_title(doc, "SECTION 173C(a)", 10)
    centered_title(doc, "CONSENT TO ACT AS DIRECTOR AND STATEMENT OF NON-DISQUALIFICATION TO ACT AS DIRECTOR", 11)
    para(doc, "Legacy display name: Form 45", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=9, after=8)
    tbl = table(doc, 2, 2, [1.55, 4.95])
    set_borders(tbl)
    cell_text(tbl.cell(0, 0), "Name of Company", bold=True)
    cell_text(tbl.cell(0, 1), "{{company.company_name}}")
    cell_text(tbl.cell(1, 0), "Company No.", bold=True)
    cell_text(tbl.cell(1, 1), "{{company.uen}}")

    para(doc, "I, {{director.full_name}}, hereby consent to act as a director of the abovenamed company with effect from {{director.appointment_date}} and declare that:", after=5)
    items = [
        "I am not below 18 years of age and I am otherwise of full legal capacity.",
        "I am not disqualified from acting as a director under the Companies Act 1967 or any other written law.",
        "I have not been convicted of any offence involving fraud or dishonesty that would disqualify me from acting as a director.",
        "I understand that I must discharge my responsibilities in the company honestly, diligently and in the best interests of the company.",
        "I understand that I must ensure the company keeps proper records and complies with its filing and statutory obligations.",
        "I have read and understood the above statements and confirm that the information provided in this form is true and complete.",
    ]
    for idx, item in enumerate(items, start=1):
        para(doc, f"({idx}) {item}", after=3)

    section_heading(doc, "Director particulars")
    tbl = table(doc, 6, 2, [1.7, 4.8])
    set_borders(tbl)
    rows = [
        ("Name", "{{director.full_name}}"),
        ("Email Address", "{{director.email}}"),
        ("Contact Number", "{{director.phone}}"),
        ("Residential Address", "{{director.residential_address}}"),
        ("{{director.id_type}} No.", "{{director.id_number}}"),
        ("Nationality", "{{director.nationality}}"),
    ]
    for row_idx, (label, value) in enumerate(rows):
        cell_text(tbl.cell(row_idx, 0), label, bold=True)
        cell_text(tbl.cell(row_idx, 1), value)

    signature_line(doc, "{{director.full_name}}", "Director", "{{director.signature_date}}")
    para(doc, "Dated this {{signature.day}} day of {{signature.month_year}}.", after=2)
    save_doc(doc, "02_director_consent_legacy_form45.docx")


def build_secretary_consent() -> None:
    doc = Document()
    configure_doc(doc, "Consent to Act as Secretary")
    centered_title(doc, "THE COMPANIES ACT 1967", 12)
    centered_title(doc, "SECTION 173C(b)", 10)
    centered_title(doc, "CONSENT TO ACT AS SECRETARY", 12)
    para(doc, "Legacy display name: Form 45B", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=9, after=8)
    tbl = table(doc, 2, 2, [1.55, 4.95])
    set_borders(tbl)
    cell_text(tbl.cell(0, 0), "Name of Company", bold=True)
    cell_text(tbl.cell(0, 1), "{{company.company_name}}")
    cell_text(tbl.cell(1, 0), "Company No.", bold=True)
    cell_text(tbl.cell(1, 1), "{{company.uen}}")

    para(doc, "1. I, {{secretary.full_name}}, the undermentioned person, hereby consent to act as Secretary of the abovenamed company with effect from {{secretary.appointment_date}}.")
    para(doc, "2. Where applicable, I confirm that I am a qualified person for the purpose of acting as secretary of the Company.")

    section_heading(doc, "Secretary particulars")
    tbl = table(doc, 5, 2, [1.7, 4.8])
    set_borders(tbl)
    rows = [
        ("Name", "{{secretary.full_name}}"),
        ("Address", "{{secretary.residential_address}}"),
        ("{{secretary.id_type}} No.", "{{secretary.id_number}}"),
        ("Nationality", "{{secretary.nationality}}"),
        ("Email / Contact", "{{secretary.email}} / {{secretary.phone}}"),
    ]
    for row_idx, (label, value) in enumerate(rows):
        cell_text(tbl.cell(row_idx, 0), label, bold=True)
        cell_text(tbl.cell(row_idx, 1), value)

    signature_line(doc, "{{secretary.full_name}}", "Secretary", "{{secretary.signature_date}}")
    para(doc, "Dated this {{signature.day}} day of {{signature.month_year}}.", after=2)
    para(doc, "* Delete or adapt where inapplicable. This rebuilt template keeps the legacy Form 45B label only as a display reference.", italic=True, size=8)
    save_doc(doc, "03_secretary_consent_legacy_form45b.docx")


def build_share_certificate() -> None:
    doc = Document()
    configure_doc(doc, "Share Certificate")
    para(doc, "Counterfoil", bold=True, size=9, after=0)
    tbl = table(doc, 5, 4, [1.55, 1.65, 1.65, 1.65])
    set_borders(tbl)
    rows = [
        ["Certificate No.", "{{shareholder.certificate_no}}", "Folio in Members' Register", "{{shareholder.folio_no}}"],
        ["Name of Company", "{{company.company_name}}", "Date", "{{signature.date}}"],
        ["Registered Holder", "{{shareholder.shareholder_name}}", "Allotment / Transfer No.", "{{shareholder.allotment_transfer_no}}"],
        ["No. of Shares", "{{shareholder.shares}}", "Distinctive No.", "{{shareholder.distinctive_numbers}}"],
        ["Remarks", "{{shareholder.remarks}}", "Share Class", "{{shareholder.share_class}}"],
    ]
    for r, values in enumerate(rows):
        for c, value in enumerate(values):
            cell_text(tbl.cell(r, c), value, bold=c in (0, 2), size=8)

    para(doc, "-" * 110, align=WD_ALIGN_PARAGRAPH.CENTER, size=8, after=4)
    para(doc, "Receipt for Share Certificate", bold=True, size=9, after=0)
    tbl = table(doc, 3, 2, [2.0, 4.5])
    set_borders(tbl)
    receipt_rows = [
        ("Certificate No.", "{{shareholder.certificate_no}}"),
        ("Received this certificate covering", "{{shareholder.shares}} {{shareholder.share_class}} shares in {{company.company_name}}"),
        ("Signature of Shareholder", "__________________________________"),
    ]
    for r, (label, value) in enumerate(receipt_rows):
        cell_text(tbl.cell(r, 0), label, bold=True, size=8)
        cell_text(tbl.cell(r, 1), value, size=8)

    para(doc, "", after=8)
    centered_title(doc, "SHARE CERTIFICATE", 14)
    company_heading(doc)
    para(doc, "Certificate No.: {{shareholder.certificate_no}}", align=WD_ALIGN_PARAGRAPH.RIGHT, size=9, after=4)
    para(doc, "THIS IS TO CERTIFY THAT", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=10, after=4)
    para(doc, "{{shareholder.shareholder_name}}", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, after=2)
    para(doc, "of {{shareholder.shareholder_address}}", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, after=4)
    para(doc, "is/are the registered holder(s) of", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, after=2)
    para(doc, "{{shareholder.shares}} {{shareholder.share_class}} Share(s)", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, after=2)
    para(doc, "fully paid in the Company subject to the Constitution of the Company.", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, after=8)
    para(doc, "Dated this {{signature.day}} day of {{signature.month_year}}.", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, after=8)
    para(doc, "Given under the Common Seal of the Company on the date stated above and in the presence of:", size=9, after=8)
    tbl = table(doc, 2, 2, [3.0, 3.0])
    set_borders(tbl)
    cell_text(tbl.cell(0, 0), "Director\n\n______________________________", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(tbl.cell(0, 1), "Director / Secretary\n\n______________________________", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(tbl.cell(1, 0), "{{director.full_name}}", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(tbl.cell(1, 1), "{{secretary_or_director.full_name}}", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Note: Transfer of any portion of the shares comprised in this certificate cannot be registered unless accompanied by this certificate.", italic=True, size=8, after=0)
    save_doc(doc, "04_share_certificate.docx")


def build_secretary_service_agreement() -> None:
    doc = Document()
    configure_doc(doc, "Secretary Service Agreement")
    para(doc, "Date: {{signature.date}}", size=10)
    para(doc, "To: {{provider.name}}", size=10)
    centered_title(doc, "Re: Services Agreement", 12)
    para(doc, "Proposed Name of Company: {{company.company_name}} (the \"Company\")", bold=True)
    para(doc, "I/We, the undersigned, hereby agree as follows:")

    clauses = [
        ("Appointment of service provider", "{{company.company_name}} appoints {{provider.name}} and its related companies to provide one or more services, including company incorporation in Singapore, corporate secretarial services, employment pass application, accounting support, filing support and related corporate services. {{company.company_name}} 委任 {{provider.name}} 及其关联公司提供包括新加坡公司注册、公司秘书、准证申请、会计及相关企业服务。"),
        ("Company name and regulatory approval", "{{company.company_name}} acknowledges that the proposed name and application details are subject to approval by the relevant Singapore authorities, including ACRA. {{company.company_name}} 确认公司名称及申请资料须经新加坡相关政府机构批准。"),
        ("Director and shareholder declarations", "Each client signatory confirms that the information and declarations provided in all forms, documents and attachments are true, accurate and complete. 各签署人确认所提供的表格、文件及附件资料真实、准确及完整。"),
        ("Authority to file", "{{company.company_name}} authorises {{provider.name}}, its designated officers, employees and agents to initiate electronic filings through BizFile and other relevant government portals for the agreed services. {{company.company_name}} 授权 {{provider.name}} 及其指定人员通过 BizFile 或相关政府系统提交约定服务所需文件。"),
        ("Confirmation after incorporation", "{{company.company_name}} agrees that the ACRA business profile shall be deemed true and correct unless written notice of any error is given within seven days from the date of incorporation. {{company.company_name}} 同意如公司成立后七日内未书面提出错误，ACRA 商业档案资料视为正确。"),
        ("Commencement of responsibility", "{{provider.name}}'s responsibility for services commences only on and after the date of appointment or the accepted engagement date, unless otherwise agreed in writing. {{provider.name}} 的服务责任自委任或确认接受服务之日起开始，另有书面约定除外。"),
        ("Electronic signature", "{{company.company_name}} understands and agrees that documents may be executed by electronic signature unless wet-ink originals are specifically required. {{company.company_name}} 理解并同意文件可通过电子签名完成，除非特别要求纸质签署原件。"),
        ("No guarantee of approval", "{{company.company_name}} acknowledges that no application outcome is guaranteed, including company name approval, incorporation, change filing, pass application or renewal. {{company.company_name}} 明白任何申请结果均不作保证。"),
        ("Paid-up capital and share obligations", "{{company.company_name}} confirms that shareholders are responsible for their subscribed shares and paid-up capital as reflected in ACRA or company records. {{company.company_name}} 确认股东应按 ACRA 或公司记录承担认购股份及实缴资本责任。"),
        ("Indemnity for inaccurate information", "{{company.company_name}} and the client signatories indemnify {{provider.name}}, its officers, employees, contractors and agents against loss, expense, cost or liability arising from inaccurate, incomplete or misleading information provided by the client. {{company.company_name}} 及客户签署人同意就其提供资料不准确、不完整或误导所产生的损失向 {{provider.name}} 及相关人员作出赔偿。"),
        ("Fees and renewals", "Renewal fees, retainer fees and corporate secretarial fees for the next service period are payable in advance unless otherwise agreed. 下一服务期间的续费、年费及秘书服务费应提前支付，另有约定除外。"),
        ("Scope before payment", "{{provider.name}} is not obliged to provide secretarial services, resolutions, annual compliance work or advisory tasks before the relevant fees are paid. 在相关费用支付前，{{provider.name}} 无义务提供秘书服务、决议、年审合规或咨询工作。"),
        ("Bank documents and third-party documents", "The company secretary will not sign on behalf of other director(s) on bank documents, tenancy documents, guarantees or other third-party documents. 公司秘书不会代表其他董事签署银行文件、租赁文件、担保或第三方文件。"),
        ("Right to accept or reject", "{{provider.name}} reserves the right to accept or reject any application or service request submitted. {{provider.name}} 保留接受或拒绝任何申请或服务请求的权利。"),
        ("Termination and records", "Upon termination of services, {{provider.name}} is not obliged to furnish hard-copy records unless required by law or agreed in writing and all outstanding fees have been paid. 服务终止后，除法律要求或另有书面约定且费用已结清外，{{provider.name}} 无义务提供纸质文件。"),
        ("Share certificates", "For share transactions, {{company.company_name}} agrees that new share certificates may be issued to affected shareholders and previous affected certificates may be cancelled. 就股份交易，{{company.company_name}} 同意可向受影响股东签发新股权证书，并取消原受影响证书。"),
        ("Share certificate numbering", "Share certificate numbering may commence from the number set by {{provider.name}} or the company records, unless otherwise agreed. 股权证书编号可按 {{provider.name}} 或公司记录设定的编号开始，另有约定除外。"),
        ("Entire agreement", "This Agreement supersedes previous agreements relating to the same subject matter. 本协议取代各方就同一事项先前达成的协议。"),
        ("Governing law", "This Agreement is governed by the laws of Singapore and the parties submit to the exclusive jurisdiction of the courts of Singapore. 本协议受新加坡法律管辖，各方接受新加坡法院的专属管辖。"),
    ]
    for idx, (heading, body) in enumerate(clauses, start=1):
        section_heading(doc, f"{idx}. {heading}")
        para(doc, body)

    section_heading(doc, "Signatures")
    signature_line(doc, "{{provider.authorised_signatory_name}}", "For and on behalf of {{provider.name}}", "{{signature.date}}")
    para(doc, "{% for client_signatory in client_signatories %}", size=1, after=1)
    signature_line(doc, "{{client_signatory.full_name}}", "Client / authorised signatory", "{{client_signatory.signature_date}}")
    para(doc, "{% endfor %}", size=1, after=1)
    save_doc(doc, "05_secretary_service_agreement.docx")


def build_nominee_director_agreement() -> None:
    doc = Document()
    configure_doc(doc, "Nominee Director Agreement")
    centered_title(doc, "AGREEMENT FOR APPOINTMENT OF NOMINEE DIRECTOR", 12)
    centered_title(doc, "委任挂名董事协议", 11)
    para(doc, "THIS AGREEMENT is dated {{signature.date}}.", after=2)
    para(doc, "本协议订立于 {{signature.date}}。", after=6)
    para(doc, "Between", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "{{nominee_director.full_name}} ({{nominee_director.id_number}})", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    para(doc, "and", align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "{{company.company_name}} (Company Registration No. {{company.uen}}), a company incorporated in the Republic of Singapore and having its registered address at {{company.registered_office_address}} (the \"Company\").", align=WD_ALIGN_PARAGRAPH.CENTER)

    section_heading(doc, "WHEREAS / 鉴于")
    para(doc, "The Company has requested the nominee director arrangement stated herein for Singapore statutory compliance purposes. 公司要求按本协议约定提供挂名董事安排，以满足新加坡合规需要。")

    clauses = [
        ("Appointment and duties / 委任和责任", "The Nominee Director shall be appointed as a director of the Company upon the terms stated herein. The duties of the Nominee Director are limited to statutory and compliance matters required under the Companies Act 1967 and the Company's Constitution, unless otherwise agreed in writing. 挂名董事按本协议条款被委任为公司董事。除另有书面约定外，其职责限于新加坡公司法及公司章程要求的法定及合规事项。"),
        ("No management role / 不参与经营", "The Nominee Director is not required to participate in the commercial management, daily operation or business decision-making of the Company. 挂名董事无需参与公司的商业管理、日常运营或业务决策。"),
        ("Restricted signing authority / 签署权限限制", "The Nominee Director is not required to sign contracts, guarantees, bank facilities, loan documents, tenancy agreements or documents that may create personal liability, unless expressly agreed in writing. 除非明确书面同意，挂名董事无需签署合同、担保、银行融资、贷款、租赁或可能产生个人责任的文件。"),
        ("Bank account assistance / 银行开户协助", "Any bank account opening assistance is limited to reasonable compliance support. The Nominee Director is not responsible for bank approval, account operation, source of funds, business transactions or continuing bank requirements. 银行开户协助仅限合理合规支持，挂名董事不对银行批准、账户运作、资金来源、业务交易或后续银行要求负责。"),
        ("Information and disclosure / 信息披露", "The Company and client signatories must provide true, accurate and complete information about the Company, shareholders, beneficial owners, controllers, business activities and source of funds. 公司及客户签署人必须提供关于公司、股东、实益拥有人、控制人、业务及资金来源的真实、准确、完整资料。"),
        ("Indemnity / 赔偿", "The Company and client signatories indemnify the Nominee Director and {{provider.name}} against all losses, claims, liabilities, penalties, costs and expenses arising from the Company's business, instructions, omissions, inaccurate information, unlawful acts or breach of this Agreement. 公司及客户签署人同意就公司业务、指示、遗漏、不准确资料、违法行为或违反本协议所产生的损失、索赔、责任、罚款、成本及费用向挂名董事及 {{provider.name}} 作出赔偿。"),
        ("Compliance and refusal / 合规及拒绝权", "The Nominee Director and {{provider.name}} may refuse any instruction that is incomplete, unclear, suspicious, unlawful, non-compliant or outside the agreed scope. 挂名董事及 {{provider.name}} 可拒绝任何不完整、不清楚、可疑、违法、不合规或超出约定范围的指示。"),
        ("Fees and deposit / 费用及押金", "The Company shall pay nominee director fees, service fees and any agreed security deposit in accordance with the accepted quotation or invoice. 公司应按已接受的报价或账单支付挂名董事费、服务费及约定押金。"),
        ("Termination / 终止", "The nominee arrangement may be terminated by written notice, subject to completion of replacement director arrangements, outstanding filings, handover requirements and payment of outstanding fees. 挂名安排可通过书面通知终止，但须完成替任董事安排、未完成申报、交接要求及结清费用。"),
        ("Confidentiality / 保密", "Each party shall keep confidential the non-public information obtained through this Agreement except where disclosure is required by law, regulators, banks, professional advisers or agreed service providers. 各方应对因本协议取得的非公开资料保密，法律、监管、银行、专业顾问或约定服务提供者要求披露的除外。"),
        ("Governing law / 管辖法律", "This Agreement is governed by the laws of Singapore and the parties submit to the exclusive jurisdiction of the courts of Singapore. 本协议受新加坡法律管辖，各方接受新加坡法院专属管辖。"),
    ]
    for idx, (heading, body) in enumerate(clauses, start=1):
        section_heading(doc, f"{idx}. {heading}")
        para(doc, body)

    section_heading(doc, "Signatures / 签署")
    signature_line(doc, "{{nominee_director.full_name}}", "Nominee Director", "{{nominee_director.signature_date}}")
    signature_line(doc, "{{client_signatory.full_name}}", "For and on behalf of {{company.company_name}} / client signatory", "{{client_signatory.signature_date}}")
    para(doc, "Manual review note: use this template only when a nominee director arrangement is actually selected and the signing relationship is confirmed.", italic=True, size=8)
    save_doc(doc, "06_nominee_director_agreement.docx")


def build_signature_record_attachment() -> None:
    doc = Document()
    configure_doc(doc, "Signature Record Attachment")
    centered_title(doc, "SIGNATURE RECORD ATTACHMENT", 12)
    centered_title(doc, "签名记录附件", 11)
    para(doc, "Company: {{company.company_name}}", after=1)
    para(doc, "Document: {{signature_record.document_title}}", after=1)
    para(doc, "Envelope / Document ID: {{signature_record.envelope_id}}", after=1)
    para(doc, "Generated on: {{signature_record.generated_at}}", after=8)

    section_heading(doc, "Signer Summary")
    tbl = table(doc, 2, 5, [1.25, 1.3, 1.55, 1.1, 1.3])
    set_borders(tbl)
    for idx, label in enumerate(["Signer", "Email", "Party ID", "Status", "Signed At"]):
        cell_text(tbl.cell(0, idx), label, bold=True, size=8)
    values = ["{{signer.full_name}}", "{{signer.email}}", "{{signer.party_id}}", "{{signer.status}}", "{{signer.signed_at}}"]
    for idx, value in enumerate(values):
        cell_text(tbl.cell(1, idx), "{% for signer in signature_record.signers %}" + value if idx == 0 else value, size=8)
    para(doc, "{% endfor %}", size=1, after=4)

    section_heading(doc, "Audit Trail")
    tbl = table(doc, 2, 4, [1.35, 1.15, 1.35, 2.55])
    set_borders(tbl)
    for idx, label in enumerate(["Timestamp", "Action", "IP Address", "Details"]):
        cell_text(tbl.cell(0, idx), label, bold=True, size=8)
    values = ["{{event.timestamp}}", "{{event.action}}", "{{event.ip_address}}", "{{event.details}}"]
    for idx, value in enumerate(values):
        cell_text(tbl.cell(1, idx), "{% for event in signature_record.events %}" + value if idx == 0 else value, size=8)
    para(doc, "{% endfor %}", size=1, after=4)

    section_heading(doc, "Digital Fingerprint")
    para(doc, "Checksum: {{signature_record.checksum}}", after=2)
    para(doc, "This attachment is intended to be populated after electronic signing is completed and appended to the relevant signed document.", italic=True, size=8)
    save_doc(doc, "07_signature_record_attachment.docx")


def write_field_guides() -> None:
    field_json = {
        "placeholder_style": "Jinja/docxtpl style: {{company.company_name}}, {% for director in directors %} ... {% endfor %}",
        "notes": [
            "Original files are not modified.",
            "Legacy names Form 45 and Form 45B are kept as display references only.",
            "Template IDs use stable business names for rule matching.",
            "Shareholder display fields such as shareholder.shareholder_name should be derived by the website from person/corporate shareholder rows.",
            "Nominee director agreement remains manual-review because signing relationship and risk scope must be confirmed.",
        ],
        "templates": MANIFEST,
        "fields": FIELD_GROUPS,
    }
    for directory in (TEMPLATE_DIR, OUTPUT_DIR):
        (directory / "template_manifest.json").write_text(json.dumps(field_json, ensure_ascii=False, indent=2), encoding="utf-8")
        readme = [
            "# P1 Rebuilt Template Pack",
            "",
            "This folder contains rebuilt DOCX templates for the first internal website launch. The original template folder was read only; no original file was changed.",
            "",
            "## Template files",
            "",
        ]
        for item in MANIFEST:
            readme.append(f"- `{item['file_name']}` - {item['display_name']} ({item['repeat']})")
        readme.extend(
            [
                "",
                "## Placeholder convention",
                "",
                "- Single value: `{{company.company_name}}`",
                "- List loop: `{% for director in directors %}` ... `{% endfor %}`",
                "- Keep internal rule names stable; old names such as Form 45/45B are display labels only.",
                "",
                "## Fields",
                "",
            ]
        )
        for group, rows in FIELD_GROUPS.items():
            readme.append(f"### {group}")
            readme.append("")
            readme.append("| Placeholder | Meaning | Source / default |")
            readme.append("|---|---|---|")
            for placeholder, meaning, source in rows:
                readme.append(f"| `{{{{{placeholder}}}}}` | {meaning} | {source} |")
            readme.append("")
        (directory / "README_fields.md").write_text("\n".join(readme), encoding="utf-8")


def zip_output() -> None:
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(OUTPUT_DIR.iterdir()):
            if path.is_file():
                zf.write(path, arcname=path.name)


def main() -> None:
    clean_dirs()
    build_first_directors_resolution()
    build_director_consent()
    build_secretary_consent()
    build_share_certificate()
    build_secretary_service_agreement()
    build_nominee_director_agreement()
    build_signature_record_attachment()
    write_field_guides()
    zip_output()
    print(f"Generated {len(MANIFEST)} templates in {TEMPLATE_DIR}")
    print(f"Copied package to {OUTPUT_DIR}")
    print(f"Zip package: {ZIP_PATH}")


if __name__ == "__main__":
    main()
