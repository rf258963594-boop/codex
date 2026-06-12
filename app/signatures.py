from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont

from config import DATA_DIR


SIGNATURE_DIR = DATA_DIR / "signatures"
SIGNATURE_WIDTH = 680
SIGNATURE_HEIGHT = 220
SIGNATURE_DOCX_WIDTH_IN = 1.55


def clean(value: Any) -> str:
    return "" if value is None else str(value).strip()


def truthy(value: Any) -> bool:
    return str(value or "").strip().lower() in {"1", "yes", "y", "true", "on"}


def safe_slug(value: Any) -> str:
    slug = re.sub(r"[^A-Za-z0-9]+", "_", clean(value)).strip("_").lower()
    return slug[:60] or "signature"


def default_signature_text(name: str) -> str:
    parts = [part for part in re.split(r"\s+", clean(name)) if part]
    if not parts:
        return ""
    if len(parts) == 1:
        return parts[0].title()
    if len(parts) == 2:
        return f"{parts[0][0].upper()}. {parts[1].title()}"
    return f"{parts[0][0].upper()}. {parts[-1].title()}"


def signature_relative_path(path: Path) -> str:
    try:
        return str(path.resolve().relative_to(DATA_DIR.resolve()))
    except ValueError:
        return str(path.resolve())


def resolve_signature_path(value: Any) -> Path | None:
    text = clean(value)
    if not text:
        return None
    path = Path(text)
    if not path.is_absolute():
        path = DATA_DIR / path
    return path if path.exists() else None


def ensure_signature_image(text: str, name_hint: str = "", *, force: bool = True) -> str:
    signature_text = clean(text) or default_signature_text(name_hint)
    if not signature_text:
        return ""
    SIGNATURE_DIR.mkdir(parents=True, exist_ok=True)
    path = SIGNATURE_DIR / f"{safe_slug(name_hint or signature_text)}.png"
    if force or not path.exists():
        create_signature_image(signature_text, path)
    return signature_relative_path(path)


def create_signature_image(text: str, output_path: Path) -> None:
    image = Image.new("RGBA", (SIGNATURE_WIDTH, SIGNATURE_HEIGHT), (255, 255, 255, 0))
    draw = ImageDraw.Draw(image)
    font = load_signature_font(text)

    bbox = draw.textbbox((0, 0), text, font=font)
    text_width = bbox[2] - bbox[0]
    text_height = bbox[3] - bbox[1]
    x = max(18, (SIGNATURE_WIDTH - text_width) // 2)
    y = max(10, (SIGNATURE_HEIGHT - text_height) // 2 - 8)

    color = (22, 30, 40, 235)
    shadow = (22, 30, 40, 42)
    draw.text((x + 2, y + 2), text, font=font, fill=shadow)
    draw.text((x, y), text, font=font, fill=color)

    rotated = image.rotate(-2.2, resample=Image.Resampling.BICUBIC, expand=True)
    bbox = rotated.getbbox()
    cropped = rotated.crop(bbox) if bbox else rotated
    final_image = Image.new("RGBA", (SIGNATURE_WIDTH, SIGNATURE_HEIGHT), (255, 255, 255, 0))
    paste_x = max(0, (SIGNATURE_WIDTH - cropped.width) // 2)
    paste_y = max(0, (SIGNATURE_HEIGHT - cropped.height) // 2)
    final_image.alpha_composite(cropped, (paste_x, paste_y))
    final_image.save(output_path)


def load_signature_font(text: str) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        r"C:\Windows\Fonts\segoesc.ttf",
        r"C:\Windows\Fonts\FRSCRIPT.TTF",
        r"C:\Windows\Fonts\BRUSHSCI.TTF",
        r"C:\Windows\Fonts\MISTRAL.TTF",
        "/usr/share/fonts/truetype/liberation2/LiberationSerif-Italic.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Italic.ttf",
    ]
    for size in range(86, 44, -4):
        for candidate in candidates:
            path = Path(candidate)
            if not path.exists():
                continue
            try:
                font = ImageFont.truetype(str(path), size=size)
            except OSError:
                continue
            test_image = Image.new("RGBA", (SIGNATURE_WIDTH, SIGNATURE_HEIGHT), (255, 255, 255, 0))
            bbox = ImageDraw.Draw(test_image).textbbox((0, 0), text, font=font)
            if bbox[2] - bbox[0] <= SIGNATURE_WIDTH - 48:
                return font
    return ImageFont.load_default()


def apply_auto_signatures(doc, context: dict[str, Any]) -> list[str]:
    signatures = collect_auto_signatures(context)
    if not signatures:
        return []
    applied: list[str] = []
    touched: set[tuple[int, str]] = set()
    for group in paragraph_groups(doc):
        apply_signatures_to_paragraph_group(group, signatures, touched, applied)
    for table in doc.tables:
        apply_signatures_to_table(table, signatures, touched, applied)
    return applied


def collect_auto_signatures(context: dict[str, Any]) -> list[dict[str, Any]]:
    found: dict[tuple[str, str], dict[str, Any]] = {}

    def walk(value: Any) -> None:
        if isinstance(value, dict):
            if truthy(value.get("auto_signature_enabled")):
                path = resolve_signature_path(value.get("signature_image_path"))
                name = clean(value.get("full_name") or value.get("common_person_name") or value.get("display_name"))
                if path and name:
                    aliases = {
                        name,
                        clean(value.get("full_name")),
                        clean(value.get("common_person_name")),
                        clean(value.get("common_person_display_name")),
                        clean(value.get("display_name")),
                    }
                    aliases = {alias for alias in aliases if alias}
                    item = {
                        "name": name,
                        "aliases": aliases,
                        "path": path,
                    }
                    for alias in aliases:
                        found[(normalize_name(alias), str(path))] = item
            for child in value.values():
                walk(child)
        elif isinstance(value, list):
            for child in value:
                walk(child)

    walk(context)
    unique: dict[str, dict[str, Any]] = {}
    for item in found.values():
        unique[f"{item['name']}::{item['path']}"] = item
    return list(unique.values())


def normalize_name(value: Any) -> str:
    return re.sub(r"[^a-z0-9]+", "", clean(value).lower())


def paragraph_groups(container) -> list[list[Any]]:
    groups: list[list[Any]] = []
    if getattr(container, "paragraphs", None):
        groups.append(list(container.paragraphs))
    for table in getattr(container, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                groups.extend(paragraph_groups(cell))
    return groups


def apply_signatures_to_paragraph_group(
    paragraphs: list[Any],
    signatures: list[dict[str, Any]],
    touched: set[tuple[int, str]],
    applied: list[str],
) -> None:
    for idx, paragraph in enumerate(paragraphs):
        text = clean(paragraph.text)
        if not text:
            continue
        if contains_signature_cue(text):
            matching = [signature for signature in signatures if text_matches_signature(text, signature)]
            if matching:
                insert_signature_pictures_near_names(paragraph, matching, touched, applied)
            continue
        for signature in signatures:
            if not text_matches_signature(text, signature):
                continue
            target = previous_signature_paragraph(paragraphs, idx)
            if target is not None:
                insert_signature_picture(target, signature["path"], touched, applied, signature["name"])


def apply_signatures_to_table(
    table,
    signatures: list[dict[str, Any]],
    touched: set[tuple[int, str]],
    applied: list[str],
) -> None:
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell_text = clean(cell.text)
            for signature in signatures:
                if not text_matches_signature(cell_text, signature):
                    continue
                target = previous_signature_cell(table, row_idx, col_idx)
                if target is not None:
                    paragraph = first_signature_paragraph(target)
                    if paragraph is not None:
                        insert_signature_picture(paragraph, signature["path"], touched, applied, signature["name"])
            for nested_table in cell.tables:
                apply_signatures_to_table(nested_table, signatures, touched, applied)


def text_matches_signature(text: str, signature: dict[str, Any]) -> bool:
    normalized = normalize_name(text)
    if not normalized:
        return False
    for alias in signature["aliases"]:
        alias_key = normalize_name(alias)
        if len(alias_key) >= 3 and alias_key in normalized:
            return True
    return False


def contains_signature_cue(text: str) -> bool:
    normalized = clean(text).lower()
    return "signature" in normalized or bool(re.search(r"_{6,}|-{6,}", normalized))


def previous_signature_paragraph(paragraphs: list[Any], index: int) -> Any | None:
    for pos in range(index - 1, max(-1, index - 4), -1):
        paragraph = paragraphs[pos]
        if contains_signature_cue(paragraph.text):
            return paragraph
    return None


def previous_signature_cell(table, row_idx: int, col_idx: int):
    for prev_idx in range(row_idx - 1, max(-1, row_idx - 4), -1):
        prev_row = table.rows[prev_idx]
        if col_idx >= len(prev_row.cells):
            continue
        cell = prev_row.cells[col_idx]
        if contains_signature_cue(cell.text):
            return cell
    return None


def first_signature_paragraph(cell) -> Any | None:
    for paragraph in cell.paragraphs:
        if contains_signature_cue(paragraph.text):
            return paragraph
    return cell.paragraphs[0] if cell.paragraphs else None


def insert_signature_picture(
    paragraph,
    image_path: Path,
    touched: set[tuple[int, str]],
    applied: list[str],
    signer_name: str,
    *,
    keep_remaining: bool = False,
) -> None:
    marker = (id(paragraph._p), normalize_name(signer_name))
    if marker in touched:
        return
    touched.add(marker)
    original_lines = clean(paragraph.text).splitlines()
    first_line = original_lines[0] if original_lines else ""
    remaining = original_lines[1:] if keep_remaining and original_lines else []
    prefix = "Signature: " if "signature" in first_line.lower() else ""
    paragraph.clear()
    if prefix:
        paragraph.add_run(prefix)
    paragraph.add_run().add_picture(str(image_path), width=Inches(SIGNATURE_DOCX_WIDTH_IN))
    for line in remaining:
        paragraph.add_run().add_break()
        paragraph.add_run(line)
    applied.append(signer_name)


def insert_signature_pictures_near_names(
    paragraph,
    signatures: list[dict[str, Any]],
    touched: set[tuple[int, str]],
    applied: list[str],
) -> None:
    original_lines = clean(paragraph.text).splitlines()
    if not original_lines:
        return

    targets: dict[int, dict[str, Any]] = {}
    used_signature_lines: set[int] = set()
    paragraph_id = id(paragraph._p)
    for signature in signatures:
        marker = (paragraph_id, normalize_name(signature["name"]))
        if marker in touched:
            continue
        signature_idx = matching_signature_line_index(original_lines, signature, used_signature_lines)
        if signature_idx is None:
            continue
        targets[signature_idx] = signature
        used_signature_lines.add(signature_idx)
        touched.add(marker)

    if not targets:
        return

    paragraph.clear()
    for idx, line in enumerate(original_lines):
        if idx:
            paragraph.add_run().add_break()
        signature = targets.get(idx)
        if signature:
            prefix = "Signature: " if "signature" in line.lower() else ""
            if prefix:
                paragraph.add_run(prefix)
            paragraph.add_run().add_picture(str(signature["path"]), width=Inches(SIGNATURE_DOCX_WIDTH_IN))
            applied.append(signature["name"])
        else:
            paragraph.add_run(line)


def matching_signature_line_index(
    lines: list[str],
    signature: dict[str, Any],
    used_signature_lines: set[int],
) -> int | None:
    matching_name_index = next(
        (idx for idx, line in enumerate(lines) if text_matches_signature(line, signature)),
        None,
    )
    if matching_name_index is not None:
        for idx in range(matching_name_index - 1, -1, -1):
            if idx not in used_signature_lines and contains_signature_cue(lines[idx]):
                return idx
    return next(
        (idx for idx, line in enumerate(lines) if idx not in used_signature_lines and contains_signature_cue(line)),
        None,
    )
