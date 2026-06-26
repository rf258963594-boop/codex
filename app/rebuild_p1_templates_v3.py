from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import rebuild_p1_templates_v2 as v2


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DIR = ROOT / "app" / "doc_templates" / "p1_standard_v3_part1"
OUTPUT_DIR = ROOT / "outputs" / "P1_standard_templates_v3_part1"
ZIP_PATH = ROOT / "outputs" / "P1_standard_templates_v3_part1.zip"
V2_OUTPUT_DIR = ROOT / "outputs" / "P1_preserved_templates_v2"

BASE_FONT = "Arial"
EAST_ASIA_FONT = "Microsoft YaHei"
INK = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x66, 0x72, 0x85)
ACCENT = RGBColor(0x1F, 0x4D, 0x78)
LIGHT_FILL = "F3F6FA"
TABLE_BORDER = "C9D2DF"


MANIFEST = [
    {
        **item,
        "file_name": item["file_name"].replace("_preserved.docx", "_standard.docx"),
        "build_method": "v3_standardized_layout_content_preserved",
    }
    for item in v2.MANIFEST
    if item["template_id"] != "signature_record_attachment"
]

MANIFEST_OVERRIDES = {
    "first_directors_resolution": {
        "signing_logic": "All directors sign in one document using company.director_signature_blocks.",
    },
    "share_certificate": {
        "signing_logic": "Certificate is executed by the company. The director signature line is intentionally blank for the actual signing director; the secretary line uses the company secretary. The system prefers a non-nominee/client director for certificate context and uses a nominee director only as fallback.",
    },
    "secretary_service_agreement": {
        "signing_logic": "Provider signs and one client representative signs. The client representative defaults to shareholder 1, then a marked client signatory, then director 1.",
    },
    "nominee_director_agreement": {
        "signing_logic": "Nominee director signs and one client representative signs. The client representative defaults to shareholder 1.",
    },
}

for item in MANIFEST:
    item.update(MANIFEST_OVERRIDES.get(item["template_id"], {}))

MANIFEST.extend(
    [
        {
            "template_id": "return_of_allotment_form24",
            "file_name": "07_return_of_allotment_form24_standard.docx",
            "display_name": "Return of Allotment of Shares / Form 24",
            "source": "2. Company Name + FORM 24.doc/pdf",
            "build_method": "v3_standardized_layout_from_source_pdf_shareholder_fields",
            "repeat": "one_per_company_allotment_with_signature_blocks_for_each_shareholder; duplicate per shareholder if separate signing copies are required",
            "manual_review": True,
        },
        {
            "template_id": "rorc_notice_controller",
            "file_name": "08_rorc_notice_controller_standard.docx",
            "display_name": "RORC Notice to Registrable Controller",
            "source": "6. Company Name + RORC.doc/pdf",
            "build_method": "v3_standardized_layout_from_source_pdf_shareholder_controller",
            "repeat": "one_per_registrable_controller_shareholder",
            "manual_review": True,
        },
        {
            "template_id": "register_of_members",
            "file_name": "09_register_of_members_standard.docx",
            "display_name": "Register of Members",
            "source": "Generated initial statutory register from shareholder rows",
            "build_method": "v3_standardized_register_template",
            "repeat": "one_per_company",
            "manual_review": True,
        },
        {
            "template_id": "paid_up_capital_confirmation",
            "file_name": "10_paid_up_capital_confirmation_standard.docx",
            "display_name": "Confirmation of Paid-Up Capital",
            "source": "Client director paid-up capital confirmation sample",
            "build_method": "v3_standardized_layout_new_confirmation",
            "repeat": "one_per_company",
            "signing_logic": "Signed by a client-side director; the system prefers a non-nominee director and only falls back if none is available.",
            "manual_review": True,
        },
    ]
)

FIELD_NOTES = {
    **{key: value for key, value in v2.FIELD_NOTES.items() if key not in {"signature_record.*", "company.fye"}},
    "company.share_currency": "Share currency; default usually SGD.",
    "company.subscriber_share_totals": "Derived total/share-count cell used by the v3 subscriber shares table.",
    "company.director_signature_blocks": "Derived signature lines for all directors in the first directors resolution.",
    "company.share_par_value": "Legacy compatibility value; current templates prefer issued/paid-up share capital fields.",
    "company.fye_month": "Financial year end month, auto-derived from incorporation date unless overridden.",
    "company.first_financial_period_start": "First accounts period start, defaults to incorporation date and renders as a long date.",
    "company.first_financial_period_end": "First accounts period end, auto-derived as the last day of the month before the incorporation month unless overridden.",
    "client_signatory.*": "Single client-side signer. Defaults to shareholder 1, then a marked client signatory, then director 1.",
    "company.issued_share_capital": "Derived total issued share capital from shareholder rows; defaults to shares for normal 1:1 ordinary shares.",
    "company.paid_up_capital": "Derived total paid-up share capital from shareholder rows; defaults to issued share capital when blank.",
    "company.unpaid_share_capital": "Derived unpaid share capital; used as a review flag when issued capital is higher than paid-up capital.",
    "company.amount_paid_per_share": "Derived paid amount per share for Form 24.",
    "company.amount_due_per_share": "Derived amount due per share for Form 24; '-' when fully paid.",
    "shareholders[]": "Repeating shareholder rows used by Form 24 and RORC; Form 24 allottees and shareholder signature blocks repeat from this list.",
    "shareholder.date_of_birth": "Derived from the matched People row when the shareholder is an individual.",
    "company.shareholder_signature_blocks": "Derived Form 24 signature lines, one block for each shareholder/allottee.",
    "shareholder.issued_share_capital": "Shareholder-level issued share capital.",
    "shareholder.paid_up_share_capital": "Shareholder-level paid-up share capital.",
    "shareholder.unpaid_share_capital": "Shareholder-level unpaid share capital.",
    "shareholder.paid_status_text": "Share certificate wording, e.g. fully paid or partly paid details.",
    "shareholder.form24_allotment_text": "Form 24 allotment summary with shares, issued capital, paid-up capital and unpaid capital when applicable.",
    "paid_up_confirmation_signatory.*": "Client-side director context for the paid-up capital confirmation; prefers a non-nominee director.",
}

P2_DEFERRED = [
    {
        "template_id": "signature_record_attachment",
        "reason": "Online e-signature audit attachment; keep for P2 when e-signing is enabled.",
    },
    {
        "template_id": "statutory_registers",
        "reason": "Registers pack is useful later but not complete/mandatory for the first internal file generator.",
    },
]


def ensure_dirs() -> None:
    for directory in (TEMPLATE_DIR, OUTPUT_DIR):
        directory.mkdir(parents=True, exist_ok=True)
        for path in directory.iterdir():
            if path.is_file():
                path.unlink()


def ensure_v2_sources() -> None:
    if not V2_OUTPUT_DIR.exists() or not list(V2_OUTPUT_DIR.glob("*.docx")):
        v2.main()


def set_run_font(
    run,
    size: float | int | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor | None = None,
    font: str = BASE_FONT,
    east_asia_font: str = EAST_ASIA_FONT,
) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def configure_style(style, *, size: float, bold: bool = False, color: RGBColor | None = None, before: float = 0, after: float = 6, line_spacing: float = 1.1) -> None:
    style.font.name = BASE_FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = color
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def configure_doc(doc: Document, title: str, *, compact: bool = False, landscape: bool = False, page_numbers: bool = True) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(11 if landscape else 8.5)
    section.page_height = Inches(8.5 if landscape else 11)
    if landscape:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
    else:
        section.top_margin = Inches(0.85 if compact else 1.0)
        section.bottom_margin = Inches(0.85 if compact else 1.0)
        section.left_margin = Inches(0.85 if compact else 1.0)
        section.right_margin = Inches(0.85 if compact else 1.0)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    configure_style(doc.styles["Normal"], size=10.5 if compact else 11, color=INK, after=6, line_spacing=1.12)
    for name in ("Title", "Heading 1", "Heading 2", "Heading 3", "Body Text", "List Paragraph"):
        if name in doc.styles:
            style = doc.styles[name]
            if name == "Title":
                configure_style(style, size=14, bold=True, color=INK, before=0, after=8, line_spacing=1.0)
            elif name == "Heading 1":
                configure_style(style, size=12.5, bold=True, color=INK, before=12, after=6, line_spacing=1.05)
            elif name == "Heading 2":
                configure_style(style, size=11.5, bold=True, color=ACCENT, before=9, after=4, line_spacing=1.05)
            elif name == "Heading 3":
                configure_style(style, size=10.5, bold=True, color=ACCENT, before=7, after=3, line_spacing=1.05)
            else:
                configure_style(style, size=10.5 if compact else 11, color=INK, after=6, line_spacing=1.12)

    for style_name, size, after, indent in [
        ("Legal Clause", 10.25 if compact else 10.75, 6, 0.28),
        ("Legal Chinese", 10.0 if compact else 10.5, 4, 0.0),
        ("Form Note", 9.0, 4, 0.0),
        ("Signature Label", 10.0, 3, 0.0),
    ]:
        if style_name not in doc.styles:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = doc.styles[style_name]
        configure_style(style, size=size, color=MUTED if style_name == "Form Note" else INK, after=after, line_spacing=1.08 if compact else 1.12)
        if indent:
            style.paragraph_format.left_indent = Inches(indent)
            style.paragraph_format.first_line_indent = Inches(-indent)

    doc.core_properties.title = title
    doc.core_properties.author = "RSIN template rebuild v3"
    doc.core_properties.comments = "Standardized P1 template v3. Content preserved from v2/source; old source files were not modified."
    if page_numbers:
        add_footer(doc)


def add_field(paragraph, code: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = code
    run._r.append(instr)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)


def add_footer(doc: Document) -> None:
    paragraph = doc.sections[0].footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    add_field(paragraph, "PAGE")
    run = paragraph.add_run(" of ")
    set_run_font(run, size=8.5, color=MUTED)
    add_field(paragraph, "NUMPAGES")


def add_para(
    doc: Document,
    text: str = "",
    *,
    style: str | None = None,
    align: WD_ALIGN_PARAGRAPH | None = None,
    size: float | int | None = None,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    before: float | int | None = None,
    after: float | int | None = None,
    keep_with_next: bool = False,
) -> object:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if keep_with_next:
        p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_title(doc: Document, company_line: bool = True, title: str = "") -> None:
    if company_line:
        add_para(doc, "{{company.company_name}}", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True, after=2, keep_with_next=True)
        add_para(doc, "Company Registration No.: {{company.uen}}", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, after=0, keep_with_next=True)
        add_para(doc, "(Incorporated in the Republic of Singapore)", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, after=12, keep_with_next=True)
    if title:
        add_para(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True, after=14, keep_with_next=True)


def add_rule(doc: Document) -> None:
    p = add_para(doc, "", after=10)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "9AA6B2")
    borders.append(bottom)
    p_pr.append(borders)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 10.5, align: WD_ALIGN_PARAGRAPH | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=INK)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color: str = TABLE_BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_key_value_table(doc: Document, rows: list[tuple[str, str]], widths: tuple[float, float] = (1.85, 4.65)) -> object:
    table = doc.add_table(rows=len(rows), cols=2)
    table_widths(table, list(widths))
    set_table_borders(table)
    set_cell_margins(table)
    for idx, (label, value) in enumerate(rows):
        shade_cell(table.cell(idx, 0), LIGHT_FILL)
        set_cell_text(table.cell(idx, 0), label, bold=True, size=9.8)
        set_cell_text(table.cell(idx, 1), value, size=10.2)
    return table


def add_signature_table(doc: Document, parties: list[tuple[str, str, str]], *, page_break_before: bool = False) -> None:
    if page_break_before:
        doc.add_page_break()
    add_para(doc, "Execution", style="Heading 2", keep_with_next=True)
    table = doc.add_table(rows=3, cols=len(parties))
    table_widths(table, [6.5 / len(parties)] * len(parties))
    set_table_borders(table, "D8DEE8")
    set_cell_margins(table, top=120, bottom=120, start=160, end=160)
    signature_line = "Signature:\n\n____________________" if len(parties) >= 3 else "\n\nSignature: ______________________________"
    for col, (label, name, capacity) in enumerate(parties):
        shade_cell(table.cell(0, col), LIGHT_FILL)
        set_cell_text(table.cell(0, col), label, bold=True, size=10.2, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.cell(1, col), signature_line, size=10.2)
        set_cell_text(table.cell(2, col), f"Name: {name}\nCapacity: {capacity}", size=10.2)
    add_para(doc, "", after=4)


def add_section_heading(doc: Document, number: str, heading: str, *, resolved: bool = False) -> None:
    label = f"{number}.  {heading}" if number else heading
    add_para(doc, label, style="Heading 1", keep_with_next=True)
    if resolved:
        add_para(doc, "Resolved -", bold=True, after=4, keep_with_next=True)


def add_legal_clause(doc: Document, label: str, text: str, *, after: float = 6) -> None:
    p = add_para(doc, "", style="Legal Clause", after=after)
    run = p.add_run(label)
    set_run_font(run, bold=True, color=INK)
    run = p.add_run(text)
    set_run_font(run, color=INK)


def clean_source_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    replacements = {
        "{{provider.name}} 及其相关公司及相关公司": "{{provider.name}} 及相关公司",
        "{{provider.name}} 及相关公司及相关公司": "{{provider.name}} 及相关公司",
        "with in accordance": "in accordance",
        "share certificates for the affected will be cancelled": "share certificates for the affected shareholders will be cancelled",
        "All previous shares certificates": "All previous share certificates",
        "in the 12- months period": "in the 12-month period",
        "hold {{provider.name}}": "hold {{provider.name}}",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def source_paragraphs(file_name: str) -> list[str]:
    doc = Document(V2_OUTPUT_DIR / file_name)
    return [clean_source_text(p.text) for p in doc.paragraphs if clean_source_text(p.text)]


def save(doc: Document, file_name: str) -> None:
    doc.save(TEMPLATE_DIR / file_name)
    doc.save(OUTPUT_DIR / file_name)


def build_first_directors_resolution() -> None:
    doc = Document()
    configure_doc(doc, "First Directors Resolution")
    add_title(doc, title="DIRECTORS' RESOLUTION IN WRITING")
    add_rule(doc)
    add_section_heading(doc, "1", "CERTIFICATE CONFIRMING THE INCORPORATION OF COMPANY")
    add_para(doc, "Noted -", bold=True, keep_with_next=True)
    add_para(doc, "The Certificate numbered {{company.incorporation_certificate_no}} confirming Incorporation of Company on {{company.incorporation_date}} together with a copy of the Constitution of the Company.")
    add_section_heading(doc, "2", "CONFIRMATION OF FIRST DIRECTORS", resolved=True)
    add_para(doc, "That {{company.first_directors_names}} be and are hereby confirmed as the First Directors of the Company.")
    add_section_heading(doc, "3", "APPOINTMENT OF SECRETARY", resolved=True)
    add_para(doc, "That {{secretary.full_name}} is hereby appointed as Secretary of the Company with immediate effect.")
    add_para(doc, "That the Accounting and Corporate Regulatory Authority, Singapore be notified accordingly.")
    add_section_heading(doc, "4", "CONFIRMATION OF REGISTERED OFFICE AND PLACE WHERE REGISTER OF MEMBERS AND INDEX IS KEPT", resolved=True)
    add_para(doc, "That the Registered Office of the Company situated at {{company.registered_office_address}} since the date of incorporation be and is hereby confirmed.")
    add_para(doc, "That the Place where Register of Members and Index is kept situated at {{company.register_location}} be and is hereby confirmed.")
    add_section_heading(doc, "5", "ADOPTION OF COMMON SEAL", resolved=True)
    add_para(doc, "That the Seal, an impression of which is affixed to the margin of these resolutions be and is hereby adopted as the Common Seal of the Company.")
    add_section_heading(doc, "6", "ISSUE OF SUBSCRIBERS' SHARES", resolved=True)
    add_para(doc, "That the signatories to the Constitution be registered as members in respect of the shares for which they subscribed in full, in cash at {{company.share_currency}} {{company.share_par_value}} /- per share, namely: -", keep_with_next=True)
    table = doc.add_table(rows=2, cols=2)
    table_widths(table, [3.45, 3.05])
    set_table_borders(table)
    set_cell_margins(table)
    shade_cell(table.cell(0, 0), LIGHT_FILL)
    shade_cell(table.cell(0, 1), LIGHT_FILL)
    set_cell_text(table.cell(0, 0), "Name of Subscriber", bold=True)
    set_cell_text(table.cell(0, 1), "No. of shares of {{company.share_currency}} {{company.share_par_value}} /- each", bold=True)
    set_cell_text(table.cell(1, 0), "{{company.subscriber_share_lines}}")
    set_cell_text(table.cell(1, 1), "{{company.subscriber_share_totals}}")
    add_para(doc, "That share certificate be issued to the subscribers and that the Common Seal be affixed thereto in accordance with the Company's Constitution.", after=8)
    add_section_heading(doc, "7", "FINANCIAL YEAR END", resolved=True)
    add_para(doc, "That the financial year end of the Company be fixed on {{company.fye}} and the first set of accounts will be for the period from {{company.first_financial_period_start}} to {{company.first_financial_period_end}}.")
    add_para(doc, "", after=10)
    add_signature_table(
        doc,
        [("Signed by", "{{director.full_name}}", "Director")],
    )
    add_para(doc, "Dated: {{signature.date}}", style="Signature Label")
    save(doc, "01_first_directors_resolution_standard.docx")


def build_director_consent() -> None:
    doc = Document()
    configure_doc(doc, "Consent to Act as Director / Form 45")
    add_para(doc, "FORM 45", align=WD_ALIGN_PARAGRAPH.RIGHT, size=11, bold=True, after=0)
    add_title(doc, company_line=False, title="CONSENT TO ACT AS DIRECTOR AND STATEMENT OF NON-DISQUALIFICATION")
    add_key_value_table(
        doc,
        [
            ("Name of Company", "{{company.company_name}}"),
            ("Company No.", "{{company.uen}}"),
            ("Director", "{{director.full_name}}"),
            ("Effective Date", "{{director.appointment_date}}"),
        ],
    )
    add_para(doc, "I, the undermentioned person, hereby consent to act as a director of the abovenamed company with effect from {{director.appointment_date}} (date) and declare that:")
    add_section_heading(doc, "(a)", "I am not disqualified from acting as a director, in that:")
    director_items = [
        "I am not an undischarged bankrupt;",
        "I have not been convicted, whether in Singapore or elsewhere, of any offence in connection with the promotion, formation or management of a corporation or any offence involving fraud or dishonesty punishable on conviction with imprisonment for three months or more;",
        "I have not had a disqualification order made by a court under section 149 of the Companies Act;",
        "I have not had a disqualification order made by the Registrar under section 155A of the Companies Act;",
        "I have not been disqualified under section 154(1) or section 154(2) of the Companies Act;",
        "I have not been disqualified under section 155 or section 155AA of the Companies Act;",
        "I have not been debarred under section 155B of the Companies Act.",
    ]
    for idx, item in enumerate(director_items, start=1):
        add_legal_clause(doc, f"({idx}) ", item, after=4)
    add_section_heading(doc, "(b)", "I am aware of and undertake to abide by my duties, responsibilities and liabilities specified in the Act as well as under the common law where applicable, including the following key administrative and substantive duties, that is, to:")
    duty_items = [
        "use reasonable diligence in the discharge of the duties of my office;",
        "act honestly and use reasonable diligence in the discharge of the duties of my office;",
        "not make improper use of my position to gain, directly or indirectly, an advantage for myself or for any other person or to cause detriment to the company;",
        "ensure that the company complies with the relevant requirements of the Companies Act, including requirements relating to registered office, registers, records, annual returns and financial statements where applicable.",
    ]
    for idx, item in enumerate(duty_items, start=1):
        add_legal_clause(doc, f"({idx}) ", item, after=4)
    add_section_heading(doc, "(c)", "That -")
    add_para(doc, "*(i) I have read and understood the above statements; or")
    add_para(doc, "*(ii) the above statements were interpreted to me in {{director.interpretation_language}} (state language/dialect) by {{director.interpreter_name}} NRIC NO: {{director.interpreter_id_number}} before I executed this form and I confirm that the statements are true. I am also aware that I can be prosecuted in Court if I willfully give any information on this form which is false.")
    add_key_value_table(
        doc,
        [
            ("Name", "{{director.full_name}}"),
            ("Email", "{{director.email}}"),
            ("Contact Number", "{{director.phone}}"),
            ("Address", "{{director.residential_address}}"),
            ("ID / Nationality", "{{director.id_type}} No: {{director.id_number}} / {{director.nationality}}"),
        ],
    )
    add_signature_table(doc, [("Signed by", "{{director.full_name}}", "Director")])
    add_para(doc, "Dated this {{signature.day}} day of {{signature.month_year}}", style="Signature Label")
    save(doc, "02_director_consent_form45_standard.docx")


def build_secretary_consent() -> None:
    doc = Document()
    configure_doc(doc, "Consent to Act as Secretary / Form 45B")
    add_para(doc, "FORM 45B", align=WD_ALIGN_PARAGRAPH.RIGHT, size=11, bold=True, after=0)
    add_title(doc, company_line=False, title="CONSENT TO ACT AS SECRETARY")
    add_key_value_table(
        doc,
        [
            ("Name of Company", "{{company.company_name}}"),
            ("Company No.", "{{company.uen}}"),
            ("Secretary", "{{secretary.full_name}}"),
            ("Effective Date", "{{secretary.appointment_date}}"),
        ],
    )
    add_para(doc, "1. I, {{secretary.full_name}}, the undermentioned person, hereby consent to act as a secretary of the abovenamed company with effect from {{secretary.appointment_date}} (date).")
    add_para(doc, "\u20202. I am a qualified person under section 171(1AA) of the Companies Act by virtue of my being -")
    secretary_items = [
        "*(i) a secretary of a company for at least 3 years of the 5 years immediately preceding the abovementioned date of my appointment as secretary of the abovenamed company.",
        "*(ii) a public accountant registered under the Accountants Act 2004.",
        "*(iii) a member of the Institute of Singapore Chartered Accountants.",
        "*(iv) a member of the Chartered Secretaries Institute of Singapore.",
        "*(v) a member of the Association of International Accountants (Singapore Branch).",
        "*(vi) a member of the Institute of Company Accountants, Singapore.",
    ]
    for item in secretary_items:
        add_para(doc, item, style="Legal Clause", after=4)
    add_key_value_table(
        doc,
        [
            ("Name", "{{secretary.full_name}}"),
            ("Address", "{{secretary.residential_address}}"),
            ("ID / Nationality", "{{secretary.id_type}} No: {{secretary.id_number}} / {{secretary.nationality}}"),
        ],
    )
    add_signature_table(doc, [("Signed by", "{{secretary.full_name}}", "Secretary")])
    add_para(doc, "Dated this {{signature.day}} day of {{signature.month_year}}", style="Signature Label")
    add_para(doc, "\u2020 To be completed by secretaries of public companies only or by secretaries of private companies appointed under section 171(1AB) of the Act.", style="Form Note", italic=True)
    save(doc, "03_secretary_consent_form45b_standard.docx")


def build_share_certificate() -> None:
    doc = Document()
    configure_doc(doc, "Share Certificate", compact=True, landscape=True, page_numbers=False)

    def cert_para(text: str = "", *, size: float = 11, bold: bool = False, italic: bool = False, align: WD_ALIGN_PARAGRAPH | None = None, after: float = 3) -> object:
        p = add_para(doc, text, size=size, bold=bold, italic=italic, align=align, after=after)
        for run in p.runs:
            set_run_font(run, size=size, bold=bold, italic=italic, color=INK, font="Times New Roman", east_asia_font="SimSun")
        return p

    def cert_cell(cell, text: str, *, bold: bool = False, size: float = 10.5, align: WD_ALIGN_PARAGRAPH | None = None) -> None:
        set_cell_text(cell, text, bold=bold, size=size, align=align)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=size, bold=bold, color=INK, font="Times New Roman", east_asia_font="SimSun")

    cert_para("Counterfoil", align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True, after=0)
    cert_para("Certificate No.: *{{shareholder.certificate_no}}*", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, after=4)
    info = doc.add_table(rows=5, cols=4)
    table_widths(info, [1.55, 4.65, 1.6, 2.0])
    set_cell_margins(info, top=60, bottom=60, start=80, end=80)
    for row in info.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    rows = [
        ("NAME OF COMPANY", "{{company.company_name}}", "DATE", "{{signature.date}}"),
        ("REG. HOLDER", "{{shareholder.shareholder_name}}", "ALLOTMENT / TRANSFER NO.", "{{shareholder.allotment_transfer_no}}"),
        ("FOLIO IN MEMBERS' REG.", "{{shareholder.folio_no}}", "No. of Shares", "{{shareholder.shares}}"),
        ("Distinctive No.", "{{shareholder.distinctive_numbers}}", "Remarks", "{{shareholder.remarks}}"),
        ("Director", "{{director.full_name}}", "Director / Secretary", "{{secretary_or_director.full_name}}"),
    ]
    for idx, row in enumerate(rows):
        for col, text in enumerate(row):
            cert_cell(info.cell(idx, col), text, bold=col in (0, 2), size=9.5 if idx < 4 else 10.5)

    cert_para("-" * 118, align=WD_ALIGN_PARAGRAPH.CENTER, size=8, after=1)
    cert_para("Receipt for Share Certificate No.: *{{shareholder.certificate_no}}*", align=WD_ALIGN_PARAGRAPH.CENTER, size=15, bold=True, after=4)
    receipt = doc.add_table(rows=4, cols=2)
    table_widths(receipt, [1.7, 8.1])
    set_cell_margins(receipt, top=80, bottom=80, start=80, end=80)
    receipt_rows = [
        ("NAME OF COMPANY", "{{company.company_name}}"),
        ("RECEIVED THIS CERTIFICATE COVERING", "{{shareholder.shares_in_words}}"),
        ("SHARES IN THE COMPANY NUMBERED AS ON THE FACE THEREOF.", ""),
        ("Signature of Shareholder", "{{shareholder.shareholder_name}}"),
    ]
    for idx, (left, right) in enumerate(receipt_rows):
        cert_cell(receipt.cell(idx, 0), left, bold=True, size=10.5)
        cert_cell(receipt.cell(idx, 1), right, bold=idx in (1, 3), size=10.5)

    doc.add_page_break()
    cert_para("SHARE CERTIFICATE", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True, after=2)
    cert_para("=" * 120, align=WD_ALIGN_PARAGRAPH.CENTER, size=8, after=6)
    head = doc.add_table(rows=1, cols=3)
    table_widths(head, [2.0, 5.8, 2.0])
    set_cell_margins(head, top=80, bottom=80, start=80, end=80)
    cert_cell(head.cell(0, 0), "Certificate No.\n\n*{{shareholder.certificate_no}}*", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    cert_cell(
        head.cell(0, 1),
        "{{company.company_name}}\n(Incorporated in the Republic of Singapore)\nCo. Reg. No.: {{company.uen}}\nRegistered Office: {{company.registered_office_address}}",
        bold=True,
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    cert_cell(head.cell(0, 2), "No. of Shares\n\n*{{shareholder.shares}}*", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    cert_para("", after=4)
    cert_para("THIS IS TO CERTIFY THAT    {{shareholder.shareholder_name}}", size=13, bold=True, after=5)
    cert_para("of    {{shareholder.shareholder_address}}", size=12, italic=True, after=5)
    cert_para("is/are the registered holder(s) of    *{{shareholder.shares_in_words}}*", size=12, italic=True, after=5)
    cert_para("{{shareholder.share_class}} Share(s) fully paid in the Company subject to the Memorandum & Articles of Association of the Company.", size=12, italic=True, after=12)
    cert_para("Dated this {{signature.day}} day of {{signature.month_year}}", size=12, italic=True, after=12)
    cert_para("Given under the Common Seal of the Company on the date stated above and in the presence of: -", size=12, italic=True, after=10)
    sign = doc.add_table(rows=2, cols=3)
    table_widths(sign, [3.6, 2.2, 3.6])
    set_cell_margins(sign, top=100, bottom=100, start=80, end=80)
    cert_cell(sign.cell(0, 0), "\n\n____________________________\nDirector\n{{director.full_name}}", bold=True, size=11)
    cert_cell(sign.cell(0, 1), "\n\n(Common Seal)", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    cert_cell(sign.cell(0, 2), "\n\n____________________________\nDirector / Secretary\n{{secretary_or_director.full_name}}", bold=True, size=11)
    sign.cell(1, 0).merge(sign.cell(1, 2))
    cert_cell(sign.cell(1, 0), "Note: Transfer of any portion of the Shares comprised in this Certificate cannot be registered unless accompanied by this Certificate.", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    save(doc, "04_share_certificate_standard.docx")


def build_secretary_agreement() -> None:
    paras = source_paragraphs("05_secretary_service_agreement_preserved.docx")
    doc = Document()
    configure_doc(doc, "Secretary Service Agreement", compact=True)
    add_para(doc, "SERVICES AGREEMENT", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True, after=10)
    add_key_value_table(
        doc,
        [
            ("Date", "{{signature.date}}"),
            ("To", "{{provider.name}}"),
            ("Re", "Services Agreement (this \"Agreement\")"),
            ("Proposed Name of Company", "{{company.company_name}} (the \"Company\")"),
        ],
    )
    add_para(doc, "I/We, the undersigned hereby agree as follows to:", after=8)
    clause_no = 1
    for text in paras:
        if text.startswith(("Date:", "To:", "Re:", "Proposed Name of Company:", "Signed by:", "For and on behalf of", "Name:")):
            continue
        if text.startswith("{{client_signatory"):
            continue
        if text.startswith("{{provider") and "RESERVES THE RIGHT" not in text:
            continue
        if clause_no <= 21:
            add_legal_clause(doc, f"{clause_no}.  ", text, after=7)
            clause_no += 1
        else:
            fixed = text
            if fixed.startswith("Each of the provisions") and fixed.endswith(".（）"):
                fixed = fixed[:-3] + ".（本协议的每一项规定都是独立的且与其他规定截然不同，如果任何时候一项或多项规定无效、非法或不可执行，则本协议其余条款的有效性、合法性和可执行性不得以任何方式受到影响或损害，并且本协议应被解释为好像本协议从未包含此类无效、非法或不可执行的条款。）"
            if fixed.startswith("No alteration or amendment"):
                fixed = "No alteration or amendment in respect of this Agreement will be effective unless in writing and executed by each of the parties.（未经各方书面签署，对本协议的任何变更或修订均不生效。）"
            add_para(doc, fixed, style="Legal Clause", after=7)
    add_signature_table(
        doc,
        [
            ("Signed for and on behalf of {{provider.name}}", "{{provider.authorised_signatory_name}}", "Authorised Signatory"),
            ("Signed by Client", "{{client_signatory.full_name}}", "Client Signatory"),
            ("Signed by Client", "{{client_signatory_2.full_name}}", "Client Signatory"),
        ],
    )
    add_para(doc, "Additional signatory: {{client_signatory_3.full_name}}", style="Signature Label")
    save(doc, "05_secretary_service_agreement_standard.docx")


def build_nominee_director_agreement() -> None:
    source_doc = Document(V2_OUTPUT_DIR / "06_nominee_director_agreement_preserved.docx")
    doc = Document()
    configure_doc(doc, "Nominee Director Agreement", compact=True)
    for raw in [clean_source_text(p.text) for p in source_doc.paragraphs if clean_source_text(p.text)]:
        if raw in {"Between", "and"}:
            add_para(doc, raw, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
        elif raw in {"AGREEMENT FOR APPOINTMENT OF NOMINEE DIRECTOR", "委任挂名董事协议"}:
            add_para(doc, raw, align=WD_ALIGN_PARAGRAPH.CENTER, size=13 if raw.isupper() else 12, bold=True, after=4)
        elif raw in {"WHEREAS", "鉴于", "NOW IT IS HEREBY AGREED AS FOLLOWS:", "双方同意如下："} or re.match(r"^[A-Z]\. [A-Z]", raw):
            add_para(doc, raw, style="Heading 2", keep_with_next=True)
        elif re.match(r"^[\u4e00-\u9fff]+$", raw) and len(raw) < 12:
            add_para(doc, raw, style="Heading 3", keep_with_next=True)
        elif re.match(r"^\d+\.", raw):
            add_para(doc, raw, style="Legal Clause", after=4)
        elif re.match(r"^[a-z]\)", raw):
            add_para(doc, raw, style="Legal Clause", after=4)
        else:
            add_para(doc, raw, after=4)
    if source_doc.tables:
        add_signature_table(
            doc,
            [
                ("The Signature of", "{{nominee_director.full_name}}", "Nominee Director / 代理董事"),
                ("The Stamp / Signature of", "{{client_signatory.full_name}}", "Director / 董事, {{company.company_name}}"),
            ],
            page_break_before=True,
        )
    save(doc, "06_nominee_director_agreement_standard.docx")


def build_signature_record() -> None:
    doc = Document()
    configure_doc(doc, "Signature Record Attachment", compact=True)
    add_para(doc, "SIGNATURE CERTIFICATE", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True, after=4)
    add_para(doc, "Document Reference: {{signature_record.envelope_id}}", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=MUTED, after=12)
    add_key_value_table(
        doc,
        [
            ("Signer", "{{signer.full_name}}"),
            ("Email", "{{signer.email}}"),
            ("Party ID", "{{signer.party_id}}"),
            ("IP Address", "{{signer.ip_address}}"),
            ("Digital Fingerprint Checksum", "{{signer.digital_fingerprint_checksum}}"),
        ],
    )
    add_signature_table(doc, [("Electronic Signature", "{{signer.full_name}}", "Signer")])
    add_key_value_table(
        doc,
        [
            ("Signer 2", "{{signer_2.full_name}}"),
            ("Party ID", "{{signer_2.party_id}}"),
            ("IP Address", "{{signer_2.ip_address}}"),
            ("Digital Fingerprint Checksum", "{{signer_2.digital_fingerprint_checksum}}"),
        ],
    )
    add_signature_table(doc, [("Electronic Signature", "{{signer_2.full_name}}", "Signer")])
    doc.add_page_break()
    add_para(doc, "AUDIT TRAIL", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True, after=8)
    table = doc.add_table(rows=5, cols=2)
    table_widths(table, [1.6, 4.9])
    set_table_borders(table)
    set_cell_margins(table)
    headers = [("Timestamp", "Audit Activity")]
    rows = headers + [
        ("{{signature_record.audit_time_1}}", "{{signature_record.audit_line_1}}"),
        ("{{signature_record.audit_time_2}}", "{{signature_record.audit_line_2}}"),
        ("{{signature_record.audit_time_3}}", "{{signature_record.audit_line_3}}"),
        ("{{signature_record.audit_time_4}}", "{{signature_record.audit_line_4}}"),
    ]
    for idx, (left, right) in enumerate(rows):
        if idx == 0:
            shade_cell(table.cell(idx, 0), LIGHT_FILL)
            shade_cell(table.cell(idx, 1), LIGHT_FILL)
        set_cell_text(table.cell(idx, 0), left, bold=idx == 0, size=9.8)
        set_cell_text(table.cell(idx, 1), right, bold=idx == 0, size=9.8)
    add_para(doc, "This signature page forms a record of the online activity executing this contract.", style="Form Note", before=8)
    save(doc, "07_signature_record_attachment_standard.docx")


def build_form24() -> None:
    doc = Document()
    configure_doc(doc, "Return of Allotment of Shares / Form 24", compact=True, page_numbers=False)
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    title_table = doc.add_table(rows=1, cols=2)
    table_widths(title_table, [4.95, 1.35])
    set_table_borders(title_table, "000000")
    set_cell_margins(title_table, top=140, bottom=140, start=120, end=120)
    set_cell_text(
        title_table.cell(0, 0),
        "THE COMPANIES ACT\n(CHAPTER 50)\nSection 63 (1)\n\nRETURN OF ALLOTMENT OF SHARES",
        bold=True,
        size=10.6,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_cell_text(title_table.cell(0, 1), "FORM\n24\n\nFolio No.", bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "", after=8)
    add_para(doc, "Name of Company: {{company.company_name}}", size=10.4, bold=True, after=2)
    add_para(doc, "Company No.: {{company.uen}}", size=10.4, bold=True, after=10)
    add_para(doc, "The issue of the shares referred to in this return was made pursuant to a resolution passed by the members of the company on {{company.incorporation_date}}.")
    add_para(doc, "The shares referred to in this return were allotted, or are deemed to have been allotted under section 63 (6) of the Companies Act, to the allottees on the dates indicated.")
    add_para(doc, "1   Payable in cash", bold=True, keep_with_next=True)
    cash = doc.add_table(rows=4, cols=4)
    table_widths(cash, [2.55, 1.45, 1.25, 1.25])
    set_table_borders(cash)
    set_cell_margins(cash)
    cash_rows = [
        ("Class of shares", "Ordinary", "Preference", "Others"),
        ("Number of shares", "{{company.subscriber_share_totals}}", "", ""),
        ("Amount paid on each share", "{{company.share_currency}} {{company.share_par_value}}", "", ""),
        ("Amount due and payable on each share", "-", "", ""),
    ]
    for r, row in enumerate(cash_rows):
        for c, text in enumerate(row):
            if r == 0:
                shade_cell(cash.cell(r, c), "FFFFFF")
            set_cell_text(cash.cell(r, c), text, bold=c == 1 and r > 0, size=9.1, align=WD_ALIGN_PARAGRAPH.CENTER if c else None)
    add_para(doc, "2   For a consideration other than cash (*See Form 25/contract in writing)", bold=True, before=8, keep_with_next=True)
    noncash = doc.add_table(rows=4, cols=4)
    table_widths(noncash, [2.8, 1.2, 1.2, 1.2])
    set_table_borders(noncash)
    set_cell_margins(noncash)
    for r, row in enumerate(
        [
            ("Class of shares", "Ordinary", "Preference", "Others"),
            ("Number of shares", "", "", ""),
            ("Amount to be treated as paid on each share", "", "", ""),
        ]
    ):
        for c, text in enumerate(row):
            if r == 0:
                shade_cell(noncash.cell(r, c), "FFFFFF")
            set_cell_text(noncash.cell(r, c), text, bold=r == 0 or c == 0, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER if c else None)
    consideration_cell = noncash.cell(3, 0).merge(noncash.cell(3, 3))
    set_cell_text(
        consideration_cell,
        "The consideration for which the shares have been so allotted is as follows:\n\nNIL",
        bold=True,
        size=9.2,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    lodged = doc.add_table(rows=4, cols=2)
    table_widths(lodged, [3.45, 3.05])
    set_table_borders(lodged, "000000")
    set_cell_margins(lodged, top=110, bottom=110)
    set_cell_text(lodged.cell(0, 0), "Lodged in the office of the Registrar of Companies & Businesses by", bold=True, size=9.2)
    set_cell_text(lodged.cell(0, 1), "For Official Use", bold=True, size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(lodged.cell(1, 0), "Name: {{provider.lodging_officer_name}}", size=9)
    set_cell_text(lodged.cell(1, 1), "Date of Registration:", size=9)
    set_cell_text(lodged.cell(2, 0), "Address: {{provider.registered_address}}", size=9)
    set_cell_text(lodged.cell(2, 1), "Receipt No:", size=9)
    set_cell_text(lodged.cell(3, 0), "A/c No.: {{provider.account_no}}        Tel No.: {{provider.phone}}", size=9)
    set_cell_text(lodged.cell(3, 1), "Checked By:", size=9)
    doc.add_page_break()
    add_para(doc, "Form 24 Continuation Sheet 1", italic=True, after=4)
    add_para(doc, "(Please use continuation sheets if space provided is insufficient)", after=8)
    add_para(doc, "Name of Company: {{company.company_name}}", bold=True, after=2)
    add_para(doc, "Company No.: {{company.uen}}", bold=True, after=12)
    add_para(doc, "3   List of the allottees and an account of the shares allotted to them are as follows:", bold=True, keep_with_next=True)
    allottees = doc.add_table(rows=2, cols=2)
    table_widths(allottees, [4.6, 2.0])
    set_table_borders(allottees, "000000")
    set_cell_margins(allottees, top=120, bottom=120)
    set_cell_text(allottees.cell(0, 0), "(a) Name\n(b) Address\n(c) NRIC/Passport No/Registration No\n(d) Nationality/Country of Incorporation", bold=True, size=9.6)
    set_cell_text(allottees.cell(0, 1), "(e) No and class of shares allotted and consideration therefor\n\n(f) Date of allotment", bold=True, size=9.6)
    set_cell_text(
        allottees.cell(1, 0),
        "{% for shareholder in shareholders %}"
        "{{shareholder.shareholder_name}}\n\n"
        "{{shareholder.shareholder_address}}\n\n"
        "IDENTIFICATION NO.: {{shareholder.id_number}}\n\n"
        "{{shareholder.nationality}}\n\n"
        "{% endfor %}",
        bold=True,
        size=10,
    )
    set_cell_text(
        allottees.cell(1, 1),
        "{% for shareholder in shareholders %}{{shareholder.shares}} {{shareholder.share_class}} Shares\n\n{% endfor %}\nAllotted on the date of Incorporation.",
        bold=True,
        size=10,
    )
    doc.add_page_break()
    add_para(doc, "Form 24 Continuation Sheet 2", italic=True, after=8)
    add_para(doc, "Name of Company: {{company.company_name}}", bold=True, after=0)
    add_para(doc, "Company No.: {{company.uen}}", bold=True, after=12)
    add_para(doc, "4   Upon the abovementioned allotment of shares, the position of the Share Capital is as follows:", bold=True, keep_with_next=True)
    capital = doc.add_table(rows=3, cols=4)
    table_widths(capital, [2.2, 1.45, 1.45, 1.45])
    set_table_borders(capital, "000000")
    set_cell_margins(capital)
    for r, row in enumerate([("", "Ordinary", "Preference", "Others"), ("Issued Share Capital", "{{company.share_currency}} {{company.paid_up_capital}}", "-", "-"), ("Paid-up Share Capital", "{{company.share_currency}} {{company.paid_up_capital}}", "-", "-")]):
        for c, text in enumerate(row):
            if r == 0:
                shade_cell(capital.cell(r, c), "FFFFFF")
            set_cell_text(capital.cell(r, c), text, bold=r == 0 or c == 0, size=10, align=WD_ALIGN_PARAGRAPH.CENTER if c else None)
    add_para(doc, "CERTIFICATE", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True, before=10)
    add_para(doc, "I hereby certify, in relation to the abovenamed company that", after=4)
    certificate_items = [
        "(a) the company has more than 500 members;",
        "(b) the company keeps its principal shares register at {{company.register_location}};",
        "(c) the company provides reasonable accommodation and facilities for persons to inspect and take copies of its list of members and its particulars of shares transferred;",
        "(d) the shares referred to in this return were allotted for cash;",
        "(e) the shares referred to in this return were allotted for a consideration other than in cash and the number of persons to whom the shares have been allotted exceeds 500.",
    ]
    for item in certificate_items:
        add_para(doc, item, style="Legal Clause", after=2)
    add_para(doc, "Dated: {{signature.date}}", before=8)
    add_para(doc, "{{company.shareholder_signature_blocks}}", size=10.5, before=14, after=8)
    save(doc, "07_return_of_allotment_form24_standard.docx")


def build_rorc_notice() -> None:
    doc = Document()
    configure_doc(doc, "RORC Notice to Registrable Controller", compact=True)
    add_para(doc, "{{company.company_name}}", size=12, bold=True, after=0)
    add_para(doc, "(Company Reg. No. {{company.uen}})", after=0)
    add_para(doc, "{{company.registered_office_address}}", after=12)
    add_para(doc, "Date: {{signature.date}}", bold=True, after=2)
    add_para(doc, "To: {{shareholder.shareholder_name}}", bold=True, after=8)
    add_para(doc, "Dear Sir/Madam,")
    add_para(doc, "RE: NOTICE FOR REGISTRABLE CONTROLLER", bold=True, after=8)
    add_para(doc, "We know or have reasonable grounds to believe that you are a registrable controller of {{company.company_name}}.")
    add_para(doc, "This notice is sent under section 386AG(2)(a) of the Companies Act (Cap. 50) and requires you to provide the following information within 30 days after the date of this notice. Failure to provide the information required by this notice may be an offence.")
    add_section_heading(doc, "1", "Are you a registrable controller* of {{company.company_name}}?")
    add_para(doc, "*A registrable controller is defined as an individual or a legal entity that has \"significant interest\" in or \"significant control\" over the company.", style="Form Note")
    add_para(doc, "Your reply: YES. Please provide the following particulars:", bold=True)
    add_para(doc, "1) If your reply is yes and you are an individual", bold=True, italic=True, keep_with_next=True)
    individual_lines = [
        "(i) Full name: {{shareholder.shareholder_name}}",
        "(ii) Aliases, if any: NIL",
        "(iii) Residential address: {{shareholder.shareholder_address}}",
        "(iv) Nationality: {{shareholder.nationality}}",
        "(v) Identity card number or Passport number: {{shareholder.id_number}}",
        "(vi) Date of birth: {{shareholder.date_of_birth}}",
        "(vii) The date on which you became an individual controller: {{company.incorporation_date}} as per BizFile Company's Date of Incorporation.",
    ]
    for line in individual_lines:
        add_para(doc, line, style="Legal Clause", after=3)
    doc.add_page_break()
    add_para(doc, "2) If your reply is yes and the person is a legal entity", bold=True, italic=True, keep_with_next=True)
    entity_lines = [
        "(i) Company name: {{shareholder.corporate_name}}",
        "(ii) Company's unique entity number issued by the Registrar, if any: {{shareholder.corporate_registration_number}}",
        "(iii) Company's registered office: {{shareholder.corporate_registered_address}}",
        "(iv) Company's legal form: Company / corporation",
        "(v) The jurisdiction where, and statute under which, the person is formed or incorporated: {{shareholder.corporate_registration_country}}",
        "(vi) The name of the corporate entity register of the jurisdiction in which the person is formed or incorporated, if applicable: {{shareholder.corporate_registration_country}} corporate registry / as per corporate registration documents",
        "(vii) The identification number or registration number on the corporate entity register of the jurisdiction where the person is formed or incorporated, if applicable: {{shareholder.corporate_registration_number}}",
        "(viii) The date on which the person became a corporate controller: {{company.incorporation_date}}",
    ]
    for line in entity_lines:
        add_para(doc, line, style="Legal Clause", after=3)
    add_section_heading(doc, "2", "Do you know or have reasonable grounds to believe that any other person is a registrable controller of {{company.company_name}}?")
    add_para(doc, "Your reply: YES, if other qualifying shareholders/controllers exist in the shareholder records. Please provide the following particulars to the best of your knowledge:", bold=True)
    add_para(doc, "1) If your reply is yes and the person is an individual", bold=True, italic=True, keep_with_next=True)
    for line in [
        "a. Full name: Generated from other qualifying individual shareholders in the Shareholders sheet, where applicable.",
        "b. Aliases, if any:",
        "c. Residential address:",
        "d. Nationality:",
        "e. Identity card number or Passport number:",
        "f. Date of birth:",
        "g. The date on which the person became an individual controller: as per BizFile Company's Date of Incorporation.",
    ]:
        add_para(doc, line, style="Legal Clause", after=2)
    add_para(doc, "2) If your reply is yes and the person is a legal entity", bold=True, italic=True, keep_with_next=True)
    for line in [
        "a. Company name: Generated from other qualifying corporate shareholders in the Shareholders sheet, where applicable.",
        "b. Company's unique entity number issued by the Registrar, if any:",
        "c. Company's registered office:",
        "d. Company's legal form:",
        "e. The jurisdiction where, and statute under which, the person is formed or incorporated:",
        "f. The name of the corporate entity register of the jurisdiction in which the person is formed or incorporated, if applicable:",
        "g. The identification number or registration number on the corporate entity register of the jurisdiction where the person is formed or incorporated, if applicable:",
        "h. The date on which the person became a corporate controller:",
    ]:
        add_para(doc, line, style="Legal Clause", after=2)
    doc.add_page_break()
    add_para(doc, "In this notice -", before=8)
    for line in [
        "\"controller\", \"corporate controller\", \"individual controller\" and \"legal entity\" have the meanings given to them in section 386AB of the Companies Act;",
        "\"identity card\" has the meaning given to it in section 2(1) of the National Registration Act (Cap. 201);",
        "\"registrable\" has the meaning given to it in section 386AC of the Companies Act.",
    ]:
        add_para(doc, line, style="Legal Clause")
    add_para(doc, "Yours sincerely,", before=16, after=26)
    add_para(doc, "______________________________", after=2)
    add_para(doc, "{{secretary.full_name}}", bold=True, after=0)
    add_para(doc, "Secretary", bold=True, after=16)
    add_rule(doc)
    add_para(doc, "I/We confirm that the information set forth above is true and accurate.", after=28)
    add_para(doc, "______________________________", after=2)
    add_para(doc, "Name: {{shareholder.shareholder_name}}", bold=True, after=0)
    add_para(doc, "{{company.company_name}}", bold=True, after=8)
    add_para(doc, "Date: {{signature.date}}", style="Signature Label")
    save(doc, "08_rorc_notice_controller_standard.docx")


def build_paid_up_capital_confirmation() -> None:
    doc = Document()
    configure_doc(doc, "Confirmation of Paid-Up Capital")
    add_title(doc, title="CONFIRMATION OF PAID-UP CAPITAL")
    add_rule(doc)
    add_para(doc, "Date: {{signature.date}}", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, after=12)
    add_para(doc, "To Whom It May Concern", after=10)
    add_para(doc, "RE: CONFIRMATION OF PAID-UP CAPITAL", bold=True, after=10)
    add_para(
        doc,
        "I, {{paid_up_confirmation_signatory.full_name}}, sign this confirmation in my capacity as "
        "{{paid_up_confirmation_signatory.capacity}} of {{company.company_name}} (the \"Company\").",
        after=8,
    )
    add_para(
        doc,
        "I confirm that, based on the information provided to the Company and its corporate service provider "
        "for the purpose of incorporation and statutory record preparation, the issued share capital and paid-up "
        "share capital recorded below are true and correct.",
        after=8,
    )
    add_key_value_table(
        doc,
        [
            ("Company", "{{company.company_name}}"),
            ("Company Registration No.", "{{company.uen}}"),
            ("Registered Office", "{{company.registered_office_address}}"),
            ("Currency", "{{company.share_currency}}"),
            ("Issued Share Capital", "{{company.share_currency}} {{company.issued_share_capital}}"),
            ("Paid-Up Share Capital", "{{company.share_currency}} {{company.paid_up_capital}}"),
            ("Unpaid Share Capital, if any", "{{company.share_currency}} {{company.unpaid_share_capital}}"),
        ],
    )
    add_para(doc, "", after=4)
    add_para(
        doc,
        "The shareholder(s) remain responsible for paying the subscribed share capital due from them. The Company "
        "should ensure that its accounting records and bank records properly reflect the paid-up share capital.",
        after=8,
    )
    add_para(
        doc,
        "This confirmation is prepared for the Company's incorporation file and internal corporate secretarial records. "
        "It does not replace supporting accounting records, bank records or any other evidence that may be required for "
        "audit, tax, bank or regulatory review.",
        after=18,
    )
    add_signature_table(
        doc,
        [("Signed by", "{{paid_up_confirmation_signatory.full_name}}", "{{paid_up_confirmation_signatory.capacity}}")],
        page_break_before=True,
    )
    add_para(doc, "Dated: {{signature.date}}", style="Signature Label")
    save(doc, "10_paid_up_capital_confirmation_standard.docx")


def build_statutory_registers() -> None:
    doc = Document()
    configure_doc(doc, "Statutory Registers", compact=True)
    add_title(doc, title="STATUTORY REGISTERS")
    add_rule(doc)
    add_key_value_table(
        doc,
        [
            ("Company", "{{company.company_name}}"),
            ("Company No.", "{{company.uen}}"),
            ("Registered Office", "{{company.registered_office_address}}"),
            ("Register Location", "{{company.register_location}}"),
            ("Date Prepared", "{{signature.date}}"),
        ],
    )
    register_specs = [
        ("Register of Members", ["Member Name", "ID / UEN", "Address", "Shares", "Certificate No."], "{{registers.members_rows}}"),
        ("Register of Directors", ["Director Name", "ID No.", "Nationality", "Address", "Appointment Date"], "{{registers.directors_rows}}"),
        ("Register of Secretaries", ["Secretary Name", "ID No.", "Address", "Appointment Date", "Status"], "{{registers.secretaries_rows}}"),
        ("Register of Registrable Controllers", ["Controller Name", "Type", "ID / UEN", "Address", "Control Start Date"], "{{registers.controllers_rows}}"),
        ("Register of Share Certificates", ["Certificate No.", "Holder", "No. of Shares", "Distinctive No.", "Issue Date"], "{{registers.share_certificates_rows}}"),
    ]
    for title, headers, row_placeholder in register_specs:
        add_para(doc, title, style="Heading 1", keep_with_next=True)
        table = doc.add_table(rows=2, cols=len(headers))
        table_widths(table, [6.5 / len(headers)] * len(headers))
        set_table_borders(table)
        set_cell_margins(table, top=80, bottom=80, start=80, end=80)
        for col, header in enumerate(headers):
            shade_cell(table.cell(0, col), LIGHT_FILL)
            set_cell_text(table.cell(0, col), header, bold=True, size=8.7, align=WD_ALIGN_PARAGRAPH.CENTER)
        table.cell(1, 0).merge(table.cell(1, len(headers) - 1))
        set_cell_text(table.cell(1, 0), row_placeholder, size=9.3)
    save(doc, "10_statutory_registers_standard.docx")


def write_guides() -> None:
    guide = {
        "version": "v3_standard_layout",
        "principle": "Preserve content from v2/source templates while rebuilding layout with a clean, consistent business-document style. No clause is intentionally summarized or removed.",
        "design_preset": {
            "base": "standard_business_brief",
            "named_override": "legal template override: Arial, 10.5-11 pt body, clean tables/signature blocks, no floating objects or manual-space layout.",
            "page": "US Letter portrait for standard documents; Share Certificate uses a landscape certificate layout.",
        },
        "quality_changes": [
            "Rebuilt the visual structure instead of carrying forward old Word/WPS floating objects, hidden white text, manual tabs and sample arrows.",
            "Restored Secretary Agreement clauses that were present but hidden in white text in the old source.",
            "Used table-based fields for forms, share certificate data and signature blocks to reduce misalignment.",
            "Added Part 1 Form 24 and RORC templates based on the old registration pack and the shareholder rows, without requiring separate manual allottee/controller fields.",
            "Form 24 now makes the repeated shareholder/allottee area and repeated shareholder signature area explicit in the template.",
            "RORC was returned to a source-like notice/list/signature structure instead of a compressed table summary.",
            "Rebuilt Share Certificate as a single-page certificate original in a centered landscape layout with a restrained red certificate border and a lower fixed signature band; counterfoil/receipt pages are deferred as optional P2 support.",
            "Removed Chinese company-name/address placeholders and interpreter placeholders from P1 templates because they are not part of the current MVP input fields.",
            "Kept legal/commercial wording from v2/source except for obvious typographical, duplicate-wording and layout-related cleanup.",
        ],
        "manual_review_flags": [
            "The templates are operational drafting templates, not legal advice.",
            "Share certificate design is standardized for generation; compare with your preferred house style before production rollout.",
            "Form 24 and RORC are included as Part 1 because requested; they use Shareholders/People data and should still receive legal/status review before cloud rollout.",
            "Signature record attachment and Statutory Registers are deferred to P2.",
            "If the final cloud system uses HTML-to-PDF for certificate-style documents, this DOCX layout can be used as the visual reference.",
        ],
        "p2_deferred": P2_DEFERRED,
        "templates": MANIFEST,
        "fields": FIELD_NOTES,
    }
    readme = [
        "# P1 Standard Template Pack v3",
        "",
        "This pack is the standardized-layout version. It uses v2/source content but rebuilds the Word layout so the files are easier to batch-generate and visually review.",
        "",
        "## Rules",
        "",
        "- Preserve clauses and signed-document content; do not summarize agreement text.",
        "- Use clean Word paragraphs/tables instead of manual spacing, floating objects, hidden text or old sample artifacts.",
        "- Keep placeholders stable and readable.",
        "- Treat unresolved legal or commercial judgment as review notes, not silent deletions.",
        "",
        "## Files",
        "",
    ]
    for item in MANIFEST:
        readme.append(f"- `{item['file_name']}` - {item['display_name']}")
    readme.extend(["", "## Deferred To P2", ""])
    for item in P2_DEFERRED:
        readme.append(f"- `{item['template_id']}` - {item['reason']}")
    readme.extend(["", "## Main Placeholders", ""])
    for key, note in FIELD_NOTES.items():
        readme.append(f"- `{{{{{key}}}}}`: {note}")

    for directory in (TEMPLATE_DIR, OUTPUT_DIR):
        (directory / "template_manifest_v3.json").write_text(json.dumps(guide, ensure_ascii=False, indent=2), encoding="utf-8")
        (directory / "README_v3_standard_layout.md").write_text("\n".join(readme), encoding="utf-8")


def zip_output() -> None:
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(OUTPUT_DIR.iterdir()):
            if path.is_file():
                zf.write(path, arcname=path.name)


def main() -> None:
    ensure_v2_sources()
    ensure_dirs()
    build_first_directors_resolution()
    build_director_consent()
    build_secretary_consent()
    build_share_certificate()
    build_secretary_agreement()
    build_nominee_director_agreement()
    build_form24()
    build_rorc_notice()
    build_paid_up_capital_confirmation()
    write_guides()
    zip_output()
    print(f"Generated standard v3 templates in {TEMPLATE_DIR}")
    print(f"Copied package to {OUTPUT_DIR}")
    print(f"Zip package: {ZIP_PATH}")


if __name__ == "__main__":
    main()
