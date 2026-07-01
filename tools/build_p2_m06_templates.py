from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DIR = ROOT / "app" / "doc_templates" / "p2_standard_v1"
OUTPUT_DIR = ROOT / "outputs" / "P2_standard_templates_v1"

TEMPLATES = {
    "directors_resolution": "M06_strike_off_directors_resolution_standard.docx",
    "shareholder_consent": "M06_strike_off_shareholder_consent_standard.docx",
    "director_declaration": "M06_strike_off_director_declaration_standard.docx",
}

BASE_FONT = "Calibri"
INK = "111827"
MUTED = "475569"


def ensure_dirs() -> None:
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def configure_doc(doc: Document, title: str, subject: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BASE_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = None
    doc.core_properties.title = title
    doc.core_properties.subject = subject


def add_para(doc: Document, text: str = "", *, bold: bool = False, size: float = 10.5, align=None, after: int = 6):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = 1.08
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.name = BASE_FONT
    run.font.size = Pt(size)
    return para


def add_company_header(doc: Document, title: str) -> None:
    add_para(doc, "{{company.company_name}}", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    add_para(doc, "Company Registration No. {{company.uen}}", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(doc, "(Incorporated in the Republic of Singapore)", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(doc, '(the "Company")', size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    add_para(doc, title, bold=True, size=12.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)


def clause_title(doc: Document, text: str) -> None:
    add_para(doc, text, bold=True, size=11.2, after=4)


def clause_para(doc: Document, text: str) -> None:
    add_para(doc, text, size=10.5, after=7)


def numbered_para(doc: Document, number: str, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.first_line_indent = Inches(-0.25)
    para.paragraph_format.space_after = Pt(5)
    para.paragraph_format.line_spacing = 1.06
    run = para.add_run(f"{number} {text}")
    run.font.name = BASE_FONT
    run.font.size = Pt(10.5)


def add_signature_text(doc: Document, heading: str, placeholder: str) -> None:
    add_para(doc, heading, bold=True, size=10.5, after=14)
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.line_spacing = 1.05
    run = para.add_run(placeholder)
    run.font.name = BASE_FONT
    run.font.size = Pt(10.5)


def build_directors_resolution() -> Document:
    doc = Document()
    configure_doc(doc, "M06 Strike-off Directors Resolution", "P2 strike-off directors resolution template")
    add_company_header(doc, "DIRECTORS' RESOLUTION IN WRITING MADE PURSUANT TO THE COMPANY'S CONSTITUTION")

    clause_title(doc, "STRIKING OFF THE COMPANY")
    clause_para(doc, "RESOLVED:")
    clause_para(
        doc,
        "That the Company hereby agrees to apply to the Accounting and Corporate Regulatory Authority to strike the Company off pursuant to Section 344 of the Companies Act 1967 as {{strike.business_status_clause}}.",
    )
    clause_para(doc, "Further, it was noted that:")
    for letter, text in [
        ("a)", "the Company has no assets and no liabilities, whether current, future or contingent;"),
        ("b)", "the directors have obtained or will obtain the written consent of all or the requisite majority of the shareholders;"),
        ("c)", "the Company has submitted the last set of audited accounts or the latest unaudited balance sheet, where applicable;"),
        ("d)", "the Company has no outstanding tax liabilities with the Inland Revenue Authority of Singapore;"),
        ("e)", "the Company has no outstanding employers' CPF contributions owing to the Central Provident Fund Board;"),
        ("f)", "the Company has no outstanding charges in the Register of Charges; and"),
        ("g)", "the Company is not involved in any court proceedings, whether within or outside Singapore."),
    ]:
        numbered_para(doc, letter, text)

    clause_title(doc, "AUTHORITY TO LODGE RETURNS")
    clause_para(
        doc,
        "RESOLVED that upon the written consent for striking off the Company being obtained from all or the requisite majority of the shareholders, the Company Secretary or authorised filing agent be authorised to lodge the necessary application and returns with the Accounting and Corporate Regulatory Authority.",
    )
    add_para(doc, "Dated this {{strike.document_date}}", after=16)
    add_signature_text(doc, "DIRECTOR(S)", "{{strike.director_signature_blocks}}")
    return doc


def build_shareholder_consent() -> Document:
    doc = Document()
    configure_doc(doc, "M06 Strike-off Shareholder Consent", "P2 strike-off shareholder consent template")
    add_para(doc, "ACCOUNTING AND CORPORATE REGULATORY AUTHORITY", bold=True, size=11, after=2)
    add_para(doc, "10 Anson Road #05-01/15", size=10.5, after=1)
    add_para(doc, "International Plaza", size=10.5, after=1)
    add_para(doc, "Singapore 079903", size=10.5, after=14)
    add_para(doc, "Dear Sir / Madam", after=12)
    add_para(doc, "{{company.company_name}}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(doc, "(Company Registration No. {{company.uen}})", align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    add_para(doc, "CONSENT TO STRIKE OFF THE COMPANY", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=16)
    clause_para(doc, "I am a shareholder holding {{strike_shareholder.shares_text}} {{strike_shareholder.share_class}} share(s) in the abovenamed Company.")
    clause_para(doc, "The Company {{strike.business_status_sentence}} and I do not intend to do any business through this Company in the future.")
    clause_para(doc, "I agree to the Company's application to be struck off from the register.")
    add_para(doc, "Yours sincerely", after=26)
    add_signature_text(doc, "", "{{strike_shareholder.signature_block}}")
    return doc


def build_director_declaration() -> Document:
    doc = Document()
    configure_doc(doc, "M06 Strike-off Director Declaration", "P2 strike-off director declaration template")
    add_para(doc, "DECLARATION", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, after=16)
    clause_para(
        doc,
        "I, {{strike_declarant.full_name}}, {{strike_declarant.id_label}} of {{strike_declarant.address}}, do solemnly and sincerely declare as follows:",
    )
    numbered_para(doc, "1.", "I am a director of {{company.company_name}} (the \"Company\"), Company Registration No. {{company.uen}}.")
    numbered_para(doc, "2.", "The Company {{strike.business_status_sentence}}.")
    numbered_para(doc, "3.", "The Company has no assets.")
    numbered_para(doc, "4.", "The Company has no liabilities.")
    numbered_para(doc, "5.", "The Company has no contingent assets and no contingent liabilities.")
    numbered_para(doc, "6.", "The Company has no outstanding tax liabilities owing to the Inland Revenue Authority of Singapore and is not indebted to any other Government department.")
    numbered_para(doc, "7.", "The Company has no outstanding charges in the Register of Charges.")
    numbered_para(doc, "8.", "The Company is not involved in or threatened with legal proceedings whether in or outside Singapore.")
    numbered_para(doc, "9.", "I have obtained or will obtain the consent of all or the requisite majority of the shareholders for the application to the Accounting and Corporate Regulatory Authority to consider striking the name of the Company off the register under Section 344 of the Companies Act 1967.")
    numbered_para(doc, "10.", "I verily believe that this is a proper case for the Accounting and Corporate Regulatory Authority to consider striking the name of the Company off the register under Section 344 of the Companies Act 1967.")
    clause_para(doc, "and I make this solemn declaration conscientiously believing the same to be true.")
    add_para(doc, "Declared at Singapore", after=2)
    add_para(doc, "this {{strike.document_date}}", after=22)
    add_signature_text(doc, "", "{{strike_declarant.signature_block}}")
    return doc


def field_map_text() -> str:
    return """# M06 Strike-off Package - Field Map

M06 is the P2 strike-off / company closure package.

## Generated files

| File | Trigger | Signer | Notes |
|---|---|---|---|
| Directors' Resolution for Strike-off | `strike_off_required = Yes` | Director signer(s) | Approves the strike-off application and authorises filing. |
| Shareholder Consent to Strike-off | `strike_off_required = Yes` | Each member/shareholder signer | One consent file per member/shareholder signer. |
| Director Declaration | `strike_off_required = Yes` | Declaration signer, default first director signer | Confirms no assets, liabilities, charges, proceedings and shareholder consent. |

## Primary input fields

| Field | Source | Used in |
|---|---|---|
| `company.company_name` | P2 one-page sheet / web form | All M06 documents |
| `company.uen` | P2 one-page sheet / web form | All M06 documents |
| `company.default_document_date` | P2 one-page sheet / web form | Document and declaration date |
| `company.director_signer_names` | P2 one-page sheet / web form | Directors' resolution and default declaration signer |
| `company.member_signer_names` | P2 one-page sheet / web form | Shareholder consent letters |
| `shareholdings.shareholder_name` / `shareholdings.shares` | P2 shareholder status sheet, if available | Shareholder consent shareholding line |
| `strike_off_required` | P2 one-page sheet / web form | Enables M06 |
| `strike_off_cessation_date` | P2 one-page sheet / web form | If blank, wording says the company has not commenced business since incorporation. If filled, wording says the company ceased business from this date. |
| `strike_off_declaration_signer_name` | P2 one-page sheet / web form | Optional override for declaration signer |

## Review guardrail

M06 prepares signing documents only. It does not verify ACRA, IRAS, CPF, charge register, bank account, creditor or litigation status. These remain manual review points before filing.
"""


def manifest() -> dict[str, object]:
    return {
        "version": "P2_M06_v0.1",
        "package": "M06 Strike-off Package",
        "templates": TEMPLATES,
        "default_files": list(TEMPLATES.keys()),
    }


def save_template(doc: Document, filename: str) -> None:
    doc.save(TEMPLATE_DIR / filename)
    doc.save(OUTPUT_DIR / filename)


def main() -> None:
    ensure_dirs()
    save_template(build_directors_resolution(), TEMPLATES["directors_resolution"])
    save_template(build_shareholder_consent(), TEMPLATES["shareholder_consent"])
    save_template(build_director_declaration(), TEMPLATES["director_declaration"])
    (TEMPLATE_DIR / "M06_field_map.md").write_text(field_map_text(), encoding="utf-8")
    (OUTPUT_DIR / "M06_field_map.md").write_text(field_map_text(), encoding="utf-8")
    (OUTPUT_DIR / "M06_manifest.json").write_text(json.dumps(manifest(), ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
