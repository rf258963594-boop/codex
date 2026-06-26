from __future__ import annotations

import json
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from build_p2_m03_templates import (
    OUTPUT_DIR,
    TEMPLATE_DIR,
    add_company_header,
    add_para,
    add_table,
    clause_para,
    clause_title,
    configure_doc,
    marker,
    set_cell_margins,
    set_cell_text,
    set_row_cant_split,
    set_table_borders,
    set_table_width,
)


TEMPLATES = {
    "agm_package": "M05_agm_documents_package_standard.docx",
    "annual_return_package": "M05_annual_return_authorisation_package_standard.docx",
    "checklist": "M05_annual_review_checklist_standard.docx",
}

VERSION = "P2_M05_v0.3"


def ensure_dirs() -> None:
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for directory in (TEMPLATE_DIR, OUTPUT_DIR):
        for path in directory.glob("M05_*"):
            if path.is_file():
                path.unlink()


def add_signature_rows(doc: Document, expr: str, heading: str) -> None:
    add_para(doc, heading, bold=True, size=10.5, after=4, keep_with_next=True)
    marker(doc, f"[[REPEAT {expr}[]]]")
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [4680, 4680])
    set_table_borders(table, "FFFFFF")
    set_cell_margins(table, top=40, bottom=60, start=0, end=160)
    set_row_cant_split(table.rows[0])
    set_cell_text(table.cell(0, 0), "{{sigrow.left_block}}", size=10.3)
    set_cell_text(table.cell(0, 1), "{{sigrow.right_block}}", size=10.3)


def conditional_start(doc: Document, expr: str) -> None:
    marker(doc, f"[[IF {expr}]]")


def conditional_end(doc: Document) -> None:
    marker(doc, "[[END IF]]")


def add_attendance_table(doc: Document) -> None:
    marker(doc, "[[REPEAT m05.attendance_rows[]]]")
    add_table(
        doc,
        ["Name", "Capacity", "Signature"],
        [["{{attendee.full_name}}", "{{attendee.capacity}}", ""]],
        [3300, 2300, 3760],
    )


def add_checklist_table(doc: Document) -> None:
    marker(doc, "[[REPEAT m05.checklist_items[]]]")
    add_table(
        doc,
        ["Item", "Status", "Review note"],
        [["{{check.item}}", "{{check.status}}", "{{check.note}}"]],
        [2350, 1450, 5560],
    )


def add_directors_resolution(doc: Document) -> None:
    add_company_header(doc, "DIRECTORS' RESOLUTIONS IN WRITING")
    add_para(
        doc,
        "Pursuant to the Constitution of the Company",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=10.5,
        after=12,
        keep_with_next=True,
    )
    clause_para(
        doc,
        "The undersigned, being the director(s) of the Company, hereby pass the following resolutions in writing. "
        "The resolutions shall have the same force and effect as if they had been passed at a meeting of the Board of Directors duly convened and held.",
    )
    clause_title(doc, "{{m05.board_accounts_title}}")
    clause_para(doc, "{{m05.board_accounts_resolution_1}}")
    clause_para(doc, "{{m05.board_accounts_resolution_2}}")
    clause_title(doc, "Annual Review Method")
    clause_para(doc, "{{m05.board_meeting_resolution}}")
    clause_para(doc, "{{m05.board_documents_resolution}}")
    clause_title(doc, "Annual Return and Statutory Declarations")
    clause_para(
        doc,
        "RESOLVED - That {{m05.ar_signer_name}}, acting as {{m05.ar_signer_capacity}}, be authorised to sign the Annual Return, "
        "Section 197 certificate, applicable statutory statement, filing authorisation and related statutory declarations, and to authorise the lodgement "
        "of the same with ACRA / BizFile."
    )
    clause_para(doc, "{{m05.directors_fee_text}}")
    clause_para(doc, "{{m05.directors_remuneration_text}}")
    add_para(doc, "Dated this {{signature.day_ordinal}} day of {{signature.month_year}}", style="Signature Block", before=12, after=10)
    add_signature_rows(doc, "m05.director_signature_rows", "DIRECTOR(S)")


def add_notice_of_agm(doc: Document) -> None:
    doc.add_page_break()
    add_company_header(doc, "NOTICE OF AN ANNUAL GENERAL MEETING OF MEMBERS")
    clause_para(
        doc,
        "NOTICE IS HEREBY GIVEN that an Annual General Meeting of the members of the Company will be held at {{m05.agm_place}} "
        "on {{m05.agm_date}} at {{m05.agm_time}} for the purpose of transacting the following business.",
    )
    clause_title(doc, "Agenda")
    clause_para(doc, "1. {{m05.agm_accounts_business}}")
    clause_para(doc, "2. To approve the directors' fees and remuneration for the financial year, where applicable.")
    clause_para(doc, "3. To note and approve the non-appointment of auditors for the ensuing year where the Company qualifies for audit exemption.")
    clause_para(doc, "4. To authorise the filing of the Annual Return and related declarations with ACRA / BizFile.")
    clause_para(doc, "5. To transact any other ordinary business which may properly be transacted at an Annual General Meeting.")
    add_para(doc, "By Order of the Board", style="Signature Block", before=14, after=18)
    add_para(doc, "_____________________________________", style="Signature Block", after=2)
    add_para(doc, "{{m05.notice_issuer_name}}", style="Signature Block", after=0)
    add_para(doc, "{{m05.notice_issuer_capacity}}", style="Signature Block", after=8)
    add_para(doc, "Dated this {{signature.day_ordinal}} day of {{signature.month_year}}", style="Signature Block", after=10)
    add_para(
        doc,
        "Notes: A member entitled to attend and vote at the meeting is entitled to appoint a proxy to attend and vote on his/her behalf. "
        "A proxy need not be a member of the Company. The instrument appointing a proxy should be deposited at the registered office of the Company "
        "or delivered to the Company before the meeting.",
        style="Small Note",
    )


def add_shorter_notice(doc: Document) -> None:
    doc.add_page_break()
    add_company_header(doc, "AGREEMENT BY MEMBERS TO SHORTER NOTICE")
    clause_para(
        doc,
        "Pursuant to Section 177(3)(a) of the Companies Act 1967 and the Constitution of the Company, the undersigned member(s) of the Company "
        "hereby agree to the Annual General Meeting being held on {{m05.agm_date}} at {{m05.agm_time}}, notwithstanding that it may be called by "
        "notice shorter than the period otherwise required."
    )
    clause_para(
        doc,
        "The undersigned member(s) further confirm that they have received, or agree to receive, the financial statements, report of the directors "
        "and related documents for consideration at the Annual General Meeting."
    )
    add_para(doc, "Dated this {{signature.day_ordinal}} day of {{signature.month_year}}", style="Signature Block", before=10, after=10)
    add_signature_rows(doc, "m05.member_signature_rows", "MEMBER(S)")


def add_proxy_form(doc: Document) -> None:
    doc.add_page_break()
    add_company_header(doc, "INSTRUMENT APPOINTING A PROXY")
    clause_para(
        doc,
        "I/We, being a member of the Company, hereby appoint the following person as my/our proxy to attend and vote for me/us and on my/our behalf "
        "at the Annual General Meeting of the Company to be held on {{m05.agm_date}} at {{m05.agm_time}}, and at any adjournment thereof."
    )
    add_table(
        doc,
        ["Member", "Proxy appointed", "Meeting"],
        [["{{m05.member_names_text}}", "", "{{m05.agm_date}} at {{m05.agm_time}}"]],
        [3000, 3000, 3360],
    )
    clause_para(
        doc,
        "Unless otherwise instructed, the proxy may vote or abstain from voting as he/she thinks fit on any resolution properly put before the meeting."
    )
    add_para(doc, "Dated this {{signature.day_ordinal}} day of {{signature.month_year}}", style="Signature Block", before=12, after=10)
    add_signature_rows(doc, "m05.member_signature_rows", "MEMBER(S)")


def add_attendance_and_minutes(doc: Document) -> None:
    doc.add_page_break()
    add_company_header(doc, "ATTENDANCE SHEET")
    clause_para(doc, "Attendance at the Annual General Meeting of the Company held on {{m05.agm_date}} at {{m05.agm_time}} at {{m05.agm_place}}.")
    add_attendance_table(doc)

    doc.add_page_break()
    add_company_header(doc, "MINUTES OF AN ANNUAL GENERAL MEETING")
    add_table(
        doc,
        ["Date and time", "Place", "Chairperson"],
        [["{{m05.agm_date}} at {{m05.agm_time}}", "{{m05.agm_place}}", "{{m05.chairperson_name}}"]],
        [2500, 5160, 1700],
    )
    clause_para(
        doc,
        "The Chairperson noted that a quorum was present and declared the meeting open. The notice convening the meeting and, where applicable, "
        "the agreement by members to shorter notice were tabled and taken as read."
    )
    clause_title(doc, "{{m05.minutes_accounts_title}}")
    clause_para(doc, "{{m05.minutes_accounts_resolution}}")
    clause_title(doc, "Accounts / Audit Status Review")
    clause_para(doc, "{{m05.audit_statement_a}}")
    clause_title(doc, "Directors' Fees and Remuneration")
    clause_para(doc, "{{m05.directors_fee_text}}")
    clause_para(doc, "{{m05.directors_remuneration_text}}")
    clause_title(doc, "Annual Return")
    clause_para(
        doc,
        "RESOLVED - That the authorised director and/or authorised representative be authorised to sign and lodge the Annual Return and related "
        "declarations with ACRA / BizFile for the financial year ended {{m05.fye_date_upper}}."
    )
    clause_title(doc, "Closure")
    clause_para(doc, "There being no further business, the meeting was declared closed.")
    add_para(doc, "", after=18)
    add_para(doc, "_____________________________________", style="Signature Block", after=2)
    add_para(doc, "{{m05.chairperson_name}}", style="Signature Block", after=0)
    add_para(doc, "Chairperson", style="Signature Block", after=0)


def add_written_annual_review_documents(doc: Document) -> None:
    doc.add_page_break()
    add_company_header(doc, "{{m05.written_resolution_title}}")
    clause_para(doc, "{{m05.written_resolution_body}}")
    clause_title(doc, "{{m05.minutes_accounts_title}}")
    clause_para(doc, "RESOLVED - {{m05.written_resolution_accounts}}")
    clause_title(doc, "Annual Return")
    clause_para(
        doc,
        "RESOLVED - That {{m05.ar_signer_name}}, acting as {{m05.ar_signer_capacity}}, be authorised to sign and lodge the Annual Return "
        "and related declarations with ACRA / BizFile for the financial year ended {{m05.fye_date_upper}}."
    )
    clause_title(doc, "Accounts / Audit Status")
    clause_para(doc, "{{m05.audit_statement_a}}")
    add_para(doc, "Dated this {{signature.day_ordinal}} day of {{signature.month_year}}", style="Signature Block", before=12, after=10)
    add_signature_rows(doc, "m05.member_signature_rows", "MEMBER(S)")


def build_agm_package() -> Document:
    doc = Document()
    configure_doc(doc, "M05 AGM Documents Package", "P2 annual review AGM documents")
    add_directors_resolution(doc)
    conditional_start(doc, "m05.use_agm_meeting_documents")
    add_notice_of_agm(doc)
    add_shorter_notice(doc)
    add_proxy_form(doc)
    add_attendance_and_minutes(doc)
    conditional_end(doc)
    conditional_start(doc, "m05.use_written_annual_review_documents")
    add_written_annual_review_documents(doc)
    conditional_end(doc)
    return doc


def add_annual_return_review_summary(doc: Document) -> None:
    add_company_header(doc, "ANNUAL RETURN REVIEW SUMMARY")
    add_table(
        doc,
        ["Company name", "UEN", "Registered office"],
        [["{{company.company_name}}", "{{company.uen}}", "{{company.registered_office_address}}"]],
        [3100, 1800, 4460],
    )
    add_table(
        doc,
        ["FYE", "AGM date", "Financial statements date", "Accounts status"],
        [["{{m05.fye_date}}", "{{m05.agm_date}}", "{{m05.financial_statement_date_display}}", "{{m05.accounts_status_label}}"]],
        [1850, 1850, 2600, 2060],
    )
    add_table(
        doc,
        ["Issued shares", "Issued share capital", "Paid-up capital", "Currency"],
        [["{{company.total_issued_shares}}", "{{company.issued_share_capital}}", "{{company.paid_up_capital}}", "{{company.currency}}"]],
        [2200, 2700, 2400, 960],
    )
    clause_para(
        doc,
        "The above particulars are prepared for annual return review. The director(s) and authorised signatory should check the Company's BizFile "
        "profile, officers, shareholders, share capital, registered office, business activities and registers before the Annual Return is lodged."
    )


def add_section_197_certificate(doc: Document) -> None:
    doc.add_page_break()
    add_company_header(doc, "CERTIFICATE BY A COMPANY LIMITED BY SHARES UNDER SECTION 197(1)")
    clause_para(doc, "I, the undermentioned officer of the abovementioned Company, hereby certify that:")
    clause_para(
        doc,
        "a) I have verified that the summary of return by a company having a share capital, as reflected in the records of the Registrar and/or "
        "the Company's statutory records, is accurate and up to date as at {{m05.ar_as_at_date}};"
    )
    clause_para(
        doc,
        "b) I have inspected, or caused to be inspected, the Register of Members, Register of Directors, Register of Secretaries, Register of "
        "Registrable Controllers, Register of Nominee Directors and other statutory records relevant to the Annual Return;"
    )
    clause_para(
        doc,
        "c) the Company is a private company limited by shares and has not issued any invitation to the public to subscribe for any shares or "
        "debentures of the Company; and"
    )
    clause_para(
        doc,
        "d) the number of members of the Company is within the limit applicable to a private company, subject to final verification against the "
        "Company's statutory records."
    )
    add_para(doc, "Dated this {{signature.day_ordinal}} day of {{signature.month_year}}", style="Signature Block", before=10, after=18)
    add_para(doc, "_____________________________________", style="Signature Block", after=2)
    add_para(doc, "{{m05.ar_signer_name}}", style="Signature Block", after=0)
    add_para(doc, "{{m05.ar_signer_capacity}}", style="Signature Block", after=0)


def add_audit_exemption_statement(doc: Document) -> None:
    doc.add_page_break()
    add_company_header(doc, "{{m05.audit_statement_title}}")
    clause_para(doc, "{{m05.audit_statement_intro}}")
    clause_para(doc, "{{m05.audit_statement_a}}")
    clause_para(doc, "{{m05.audit_statement_b}}")
    clause_para(doc, "{{m05.audit_statement_c}}")
    clause_para(doc, "{{m05.audit_statement_d}}")
    add_para(doc, "Dated this {{signature.day_ordinal}} day of {{signature.month_year}}", style="Signature Block", before=10, after=10)
    add_signature_rows(doc, "m05.director_signature_rows", "DIRECTOR(S)")


def add_ar_authorisation_letter(doc: Document) -> None:
    doc.add_page_break()
    add_para(doc, "{{m05.agm_date}}", align=WD_ALIGN_PARAGRAPH.RIGHT, after=14)
    add_para(doc, "{{provider.name}}", bold=True, after=2)
    add_para(doc, "{{provider.registered_address}}", after=14)
    add_para(doc, "Dear Sirs", after=10)
    add_para(doc, "FILING OF ANNUAL RETURN - FINANCIAL YEAR ENDED {{m05.fye_date_upper}}", bold=True, after=10)
    clause_para(
        doc,
        "I/We, the director(s), member(s) and/or authorised signatory of {{company.company_name}} (UEN: {{company.uen}}), hereby authorise "
        "{{provider.name}} and its designated officers or agents to prepare, initiate and lodge the Annual Return and related electronic filings "
        "with ACRA / BizFile for the financial year ended {{m05.fye_date_upper}}."
    )
    clause_para(
        doc,
        "I/We declare that the particulars of the Company and the information provided for the Annual Return are, to the best of my/our knowledge "
        "and belief, true, complete and up to date."
    )
    clause_para(
        doc,
        "I/We confirm that responsibility for the accuracy and completeness of the Annual Return, statutory registers, accounting records and "
        "supporting information remains with the Company and its directors."
    )
    add_para(doc, "Yours faithfully,", after=18)
    add_para(doc, "_____________________________________", style="Signature Block", after=2)
    add_para(doc, "{{m05.ar_signer_name}}", style="Signature Block", after=0)
    add_para(doc, "{{m05.ar_signer_capacity}}", style="Signature Block", after=0)


def add_management_representation(doc: Document) -> None:
    doc.add_page_break()
    add_company_header(doc, "{{m05.management_representation_title}}")
    clause_para(doc, "{{m05.management_representation_intro}}")
    clause_para(doc, "{{m05.management_representation_1}}")
    clause_para(doc, "{{m05.management_representation_2}}")
    clause_para(doc, "{{m05.management_representation_3}}")
    clause_para(doc, "{{m05.management_representation_4}}")
    clause_para(doc, "{{m05.management_representation_5}}")
    clause_para(doc, "{{m05.management_representation_6}}")
    clause_para(doc, "{{m05.management_representation_7}}")
    add_para(doc, "Yours faithfully,", after=18)
    add_signature_rows(doc, "m05.director_signature_rows", "DIRECTOR(S)")


def build_annual_return_package() -> Document:
    doc = Document()
    configure_doc(doc, "M05 Annual Return Authorisation Package", "P2 annual return authorisation and declarations")
    add_annual_return_review_summary(doc)
    add_section_197_certificate(doc)
    add_audit_exemption_statement(doc)
    add_ar_authorisation_letter(doc)
    add_management_representation(doc)
    return doc


def build_checklist() -> Document:
    doc = Document()
    configure_doc(doc, "M05 Annual Review Checklist", "P2 annual review internal checklist")
    add_company_header(doc, "ANNUAL REVIEW INTERNAL CHECKLIST")
    add_table(
        doc,
        ["FYE", "AGM date", "Annual review method", "Accounts status"],
        [["{{m05.fye_date}}", "{{m05.agm_date}}", "{{m05.agm_mode_label}}", "{{m05.accounts_status_label}}"]],
        [1800, 1800, 2500, 2260],
    )
    add_checklist_table(doc)
    add_para(
        doc,
        "This checklist is for internal review only and should not be treated as confirmation that ACRA / BizFile filing has been completed.",
        style="Small Note",
        before=8,
        after=0,
    )
    return doc


def field_map_text() -> str:
    return """# M05 Annual Review Package - Field Map

M05 is the annual review / AGM / Annual Return authorisation package for an existing Singapore company.

## Rebuild note

Version P2_M05_v0.4 uses a simplified annual review status model. The main `accounts_status` values are `active`, `dormant` and `audited`; older `non_dormant` / `unaudited` values are treated as active for compatibility.

## Generated PDFs

| File | Content | Main signer |
|---|---|---|
| Annual review signing package | AGM / written annual review documents plus Annual Return review, Section 197 certificate, dynamic audit/dormant/audited statement, AR filing authorisation and management representation | Directors, member(s), authorised signer |
| Internal checklist | Filing and review checklist | Internal only |

## Primary fields

| Field | Source |
|---|---|
| `annual_review_required` | P2 one-page sheet or Quick Annual Review sheet |
| `fye_date` | Quick Annual Review |
| `agm_date` | Quick Annual Review; defaults to document date if blank |
| `agm_time` | Quick Annual Review; defaults to 10.00 a.m. |
| `agm_place` | Quick Annual Review; defaults to registered office |
| `financial_statement_date` | Quick Annual Review; defaults to FYE |
| `accounts_status` | active / dormant / audited |
| `company_activity_status` | Optional compatibility field; normally leave blank |
| `financial_statements_type` | Optional compatibility field; normally leave blank |
| `financial_statements_required` | Optional compatibility field; normally leave blank |
| `audit_exemption_status` | Optional compatibility field; normally leave blank |
| `agm_status` | Auto / Held AGM / Dispensed with AGM / Exempt from AGM / Written resolutions |
| `acra_dormant_relevant_company` | Auto / Yes / No |
| `total_assets_under_500k` | Auto / Yes / No |
| `iras_tax_status` | Active / Dormant / Dormant waiver granted / Manual review |
| `director_signer_name` / `director_signer_names` | Quick Annual Review or company page |
| `shareholder_signer_name` / `member_signer_names` | Quick Annual Review or company page |
| `ar_authorized_signer_name` | Quick Annual Review; defaults to first director signer |
| `directors_fee` / `directors_remuneration` | Quick Annual Review |
| `management_rep_letter` | Quick Annual Review; default Yes |

## Guardrails

The package prepares signing documents and authorisations. It does not represent that ACRA / BizFile filing has already been lodged.
"""


def manifest() -> dict[str, object]:
    return {
        "version": VERSION,
        "package": "M05 Annual Review Package",
        "templates": TEMPLATES,
    }


def save_template(doc: Document, filename: str) -> None:
    output_path = OUTPUT_DIR / filename
    doc.save(output_path)
    shutil.copy2(output_path, TEMPLATE_DIR / filename)


def main() -> None:
    ensure_dirs()
    save_template(build_agm_package(), TEMPLATES["agm_package"])
    save_template(build_annual_return_package(), TEMPLATES["annual_return_package"])
    save_template(build_checklist(), TEMPLATES["checklist"])
    (OUTPUT_DIR / "M05_field_map.md").write_text(field_map_text(), encoding="utf-8")
    (OUTPUT_DIR / "M05_manifest.json").write_text(json.dumps(manifest(), ensure_ascii=False, indent=2), encoding="utf-8")
    for filename in TEMPLATES.values():
        print(TEMPLATE_DIR / filename)
    print(OUTPUT_DIR / "M05_field_map.md")


if __name__ == "__main__":
    main()
