from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DIR = ROOT / "app" / "doc_templates" / "p2_standard_v1"
OUTPUT_DIR = ROOT / "outputs" / "P2_standard_templates_v1"

TEMPLATES = {
    "resolution_package": "M02_resolution_package_transfer_in_standard.docx",
    "handover_resignation_package": "M02_handover_and_resignation_package_standard.docx",
}

BASE_FONT = "Calibri"
EAST_ASIA_FONT = "Microsoft YaHei"
INK = RGBColor(0x11, 0x18, 0x27)
MUTED = RGBColor(0x63, 0x6B, 0x78)
ACCENT = RGBColor(0x1F, 0x4D, 0x78)
LIGHT_FILL = "F2F4F7"
TABLE_BORDER = "B8C1CC"


def ensure_dirs() -> None:
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for directory in (TEMPLATE_DIR, OUTPUT_DIR):
        for path in directory.glob("M02_*"):
            if path.is_file():
                path.unlink()


def style_font(
    style,
    size: float,
    *,
    bold: bool = False,
    color: RGBColor | None = None,
    after: float = 6,
    before: float = 0,
    line_spacing: float = 1.10,
) -> None:
    style.font.name = BASE_FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    paragraph = style.paragraph_format
    paragraph.space_before = Pt(before)
    paragraph.space_after = Pt(after)
    paragraph.line_spacing = line_spacing


def set_run_font(
    run,
    size: float | None = None,
    *,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    run.font.name = BASE_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def configure_doc(doc: Document, title: str, subject: str) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    style_font(doc.styles["Normal"], 11, color=INK, after=6, line_spacing=1.10)
    style_font(doc.styles["Title"], 12.5, bold=True, color=INK, after=12, line_spacing=1.0)
    style_font(doc.styles["Heading 1"], 12, bold=True, color=INK, before=10, after=5, line_spacing=1.05)
    style_font(doc.styles["Heading 2"], 11, bold=True, color=ACCENT, before=8, after=4, line_spacing=1.05)
    style_font(doc.styles["Heading 3"], 10.5, bold=True, color=ACCENT, before=7, after=3, line_spacing=1.05)

    for style_name, size, color in [
        ("Legal Clause", 10.8, INK),
        ("Signature Block", 10.5, INK),
        ("Small Note", 8.8, MUTED),
    ]:
        if style_name not in doc.styles:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = doc.styles[style_name]
        style_font(style, size, color=color, after=4, line_spacing=1.10)

    doc.core_properties.title = title
    doc.core_properties.subject = subject
    doc.core_properties.author = "RSIN template rebuild"
    doc.core_properties.comments = "M02 P2 transfer-in package master template."
    add_footer(doc)


def add_field(paragraph, code: str) -> None:
    run = paragraph.add_run()
    field = OxmlElement("w:fldChar")
    field.set(qn("w:fldCharType"), "begin")
    run._r.append(field)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = code
    run._r.append(instr)
    field = OxmlElement("w:fldChar")
    field.set(qn("w:fldCharType"), "end")
    run._r.append(field)


def add_footer(doc: Document) -> None:
    paragraph = doc.sections[0].footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Page ")
    set_run_font(run, 8.5, color=MUTED)
    add_field(paragraph, "PAGE")
    run = paragraph.add_run(" of ")
    set_run_font(run, 8.5, color=MUTED)
    add_field(paragraph, "NUMPAGES")


def add_para(
    doc: Document,
    text: str = "",
    *,
    style: str | None = None,
    align: WD_ALIGN_PARAGRAPH | None = None,
    size: float | None = None,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    before: float | None = None,
    after: float | None = None,
    keep_with_next: bool = False,
    keep_together: bool = False,
):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    if before is not None:
        paragraph.paragraph_format.space_before = Pt(before)
    if after is not None:
        paragraph.paragraph_format.space_after = Pt(after)
    if keep_with_next:
        paragraph.paragraph_format.keep_with_next = True
    if keep_together:
        paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return paragraph


def add_company_header(doc: Document, title: str) -> None:
    add_para(doc, "{{company.company_name}}", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12.2, after=1, keep_with_next=True)
    add_para(doc, "Registration No. {{company.uen}}", align=WD_ALIGN_PARAGRAPH.CENTER, size=10.2, after=1, keep_with_next=True)
    add_para(doc, "(Incorporated in the Republic of Singapore)", align=WD_ALIGN_PARAGRAPH.CENTER, size=10.2, after=1, keep_with_next=True)
    add_para(doc, "(the \"Company\")", align=WD_ALIGN_PARAGRAPH.CENTER, size=10.2, after=12, keep_with_next=True)
    add_para(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12.4, after=14, keep_with_next=True)


def clause_title(doc: Document, title: str) -> None:
    add_para(doc, title, style="Heading 1", keep_with_next=True)


def clause_para(doc: Document, text: str) -> None:
    add_para(doc, text, style="Legal Clause")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_table_borders(table, color: str = TABLE_BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for tag in ("w:top", "w:left", "w:bottom", "w:right", "w:insideH", "w:insideV"):
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top: int = 90, bottom: int = 90, start: int = 120, end: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for tag, value in (("w:top", top), ("w:bottom", bottom), ("w:start", start), ("w:end", end)):
        element = margins.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        row_pr = row._tr.get_or_add_trPr()
        if row_pr.find(qn("w:cantSplit")) is None:
            row_pr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 9.8, align: WD_ALIGN_PARAGRAPH | None = None, fill: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=INK)
    if fill:
        shade_cell(cell, fill)


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, [2600, 6760])
    set_table_borders(table)
    set_cell_margins(table)
    for idx, (label, value) in enumerate(rows):
        set_cell_text(table.cell(idx, 0), label, bold=True, size=9.4, fill=LIGHT_FILL)
        set_cell_text(table.cell(idx, 1), value, size=9.4)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_signature_text(doc: Document, title: str, placeholder: str) -> None:
    clause_title(doc, title)
    add_para(doc, placeholder, style="Signature Block", after=8, keep_together=True)


def build_resolution_package() -> Document:
    doc = Document()
    configure_doc(doc, "M02 EGM and Board Resolutions", "P2 transfer-in EGM notice, members authority and directors implementation template")
    add_company_header(doc, "NOTICE OF EXTRAORDINARY GENERAL MEETING AND RESOLUTIONS")
    add_key_value_table(
        doc,
        [
            ("Date", "{{m02.effective_date}}"),
            ("Meeting / signing venue", "{{m02.meeting_place}}"),
            ("Existing secretarial arrangement", "{{m02.old_secretary_company}}"),
            ("New service provider", "{{m02.new_secretary_company}}"),
        ],
    )
    clause_title(doc, "A. Notice of Extraordinary General Meeting")
    clause_para(doc, "NOTICE IS HEREBY GIVEN that an Extraordinary General Meeting of the Company will be held at {{m02.meeting_place}}, or such other place as may be agreed by the members, on {{m02.effective_date}} at 10.00 a.m. for the purpose of considering and, if thought fit, passing the ordinary resolutions set out in this package, with or without modification.")
    clause_para(doc, "The matters to be considered include the Company's corporate secretarial service arrangement, the handover of statutory records, the authority for related ACRA / BizFile lodgements and, where applicable, any related appointment, resignation, removal or cessation item separately identified in the Company's instructions.")
    add_signature_text(doc, "Issued By Order of the Board", "{{m02.notice_issuer_signature_block}}")

    doc.add_page_break()
    add_company_header(doc, "CONSENT TO SHORTER NOTICE OF EXTRAORDINARY GENERAL MEETING")
    clause_para(doc, "The undersigned member(s) of the Company hereby consent to the Extraordinary General Meeting of the Company being convened and held on shorter notice for the purpose of considering the ordinary resolutions set out in the notice of meeting, notwithstanding any shorter period of notice than may otherwise be required under the Constitution of the Company or applicable law.")
    clause_para(doc, "This consent is given in connection with the Company's transfer of corporate secretarial administration, statutory records handover and related corporate records update.")
    add_signature_text(doc, "Signed by the Member(s)", "{{m02.member_signature_blocks}}")

    doc.add_page_break()
    add_company_header(doc, "MINUTES OF EXTRAORDINARY GENERAL MEETING OF THE MEMBERS")
    clause_title(doc, "Background")
    clause_para(doc, "The Member(s) have considered the Company's existing corporate secretarial service arrangement with {{m02.old_secretary_company}} and the proposed engagement of {{m02.new_secretary_company}} to support the Company's statutory records and corporate secretarial administration.")
    clause_para(doc, "The Member(s) consider it in the interests of the Company to authorise the directors and the Company's authorised signatory to take all necessary steps in connection with the transfer of corporate secretarial administration, the handover of statutory records and any relevant ACRA / BizFile arrangement.")
    clause_title(doc, "Ordinary Resolutions")
    clause_para(doc, "RESOLVED THAT the Company's existing corporate secretarial service arrangement with {{m02.old_secretary_company}} be reviewed and, where applicable, terminated or transitioned with effect from {{m02.effective_date}} or such other date as the directors may determine.")
    clause_para(doc, "RESOLVED THAT {{m02.new_secretary_company}} be engaged and authorised as the Company's corporate secretarial service provider, and that the directors and/or authorised signatory of the Company be authorised to sign all engagement, handover, termination, authorisation and related documents required for such arrangement.")
    clause_para(doc, "RESOLVED THAT the directors be authorised to request the release or handover of the Company's statutory records, update the Company's internal records, and arrange the filing or updating of information with ACRA / BizFile where applicable.")
    clause_para(doc, "RESOLVED THAT any related appointment, resignation, removal, cessation or change of officer shall be documented and effected only where the Company has the relevant consent, resignation, member authority, board approval or other proper authority required for that item.")
    clause_para(doc, "RESOLVED THAT all actions already taken by any director, member or authorised representative of the Company in connection with the matters above be and are hereby approved, confirmed and ratified.")
    add_signature_text(doc, "Signed by the Member(s)", "{{m02.member_signature_blocks}}")

    doc.add_page_break()
    add_company_header(doc, "DIRECTORS' RESOLUTIONS IN WRITING")
    clause_title(doc, "Noted")
    clause_para(doc, "The directors noted the authority given by the Member(s) in relation to the Company's transfer of corporate secretarial administration and statutory records.")
    clause_para(doc, "The directors further noted the following related officer or service arrangement items identified for this transfer-in package:")
    clause_para(doc, "{{m02.personnel_change_summary}}")
    clause_title(doc, "Resolved")
    clause_para(doc, "RESOLVED THAT {{m02.new_secretary_company}} be authorised to liaise with {{m02.old_secretary_company}}, request and receive the Company's statutory records and registers, and assist the Company with the necessary records update and filing arrangements.")
    clause_para(doc, "RESOLVED THAT any director of the Company and/or the authorised signatory named in the Company's records be authorised to sign the handover letter, service termination letter, authorisation letter, BizFile authorisation and any ancillary documents required for the matters above.")
    clause_para(doc, "RESOLVED THAT where the appointment or resignation of an individual company secretary, director or officer is required, such appointment, resignation or cessation shall be supported by the relevant consent, resignation letter, directors' resolution, members' resolution or other proper authority, as applicable.")
    clause_para(doc, "RESOLVED THAT the Company shall continue to maintain at least one director who is ordinarily resident in Singapore, and that no change shall be lodged if it would result in the Company failing to meet that requirement.")
    clause_para(doc, "RESOLVED THAT the Company's statutory registers, minute book and internal records be updated accordingly after the relevant changes have been validly effected.")
    add_signature_text(doc, "Signed by the Directors", "{{m02.director_signature_blocks}}")
    return doc


def build_handover_resignation_package() -> Document:
    doc = Document()
    configure_doc(doc, "M02 Handover and Resignation Package", "P2 transfer-in handover and optional resignation package template")
    add_para(doc, "{{m02.effective_date}}", align=WD_ALIGN_PARAGRAPH.RIGHT, after=14)
    add_para(doc, "{{m02.old_secretary_company}}", bold=True, after=2)
    add_para(doc, "Dear Sirs", after=10)
    add_para(doc, "RE: {{company.company_name}} (UEN: {{company.uen}})", bold=True, after=8)
    clause_para(doc, "We refer to the Company's corporate secretarial service arrangement and statutory records maintained by your office.")
    clause_para(doc, "Please be informed that the Company has authorised {{m02.new_secretary_company}} to assist with its corporate secretarial administration with effect from {{m02.effective_date}}. Please therefore arrange the handover of the Company's statutory records and relevant records to {{m02.new_secretary_company}} or to the Company's authorised representative.")
    clause_para(doc, "The handover should include, where applicable:")
    clause_para(doc, "{{m02.handover_request_items}}")
    clause_para(doc, "Please also confirm whether there are any outstanding filings, compliance matters, invoices or documents requiring attention before the handover is completed.")
    clause_para(doc, "We appreciate your assistance with the handover and look forward to completing the transition in an orderly and professional manner.")
    add_para(doc, "Yours faithfully,", after=20)
    add_para(doc, "{{m02.client_signature_block}}", style="Signature Block")
    return doc


def field_map_text() -> str:
    return """# M02 Transfer-In Package - Field Map

M02 is the transfer-in package for changing or transferring the corporate secretarial administration of an existing Singapore company.

## Generated files

| File | Trigger | Signer | Notes |
|---|---|---|---|
| Notice of EGM and Resolutions | `transfer_in_required = Yes` | Notice issuer, member(s), and all director signers | Contains the Notice of EGM, Consent to Shorter Notice, Members' Resolution / Minutes and Directors' Resolution in one PDF. |
| Handover and Resignation Package | `transfer_in_required = Yes` | Client signatory; resigning person(s) only if resignation letters are requested and valid | Contains the service handover letter and optional resignation letters in one PDF. |

## Primary input fields

| Field | Source | Used in |
|---|---|---|
| `company.company_name` | P2 one-page sheet | All M02 documents |
| `company.uen` | P2 one-page sheet | All M02 documents |
| `company.registered_office_address` | P2 one-page sheet | Resolution package |
| `company.default_document_date` | P2 one-page sheet | Default document/effective date |
| `company.director_signer_names` | P2 one-page sheet | Directors' resolution signatures |
| `company.member_signer_names` / `company.shareholder_signer_names` | P2 one-page sheet | Member consent/resolution signatures; supports multiple names separated by newline, comma or semicolon |
| `company.client_signatory_name` | P2 one-page sheet | Handover letter signer; if empty, the first member/shareholder signer is used |
| `transfer_in_required` | P2 one-page sheet | Enables M02 |
| `transfer_in_mode` | P2 one-page sheet | Internal review flag only; not shown in formal documents |
| `old_secretary_company` | P2 one-page sheet | Existing corporate secretarial service provider |
| `new_secretary_company` | P2 one-page sheet | New corporate secretarial service provider |
| `generate_resignation_letter` | P2 one-page sheet | Optional resignation letters |
| Personnel action `action_type` | P2 one-page sheet | Identifies related appointment/resignation items |
| Personnel action `target_name` | P2 one-page sheet | Officer/person name |
| Personnel action `effective_date` | P2 one-page sheet | Change or resignation effective date |

## Legal drafting guardrail

The new service provider is treated as the corporate secretarial service provider, not as the individual statutory company secretary. Appointment or resignation of an individual secretary should remain a separate appointment/resignation item where required.
"""


def manifest() -> dict[str, object]:
    return {
        "version": "P2_M02_v0.2",
        "package": "M02 Transfer-In Package",
        "templates": TEMPLATES,
        "default_files": [
            "resolution_package",
            "handover_resignation_package",
        ],
        "optional_sections": ["resignation_letters"],
    }


def save_template(doc: Document, filename: str) -> None:
    output_path = OUTPUT_DIR / filename
    doc.save(output_path)
    shutil.copy2(output_path, TEMPLATE_DIR / filename)


def main() -> None:
    ensure_dirs()
    save_template(build_resolution_package(), TEMPLATES["resolution_package"])
    save_template(build_handover_resignation_package(), TEMPLATES["handover_resignation_package"])
    (OUTPUT_DIR / "M02_field_map.md").write_text(field_map_text(), encoding="utf-8")
    (OUTPUT_DIR / "M02_manifest.json").write_text(json.dumps(manifest(), ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
