from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
P1_DIR = ROOT / "app" / "doc_templates" / "p1_standard_v3_part1"


def set_cell_text(cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"


def remove_extra_rows(table, keep_count: int) -> None:
    while len(table.rows) > keep_count:
        table._tbl.remove(table.rows[-1]._tr)


def remove_column(table, index: int) -> None:
    for row in table.rows:
        row._tr.remove(row.cells[index]._tc)
    grid = table._tbl.tblGrid
    if grid is not None and len(grid.gridCol_lst) > index:
        grid.remove(grid.gridCol_lst[index])


def remove_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.text = text
    for run in paragraph.runs:
        run.font.name = "Arial"


def add_paragraph(
    doc: Document,
    text: str = "",
    *,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size: float = 12,
    bold: bool = False,
    italic: bool = False,
    color: str | None = None,
    before: float = 0,
    after: float = 4,
):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return paragraph


def set_certificate_cell(cell, text: str, *, size: float = 11, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold


def add_cell_paragraph(
    cell,
    text: str = "",
    *,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size: float = 12,
    bold: bool = False,
    italic: bool = False,
    color: str | None = None,
    after: float = 3,
):
    paragraph = cell.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return paragraph


def ensure_page_break_before(paragraph) -> None:
    previous = paragraph._p.getprevious()
    if previous is not None and "<w:br" in previous.xml and 'type="page"' in previous.xml:
        return
    breaker = paragraph.insert_paragraph_before()
    breaker.add_run().add_break(WD_BREAK.PAGE)


def set_table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)


def set_table_borders_none(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_row_min_height(row, height_inches: float) -> None:
    row.height = Inches(height_inches)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def set_certificate_page_border(section, color: str = "9A1B1E") -> None:
    sect_pr = section._sectPr
    for existing in sect_pr.findall(qn("w:pgBorders")):
        sect_pr.remove(existing)

    borders = OxmlElement("w:pgBorders")
    borders.set(qn("w:offsetFrom"), "page")
    borders.set(qn("w:display"), "allPages")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "thinThickThinMediumGap")
        border.set(qn("w:sz"), "20")
        border.set(qn("w:space"), "18")
        border.set(qn("w:color"), color)
        borders.append(border)
    sect_pr.append(borders)


def update_first_directors_resolution() -> None:
    doc = Document(P1_DIR / "01_first_directors_resolution_standard.docx")
    for paragraph in doc.paragraphs:
        if "financial year end of the Company be fixed on" in paragraph.text:
            set_paragraph_text(
                paragraph,
                "That the financial year end of the Company be fixed on {{company.fye_month}} "
                "and the first set of accounts will be for the period from "
                "{{company.first_financial_period_start}} to {{company.first_financial_period_end}}.",
            )

    subscriber_table = doc.tables[0]
    set_cell_text(subscriber_table.rows[1].cells[0], "{% for shareholder in shareholders %}{{shareholder.shareholder_name}}")
    set_cell_text(subscriber_table.rows[1].cells[1], "{{shareholder.shares}}{% endfor %}")

    table = doc.tables[1]
    remove_extra_rows(table, 2)
    set_cell_text(table.rows[0].cells[0], "Signed by the Directors")
    set_cell_text(table.rows[1].cells[0], "{{company.director_signature_blocks}}")
    date_row = table.add_row()
    set_cell_text(date_row.cells[0], "Dated: {{signature.date}}")
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip().startswith("Dated:"):
            remove_paragraph(paragraph)
    doc.save(P1_DIR / "01_first_directors_resolution_standard.docx")


def update_director_consent_form45() -> None:
    doc = Document(P1_DIR / "02_director_consent_form45_standard.docx")
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip().startswith("*(i) I have read and understood the above statements"):
            set_paragraph_text(
                paragraph,
                "I have read and understood the above statements and confirm that the statements are true. "
                "I am also aware that I can be prosecuted in Court if I willfully give any information on this form which is false.",
            )
        elif "above statements were interpreted to me" in paragraph.text:
            remove_paragraph(paragraph)
    doc.save(P1_DIR / "02_director_consent_form45_standard.docx")


def update_share_certificate() -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    set_certificate_page_border(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    layout = doc.add_table(rows=4, cols=1)
    layout.autofit = False
    set_table_widths(layout, [9.9])
    set_table_borders_none(layout)
    set_row_min_height(layout.rows[0], 1.65)
    set_row_min_height(layout.rows[1], 3.05)
    set_row_min_height(layout.rows[2], 1.25)
    set_row_min_height(layout.rows[3], 0.28)

    for row in layout.rows:
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    header_cell = layout.rows[0].cells[0]
    add_cell_paragraph(header_cell, "SHARE CERTIFICATE", size=19, bold=True, color="7F1416", after=0)
    add_cell_paragraph(header_cell, "======================================================================", size=9, color="9A1B1E", after=4)

    meta = header_cell.add_table(rows=1, cols=3)
    meta.autofit = False
    set_table_widths(meta, [2.0, 5.7, 2.0])
    set_table_borders_none(meta)
    for cell in meta.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_certificate_cell(meta.rows[0].cells[0], "Certificate No.\n\n{{shareholder.certificate_no}}", size=10.5, bold=True)
    set_certificate_cell(
        meta.rows[0].cells[1],
        "{{company.company_name}}\n(Incorporated in the Republic of Singapore)\n"
        "Co. Reg. No.: {{company.uen}}\nRegistered Office: {{company.registered_office_address}}",
        size=10.5,
        bold=True,
    )
    set_certificate_cell(meta.rows[0].cells[2], "No. of Shares\n\n{{shareholder.shares}}", size=10.5, bold=True)

    body_cell = layout.rows[1].cells[0]
    add_cell_paragraph(body_cell, "THIS IS TO CERTIFY THAT", size=14, bold=True, after=1)
    add_cell_paragraph(body_cell, "{{shareholder.shareholder_name}}", size=15, bold=True, after=2)
    add_cell_paragraph(body_cell, "of", size=11.5, italic=True, after=0)
    add_cell_paragraph(body_cell, "{{shareholder.shareholder_address}}", size=11.5, italic=True, after=3)
    add_cell_paragraph(body_cell, "is/are the registered holder(s) of", size=11.5, italic=True, after=0)
    add_cell_paragraph(body_cell, "{{shareholder.shares_in_words}}", size=13, bold=True, italic=True, after=3)
    add_cell_paragraph(
        body_cell,
        "fully paid in the Company subject to the Memorandum & Articles of Association of the Company.",
        size=11.5,
        italic=True,
        after=5,
    )
    add_cell_paragraph(body_cell, "Dated this {{signature.day}} day of {{signature.month_year}}", size=11.5, italic=True, after=5)
    add_cell_paragraph(body_cell, "Executed by the Company on the date stated above in accordance with the Companies Act.", size=11.5, italic=True, after=0)

    signatures_cell = layout.rows[2].cells[0]
    signatures_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    signatures = signatures_cell.add_table(rows=1, cols=3)
    signatures.autofit = False
    set_table_widths(signatures, [3.15, 3.15, 3.15])
    set_table_borders_none(signatures)
    set_certificate_cell(signatures.rows[0].cells[0], "____________________________\nDirector\nName: ____________________", size=10.5)
    set_certificate_cell(signatures.rows[0].cells[1], "For and on behalf of\n{{company.company_name}}", size=10.5)
    set_certificate_cell(signatures.rows[0].cells[2], "____________________________\nCompany Secretary\n{{secretary.full_name}}", size=10.5)

    note_cell = layout.rows[3].cells[0]
    add_cell_paragraph(
        note_cell,
        "Note: Transfer of any portion of the Shares comprised in this Certificate cannot be registered unless accompanied by this Certificate.",
        size=8,
        after=0,
    )
    doc.save(P1_DIR / "04_share_certificate_standard.docx")


def update_secretary_service_agreement() -> None:
    doc = Document(P1_DIR / "05_secretary_service_agreement_standard.docx")
    for paragraph in list(doc.paragraphs):
        if "Additional signatory:" in paragraph.text:
            remove_paragraph(paragraph)
        if paragraph.text.strip() == "Execution":
            ensure_page_break_before(paragraph)
    table = doc.tables[1]
    if len(table.rows[0].cells) > 2:
        remove_column(table, 2)
    set_table_widths(table, [3.05, 3.05])
    set_cell_text(table.rows[0].cells[0], "Signed for and on behalf of {{provider.name}}")
    set_cell_text(table.rows[0].cells[1], "Signed by Client")
    set_cell_text(table.rows[1].cells[0], "Signature:\n\n____________________")
    set_cell_text(table.rows[1].cells[1], "Signature:\n\n____________________")
    set_cell_text(table.rows[2].cells[0], "Name: {{provider.authorised_signatory_name}}\nCapacity: Authorised Signatory")
    set_cell_text(table.rows[2].cells[1], "Name: {{client_signatory.full_name}}\nCapacity: {{client_signatory.capacity}}")
    doc.save(P1_DIR / "05_secretary_service_agreement_standard.docx")


def update_nominee_director_agreement() -> None:
    doc = Document(P1_DIR / "06_nominee_director_agreement_standard.docx")
    for paragraph in doc.paragraphs:
        if "{{company.company_name_chinese}}" in paragraph.text or "{{company.registered_office_address_chinese}}" in paragraph.text:
            set_paragraph_text(
                paragraph,
                '{{company.company_name}} (Company Registration No. {{company.uen}}), a company incorporated in the Republic of Singapore '
                'and having its registered address at {{company.registered_office_address}} ("Company").',
            )
    table = doc.tables[0]
    set_cell_text(
        table.rows[2].cells[1],
        "Name: {{client_signatory.full_name}}\nCapacity: {{client_signatory.capacity}}, {{company.company_name}}",
    )
    doc.save(P1_DIR / "06_nominee_director_agreement_standard.docx")


def update_form24() -> None:
    doc = Document(P1_DIR / "07_return_of_allotment_form24_standard.docx")
    table = doc.tables[4]
    if len(table.rows) >= 2:
        row = table.rows[1]
        set_cell_text(
            row.cells[0],
            "{% for shareholder in shareholders %}{{shareholder.shareholder_name}}\n\n"
            "{{shareholder.shareholder_address}}\n\n"
            "IDENTIFICATION NO.: {{shareholder.id_number}}\n\n"
            "{{shareholder.nationality}}",
        )
        set_cell_text(
            row.cells[1],
            "{{shareholder.shares}} {{shareholder.share_class}} Shares\n\n"
            "Allotted on the date of Incorporation.{% endfor %}",
        )
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    doc.save(P1_DIR / "07_return_of_allotment_form24_standard.docx")


def main() -> None:
    update_first_directors_resolution()
    update_director_consent_form45()
    update_share_certificate()
    update_secretary_service_agreement()
    update_nominee_director_agreement()
    update_form24()


if __name__ == "__main__":
    main()
