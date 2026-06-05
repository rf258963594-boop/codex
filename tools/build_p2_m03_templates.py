from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DIR = ROOT / "app" / "doc_templates" / "p2_standard_v1"
OUTPUT_DIR = ROOT / "outputs" / "P2_standard_templates_v1"

TEMPLATES = {
    "resolution": "M03_share_transfer_directors_resolution_standard.docx",
    "instrument": "M03_instrument_of_transfer_standard.docx",
    "certificate": "M03_updated_share_certificate_standard.docx",
    "checklist": "M03_register_and_stamp_duty_checklist_standard.docx",
}

BASE_FONT = "Calibri"
EAST_ASIA_FONT = "Microsoft YaHei"
INK = RGBColor(0x11, 0x18, 0x27)
MUTED = RGBColor(0x63, 0x6B, 0x78)
ACCENT = RGBColor(0x1F, 0x4D, 0x78)
RED = RGBColor(0x9F, 0x12, 0x12)
LIGHT_FILL = "F2F4F7"
TABLE_BORDER = "B8C1CC"


def ensure_dirs() -> None:
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for directory in (TEMPLATE_DIR, OUTPUT_DIR):
        for path in directory.glob("M03_*"):
            if path.is_file():
                path.unlink()


def style_font(
    style,
    size: float,
    *,
    bold: bool = False,
    color: RGBColor | None = None,
    after: float = 6,
    before: float = 0,
    line_spacing: float = 1.10,
) -> None:
    style.font.name = BASE_FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    style.font.size = Pt(size)
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    paragraph = style.paragraph_format
    paragraph.space_before = Pt(before)
    paragraph.space_after = Pt(after)
    paragraph.line_spacing = line_spacing


def set_run_font(
    run,
    size: float | None = None,
    *,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    run.font.name = BASE_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def configure_doc(doc: Document, title: str, subject: str) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    style_font(doc.styles["Normal"], 11, color=INK, after=6, line_spacing=1.10)
    style_font(doc.styles["Title"], 12.5, bold=True, color=INK, after=12, line_spacing=1.0)
    style_font(doc.styles["Heading 1"], 12, bold=True, color=INK, before=10, after=5, line_spacing=1.05)
    style_font(doc.styles["Heading 2"], 11, bold=True, color=ACCENT, before=8, after=4, line_spacing=1.05)

    for style_name, size, color in [
        ("Legal Clause", 10.8, INK),
        ("Signature Block", 10.5, INK),
        ("Small Note", 8.8, MUTED),
    ]:
        if style_name not in doc.styles:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = doc.styles[style_name]
        style_font(style, size, color=color, after=4, line_spacing=1.10)

    doc.core_properties.title = title
    doc.core_properties.subject = subject
    doc.core_properties.author = "RSIN template rebuild"
    doc.core_properties.comments = "M03 P2 share transfer package master template."
    add_footer(doc)


def configure_certificate_doc(doc: Document, title: str, subject: str) -> None:
    configure_doc(doc, title, subject)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    for paragraph in section.footer.paragraphs:
        paragraph.clear()


def add_field(paragraph, code: str) -> None:
    run = paragraph.add_run()
    field = OxmlElement("w:fldChar")
    field.set(qn("w:fldCharType"), "begin")
    run._r.append(field)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = code
    run._r.append(instr)
    field = OxmlElement("w:fldChar")
    field.set(qn("w:fldCharType"), "end")
    run._r.append(field)


def add_footer(doc: Document) -> None:
    paragraph = doc.sections[0].footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Page ")
    set_run_font(run, 8.5, color=MUTED)
    add_field(paragraph, "PAGE")
    run = paragraph.add_run(" of ")
    set_run_font(run, 8.5, color=MUTED)
    add_field(paragraph, "NUMPAGES")


def add_para(
    doc: Document,
    text: str = "",
    *,
    style: str | None = None,
    align: WD_ALIGN_PARAGRAPH | None = None,
    size: float | None = None,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
    before: float | None = None,
    after: float | None = None,
    keep_with_next: bool = False,
    keep_together: bool = False,
):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    if before is not None:
        paragraph.paragraph_format.space_before = Pt(before)
    if after is not None:
        paragraph.paragraph_format.space_after = Pt(after)
    if keep_with_next:
        paragraph.paragraph_format.keep_with_next = True
    if keep_together:
        paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return paragraph


def add_company_header(doc: Document, title: str) -> None:
    add_para(doc, "{{company.company_name}}", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12.2, after=1, keep_with_next=True)
    add_para(doc, "Company Registration No. {{company.uen}}", align=WD_ALIGN_PARAGRAPH.CENTER, size=10.2, after=1, keep_with_next=True)
    add_para(doc, "(Incorporated in the Republic of Singapore)", align=WD_ALIGN_PARAGRAPH.CENTER, size=10.2, after=1, keep_with_next=True)
    add_para(doc, "(the \"Company\")", align=WD_ALIGN_PARAGRAPH.CENTER, size=10.2, after=12, keep_with_next=True)
    add_para(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12.4, after=14, keep_with_next=True)


def clause_title(doc: Document, title: str) -> None:
    add_para(doc, title, style="Heading 1", keep_with_next=True)


def clause_para(doc: Document, text: str) -> None:
    add_para(doc, text, style="Legal Clause")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_table_borders(table, color: str = TABLE_BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=90, bottom=90, start=110, end=110) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for key, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        element = margins.find(qn(f"w:{key}"))
        if element is None:
            element = OxmlElement(f"w:{key}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=False):
            set_cell_width(cell, width)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 9.6, align: WD_ALIGN_PARAGRAPH | None = None, color: RGBColor | None = None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.clear()
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def marker(doc: Document, text: str) -> None:
    add_para(doc, text, size=1, color=RGBColor(0xFF, 0xFF, 0xFF), after=0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    set_table_width(table, widths)
    set_table_borders(table)
    set_cell_margins(table)
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        shade_cell(cell, LIGHT_FILL)
        set_cell_text(cell, header, bold=True, size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            set_cell_text(table.cell(row_idx, col_idx), value, size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_transfer_table(doc: Document) -> None:
    marker(doc, "[[REPEAT m03.transfers[]]]")
    add_table(
        doc,
        ["Transferor", "Transferee", "Class", "No. of shares", "Transfer date", "Consideration / certificates"],
        [[
            "{{transfer.transferor_name}}",
            "{{transfer.transferee_name}}",
            "{{transfer.share_class}}",
            "{{transfer.shares_transferred}}",
            "{{transfer.transfer_date}}",
            "{{transfer.consideration_and_certificate_text}}",
        ]],
        [1580, 1580, 1100, 1160, 1260, 2680],
    )


def add_party_signature_table(doc: Document) -> None:
    marker(doc, "[[REPEAT m03.transfer_parties[]]]")
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [4680, 4680])
    set_table_borders(table, "FFFFFF")
    set_cell_margins(table, top=40, bottom=60, start=0, end=160)
    set_row_cant_split(table.rows[0])
    set_cell_text(table.cell(0, 0), "\n______________________________\n{{party.full_name}}\n{{party.capacity}}", size=10.2)
    set_cell_text(table.cell(0, 1), "{{party.id_line}}\n{{party.address_line}}", size=8.8, color=MUTED)


def add_director_signature_table(doc: Document) -> None:
    marker(doc, "[[REPEAT m03.director_signature_rows[]]]")
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [4680, 4680])
    set_table_borders(table, "FFFFFF")
    set_cell_margins(table, top=40, bottom=60, start=0, end=160)
    set_row_cant_split(table.rows[0])
    set_cell_text(table.cell(0, 0), "{{sigrow.left_block}}", size=10.3)
    set_cell_text(table.cell(0, 1), "{{sigrow.right_block}}", size=10.3)


def add_checklist_table(doc: Document) -> None:
    marker(doc, "[[REPEAT m03.checklist_items[]]]")
    add_table(
        doc,
        ["Item", "Status", "Review note"],
        [["{{check.item}}", "{{check.status}}", "{{check.note}}"]],
        [2350, 1450, 5560],
    )


def build_resolution_docx() -> Document:
    doc = Document()
    configure_doc(doc, "M03 Share Transfer Directors' Resolution", "P2 share transfer directors' resolution")
    add_company_header(doc, "DIRECTORS' RESOLUTIONS IN WRITING")
    clause_para(doc, "The undersigned, being the director(s) of the Company, hereby pass the following resolutions in writing pursuant to the Constitution of the Company.")
    clause_para(doc, "RESOLVED: -")
    clause_title(doc, "APPROVAL OF TRANSFER OF SHARES")
    clause_para(doc, "That the transfer of the following shares in the Company be and is hereby approved:")
    add_transfer_table(doc)
    clause_para(doc, "That the relevant share certificate(s) be cancelled and/or issued as required, and that the Register of Members and related corporate records of the Company be updated accordingly.")
    clause_para(doc, "That any director of the Company and/or the company secretary and/or {{provider.name}} be authorised to prepare, sign, lodge, amend and submit all necessary documents, notices and electronic transactions with ACRA and any other relevant authority to give effect to the above resolutions.")
    add_para(doc, "Dated this {{signature.day_ordinal}} day of {{signature.month_year}}", style="Signature Block", before=10, after=10)
    add_para(doc, "DIRECTOR(S)", bold=True, size=10.5, after=4, keep_with_next=True)
    add_director_signature_table(doc)
    return doc


def build_instrument_docx() -> Document:
    doc = Document()
    configure_doc(doc, "M03 Instrument of Transfer", "P2 share transfer instrument")
    add_company_header(doc, "INSTRUMENT OF TRANSFER")
    clause_para(doc, "The transferor(s) and transferee(s) named in the Schedule below agree to the transfer of the shares described in the Schedule, subject to the Constitution of the Company and all applicable legal, filing and stamp duty requirements.")
    clause_para(doc, "For the consideration stated in the Schedule or otherwise agreed between the relevant transferor and transferee, the transferor transfers to the transferee the legal and beneficial interest in the shares specified below.")
    clause_title(doc, "SCHEDULE OF TRANSFER")
    add_transfer_table(doc)
    clause_para(doc, "The parties acknowledge that the Company may update its Register of Members, cancel and/or issue share certificate(s), and complete the relevant statutory records only after the required documents and review items have been completed.")
    add_para(doc, "Dated this {{signature.day_ordinal}} day of {{signature.month_year}}", style="Signature Block", before=10, after=10)
    add_para(doc, "TRANSFEROR(S) AND TRANSFEREE(S)", bold=True, size=10.5, after=4, keep_with_next=True)
    add_party_signature_table(doc)
    return doc


def add_certificate_border_table(doc: Document) -> None:
    add_landscape_certificate_box(doc)


def add_landscape_certificate_box(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [14400])
    set_table_borders(table, "9F1212")
    set_cell_margins(table, top=340, bottom=300, start=420, end=420)
    table.rows[0].height = Inches(7.35)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def cell_para(text: str, size: float, *, bold: bool = False, color: RGBColor = INK, after: float = 6) -> None:
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(after)
        run = paragraph.add_run(text)
        set_run_font(run, size, bold=bold, color=color)

    first = cell.paragraphs[0]
    first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first.paragraph_format.space_after = Pt(4)
    run = first.add_run("SHARE CERTIFICATE")
    set_run_font(run, 24, bold=True, color=RED)

    cell_para("{{company.company_name}}", 15, bold=True, after=4)
    cell_para("Company Registration No. {{company.uen}}", 10.8, after=2)
    cell_para("Incorporated in the Republic of Singapore", 10.5, color=MUTED, after=12)
    cell_para("Certificate No.: {{certificate.certificate_no}}", 12, bold=True, color=RED, after=14)
    cell_para("This is to certify that", 11.5, color=MUTED, after=2)
    cell_para("{{certificate.holder_name}}", 18, bold=True, after=2)
    cell_para("is the registered holder of", 11.5, color=MUTED, after=2)
    cell_para("{{certificate.shares_text}} {{certificate.share_class}} share(s)", 18, bold=True, after=2)
    cell_para("in the Company, {{certificate.paid_status_text}}.", 12.2, after=12)
    cell_para("Issued on {{certificate.issue_date}}     {{certificate.source_reference}}", 10.6, color=MUTED, after=20)

    sig_table = cell.add_table(rows=2, cols=2)
    set_table_width(sig_table, [6100, 6100])
    set_table_borders(sig_table, "FFFFFF")
    set_cell_margins(sig_table, top=80, bottom=40, start=160, end=160)
    set_cell_text(sig_table.cell(0, 0), "Director", bold=True, size=10.8, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(sig_table.cell(0, 1), "Company Secretary", bold=True, size=10.8, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(sig_table.cell(1, 0), "\n______________________________", size=10.8, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(sig_table.cell(1, 1), "\n______________________________", size=10.8, align=WD_ALIGN_PARAGRAPH.CENTER)

    cell_para("This certificate is issued subject to the Constitution of the Company and the Company's register of members.", 8.8, color=MUTED, after=0)


def build_certificate_docx() -> Document:
    doc = Document()
    configure_certificate_doc(doc, "M03 Updated Share Certificate", "P2 updated share certificate after share transfer")
    add_landscape_certificate_box(doc)
    return doc


def build_checklist_docx() -> Document:
    doc = Document()
    configure_doc(doc, "M03 Register and Stamp Duty Checklist", "P2 internal share transfer checklist")
    add_company_header(doc, "INTERNAL REGISTER AND STAMP DUTY REVIEW CHECKLIST")
    clause_para(doc, "This checklist is for internal review before releasing the share transfer package or updating the Company's statutory records.")
    add_checklist_table(doc)
    clause_title(doc, "TRANSFER SUMMARY")
    add_transfer_table(doc)
    add_para(doc, "Prepared by: ______________________________", style="Signature Block", before=12, after=6)
    add_para(doc, "Reviewed by: ______________________________", style="Signature Block", after=6)
    return doc


def field_map_markdown() -> str:
    return """# M03 Share Transfer Package - Field Map

M03 is the P2 package for ordinary share transfers. It does not generate Form 24; Form 24 is reserved for share allotment / capital increase workflows.

## Generated files

| File | Trigger | Signer | Notes |
|---|---|---|---|
| Share Transfer Directors' Resolution | At least one active share transfer row | All director signers | Approves transfer and authorises records / filing updates. |
| Instrument of Transfer | At least one active share transfer row | Transferor(s) and transferee(s) | Uses the share transfer schedule from the P2 one-page sheet. |
| Updated Share Certificate | Per transferee where `generate_new_certificate` is not No | Director signature line + Company Secretary | Certificate number should be checked before signing. |
| Register and Stamp Duty Checklist | At least one active share transfer row | Internal review | Internal checklist for EROM/Register, certificates and stamp duty / NAV review. |

## Primary input fields

| Field | Source | Used in |
|---|---|---|
| `company.company_name` | P2 one-page sheet | All M03 documents |
| `company.uen` | P2 one-page sheet | All M03 documents |
| `company.registered_office_address` | P2 one-page sheet | Background and instrument context |
| `company.default_document_date` | P2 one-page sheet | Default transfer / document date |
| `company.director_signer_names` | P2 one-page sheet | Directors' resolution signatures |
| `share_transfers[].transferor_name` | Share transfer table | Resolution, instrument, checklist |
| `share_transfers[].transferee_name` | Share transfer table | Resolution, instrument, certificate, checklist |
| `share_transfers[].shares_transferred` | Share transfer table | Resolution, instrument, certificate |
| `share_transfers[].share_class` | Share transfer table | Defaults to Ordinary |
| `share_transfers[].transfer_date` | Share transfer table | Defaults to company document date |
| `share_transfers[].consideration_amount` / `consideration_basis` | Share transfer table | Instrument and checklist |
| `share_transfers[].old_certificate_no` / `new_certificate_no` | Share transfer table | Certificate and checklist |
| `share_transfers[].stamp_duty_review` | Share transfer table | Internal checklist warning |

## Guardrails

- Ordinary share transfer does not generate Form 24.
- Stamp duty / NAV / actual consideration is a manual review point, not an automatic legal conclusion.
- If the transferor or transferee is a corporation, authorised representative details should be manually reviewed.
- RORC / controller changes should be reviewed separately when the post-transfer holding changes control.
"""


def build_all() -> list[Path]:
    ensure_dirs()
    docs = {
        TEMPLATES["resolution"]: build_resolution_docx(),
        TEMPLATES["instrument"]: build_instrument_docx(),
        TEMPLATES["certificate"]: build_certificate_docx(),
        TEMPLATES["checklist"]: build_checklist_docx(),
    }
    paths: list[Path] = []
    for name, doc in docs.items():
        path = TEMPLATE_DIR / name
        doc.save(path)
        shutil.copy2(path, OUTPUT_DIR / name)
        paths.append(path)
    field_map = OUTPUT_DIR / "M03_field_map.md"
    field_map.write_text(field_map_markdown(), encoding="utf-8")
    paths.append(field_map)
    return paths


if __name__ == "__main__":
    for output in build_all():
        print(output)
