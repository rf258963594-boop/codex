from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from build_p2_m03_templates import (
    MUTED,
    OUTPUT_DIR,
    TEMPLATE_DIR,
    add_company_header,
    add_para,
    add_table,
    clause_para,
    clause_title,
    configure_doc,
)


ROOT = Path(__file__).resolve().parents[1]
P1_TEMPLATE_DIR = ROOT / "app" / "doc_templates" / "p1_standard_v3_part1"
P2_TEMPLATE_DIR = ROOT / "app" / "doc_templates" / "p2_standard_v1"

P1_TEMPLATE = "09_register_of_members_standard.docx"
M03_TEMPLATE = "M03_register_of_members_update_standard.docx"
M04_TEMPLATE = "M04_register_of_members_update_standard.docx"


def configure_register_doc(doc: Document, title: str, subject: str) -> None:
    configure_doc(doc, title, subject)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)


def add_register_detail_table(doc: Document) -> None:
    add_table(
        doc,
        ["Item", "Particulars"],
        [
            ["Registered office", "{{company.registered_office_address}}"],
            ["Register location", "{{register.location}}"],
            ["Class of shares", "{{register.share_class}}"],
            ["Currency", "{{register.currency}}"],
            ["Total shares in this record", "{{register.total_shares}}"],
            ["Total paid-up amount in this record", "{{register.total_paid}}"],
            ["Prepared / effective date", "{{register.prepared_date}}"],
            ["Transaction type", "{{register.transaction_type}}"],
        ],
        [3000, 9600],
    )


def add_register_entries_table(doc: Document) -> None:
    add_table(
        doc,
        [
            "Folio / Ref.",
            "Member",
            "ID / Reg. No.",
            "Address",
            "Shares",
            "Class",
            "Cert. No.",
            "Date / Remarks",
        ],
        [
            [
                "{% for entry in register.entries %}{{entry.folio_no}}{% endfor %}",
                "{{entry.member_name}}",
                "{{entry.id_number}}",
                "{{entry.address}}",
                "{{entry.shares}}",
                "{{entry.share_class}}",
                "{{entry.certificate_no}}",
                "{{entry.entry_date}}\n{{entry.paid_status_text}}\n{{entry.remarks}}",
            ]
        ],
        [850, 2000, 1300, 2200, 900, 1100, 1100, 3150],
    )


def add_preparation_table(doc: Document) -> None:
    add_table(
        doc,
        ["Prepared by", "Date", "Status"],
        [["{{provider.name}}", "{{register.prepared_date}}", "For internal register update / statutory records review"]],
        [3000, 2000, 7600],
    )


def build_register_template(path: Path, *, title: str, subject: str) -> None:
    doc = Document()
    configure_register_doc(doc, title, subject)
    add_company_header(doc, "{{register.title}}")
    add_para(doc, "{{register.subtitle}}", align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5, color=MUTED, after=14)

    clause_title(doc, "1. COMPANY AND REGISTER DETAILS")
    add_register_detail_table(doc)

    clause_title(doc, "2. MEMBER ENTRIES")
    clause_para(
        doc,
        "The following entries are prepared from the information provided for this document package. "
        "They should be checked against the signed source documents and final statutory records.",
    )
    add_register_entries_table(doc)

    clause_title(doc, "3. RECORD NOTE")
    clause_para(doc, "{{register.note}}")

    clause_title(doc, "4. PREPARATION")
    add_preparation_table(doc)
    add_para(
        doc,
        "This template is generated for document preparation and internal statutory records maintenance. "
        "It is not a confirmation that ACRA / BizFile filing, stamp duty review or final statutory register update has been completed.",
        style="Small Note",
        after=0,
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> None:
    build_register_template(
        P1_TEMPLATE_DIR / P1_TEMPLATE,
        title="P1 Register of Members",
        subject="Initial register of members after incorporation",
    )
    build_register_template(
        P2_TEMPLATE_DIR / M03_TEMPLATE,
        title="M03 Register of Members Update",
        subject="Register of members update record for share transfer",
    )
    build_register_template(
        P2_TEMPLATE_DIR / M04_TEMPLATE,
        title="M04 Register of Members Update",
        subject="Register of members update record for share allotment",
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(P2_TEMPLATE_DIR / M03_TEMPLATE, OUTPUT_DIR / M03_TEMPLATE)
    shutil.copy2(P2_TEMPLATE_DIR / M04_TEMPLATE, OUTPUT_DIR / M04_TEMPLATE)
    print(P1_TEMPLATE_DIR / P1_TEMPLATE)
    print(P2_TEMPLATE_DIR / M03_TEMPLATE)
    print(P2_TEMPLATE_DIR / M04_TEMPLATE)


if __name__ == "__main__":
    main()
