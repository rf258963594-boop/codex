from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from build_p2_m03_templates import (
    MUTED,
    OUTPUT_DIR,
    add_company_header,
    add_para,
    add_table,
    clause_para,
    clause_title,
    marker,
)
from build_register_of_members_templates import configure_register_doc


ROOT = Path(__file__).resolve().parents[1]
P1_TEMPLATE_DIR = ROOT / "app" / "doc_templates" / "p1_standard_v3_part1"
TEMPLATE_NAME = "11_statutory_registers_package_standard.docx"


def add_register_overview(doc: Document) -> None:
    clause_title(doc, "1. COMPANY AND REGISTER DETAILS")
    add_table(
        doc,
        ["Item", "Particulars"],
        [
            ["Registered office", "{{company.registered_office_address}}"],
            ["Register location", "{{m06.location}}"],
            ["Prepared / effective date", "{{m06.prepared_date}}"],
            ["Share class", "{{register.share_class}}"],
            ["Currency", "{{register.currency}}"],
            ["Package basis", "{{m06.subtitle}}"],
        ],
        [3200, 9800],
    )


def add_member_register(doc: Document) -> None:
    clause_title(doc, "2. REGISTER OF MEMBERS")
    clause_para(
        doc,
        "The following entries record the initial members and shareholdings prepared from the incorporation share allotment information.",
    )
    marker(doc, "[[REPEAT m06.members[]]]")
    add_table(
        doc,
        ["Folio", "Member", "ID / Reg. No.", "Address", "Shares", "Class", "Cert. No.", "Date / Remarks"],
        [
            [
                "{{member.folio_no}}",
                "{{member.member_name}}",
                "{{member.id_number}}",
                "{{member.address}}",
                "{{member.shares}}",
                "{{member.share_class}}",
                "{{member.certificate_no}}",
                "{{member.entry_date}}\n{{member.paid_status_text}}\n{{member.remarks}}",
            ]
        ],
        [800, 1900, 1300, 2600, 850, 1000, 900, 3650],
    )


def add_officer_register(doc: Document, title: str, expr: str, variable: str) -> None:
    clause_title(doc, title)
    marker(doc, f"[[REPEAT {expr}[]]]")
    add_table(
        doc,
        ["Name", "ID", "Nationality", "Address", "DOB", "Appointment", "Cessation", "Remarks"],
        [
            [
                f"{{{{{variable}.name}}}}",
                f"{{{{{variable}.id_type}}}} {{{{{variable}.id_number}}}}",
                f"{{{{{variable}.nationality}}}}",
                f"{{{{{variable}.address}}}}",
                f"{{{{{variable}.date_of_birth}}}}",
                f"{{{{{variable}.appointment_date}}}}",
                f"{{{{{variable}.cessation_date}}}}",
                f"{{{{{variable}.remarks}}}}",
            ]
        ],
        [1800, 1500, 1200, 3300, 950, 1400, 1200, 1650],
    )


def add_controller_register(doc: Document) -> None:
    clause_title(doc, "5. REGISTER OF REGISTRABLE CONTROLLERS")
    clause_para(
        doc,
        "This section is prepared from the shareholder/controller information available in the worksheet and should be checked against the Company's RORC records.",
    )
    marker(doc, "[[REPEAT m06.controllers[]]]")
    add_table(
        doc,
        ["Controller", "Type", "ID / Reg. No.", "Address", "Basis", "Entry date", "Status / Remarks"],
        [
            [
                "{{controller.name}}",
                "{{controller.controller_type}}",
                "{{controller.id_number}}",
                "{{controller.address}}",
                "{{controller.basis}}",
                "{{controller.entry_date}}",
                "{{controller.status}}\n{{controller.remarks}}",
            ]
        ],
        [1900, 900, 1400, 3100, 2800, 1300, 1600],
    )


def add_nominee_director_register(doc: Document) -> None:
    clause_title(doc, "6. REGISTER OF NOMINEE DIRECTORS")
    marker(doc, "[[REPEAT m06.nominee_directors[]]]")
    add_table(
        doc,
        ["Nominee director", "ID", "Address", "Appointment", "Nominator", "Nominator details", "Basis / Remarks"],
        [
            [
                "{{nominee.name}}",
                "{{nominee.id_number}}",
                "{{nominee.address}}",
                "{{nominee.appointment_date}}",
                "{{nominee.nominator_name}}",
                "{{nominee.nominator_id_number}}\n{{nominee.nominator_address}}",
                "{{nominee.basis}}\n{{nominee.remarks}}",
            ]
        ],
        [1800, 1200, 2600, 1300, 1800, 2300, 2000],
    )


def add_nominee_shareholder_register(doc: Document) -> None:
    add_para(doc, "[[IF m06.nominee_shareholders]]", size=1, color=MUTED, after=0)
    clause_title(doc, "7. REGISTER OF NOMINEE SHAREHOLDERS")
    marker(doc, "[[REPEAT m06.nominee_shareholders[]]]")
    add_table(
        doc,
        ["Nominee shareholder", "ID", "Nominator / Beneficial owner", "Nominator details", "Shares", "Entry date", "Remarks"],
        [
            [
                "{{nominee.nominee_name}}",
                "{{nominee.nominee_id_number}}",
                "{{nominee.nominator_name}}",
                "{{nominee.nominator_id_number}}\n{{nominee.nominator_address}}",
                "{{nominee.shares}} {{nominee.share_class}}",
                "{{nominee.entry_date}}",
                "{{nominee.remarks}}",
            ]
        ],
        [1900, 1300, 2300, 2600, 1350, 1200, 2350],
    )
    add_para(doc, "[[END IF]]", size=1, color=MUTED, after=0)


def add_certificate_register(doc: Document) -> None:
    clause_title(doc, "8. REGISTER OF SHARE CERTIFICATES")
    marker(doc, "[[REPEAT m06.share_certificates[]]]")
    add_table(
        doc,
        ["Cert. No.", "Holder", "Shares / Class", "Paid amount", "Issue date", "Status", "Remarks"],
        [
            [
                "{{certificate.certificate_no}}",
                "{{certificate.holder_name}}",
                "{{certificate.shares}} {{certificate.share_class}}",
                "{{certificate.paid_amount}}",
                "{{certificate.issue_date}}",
                "{{certificate.status}}",
                "{{certificate.remarks}}",
            ]
        ],
        [1000, 2400, 1600, 1600, 1300, 1100, 4000],
    )


def add_review_notes(doc: Document) -> None:
    clause_title(doc, "9. PREPARATION NOTES")
    marker(doc, "[[REPEAT m06.review_notes[]]]")
    add_table(
        doc,
        ["Review item", "Note"],
        [["{{note.item}}", "{{note.note}}"]],
        [2600, 10400],
    )
    add_para(
        doc,
        "This package is a working statutory register pack for document preparation and internal records maintenance. "
        "It is not a confirmation that any ACRA / BizFile filing has been completed.",
        style="Small Note",
        after=0,
    )


def build_template(path: Path) -> None:
    doc = Document()
    configure_register_doc(doc, "M06 Statutory Registers Package", "Initial statutory registers package")
    add_company_header(doc, "{{m06.title}}")
    add_para(doc, "{{m06.subtitle}}", align=WD_ALIGN_PARAGRAPH.CENTER, size=10.2, color=MUTED, after=12)

    add_register_overview(doc)
    add_member_register(doc)
    add_officer_register(doc, "3. REGISTER OF DIRECTORS", "m06.directors", "officer")
    add_officer_register(doc, "4. REGISTER OF SECRETARIES", "m06.secretaries", "officer")
    add_controller_register(doc)
    add_nominee_director_register(doc)
    add_nominee_shareholder_register(doc)
    add_certificate_register(doc)
    add_review_notes(doc)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> None:
    path = P1_TEMPLATE_DIR / TEMPLATE_NAME
    build_template(path)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(path, OUTPUT_DIR / TEMPLATE_NAME)
    print(path)


if __name__ == "__main__":
    main()
