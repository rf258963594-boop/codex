from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_p2_m03_templates import (
    ACCENT,
    INK,
    LIGHT_FILL,
    MUTED,
    RED,
    TEMPLATE_DIR,
    OUTPUT_DIR,
    add_company_header,
    add_para,
    add_table,
    clause_para,
    clause_title,
    add_landscape_certificate_box,
    configure_certificate_doc,
    configure_doc,
    marker,
    set_cell_margins,
    set_cell_text,
    set_row_cant_split,
    set_run_font,
    set_table_borders,
    set_table_width,
)


TEMPLATES = {
    "authority": "M04_s161_members_authority_standard.docx",
    "resolution": "M04_allotment_directors_resolution_standard.docx",
    "application": "M04_share_application_standard.docx",
    "certificate": "M04_share_certificate_standard.docx",
    "form24": "M04_return_of_allotment_form24_standard.docx",
    "checklist": "M04_register_update_checklist_standard.docx",
}


def ensure_dirs() -> None:
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for directory in (TEMPLATE_DIR, OUTPUT_DIR):
        for name in TEMPLATES.values():
            path = directory / name
            if path.is_file():
                path.unlink()


def add_allotment_table(doc: Document) -> None:
    marker(doc, "[[REPEAT m04.allotments[]]]")
    add_table(
        doc,
        ["Allottee", "Class", "Shares", "Issued capital", "Paid-up capital", "Allotment date"],
        [[
            "{{allotment.allottee_name}}",
            "{{allotment.share_class}}",
            "{{allotment.shares_allotted}}",
            "{{allotment.issued_share_capital_text}}",
            "{{allotment.paid_up_share_capital_text}}",
            "{{allotment.allotment_date}}",
        ]],
        [2050, 950, 1200, 1650, 1650, 1860],
        body_alignments=[
            None,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
        ],
    )


def add_form24_allottee_table(doc: Document) -> None:
    marker(doc, "[[REPEAT m04.allotments[]]]")
    add_table(
        doc,
        ["Name / ID", "Address", "Nationality / Country", "Shares allotted and consideration", "Date"],
        [[
            "{{allotment.allottee_form24_name_id}}",
            "{{allotment.allottee_address}}",
            "{{allotment.nationality}}",
            "{{allotment.form24_allotment_text}}",
            "{{allotment.allotment_date}}",
        ]],
        [1600, 1950, 1900, 2310, 1600],
        body_alignments=[
            None,
            None,
            WD_ALIGN_PARAGRAPH.CENTER,
            None,
            WD_ALIGN_PARAGRAPH.CENTER,
        ],
    )


def add_signature_rows(doc: Document, expr: str, heading: str) -> None:
    add_para(doc, heading, bold=True, size=10.5, after=4, keep_with_next=True)
    marker(doc, f"[[REPEAT {expr}[]]]")
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [4680, 4680])
    set_table_borders(table, "FFFFFF")
    set_cell_margins(table, top=40, bottom=60, start=0, end=160)
    set_row_cant_split(table.rows[0])
    set_cell_text(table.cell(0, 0), "{{sigrow.left_block}}", size=10.3)
    set_cell_text(table.cell(0, 1), "{{sigrow.right_block}}", size=10.3)


def add_checklist_table(doc: Document) -> None:
    marker(doc, "[[REPEAT m04.checklist_items[]]]")
    add_table(
        doc,
        ["Item", "Status", "Review note"],
        [["{{check.item}}", "{{check.status}}", "{{check.note}}"]],
        [2350, 1450, 5560],
    )


def build_authority_docx() -> Document:
    doc = Document()
    configure_doc(doc, "M04 Section 161 Members' Authority", "P2 share allotment members' authority")
    add_company_header(doc, "MEMBERS' WRITTEN RESOLUTION / SECTION 161 AUTHORITY")
    clause_para(doc, "The undersigned member(s) of the Company hereby pass the following ordinary resolution in writing, subject to the Constitution of the Company and the Companies Act.")
    clause_title(doc, "AUTHORITY TO ISSUE SHARES")
    clause_para(doc, "That pursuant to Section 161 of the Companies Act, the directors of the Company be and are hereby authorised to allot and issue shares in the Company at any time and upon such terms and conditions and for such purposes as the directors may, in their absolute discretion, deem fit.")
    clause_para(doc, "That the directors be authorised to approve and give effect to the following proposed allotment(s), and to do all acts and things necessary or expedient in connection with such allotment(s):")
    add_allotment_table(doc)
    clause_para(doc, "That any director and/or the company secretary and/or {{provider.name}} be authorised to prepare, sign, lodge, amend and submit all documents, notices and electronic transactions necessary to give effect to the above resolutions.")
    add_para(doc, "Dated this {{signature.day_ordinal}} day of {{signature.month_year}}", style="Signature Block", before=10, after=10)
    add_signature_rows(doc, "m04.member_signature_rows", "MEMBER(S) / AUTHORISED SIGNATORY")
    return doc


def build_resolution_docx() -> Document:
    doc = Document()
    configure_doc(doc, "M04 Allotment Directors' Resolution", "P2 share allotment directors' resolution")
    add_company_header(doc, "DIRECTORS' RESOLUTIONS IN WRITING")
    clause_para(doc, "The undersigned, being the director(s) of the Company, hereby pass the following resolutions in writing pursuant to the Constitution of the Company.")
    clause_title(doc, "ALLOTMENT OF SHARES")
    clause_para(doc, "That pursuant to the authority given by the member(s) of the Company, the following allotment(s) of shares be and are hereby approved:")
    add_allotment_table(doc)
    clause_para(doc, "That the relevant share certificate(s) be issued, the Register of Members be updated, and the required statutory returns and electronic lodgements be prepared and submitted to ACRA as required.")
    clause_para(doc, "That any director of the Company and/or the company secretary and/or {{provider.name}} be authorised to prepare, sign, lodge, amend and submit all documents and electronic transactions necessary to give effect to these resolutions.")
    add_para(doc, "Dated this {{signature.day_ordinal}} day of {{signature.month_year}}", style="Signature Block", before=10, after=10)
    add_signature_rows(doc, "m04.director_signature_rows", "DIRECTOR(S)")
    return doc


def build_application_docx() -> Document:
    doc = Document()
    configure_doc(doc, "M04 Share Application", "P2 share allotment application")
    add_para(doc, "Date: {{allotment.allotment_date}}", align=WD_ALIGN_PARAGRAPH.RIGHT, size=10.5, after=14)
    add_para(doc, "The Board of Directors", size=10.5, after=2)
    add_para(doc, "{{company.company_name}}", bold=True, size=10.5, after=2)
    add_para(doc, "Company Registration No. {{company.uen}}", size=10.5, after=2)
    add_para(doc, "{{company.registered_office_address}}", size=10.5, after=14)
    add_para(doc, "Dear Sirs", size=10.5, after=10)
    add_para(doc, "APPLICATION FOR SHARES", bold=True, size=12, after=10)
    clause_para(doc, "I/We hereby apply for an allotment of shares in the above named Company and provide the following information:")
    add_table(
        doc,
        ["Field", "Details"],
        [
            ["Applicant", "{{allotment.allottee_name}}"],
            ["ID / Registration No.", "{{allotment.allottee_id_number}}"],
            ["Address", "{{allotment.allottee_address}}"],
            ["No. of shares", "{{allotment.shares_allotted}}"],
            ["Class of shares", "{{allotment.share_class}}"],
            ["Amount paid per share", "{{allotment.amount_paid_per_share_text}}"],
            ["Total paid-up capital", "{{allotment.paid_up_share_capital_text}}"],
            ["Consideration", "{{allotment.consideration_text}}"],
        ],
        [2450, 6910],
    )
    clause_para(doc, "I/We agree to accept the shares which may be allotted to me/us subject to the Constitution of the Company and authorise the Company to enter my/our name in the Register of Members in respect of the shares so allotted.")
    clause_para(doc, "This application is irrevocable unless otherwise agreed in writing by the Company.")
    add_para(doc, "Yours faithfully,", style="Signature Block", before=12, after=26)
    add_para(doc, ".....................................................", style="Signature Block", after=2)
    add_para(doc, "{{allotment.allottee_name}}", style="Signature Block", after=0)
    return doc


def add_certificate_box(doc: Document) -> None:
    add_p1_style_certificate_box(doc)


def set_table_borders_style(table, *, color: str = "9F1212", val: str = "double", size: str = "8") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), val)
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_times_run(run, size: float, *, bold: bool = False, italic: bool = False, color: RGBColor = INK) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def clear_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        paragraph.clear()


def add_cert_para(cell, text: str, size: float, *, bold: bool = False, italic: bool = False, color: RGBColor = INK, after: float = 4, align=WD_ALIGN_PARAGRAPH.CENTER):
    paragraph = cell.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_times_run(run, size, bold=bold, italic=italic, color=color)
    return paragraph


def set_cert_cell(cell, lines: list[tuple[str, float, bool, bool, RGBColor]], *, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    clear_cell(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    for idx, (text, size, bold, italic, color) in enumerate(lines):
        paragraph = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        paragraph.alignment = align
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(text)
        set_times_run(run, size, bold=bold, italic=italic, color=color)


def add_p1_style_certificate_box(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [14400])
    set_table_borders_style(table, color="9F1212", val="double", size="10")
    set_cell_margins(table, top=330, bottom=220, start=330, end=330)
    table.rows[0].height = Inches(7.35)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    clear_cell(cell)
    first = cell.paragraphs[0]
    first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first.paragraph_format.space_after = Pt(0)
    run = first.add_run("SHARE CERTIFICATE")
    set_times_run(run, 22, bold=True, color=RED)

    underline = cell.add_paragraph()
    underline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    underline.paragraph_format.space_after = Pt(8)
    run = underline.add_run("=" * 42)
    set_times_run(run, 8, color=RED)

    top = cell.add_table(rows=1, cols=3)
    set_table_width(top, [3000, 8400, 3000])
    set_table_borders(top, "FFFFFF")
    set_cell_margins(top, top=0, bottom=0, start=80, end=80)
    set_cert_cell(
        top.cell(0, 0),
        [
            ("Certificate No.", 12.5, False, False, INK),
            ("", 4, False, False, INK),
            ("{{certificate.certificate_no}}", 12.5, False, False, INK),
        ],
    )
    set_cert_cell(
        top.cell(0, 1),
        [
            ("{{company.company_name}}", 14, False, False, INK),
            ("(Incorporated in the Republic of Singapore)", 12, False, False, INK),
            ("Co. Reg. No.: {{company.uen}}", 12, False, False, INK),
            ("Registered Office: {{company.registered_office_address}}", 11.5, False, False, INK),
        ],
    )
    set_cert_cell(
        top.cell(0, 2),
        [
            ("No. of Shares", 12.5, False, False, INK),
            ("", 4, False, False, INK),
            ("{{certificate.shares_text}}", 12.5, False, False, INK),
        ],
    )

    spacer = cell.add_paragraph()
    spacer.paragraph_format.space_after = Pt(18)
    add_cert_para(cell, "THIS IS TO CERTIFY THAT", 15, bold=True, after=1)
    add_cert_para(cell, "{{certificate.holder_name}}", 15, bold=True, after=1)
    add_cert_para(cell, "of", 12, italic=True, after=1)
    add_cert_para(cell, "{{certificate.holder_address}}", 11.5, italic=True, after=1)
    add_cert_para(cell, "is/are the registered holder(s) of", 11.5, italic=True, after=1)
    add_cert_para(cell, "{{certificate.shares_in_words}}", 14, bold=True, italic=True, after=2)
    add_cert_para(
        cell,
        "{{certificate.paid_status_line}} in the Company subject to the Constitution of the Company.",
        11.5,
        italic=True,
        after=5,
    )
    add_cert_para(cell, "Dated this {{signature.day_ordinal}} day of {{signature.month_year}}", 11.5, italic=True, after=6)
    add_cert_para(cell, "Executed by the Company on the date stated above in accordance with the Companies Act.", 11.5, italic=True, after=22)

    sig_table = cell.add_table(rows=1, cols=3)
    set_table_width(sig_table, [4300, 5700, 4300])
    set_table_borders(sig_table, "FFFFFF")
    set_cell_margins(sig_table, top=0, bottom=0, start=80, end=80)
    set_cert_cell(
        sig_table.cell(0, 0),
        [
            ("______________________________", 11, False, False, INK),
            ("Director", 11.5, False, False, INK),
            ("Name: __________________________", 10.5, False, False, INK),
        ],
    )
    set_cert_cell(
        sig_table.cell(0, 1),
        [
            ("For and on behalf of", 11.5, False, False, INK),
            ("{{company.company_name}}", 11.5, False, False, INK),
        ],
    )
    set_cert_cell(
        sig_table.cell(0, 2),
        [
            ("______________________________", 11, False, False, INK),
            ("Company Secretary", 11.5, False, False, INK),
        ],
    )

    note = cell.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(24)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("Note: Transfer of any portion of the Shares comprised in this Certificate cannot be registered unless accompanied by this Certificate.")
    set_times_run(run, 8, color=INK)


def build_certificate_docx() -> Document:
    doc = Document()
    configure_certificate_doc(doc, "M04 Share Certificate", "P2 share certificate after allotment")
    add_certificate_box(doc)
    return doc


def build_form24_docx() -> Document:
    doc = Document()
    configure_doc(doc, "M04 Return of Allotment / Form 24", "P2 return of allotment")
    add_company_header(doc, "RETURN OF ALLOTMENT OF SHARES / FORM 24")
    clause_para(doc, "The issue of the shares referred to in this return was made pursuant to the member(s)' resolution / authority dated {{m04.authority_date}}.")
    clause_para(doc, "The shares referred to in this return were allotted, or are deemed to have been allotted, to the allottee(s) on the date(s) indicated below.")
    clause_title(doc, "1. PAYABLE IN CASH")
    add_table(
        doc,
        ["Class of shares", "Number of shares", "Amount paid on each share", "Amount due and payable on each share"],
        [["{{m04.share_class_summary}}", "{{m04.total_shares_allotted}}", "{{m04.amount_paid_per_share_summary}}", "{{m04.amount_due_per_share_summary}}"]],
        [2300, 1900, 2450, 2710],
    )
    clause_title(doc, "2. LIST OF ALLOTTEES")
    add_form24_allottee_table(doc)
    clause_title(doc, "3. SHARE CAPITAL POSITION AFTER ALLOTMENT")
    add_table(
        doc,
        ["Share class", "Issued share capital", "Paid-up share capital", "Review note"],
        [["{{m04.share_class_summary}}", "{{m04.post_issued_share_capital_text}}", "{{m04.post_paid_up_capital_text}}", "{{m04.capital_review_note}}"]],
        [1800, 2200, 2200, 3160],
    )
    clause_title(doc, "CERTIFICATE")
    clause_para(doc, "I/We certify that the above particulars are prepared for lodgement and records update in relation to the allotment of shares described above. The final ACRA / BizFile entry and supporting registers should be checked before filing completion.")
    add_para(doc, "Dated this {{signature.day_ordinal}} day of {{signature.month_year}}", style="Signature Block", before=10, after=10)
    add_signature_rows(doc, "m04.director_signature_rows", "DIRECTOR(S) / AUTHORISED SIGNATORY")
    return doc


def build_checklist_docx() -> Document:
    doc = Document()
    configure_doc(doc, "M04 Register Update Checklist", "P2 internal share allotment checklist")
    add_company_header(doc, "INTERNAL ALLOTMENT AND REGISTER UPDATE CHECKLIST")
    clause_para(doc, "This checklist is for internal review before releasing the allotment package or updating the Company's statutory records.")
    add_checklist_table(doc)
    clause_title(doc, "ALLOTMENT SUMMARY")
    add_allotment_table(doc)
    add_para(doc, "Prepared by: ______________________________", style="Signature Block", before=12, after=6)
    add_para(doc, "Reviewed by: ______________________________", style="Signature Block", after=6)
    return doc


def field_map_markdown() -> str:
    return """# M04 Share Allotment Package - Field Map

M04 is the P2 package for share allotment / capital increase. Form 24 belongs here, not in ordinary share transfer M03.

## Generated files

| File | Trigger | Signer | Notes |
|---|---|---|---|
| Section 161 / Members' Authority | At least one active allotment row | Member(s) / authorised signatory | Gives directors authority to allot and issue shares. |
| Allotment Directors' Resolution | At least one active allotment row | All director signers | Approves the specific allotment(s). |
| Share Application | One per allottee | Allottee / authorised signatory | Applicant agrees to accept new shares. |
| Share Certificate | One per allottee unless certificate is disabled | Director line + Company Secretary | Certificate number and paid-up status require review. |
| Return of Allotment / Form 24 | At least one active allotment row | Director(s) / authorised signatory | Prepared as signing / lodgement support; ACRA entry still needs review. |
| Register Update Checklist | At least one active allotment row | Internal review | Internal reminder for register, ACRA, certificate and capital checks. |

## Primary input fields

| Field | Source | Used in |
|---|---|---|
| `company.company_name` / `company.uen` / `company.registered_office_address` | P2 one-page sheet | All M04 documents |
| `company.default_document_date` | P2 one-page sheet | Default authority / allotment date |
| `company.director_signer_names` | P2 one-page sheet | Directors' resolution and Form 24 signatures |
| `company.member_signer_names` | P2 one-page sheet | Members' authority signatures |
| `share_allotments[].allottee_name` | Allotment table | All M04 documents |
| `share_allotments[].shares_allotted` | Allotment table | All M04 documents |
| `share_allotments[].issued_share_capital` | Allotment table | Form 24 and certificate review |
| `share_allotments[].total_paid` / `paid_up_share_capital` | Allotment table | Form 24, application and certificate |
| `share_allotments[].allotment_date` / `authority_date` | Allotment table | Resolutions, Form 24 and application |
| `share_allotments[].certificate_no` | Allotment table | Share certificate |

## Guardrails

- If capital totals are blank, M04 can still generate but the Form 24 capital summary is marked for manual review.
- Section 161 / members' authority is generated as the standard first-version route.
- ACRA / BizFile lodgement is not automatically completed by this package.
"""


def build_all() -> list[Path]:
    ensure_dirs()
    docs = {
        TEMPLATES["authority"]: build_authority_docx(),
        TEMPLATES["resolution"]: build_resolution_docx(),
        TEMPLATES["application"]: build_application_docx(),
        TEMPLATES["certificate"]: build_certificate_docx(),
        TEMPLATES["form24"]: build_form24_docx(),
        TEMPLATES["checklist"]: build_checklist_docx(),
    }
    paths: list[Path] = []
    for name, doc in docs.items():
        path = TEMPLATE_DIR / name
        doc.save(path)
        shutil.copy2(path, OUTPUT_DIR / name)
        paths.append(path)
    field_map = OUTPUT_DIR / "M04_field_map.md"
    field_map.write_text(field_map_markdown(), encoding="utf-8")
    paths.append(field_map)
    return paths


if __name__ == "__main__":
    for output in build_all():
        print(output)
