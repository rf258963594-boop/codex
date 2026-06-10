from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
P1_DIR = ROOT / "app" / "doc_templates" / "p1_standard_v3_part1"


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    style = paragraph.style
    first = paragraph.runs[0] if paragraph.runs else None
    paragraph.clear()
    paragraph.style = style
    run = paragraph.add_run(text)
    if first is not None:
        run.bold = first.bold
        run.italic = first.italic
        run.underline = first.underline
        if first.font.name:
            run.font.name = first.font.name
        if first.font.size:
            run.font.size = first.font.size
        if first.font.color and first.font.color.rgb:
            run.font.color.rgb = first.font.color.rgb


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._p
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def insert_paragraph_before(paragraph: Paragraph, text: str) -> Paragraph:
    new_element = OxmlElement("w:p")
    paragraph._p.addprevious(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    new_paragraph.add_run(text)
    return new_paragraph


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    new_paragraph.add_run(text)
    return new_paragraph


def find_paragraph(doc: Document, startswith: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(startswith):
            return paragraph
    return None


def first_paragraph(doc: Document, startswith: str) -> Paragraph:
    paragraph = find_paragraph(doc, startswith)
    if paragraph is not None:
        return paragraph
    raise ValueError(f"Cannot find paragraph starting with: {startswith}")


def update_first_directors_resolution() -> None:
    path = P1_DIR / "01_first_directors_resolution_standard.docx"
    doc = Document(path)
    paragraph = find_paragraph(doc, "The Certificate numbered")
    if paragraph is not None:
        set_paragraph_text(
            paragraph,
            "The electronic certificate confirming Incorporation of Company on {{company.incorporation_date}} "
            "together with a copy of the Constitution of the Company was noted.",
        )
    doc.save(path)


def update_form24() -> None:
    path = P1_DIR / "07_return_of_allotment_form24_standard.docx"
    doc = Document(path)
    certify = find_paragraph(doc, "I hereby certify")
    if certify is not None:
        set_paragraph_text(certify, "I hereby certify, in relation to the abovenamed company, that:")
    para_a = find_paragraph(doc, "(a) the company has more than 500 members")
    if para_a is not None:
        set_paragraph_text(para_a, "(a) the company keeps its principal shares register at {{company.register_location}};")
    para_b = find_paragraph(doc, "(b) the company keeps its principal shares register")
    if para_b is not None:
        set_paragraph_text(para_b, "(b) the shares referred to in this return were allotted for cash; and")
    para_c = find_paragraph(doc, "(c) the company provides reasonable accommodation")
    if para_c is not None:
        set_paragraph_text(para_c, "(c) the particulars stated in this return are true and correct to the best of my knowledge.")
    for prefix in [
        "(d) the shares referred to in this return were allotted for cash",
        "(e) the shares referred to in this return were allotted for a consideration other than in cash",
    ]:
        paragraph = find_paragraph(doc, prefix)
        if paragraph is not None:
            remove_paragraph(paragraph)
    doc.save(path)


def marker_exists(doc: Document, marker: str) -> bool:
    return any(paragraph.text.strip() == marker for paragraph in doc.paragraphs)


def has_page_break(paragraph: Paragraph) -> bool:
    for run in paragraph.runs:
        for br in run._r.iter(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return True
    return False


def remove_page_break_before_corporate_block(doc: Document) -> None:
    paragraphs = list(doc.paragraphs)
    for idx, paragraph in enumerate(paragraphs[:-1]):
        if paragraph.text.strip() or not has_page_break(paragraph):
            continue
        next_text = ""
        for next_paragraph in paragraphs[idx + 1 :]:
            next_text = next_paragraph.text.strip()
            if next_text:
                break
        if next_text == "[[IF shareholder.is_corporate]]" or next_text.startswith("2) If your reply is yes and the person is a legal entity"):
            remove_paragraph(paragraph)


def update_rorc() -> None:
    path = P1_DIR / "08_rorc_notice_controller_standard.docx"
    doc = Document(path)
    if not marker_exists(doc, "[[IF shareholder.is_individual]]"):
        individual_start = first_paragraph(doc, "1) If your reply is yes and you are an individual")
        individual_end = first_paragraph(doc, "(vii) The date on which you became an individual controller")
        insert_paragraph_before(individual_start, "[[IF shareholder.is_individual]]")
        insert_paragraph_after(individual_end, "[[END IF]]")
    if not marker_exists(doc, "[[IF shareholder.is_corporate]]"):
        corporate_start = first_paragraph(doc, "2) If your reply is yes and the person is a legal entity")
        corporate_end = first_paragraph(doc, "(viii) The date on which the person became a corporate controller")
        insert_paragraph_before(corporate_start, "[[IF shareholder.is_corporate]]")
        insert_paragraph_after(corporate_end, "[[END IF]]")
    remove_page_break_before_corporate_block(doc)
    doc.save(path)


def main() -> None:
    update_first_directors_resolution()
    update_form24()
    update_rorc()
    print("P1 template logic updated.")


if __name__ == "__main__":
    main()
