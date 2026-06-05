from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
TARGET_DIRS = [
    ROOT / "app" / "doc_templates" / "p1_standard_v3_part1",
    ROOT / "outputs" / "P1_standard_templates_v3_part1",
]

REPLACEMENTS = {
    "01_first_directors_resolution_standard.docx": {
        "That the signatories to the Constitution be registered as members in respect of the shares for which they subscribed in full, in cash at {{company.share_currency}} {{company.share_par_value}} /- per share, namely: -":
            "That the signatories to the Constitution be registered as members in respect of the shares for which they subscribed in cash, namely: -",
        "No. of shares of {{company.share_currency}} {{company.share_par_value}} /- each":
            "No. of shares / issued and paid-up share capital",
    },
    "04_share_certificate_standard.docx": {
        "{{shareholder.share_class}} Share(s) fully paid in the Company subject to the Memorandum & Articles of Association of the Company.":
            "{{shareholder.share_class}} Share(s) {{shareholder.paid_status_text}} in the Company subject to the Memorandum & Articles of Association of the Company.",
        "fully paid in the Company subject to the Memorandum & Articles of Association of the Company.":
            "{{shareholder.paid_status_text}} in the Company subject to the Memorandum & Articles of Association of the Company.",
    },
    "07_return_of_allotment_form24_standard.docx": {
        "{{company.share_currency}} {{company.share_par_value}}":
            "{{company.share_currency}} {{company.amount_paid_per_share}}",
        "Amount due and payable on each share | - |  | ":
            "Amount due and payable on each share | {{company.share_currency}} {{company.amount_due_per_share}} |  | ",
        "{% for shareholder in shareholders %}{{shareholder.shares}} {{shareholder.share_class}} Shares\n\n{% endfor %}\nAllotted on the date of Incorporation.":
            "{% for shareholder in shareholders %}{{shareholder.form24_allotment_text}}\n\n{% endfor %}\nAllotted on the date of Incorporation.",
        "{% for shareholder in shareholders %}{{shareholder.shareholder_name}}\n\n{{shareholder.shareholder_address}}\n\nIDENTIFICATION NO.: {{shareholder.id_number}}\n\n{{shareholder.nationality}} | {{shareholder.shares}} {{shareholder.share_class}} Shares\n\nAllotted on the date of Incorporation.{% endfor %}":
            "{% for shareholder in shareholders %}{{shareholder.shareholder_name}}\n\n{{shareholder.shareholder_address}}\n\nIDENTIFICATION NO.: {{shareholder.id_number}}\n\n{{shareholder.nationality}} | {{shareholder.form24_allotment_text}}\n\nAllotted on the date of Incorporation.{% endfor %}",
        "Issued Share Capital | {{company.share_currency}} {{company.paid_up_capital}} | - | -":
            "Issued Share Capital | {{company.share_currency}} {{company.issued_share_capital}} | - | -",
    },
}


def set_text(paragraph, text: str) -> None:
    if paragraph.runs:
        first = paragraph.runs[0]
        for run in paragraph.runs[1:]:
            run._element.getparent().remove(run._element)
        first.text = text
    else:
        paragraph.add_run(text)


def replace_in_paragraph(paragraph, replacements: dict[str, str]) -> int:
    original = paragraph.text
    updated = original
    for old, new in replacements.items():
        if old in updated:
            updated = updated.replace(old, new)
    if updated != original:
        set_text(paragraph, updated)
        return 1
    return 0


def replace_in_docx(path: Path, replacements: dict[str, str]) -> int:
    doc = Document(path)
    count = 0
    for paragraph in doc.paragraphs:
        count += replace_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            for old, new in replacements.items():
                if old == row_text:
                    values = new.split(" | ")
                    for idx, value in enumerate(values[: len(row.cells)]):
                        set_text(row.cells[idx].paragraphs[0], value)
                    count += 1
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    count += replace_in_paragraph(paragraph, replacements)
    doc.save(path)
    return count


def main() -> None:
    results = {}
    for directory in TARGET_DIRS:
        for filename, replacements in REPLACEMENTS.items():
            path = directory / filename
            if path.exists():
                results[str(path)] = replace_in_docx(path, replacements)
    for path, count in results.items():
        print(f"{count}\t{path}")


if __name__ == "__main__":
    main()
