from __future__ import annotations

import json
import re
import sys
from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
APP_DIR = ROOT / "app"
OUT_DIR = ROOT / "outputs"
GENERATED_DIR = APP_DIR / "generated"

sys.path.insert(0, str(ROOT / "tools"))
sys.path.insert(0, str(APP_DIR))

from build_maintenance_one_page_template_v6 import build_workbook  # noqa: E402
from db import connect  # noqa: E402
from doc_generator import generate_p2_m01_pdf_package  # noqa: E402
from doc_render import pdf_info  # noqa: E402
from excel_parser import parse_excel  # noqa: E402
from rules import suggest_maintenance  # noqa: E402


JOB_CODE = "M01-V7-ONEPAGE-STRESS-V3"
WORKBOOK_PATH = OUT_DIR / "M01_v7_onepage_stress_input.xlsx"
AUDIT_PATH = OUT_DIR / "M01_v7_onepage_stress_audit_report.json"
DELIVERY_PATH = OUT_DIR / "M01_v7_onepage_stress_delivery_package.zip"


COMMON_PEOPLE = [
    {
        "display_name": "TEST NOMINEE DIRECTOR A",
        "default_role": "Nominee Director",
        "id_type": "NRIC",
        "id_number": "S7654321Z",
        "nationality": "Singaporean",
        "residential_address": "20 CECIL STREET, #12-02 PLUS, SINGAPORE 049705",
        "email": "test.nominee.director@example.com",
        "phone": "+65 9888 0001",
        "is_local_resident_director": 1,
        "notes": "V7 one-page stress test simulated common person.",
    },
    {
        "display_name": "TEST COMPANY SECRETARY A",
        "default_role": "Secretary",
        "id_type": "NRIC",
        "id_number": "S8765432Y",
        "nationality": "Singaporean",
        "residential_address": "111 NORTH BRIDGE ROAD, #29-06A PENINSULA PLAZA, SINGAPORE 179098",
        "email": "test.secretary@example.com",
        "phone": "+65 9888 0002",
        "is_local_resident_director": 0,
        "notes": "V7 one-page stress test simulated common person.",
    },
]


def clean(value) -> str:
    return "" if value is None else str(value).strip()


def header_key(value) -> str:
    raw = clean(value)
    if not raw:
        return ""
    for part in reversed(re.split(r"[\r\n]+", raw)):
        part = part.strip().strip("()[] ")
        if re.fullmatch(r"[A-Za-z][A-Za-z0-9_]*", part):
            return part
    compact = re.sub(r"[^A-Za-z0-9_]", "", raw)
    return compact if re.fullmatch(r"[A-Za-z][A-Za-z0-9_]*", compact) else raw


def upsert_common_people() -> list[dict]:
    with connect() as conn:
        for person in COMMON_PEOPLE:
            existing = conn.execute(
                "SELECT id FROM common_people WHERE display_name = ?",
                (person["display_name"],),
            ).fetchone()
            values = (
                person["default_role"],
                person["id_type"],
                person["id_number"],
                person["nationality"],
                person["residential_address"],
                person["email"],
                person["phone"],
                person["is_local_resident_director"],
                1,
                person["notes"],
                person["display_name"],
            )
            if existing:
                conn.execute(
                    """
                    UPDATE common_people
                    SET default_role = ?, id_type = ?, id_number = ?, nationality = ?,
                        residential_address = ?, email = ?, phone = ?,
                        is_local_resident_director = ?, active = ?, notes = ?
                    WHERE display_name = ?
                    """,
                    values,
                )
            else:
                conn.execute(
                    """
                    INSERT INTO common_people
                    (default_role, id_type, id_number, nationality, residential_address,
                     email, phone, is_local_resident_director, active, notes, display_name)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    values,
                )
        rows = conn.execute(
            "SELECT * FROM common_people WHERE display_name IN (?, ?) ORDER BY display_name",
            tuple(person["display_name"] for person in COMMON_PEOPLE),
        ).fetchall()
        return [dict(row) for row in rows]


def load_common_people() -> dict[str, dict]:
    with connect() as conn:
        rows = conn.execute("SELECT * FROM common_people WHERE active = 1").fetchall()
        return {row["display_name"]: dict(row) for row in rows}


def set_vertical_values(ws, values: dict[str, object]) -> None:
    for row in range(1, ws.max_row + 1):
        key = clean(ws.cell(row, 3).value)
        if key in values:
            ws.cell(row, 4).value = values[key]


def set_table(ws, required_keys: set[str], records: list[dict[str, object]]) -> None:
    header_row = None
    headers: list[str] = []
    for row in range(1, ws.max_row + 1):
        candidate = [header_key(ws.cell(row, col).value) for col in range(1, ws.max_column + 1)]
        if required_keys.issubset({key for key in candidate if key}):
            header_row = row
            headers = candidate
            break
    if not header_row:
        raise ValueError(f"Could not locate table with keys: {sorted(required_keys)}")
    for offset, record in enumerate(records, start=1):
        row = header_row + offset
        for col, key in enumerate(headers, start=1):
            if key in record:
                ws.cell(row, col).value = record[key]


def build_stress_workbook() -> Path:
    wb = build_workbook(False)
    main = wb.worksheets[0]
    annual = wb.worksheets[1]

    old_address = "8 TEMASEK BOULEVARD, #35-01 SUNTEC TOWER THREE, SINGAPORE 038988"
    new_address = (
        "71 ROBINSON ROAD, #14-01, SINGAPORE 068895, "
        "CARE OF RSIN GROUP CORPORATE SECRETARY DEPARTMENT"
    )
    set_vertical_values(
        main,
        {
            "company_name": "V7 ONE PAGE STRESS TEST PTE. LTD.",
            "uen": "202688888M",
            "registered_office_address": old_address,
            "default_document_date": "03/06/2026",
            "director_signer_names": "ALICE TAN, BOB LIM",
            "client_signatory_name": "ALICE TAN",
            "business_order_id": JOB_CODE,
            "source_type": "AI + Manual stress test",
            "source_file_id": "One-page v7 stress data",
            "prepared_by": "Codex v7 stress test",
            "change_registered_office_required": "Yes",
            "new_registered_office_address": new_address,
            "change_business_activity_required": "Yes",
            "new_primary_ssic": "62011",
            "new_primary_activity": "DEVELOPMENT OF SOFTWARE AND APPLICATIONS EXCEPT GAMES",
            "new_secondary_ssic": "63119",
            "new_secondary_activity": "DATA PROCESSING, HOSTING AND RELATED ACTIVITIES",
            "change_fye_required": "Yes",
            "old_fye": "31/12/2025",
            "new_fye": "30/06/2026",
            "next_accounts_period_start": "01/01/2026",
            "next_accounts_period_end": "30/06/2026",
            "transfer_in_required": "No",
            "generate_resignation_letter": "No",
            "annual_review_required": "Yes",
            "fye_date": "31/12/2025",
            "agm_date": "30/06/2026",
            "annual_review_remarks": "Stress test annual review enabled. Formal annual package is currently preview-only.",
            "notes": "One-page v7 stress test with comma-separated director signers, common people, resignation letter flags and annual review.",
        },
    )

    set_table(
        main,
        {"generate", "action_type", "target_name"},
        [
            {
                "generate": "Yes",
                "action_type": "resign_director",
                "target_name": "BOB LIM",
                "effective_date": "03/06/2026",
                "resignation_letter": "Yes",
                "remarks": "Director resignation letter requested; separate output package.",
            },
            {
                "generate": "Yes",
                "action_type": "appoint_director",
                "target_name": "TEST NOMINEE DIRECTOR A",
                "effective_date": "03/06/2026",
                "resignation_letter": "No",
                "remarks": "Appoint common nominee director from backend common people.",
            },
            {
                "generate": "Yes",
                "action_type": "resign_secretary",
                "target_name": "OLD SECRETARY ONG",
                "effective_date": "03/06/2026",
                "resignation_letter": "Yes",
                "remarks": "Secretary resignation letter requested; separate output package.",
            },
            {
                "generate": "Yes",
                "action_type": "appoint_secretary",
                "target_name": "TEST COMPANY SECRETARY A",
                "effective_date": "03/06/2026",
                "resignation_letter": "No",
                "remarks": "Appoint common secretary from backend common people.",
            },
        ],
    )
    set_table(
        main,
        {"generate", "field_label", "old_value", "new_value"},
        [
            {
                "generate": "Yes",
                "target_name": "ALICE TAN",
                "field_label": "ID type",
                "old_value": "Passport",
                "new_value": "NRIC",
                "effective_date": "03/06/2026",
                "remarks": "Personal particulars update: ID type.",
            },
            {
                "generate": "Yes",
                "target_name": "ALICE TAN",
                "field_label": "ID number",
                "old_value": "E12345678",
                "new_value": "S1111111B",
                "effective_date": "03/06/2026",
                "remarks": "Personal particulars update: IC / ID number.",
            },
            {
                "generate": "Yes",
                "target_name": "ALICE TAN",
                "field_label": "Residential address",
                "old_value": "1 RAFFLES PLACE, #20-61 ONE RAFFLES PLACE, SINGAPORE 048616",
                "new_value": "9 STRAITS VIEW, #12-07 MARINA ONE WEST TOWER, SINGAPORE 018937",
                "effective_date": "03/06/2026",
                "remarks": "Personal particulars update: local residential address.",
            },
            {
                "generate": "Yes",
                "target_name": "ALICE TAN",
                "field_label": "Email",
                "old_value": "alice.old@example.com",
                "new_value": "alice.updated@example.com",
                "effective_date": "03/06/2026",
                "remarks": "Personal particulars update: email.",
            },
            {
                "generate": "Yes",
                "target_name": "ALICE TAN",
                "field_label": "Phone",
                "old_value": "+65 8000 0001",
                "new_value": "+65 8123 4567",
                "effective_date": "03/06/2026",
                "remarks": "Personal particulars update: phone.",
            },
            {
                "generate": "Yes",
                "target_name": "ALICE TAN",
                "field_label": "Nationality",
                "old_value": "Chinese",
                "new_value": "Singaporean",
                "effective_date": "03/06/2026",
                "remarks": "Personal particulars update: nationality.",
            },
            {
                "generate": "Yes",
                "target_name": "BOB LIM",
                "field_label": "Name",
                "old_value": "BOB OLD NAME",
                "new_value": "BOB LIM",
                "effective_date": "03/06/2026",
                "remarks": "Personal particulars update: name.",
            },
            {
                "generate": "Yes",
                "target_name": "BOB LIM",
                "field_label": "Residential address",
                "old_value": "88 MARKET STREET, #10-01, SINGAPORE 048948",
                "new_value": "6 SHENTON WAY, #33-00 OUE DOWNTOWN, SINGAPORE 068809",
                "effective_date": "03/06/2026",
                "remarks": "Personal particulars update: address for second officer.",
            },
        ],
    )
    set_table(
        main,
        {"generate", "transferor_name", "transferee_name", "shares_transferred"},
        [
            {
                "generate": "Yes",
                "transferor_name": "ALICE TAN",
                "transferee_name": "DAVID TRANSFEREE",
                "shares_transferred": 120,
                "share_class": "Ordinary",
                "transfer_date": "03/06/2026",
                "consideration_basis": "internal_paid_up_basis",
                "remarks": "Share transfer preview and M01 approval test.",
            }
        ],
    )
    set_table(
        main,
        {"generate", "allottee_name", "shares_allotted"},
        [
            {
                "generate": "Yes",
                "allottee_name": "ERIN ALLOTTEE",
                "shares_allotted": 300,
                "share_class": "Ordinary",
                "amount_paid_per_share": 1,
                "total_paid": 300,
                "allotment_date": "03/06/2026",
                "remarks": "Allotment preview and M01 approval test.",
            }
        ],
    )

    set_vertical_values(
        annual,
        {
            "annual_review_required": "Yes",
            "fye_date": "31/12/2025",
            "agm_date": "30/06/2026",
            "agm_time": "10.00 a.m.",
            "agm_place": old_address,
            "agm_route": "ordinary_agm",
            "accounts_status": "active",
            "financial_statement_date": "31/12/2025",
            "financial_year_start": "01/01/2025",
            "director_signer_name": "ALICE TAN",
            "shareholder_signer_name": "ALICE TAN",
            "ar_authorized_signer_name": "ALICE TAN",
            "directors_fee": "0",
            "directors_remuneration": "0",
            "shorter_notice_consent": "Auto",
            "management_rep_letter": "Yes",
            "remarks": "Annual review stress test. Current website stage recognises this as annual package preview.",
        },
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    wb.save(WORKBOOK_PATH)
    return WORKBOOK_PATH


def extract_docx_text(docx_path: Path) -> str:
    doc = Document(docx_path)
    chunks: list[str] = []
    for paragraph in doc.paragraphs:
        chunks.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            chunks.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(chunks)


def audit(parsed: dict, common_rows: list[dict], pdf_zip: Path) -> dict:
    docx_dir = GENERATED_DIR / f"{JOB_CODE}_P2_M01_docs"
    pdf_dir = GENERATED_DIR / f"{JOB_CODE}_P2_M01_pdf"
    summary_path = docx_dir / "generation_summary.json"
    summary = json.loads(summary_path.read_text(encoding="utf-8"))
    docx_files = sorted(docx_dir.glob("*.docx"))
    pdf_files = sorted(pdf_dir.glob("*.pdf"))
    text = extract_docx_text(docx_files[0])
    preview = suggest_maintenance(parsed)
    with ZipFile(pdf_zip, "r") as zf:
        zip_entries = zf.namelist()
    expected = {
        "comma_separated_signer_1": "ALICE TAN",
        "comma_separated_signer_2": "BOB LIM",
        "common_nominee_name": "TEST NOMINEE DIRECTOR A",
        "common_secretary_name": "TEST COMPANY SECRETARY A",
        "resigning_director": "BOB LIM",
        "resigning_secretary": "OLD SECRETARY ONG",
        "personal_multi_officer_heading": "the following person(s)",
        "personal_alice_id_type": "ALICE TAN - ID type",
        "personal_alice_ic": "S1111111B",
        "personal_alice_address": "9 STRAITS VIEW",
        "personal_alice_email": "alice.updated@example.com",
        "personal_alice_phone": "+65 8123 4567",
        "personal_alice_nationality": "Singaporean",
        "personal_bob_name": "BOB LIM - Name",
        "personal_bob_address": "6 SHENTON WAY",
        "share_transfer": "DAVID TRANSFEREE",
        "share_allotment": "ERIN ALLOTTEE",
        "acra": "ACRA",
    }
    lower_text = text.lower()
    checks = {key: value.lower() in lower_text for key, value in expected.items()}
    removed_checks = {
        "business_existing_acra_column_removed": "Existing ACRA record" not in text,
        "nominee_ic_removed_from_m01": "S7654321Z" not in text,
        "secretary_ic_removed_from_m01": "S8765432Y" not in text,
        "nominee_address_removed_from_m01": "20 CECIL STREET" not in text,
        "secretary_address_removed_from_m01": "111 NORTH BRIDGE ROAD" not in text,
    }
    preview_files = preview.get("files", [])
    resignation_preview = [item for item in preview_files if "Resignation Letter" in item.get("name", "")]
    annual_preview = [item for item in preview_files if item.get("package") == "年审包"]
    return {
        "input_workbook": str(WORKBOOK_PATH),
        "common_people_written": common_rows,
        "job_code": JOB_CODE,
        "docx_files": [str(path) for path in docx_files],
        "pdf_files": [str(path) for path in pdf_files],
        "pdf_zip": str(pdf_zip),
        "pdf_zip_entries": zip_entries,
        "generation_summary": summary,
        "preview_summary": preview.get("summary", {}),
        "preview_files": preview_files,
        "resignation_letter_preview_count": len(resignation_preview),
        "annual_review_preview_count": len(annual_preview),
        "expected_text_checks": checks,
        "missing_expected_text": [key for key, ok in checks.items() if not ok],
        "removed_field_checks": removed_checks,
        "failed_removed_field_checks": [key for key, ok in removed_checks.items() if not ok],
        "unresolved_placeholder_hits": re.findall(r"\{\{|\}\}|\[\[|\]\]|\bNone\b|\bnull\b", text, flags=re.IGNORECASE),
        "m01_pdf_info": pdf_info(pdf_files[0]) if pdf_files else {},
        "important_limitations": [
            "M01 PDF is generated as a formal ordinary directors' resolution.",
            "Resignation letters are recognised in preview but are not generated by M01.",
            "Annual review files are recognised in preview but the annual package generator is not connected yet.",
        ],
    }


def build_delivery(paths: list[Path]) -> Path:
    if DELIVERY_PATH.exists():
        DELIVERY_PATH.unlink()
    with ZipFile(DELIVERY_PATH, "w") as zf:
        for path in paths:
            if path.exists():
                zf.write(path, arcname=path.name)
    return DELIVERY_PATH


def main() -> None:
    common_rows = upsert_common_people()
    workbook_path = build_stress_workbook()
    parsed = parse_excel(workbook_path, load_common_people())
    pdf_zip = generate_p2_m01_pdf_package(parsed, JOB_CODE)
    report = audit(parsed, common_rows, pdf_zip)
    AUDIT_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    docx_paths = [Path(path) for path in report["docx_files"]]
    pdf_paths = [Path(path) for path in report["pdf_files"]]
    build_delivery([WORKBOOK_PATH, AUDIT_PATH, pdf_zip, *docx_paths, *pdf_paths])
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
